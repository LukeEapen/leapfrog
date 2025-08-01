from flask import Flask, render_template, request, jsonify, session
import openai
from openai import OpenAI
import os
import time
import logging
import tiktoken
import asyncio
import aiohttp
import threading
import string
import traceback
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
import hashlib
import json
from cachetools import TTLCache
import tenacity
import csv
import io

# Vector Database and RAG imports
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Document processing imports
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
    logger.info("python-docx library available for DOCX file processing")
except ImportError:
    logger.warning("python-docx not available. DOCX files will be processed as text extraction fallback.")
    DOCX_AVAILABLE = False

# Initialize Vector Database and RAG components
logger.info("Initializing Vector Database and RAG components...")

# Initialize ChromaDB client
chroma_client = chromadb.PersistentClient(path="./vector_db")

# Initialize embedding model for RAG
try:
    embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    logger.info("Embedding model loaded successfully")
except Exception as e:
    logger.warning(f"Failed to load embedding model: {e}. RAG features may be limited.")
    embedding_model = None

# Create or get collection for PRD documents
try:
    prd_collection = chroma_client.get_or_create_collection(
        name="prd_documents",
        metadata={"description": "PRD and documentation storage for RAG"}
    )
    logger.info("ChromaDB collection initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize ChromaDB collection: {e}")
    prd_collection = None

app = Flask(__name__)

# Configure Flask app
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'dev-secret-key-change-in-production')
# In production, make sure to set FLASK_SECRET_KEY environment variable

# Performance optimizations
THREAD_POOL_SIZE = 12  # Increased for more parallelism (tune as needed for your hardware)
MAX_CACHE_SIZE = 1000
CACHE_TTL = 3600  # 1 hour

# Initialize thread pool for parallel processing
executor = ThreadPoolExecutor(max_workers=THREAD_POOL_SIZE)

# Utility for parallel agent calls
def run_agents_in_parallel(agent_tasks):
    """
    agent_tasks: dict of {key: (func, args)}
    Returns: dict of {key: result}
    """
    import time
    results = {}
    start = time.time()
    with ThreadPoolExecutor(max_workers=THREAD_POOL_SIZE) as executor:
        future_to_key = {executor.submit(func, *args): key for key, (func, args) in agent_tasks.items()}
        for future in as_completed(future_to_key):
            key = future_to_key[future]
            try:
                results[key] = future.result()
            except Exception as exc:
                results[key] = f"Error: {exc}"
    logger.info(f"[PERF] Parallel agent batch ({list(agent_tasks.keys())}) took {time.time() - start:.2f}s")
    return results

# Cache for responses (in-memory with TTL)
response_cache = TTLCache(maxsize=MAX_CACHE_SIZE, ttl=CACHE_TTL)

logger.info("Flask application starting up with performance optimizations...")
logger.info(f"OpenAI API key configured: {'Yes' if os.getenv('OPENAI_API_KEY') else 'No'}")
logger.info(f"Thread pool size: {THREAD_POOL_SIZE}")
logger.info(f"Cache size: {MAX_CACHE_SIZE} items, TTL: {CACHE_TTL}s")

@app.route("/prd_parser", methods=["GET"])
def user_story_input():
    logger.info("GET request received for user story input page")
    try:
        logger.info("Rendering poc2_user_story_input.html template")
        return render_template("poc2_user_story_input.html")
    except Exception as e:
        logger.error(f"Error rendering user story input template: {str(e)}")
        raise

@app.route("/", methods=["GET"])
def home():
    """Home page with navigation to all features."""
    logger.info("GET request received for home page")
    
    html_content = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>RAG-Enhanced PRD to Epic Generator</title>
        <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css"/>
        <style>
            body { 
                font-family: Arial, sans-serif; 
                margin: 20px; 
                background-color: #2f323a; 
                color: #fff;
            }
            .container { 
                max-width: 1200px; 
                margin: 0 auto; 
                background-color: #2f323a; 
                padding: 30px; 
                border-radius: 10px; 
            }
            .header-bar {
                background-color: #d0021b;
                color: white;
                text-align: center;
                font-weight: bold;
                padding: 15px;
                border-radius: 6px;
                margin-bottom: 20px;
            }
            .header { 
                text-align: center; 
                margin-bottom: 30px; 
                color: #fff;
            }
            .header h1 {
                color: #fff;
                margin-bottom: 15px;
            }
            .header p {
                color: #ccc;
                font-size: 1.1rem;
            }
            .feature-grid { 
                display: grid; 
                grid-template-columns: 1fr 1fr 1fr; 
                gap: 20px; 
                margin-top: 30px; 
            }
            @media (max-width: 1024px) {
                .feature-grid {
                    grid-template-columns: 1fr 1fr;
                }
            }
            @media (max-width: 768px) {
                .feature-grid {
                    grid-template-columns: 1fr;
                }
            }
            .feature-card { 
                background-color: #f8f9fa; 
                color: #000;
                padding: 20px; 
                border-radius: 8px; 
                border: 1px solid #dee2e6;
                transition: transform 0.2s ease;
            }
            .feature-card:hover {
                transform: translateY(-2px);
                box-shadow: 0 4px 8px rgba(0,0,0,0.2);
            }
            .feature-card h3 { 
                color: #d0021b; 
                margin-top: 0; 
                font-weight: bold;
            }
            .feature-card p, .feature-card ul {
                color: #333;
            }
            .btn { 
                display: inline-block; 
                padding: 12px 24px; 
                background-color: #d0021b; 
                color: white; 
                text-decoration: none; 
                border-radius: 5px; 
                margin: 10px 0; 
                font-weight: bold;
                transition: background-color 0.2s ease;
            }
            .btn:hover { 
                background-color: #b0011a; 
                color: white;
                text-decoration: none;
            }
            .btn-secondary { 
                background-color: #444; 
            }
            .btn-secondary:hover { 
                background-color: #333; 
            }
            .status { 
                background-color: #444; 
                border: 1px solid #555; 
                color: #fff; 
                padding: 15px; 
                border-radius: 8px; 
                margin-bottom: 20px; 
                text-align: center;
                font-weight: bold;
            }
            .footer-text {
                margin-top: 30px; 
                text-align: center; 
                color: #ccc;
                font-style: italic;
            }
            .btn-danger { 
                background-color: #dc3545; 
            }
            .btn-danger:hover { 
                background-color: #c82333; 
            }
            .btn-danger:disabled { 
                background-color: #6c757d; 
                cursor: not-allowed;
            }
            .clear-status {
                margin-top: 10px;
                padding: 10px;
                border-radius: 4px;
                display: none;
            }
            .clear-status.success {
                background-color: #d4edda;
                color: #155724;
                border: 1px solid #c3e6cb;
            }
            .clear-status.error {
                background-color: #f8d7da;
                color: #721c24;
                border: 1px solid #f5c6cb;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header-bar">
                RAG-Enhanced PRD to Epic Generator
            </div>
            
            <div class="header">
                <h1>Transform PRDs into Comprehensive Epics & User Stories</h1>
                <p>Advanced RAG technology for intelligent document processing and epic generation</p>
                <div class="status">
                    Vector Database Active | RAG Processing Ready | Epic Generator Online
                </div>
            </div>
            
            <div class="feature-grid">
                <div class="feature-card">
                    <h3>Generate Epics & User Stories</h3>
                    <p>Upload your PRD and additional documentation to generate comprehensive epics and user stories using RAG-enhanced processing.</p>
                    <a href="/prd_parser" class="btn">Start Processing</a>
                </div>
                
                <div class="feature-card">
                    <h3>Vector Database Status</h3>
                    <p>View all documents stored in the vector database, including chunks, metadata, and content previews.</p>
                    <a href="/vector-db-status" class="btn btn-secondary">View Database</a>
                </div>
                
                <div class="feature-card">
                    <h3>Search Vector Database</h3>
                    <p>Search through stored documents using semantic similarity to find relevant content for any query.</p>
                    <a href="/vector-db-search" class="btn btn-secondary">Search Database</a>
                </div>
                
                <div class="feature-card">
                    <h3>Clear Vector Database</h3>
                    <p>Remove all stored documents from the vector database. This will clear all PRD documents, summaries, and chunks. <strong>This action cannot be undone.</strong></p>
                    <button onclick="clearVectorDatabase()" class="btn btn-danger" id="clearDbBtn">
                        🗑️ Clear All Documents
                    </button>
                    <div id="clearStatus" class="clear-status"></div>
                </div>
                
                <div class="feature-card">
                    <h3>RAG Features</h3>
                    <p>This system uses Retrieval Augmented Generation with:</p>
                    <ul>
                        <li>ChromaDB vector storage</li>
                        <li>SentenceTransformers embeddings</li>
                        <li>Intelligent document chunking</li>
                        <li>Semantic content retrieval</li>
                        <li>Single-agent optimization</li>
                    </ul>
                </div>
                
                <div class="feature-card">
                    <h3>🗺️ System Mapping</h3>
                    <p>Create and manage system mappings with an intuitive drag-and-drop interface. Map systems for technical architecture and deployment planning.</p>
                    <ul>
                        <li>Drag & drop system selection</li>
                        <li>Custom mapping creation</li>
                        <li>CSV export functionality</li>
                        <li>AI-powered mapping assistance</li>
                    </ul>
                    <a href="/system-mapping" class="btn">🏗️ Create System Map</a>
                </div>
            </div>
            
            <div class="footer-text">
                <p>PRD Parser Agent replaced with RAG summaries | 50%+ faster processing | Reduced token costs</p>
            </div>
        </div>
        
        <script>
            async function clearVectorDatabase() {
                const clearBtn = document.getElementById('clearDbBtn');
                const statusDiv = document.getElementById('clearStatus');
                
                // Confirm action
                if (!confirm('Are you sure you want to clear ALL documents from the vector database? This action cannot be undone.')) {
                    return;
                }
                
                // Show loading state
                clearBtn.disabled = true;
                clearBtn.innerHTML = '🔄 Clearing Database...';
                statusDiv.style.display = 'none';
                
                try {
                    const response = await fetch('/vector-db-clear', {
                        method: 'POST',
                        headers: {
                            'Content-Type': 'application/json'
                        }
                    });
                    
                    const result = await response.json();
                    
                    if (result.status === 'success') {
                        statusDiv.innerHTML = `✅ ${result.message}`;
                        statusDiv.className = 'clear-status success';
                        statusDiv.style.display = 'block';
                    } else {
                        statusDiv.innerHTML = `❌ Error: ${result.message}`;
                        statusDiv.className = 'clear-status error';
                        statusDiv.style.display = 'block';
                    }
                    
                } catch (error) {
                    console.error('Error clearing vector database:', error);
                    statusDiv.innerHTML = `❌ Network error: ${error.message}`;
                    statusDiv.className = 'clear-status error';
                    statusDiv.style.display = 'block';
                } finally {
                    // Reset button state
                    clearBtn.disabled = false;
                    clearBtn.innerHTML = '🗑️ Clear All Documents';
                    
                    // Hide status after 5 seconds
                    setTimeout(() => {
                        statusDiv.style.display = 'none';
                    }, 5000);
                }
            }
        </script>
    </body>
    </html>
    """
    
    return html_content, 200, {'Content-Type': 'text/html'}

@app.route("/user-story-upload", methods=["POST"])
def process_user_story():
    start_time = time.time()
    logger.info("[PERF] process_user_story: started")
    logger.info("POST request received for user story processing")
    
    # Collect input
    logger.info("Collecting form data and file uploads")
    context = request.form.get("context", "")
    prd_file = request.files.get("prd_file")
    additional_docs = request.files.get("additional_docs")
    
    logger.info(f"Context provided: {'Yes' if context else 'No'}")
    logger.info(f"PRD file uploaded: {'Yes' if prd_file else 'No'}")
    logger.info(f"Additional docs uploaded: {'Yes' if additional_docs else 'No'}")
    
    if prd_file:
        logger.info(f"PRD file name: {prd_file.filename}")
    if additional_docs:
        logger.info(f"Additional docs file name: {additional_docs.filename}")

    # Parallel file reading for better performance
    logger.info("Reading uploaded files in parallel")
    
    def read_files_parallel():
        with ThreadPoolExecutor(max_workers=2) as file_executor:
            prd_future = file_executor.submit(safe_read, prd_file) if prd_file else None
            docs_future = file_executor.submit(safe_read, additional_docs) if additional_docs else None
            
            prd_content = prd_future.result() if prd_future else ""
            docs_content = docs_future.result() if docs_future else ""
            
            return prd_content, docs_content
    
    t0 = time.time()
    prd_content, docs_content = read_files_parallel()
    logger.info(f"[PERF] File read parallel: {time.time() - t0:.2f}s")
    
    # Debug file content
    if prd_content:
        debug_file_content(prd_content, prd_file.filename if prd_file else "prd_content")
    if docs_content:
        debug_file_content(docs_content, additional_docs.filename if additional_docs else "additional_docs")    # Store documents in vector database and create RAG summaries
    logger.info("Processing documents with RAG enhancement")
    
    # Process PRD with RAG if content is substantial and valid
    t1 = time.time()
    if is_valid_content(prd_content):
        if len(prd_content) > 5000:
            logger.info("Creating RAG-enhanced PRD summary")
            prd_content = create_rag_summary(prd_content, prd_file.filename if prd_file else "prd_content", max_summary_length=25000)
        elif len(prd_content) > 30000:
            logger.info("Large PRD detected - using traditional optimization")
            prd_content = optimize_prd_content(prd_content, max_length=40000)
        else:
            logger.info("PRD content is valid but small, using as-is")
    else:
        logger.warning(f"PRD content is invalid for RAG processing: {prd_content[:100] if prd_content else 'None'}")
        prd_content = "No valid PRD content available - please check file format and encoding."
    logger.info(f"[PERF] PRD content processing: {time.time() - t1:.2f}s")
    
    # Process additional docs with RAG
    t2 = time.time()
    if is_valid_content(docs_content):
        if len(docs_content) > 3000:
            logger.info("Creating RAG-enhanced docs summary")
            docs_content = create_rag_summary(docs_content, additional_docs.filename if additional_docs else "additional_docs", max_summary_length=10000)
        elif len(docs_content) > 15000:
            logger.info("Large additional docs detected - truncating for faster processing")
            docs_content = docs_content[:15000] + "...\n[Content truncated for performance]"
        else:
            logger.info("Additional docs content is valid but small, using as-is")
    else:
        logger.warning(f"Additional docs content is invalid: {docs_content[:100] if docs_content else 'None'}")
        docs_content = "No valid additional documentation available."
    logger.info(f"[PERF] Docs content processing: {time.time() - t2:.2f}s")
    
    logger.info(f"PRD content length (post-RAG): {len(prd_content)} characters")
    logger.info(f"Additional docs content length (post-RAG): {len(docs_content)} characters")
    logger.info("Combining context and RAG-enhanced content")
    t3 = time.time()
    
    # Create enhanced context for Epic Generator
    enhanced_context = f"""
{context}

RAG-Enhanced PRD Analysis:
{prd_content}

Additional Context:
{docs_content}

Instructions: The above content has been intelligently extracted and summarized using RAG (Retrieval Augmented Generation). 
It contains the most relevant requirements, user stories, and business objectives from the original documents.
Use this curated information to generate comprehensive epics and user stories.
"""
    logger.info(f"[PERF] Context assembly: {time.time() - t3:.2f}s")
    logger.info(f"Enhanced context length: {len(enhanced_context)} characters")
    
    # Check cache first with more granular caching
    t4 = time.time()
    prompt_hash = get_cache_key(enhanced_context)
    cached_result = response_cache.get(prompt_hash)
    if cached_result:
        logger.info("Cache hit! Returning cached response")
        processing_time = time.time() - start_time
        logger.info(f"[PERF] Total processing time (cached): {processing_time:.2f} seconds")
        return render_template("poc2_epic_story_screen.html", epics=cached_result)
    logger.info(f"[PERF] Cache check: {time.time() - t4:.2f}s")
      # Log token count for the enhanced context
    t5 = time.time()
    context_tokens = count_tokens(enhanced_context, "gpt-4o")
    logger.info(f"Enhanced context token count: {context_tokens:,} tokens")
    logger.info(f"[PERF] Token count: {time.time() - t5:.2f}s")
    
    # Check if context is within token limits
    if context_tokens > 120000:
        logger.warning(f"Enhanced context token count ({context_tokens:,}) is approaching the 128k limit!")
        # Further optimize if needed
        enhanced_context = enhanced_context[:100000] + "\n[Content truncated due to token limits]"
        context_tokens = count_tokens(enhanced_context, "gpt-4o")
        logger.info(f"Truncated context token count: {context_tokens:,} tokens")
    else:
        logger.info(f"Enhanced context is within safe token limits ({context_tokens:,}/128,000 tokens)")    # Skip PRD Parser Agent - RAG has already done the parsing and summarization
    logger.info("****************Skipping PRD Parser - Using RAG Summary Directly")
    logger.info("Starting Epic Generator with RAG-enhanced content")
    
    # Print Enhanced Context being sent to Epic Agent
    logger.info("=" * 80)
    logger.info("ENHANCED CONTEXT INPUT TO EPIC AGENT:")
    logger.info("=" * 80)
    logger.info(enhanced_context[:3000] + "..." if len(enhanced_context) > 3000 else enhanced_context)
    logger.info("=" * 80)
    
    try:
        # Directly use Epic Generator with RAG-enhanced content
        t6 = time.time()
        start_epic_time = time.time()
        epic_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", enhanced_context)
        epic_processing_time = time.time() - start_epic_time
        logger.info(f"[PERF] Epic agent call: {epic_processing_time:.2f}s")
        logger.info("################Epic Generator response received")
        # Print Epic Agent Response
        logger.info("=" * 80)
        logger.info("EPIC AGENT RESPONSE OUTPUT:")
        logger.info("=" * 80)
        logger.info(epic_response)
        logger.info("=" * 80)
        # Log token usage for Epic Generator interaction
        log_token_usage(enhanced_context, epic_response, model="gpt-4o", context="RAG-Enhanced Epic Generator")
        final_output = epic_response
        logger.info(f"Final output length: {len(final_output)} characters")
        # Cache the result for future requests
        response_cache[prompt_hash] = final_output
        logger.info("Response cached for future requests")
        processing_time = time.time() - start_time
        logger.info(f"[PERF] Total RAG-optimized processing time: {processing_time:.2f} seconds")
        # Calculate time savings
        estimated_traditional_time = processing_time * 2  # Estimate of traditional two-agent approach
        time_saved = estimated_traditional_time - processing_time
        logger.info(f"Estimated time saved by RAG optimization: {time_saved:.2f} seconds")
        logger.info("Rendering epic story screen with RAG-optimized output")
        return render_template("poc2_epic_story_screen.html", epics=final_output)
    except Exception as e:
        logger.error(f"Error in RAG-optimized processing: {str(e)}")
        raise

def process_large_prompt_chunked(prompt, start_time):
    """Process large prompts using RAG-enhanced chunking for better performance."""
    logger.info("Processing large prompt using RAG-enhanced chunked approach")
    
    # Check cache first
    prompt_hash = get_cache_key(prompt)
    cached_result = response_cache.get(prompt_hash)
    if cached_result:
        logger.info("Cache hit! Returning cached response")
        processing_time = time.time() - start_time
        logger.info(f"Total processing time (cached): {processing_time:.2f} seconds")
        return render_template("poc2_epic_story_screen.html", epics=cached_result)
    
    # Split prompt into sections and extract content
    sections = prompt.split('\n\n')
    context = sections[0] if sections else ""
    prd_section = ""
    docs_section = ""
    
    # Extract PRD and Docs sections
    for i, section in enumerate(sections):
        if 'PRD:' in section or 'RAG-Enhanced PRD Analysis:' in section:
            prd_section = '\n\n'.join(sections[i:i+15])  # Take next 15 sections for PRD
            break
    
    for i, section in enumerate(sections):
        if 'Additional Docs:' in section or 'Additional Context:' in section:
            docs_section = '\n\n'.join(sections[i:])  # Take remaining for docs
            break
    
    # Create RAG-optimized prompt for direct Epic Generator use
    rag_optimized_prompt = f"""
{context}

RAG-Enhanced PRD Analysis (Chunked Processing):
{prd_section[:35000]}

Additional Context:
{docs_section[:12000]}

Instructions: The above content has been intelligently processed and chunked for optimal performance. 
It contains the most relevant requirements, user stories, and business objectives from the original documents.
Generate comprehensive epics and user stories directly from this curated information.
"""
    
    logger.info(f"RAG-optimized chunked prompt length: {len(rag_optimized_prompt)} characters (reduced from {len(prompt)})")      # Log token count for the optimized prompt
    optimized_tokens = count_tokens(rag_optimized_prompt, "gpt-4o")
    logger.info(f"RAG-optimized chunked prompt token count: {optimized_tokens:,} tokens")
    
    try:
        # Skip PRD Parser - Use Epic Generator directly with RAG-enhanced content
        logger.info("****************Skipping PRD Parser in chunked processing - Using RAG-enhanced content directly")
        
        start_epic_time = time.time()
        epic_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", rag_optimized_prompt)
        epic_processing_time = time.time() - start_epic_time
        
        logger.info("Epic Generator (chunked) response received")
        logger.info(f"Epic Generator processing time: {epic_processing_time:.2f} seconds")
          # Log token usage for Epic Generator interaction
        log_token_usage(rag_optimized_prompt, epic_response, model="gpt-4o", context="RAG-Enhanced Epic Generator (Chunked)")
        
        final_output = epic_response
        
        # Cache the result
        response_cache[prompt_hash] = final_output
        logger.info("Chunked response cached for future requests")
        
        processing_time = time.time() - start_time
        logger.info(f"Total RAG-enhanced chunked processing time: {processing_time:.2f} seconds")
        
        # Calculate time savings compared to traditional two-agent approach
        estimated_traditional_time = processing_time * 2.5  # Estimate of traditional approach with chunking
        time_saved = estimated_traditional_time - processing_time
        logger.info(f"Estimated time saved by RAG chunked optimization: {time_saved:.2f} seconds")
        
        return render_template("poc2_epic_story_screen.html", epics=final_output)
        
    except Exception as e:
        logger.error(f"Error in RAG-enhanced chunked processing: {str(e)}")
        raise

def get_cache_key(prompt):
    """Generate a cache key from the prompt."""
    return hashlib.md5(prompt.encode('utf-8')).hexdigest()

def store_document_in_vector_db(content, filename, doc_type="prd"):
    """Store document content in vector database for RAG retrieval."""
    if not prd_collection or not embedding_model:
        logger.warning("Vector DB or embedding model not available. Skipping storage.")
        return None
    
    try:
        logger.info(f"Storing {doc_type} document in vector database: {filename}")
        
        # Split content into chunks for better retrieval
        chunks = split_document_into_chunks(content, chunk_size=1000, overlap=200)
        logger.info(f"Document split into {len(chunks)} chunks")
        
        # Generate embeddings for each chunk
        chunk_embeddings = embedding_model.encode(chunks)
        
        # Create metadata for each chunk
        documents = []
        metadatas = []
        ids = []
        
        doc_id = hashlib.md5(f"{filename}_{doc_type}".encode()).hexdigest()
        
        for i, (chunk, embedding) in enumerate(zip(chunks, chunk_embeddings)):
            chunk_id = f"{doc_id}_chunk_{i}"
            
            documents.append(chunk)
            metadatas.append({
                "filename": filename,
                "doc_type": doc_type,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "timestamp": datetime.now().isoformat()
            })
            ids.append(chunk_id)
        
        # Store in ChromaDB
        prd_collection.add(
            documents=documents,
            metadatas=metadatas,
            ids=ids
        )
        
        logger.info(f"Successfully stored {len(chunks)} chunks in vector database")
        return doc_id
        
    except Exception as e:
        logger.error(f"Error storing document in vector DB: {str(e)}")
        return None

def split_document_into_chunks(content, chunk_size=1000, overlap=200):
    """Split document content into overlapping chunks for better retrieval."""
    if len(content) <= chunk_size:
        return [content]
    
    chunks = []
    start = 0
    
    while start < len(content):
        end = start + chunk_size
        
        # Try to break at sentence boundaries
        if end < len(content):
            # Look for sentence endings within overlap range
            for punct in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                last_punct = content.rfind(punct, start, end)
                if last_punct > start:
                    end = last_punct + len(punct)
                    break
        
        chunk = content[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
        
        # Avoid infinite loop
        if start >= end:
            start = end
    
    return chunks

def retrieve_relevant_content(query, top_k=5, doc_type=None):
    """Retrieve relevant content from vector database using RAG."""
    if not prd_collection or not embedding_model:
        logger.warning("Vector DB or embedding model not available. Skipping RAG retrieval.")
        return []
    
    try:
        logger.info(f"Retrieving relevant content for query (top {top_k} results)")
        
        # Create query filter if doc_type is specified
        where_filter = {"doc_type": doc_type} if doc_type else None
        
        # Query the collection
        results = prd_collection.query(
            query_texts=[query],
            n_results=top_k,
            where=where_filter
        )
        
        relevant_chunks = []
        if results['documents'] and results['documents'][0]:
            for i, (doc, metadata) in enumerate(zip(results['documents'][0], results['metadatas'][0])):
                relevant_chunks.append({
                    'content': doc,
                    'metadata': metadata,
                    'relevance_rank': i + 1
                })
                
        logger.info(f"Retrieved {len(relevant_chunks)} relevant chunks")
        return relevant_chunks
        
    except Exception as e:
        logger.error(f"Error retrieving from vector DB: {str(e)}")
        return []

def create_rag_summary(content, filename, max_summary_length=5000):
    """Create a RAG-enhanced summary of the document content with size optimization."""
    logger.info(f"Creating RAG-enhanced summary for {filename}")
    
    # For very large files, skip expensive RAG processing and use intelligent fallback
    if len(content) > 50000:  # 50KB threshold
        logger.info(f"File {filename} is large ({len(content):,} chars), using intelligent fallback instead of RAG")
        return create_intelligent_summary_fallback(content, filename, max_summary_length)
    
    # Store the document in vector DB
    doc_id = store_document_in_vector_db(content, filename)
    
    if not doc_id:
        # Fallback to intelligent content extraction if vector DB fails
        logger.warning("Vector DB storage failed, using intelligent content extraction")
        return create_intelligent_summary_fallback(content, filename, max_summary_length)
    
    # Create comprehensive summary queries to extract key information
    summary_queries = [
        "requirements and functional specifications",
        "user stories and acceptance criteria", 
        "business objectives and goals",
        "system constraints and dependencies",
        "key features and capabilities",
        "technical architecture and design",
        "data flow and integration requirements",
        "security and compliance requirements",
        "performance and scalability requirements",
        "user interface and user experience requirements"
    ]
    
    # Retrieve relevant content for each query
    summary_parts = []
    for query in summary_queries:
        relevant_chunks = retrieve_relevant_content(query, top_k=2, doc_type="prd")
        
        if relevant_chunks:
            # Take the most relevant chunk for each query
            best_chunk = relevant_chunks[0]['content']
            summary_parts.append(f"[{query.title()}]\n{best_chunk}\n")
    
    # Combine summary parts with intelligent ordering
    ordered_summary = order_summary_parts(summary_parts)
    final_summary = "\n\n".join(ordered_summary)
    
    # Truncate to max length
    if len(final_summary) > max_summary_length:
        logger.info(f"Final summary length ({len(final_summary)}) exceeds max length ({max_summary_length}), truncating")
        final_summary = final_summary[:max_summary_length] + "... [Content truncated]"
    
    logger.info(f"RAG summary created, length: {len(final_summary)} characters")
    
    # Print RAG Summary for debugging
    logger.info("=" * 80)
    logger.info("RAG SUMMARY GENERATED:")
    logger.info("=" * 80)
    logger.info(final_summary[:2000] + "..." if len(final_summary) > 2000 else final_summary)
    logger.info("=" * 80)
    
    return final_summary.strip()

def order_summary_parts(summary_parts):
    """Order summary parts based on importance and relevance."""
    logger.info("Ordering summary parts by importance")
    
    # Simple heuristic: prioritize sections with more details and specific requirements
    def section_score(section):
        score = 0
        # More weight to user stories and requirements
        if "user story" in section.lower():
            score += 10
        if "requirement" in section.lower():
            score += 8
        # Add score for each bullet point or numbered item
        score += section.count('•') * 2
        score += section.count('-') * 2
        score += section.count('1.') * 2
        score += section.count('2.') * 2
        return score
    
    # Sort sections by computed score
    ordered_sections = sorted(summary_parts, key=section_score, reverse=True)
    logger.info(f"Ordered {len(ordered_sections)} summary parts by importance")
    return ordered_sections

def extract_docx_text(file):
    """Extract text content from a DOCX file."""
    try:
        if not DOCX_AVAILABLE:
            logger.warning(f"Cannot process DOCX file {file.filename}: python-docx not available")
            return "[DOCX file detected but python-docx library not available. Please install python-docx to process Word documents.]"
        
        logger.info(f"Extracting text from DOCX file: {file.filename}")
        
        # Create a temporary BytesIO object from the file
        from io import BytesIO
        file_content = file.read()
        file.seek(0)  # Reset file pointer in case it's needed again
        
        # Load the document using python-docx
        doc = DocxDocument(BytesIO(file_content))
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        final_text = "\n\n".join(text_content)
        
        if final_text.strip():
            logger.info(f"Successfully extracted {len(final_text)} characters from DOCX file {file.filename}")
            return final_text
        else:
            logger.warning(f"No text content found in DOCX file {file.filename}")
            return "[DOCX file processed but no readable text content found]"
            
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {file.filename}: {str(e)}")
        return f"[Error processing DOCX file: {str(e)}]"

def safe_read(file):
    """Safely read file content with proper error handling and encoding detection."""
    try:
        logger.debug(f"Attempting to read file: {file.filename if file else 'None'}")
        
        # Check if this is a DOCX file and handle it specially
        if file and file.filename and file.filename.lower().endswith('.docx'):
            logger.info(f"Detected DOCX file: {file.filename}")
            return extract_docx_text(file)
        
        # Read the raw bytes first
        raw_content = file.read()
        logger.debug(f"Read {len(raw_content)} bytes from file: {file.filename}")
        
        # Check if this looks like a DOCX file even if extension is wrong
        if raw_content.startswith(b'PK\x03\x04') and b'[Content_Types].xml' in raw_content:
            logger.info(f"Detected DOCX format in file: {file.filename}")
            file.seek(0)  # Reset file pointer for DOCX processing
            return extract_docx_text(file)
        
        # Try multiple encodings for regular text files
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                content = raw_content.decode(encoding)
                logger.info(f"Successfully decoded file {file.filename} using {encoding} encoding, length: {len(content)} characters")
                
                # Validate that we have meaningful content
                if len(content.strip()) > 0:
                    return content
                else:
                    logger.warning(f"File {file.filename} appears to be empty or whitespace only")
                    return "[File appears to be empty]"
                    
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try to extract text from binary content
        logger.warning(f"All standard encodings failed for {file.filename}, attempting binary extraction")
        try:
            # Try to extract printable ASCII characters
            import string
            printable_chars = set(string.printable)
            extracted_text = ''.join(char for char in raw_content.decode('latin-1') if char in printable_chars)
            
            if len(extracted_text.strip()) > 50:  # If we got some meaningful text
                logger.info(f"Extracted {len(extracted_text)} printable characters from {file.filename}")
                return extracted_text
            else:
                logger.error(f"Unable to extract meaningful text from {file.filename}")
                return f"[Unable to decode file - tried encodings: {', '.join(encodings)}]"
                
        except Exception as extract_error:
            logger.error(f"Binary extraction failed for {file.filename}: {str(extract_error)}")
            return f"[File reading failed - may be binary or corrupted]"
            
    except Exception as e:
        logger.error(f"Error reading file {file.filename}: {str(e)}")
        return f"[Error reading file: {str(e)}]"

def optimize_prd_content(prd_content, max_length=40000):
    """Optimize PRD content by extracting key sections and reducing verbosity."""
    if len(prd_content) <= max_length:
        return prd_content
    
    logger.info(f"Optimizing PRD content from {len(prd_content)} to ~{max_length} characters")
    
    # Split into sections and prioritize important ones
    sections = prd_content.split('\n\n')
    important_keywords = [
        'requirement', 'feature', 'user story', 'acceptance criteria', 
        'functional', 'non-functional', 'business rule', 'constraint',
        'objective', 'goal', 'scope', 'assumption', 'dependency'
    ]
    
    # Score sections based on importance
    scored_sections = []
    for section in sections:
        score = 0
        section_lower = section.lower()
        
        # Score based on keywords
        for keyword in important_keywords:
            score += section_lower.count(keyword) * 10
        
        # Prefer sections with structured content
        score += section.count('•') * 5
        score += section.count('-') * 3
        score += section.count('1.') * 5
        
        scored_sections.append((score, section))
    
    # Sort by score and take top sections
    scored_sections.sort(key=lambda x: x[0], reverse=True)
    
    optimized_content = ""
    for score, section in scored_sections:
        if len(optimized_content) + len(section) <= max_length:
            optimized_content += section + "\n\n"
        else:
            break
    
    logger.info(f"Optimized PRD content to {len(optimized_content)} characters")
    return optimized_content.strip()

@lru_cache(maxsize=100)
def count_tokens(text, model="gpt-4o"):
    """Count tokens in text using tiktoken. Cached for performance."""
    try:
        # Map gpt-4o and gpt-4o to gpt-4 for tiktoken compatibility
        tiktoken_model = "gpt-4" if model in ["gpt-4-turbo", "gpt-4o"] else model
        encoding = tiktoken.encoding_for_model(tiktoken_model)
        return len(encoding.encode(text))
    except Exception as e:
        logger.warning(f"Error counting tokens: {e}, using fallback estimate")
        # Fallback: rough estimate of 4 characters per token
        return len(text) // 4

def log_token_usage(prompt_text, response_text, model="gpt-4o", context=""):
    """Log detailed token usage and cost estimation."""
    prompt_tokens = count_tokens(prompt_text, model)
    response_tokens = count_tokens(response_text, model)
    total_tokens = prompt_tokens + response_tokens
    
    # Cost estimation for GPT-4 Turbo (approximate rates)
    input_cost_per_1k = 0.01   # $0.01 per 1K input tokens
    output_cost_per_1k = 0.03  # $0.03 per 1K output tokens
    
    input_cost = (prompt_tokens / 1000) * input_cost_per_1k
    output_cost = (response_tokens / 1000) * output_cost_per_1k
    total_cost = input_cost + output_cost
    
    logger.info(f"Token Usage - {context}")
    logger.info(f"  Input tokens: {prompt_tokens:,}")
    logger.info(f"  Output tokens: {response_tokens:,}")
    logger.info(f"  Total tokens: {total_tokens:,}")
    logger.info(f"  Estimated cost: ${total_cost:.4f}")

def process_csv_to_vector_db(csv_file):
    """Process CSV file and store it in the vector database."""
    try:
        logger.info(f"Processing CSV file: {csv_file.filename}")
        
        # Read CSV content with proper encoding handling
        csv_content = csv_file.read().decode('utf-8-sig')  # utf-8-sig handles BOM
        csv_file.seek(0)  # Reset file pointer
        
        # Clean up the content - remove non-breaking spaces and other problematic characters
        csv_content = csv_content.replace('\xa0', ' ').replace('\ufeff', '')
        
        # Parse CSV
        csv_reader = csv.DictReader(io.StringIO(csv_content))
        rows = list(csv_reader)
        
        if not rows:
            logger.warning("CSV file is empty or has no data rows")
            return "CSV file is empty"
        
        # Clean column names
        cleaned_rows = []
        for row in rows:
            cleaned_row = {}
            for key, value in row.items():
                # Clean column names and values
                clean_key = key.strip().replace('\xa0', ' ').replace('\ufeff', '') if key else 'Unknown'
                clean_value = value.strip().replace('\xa0', ' ').replace('\ufeff', '') if value else ''
                cleaned_row[clean_key] = clean_value
            cleaned_rows.append(cleaned_row)
        
        rows = cleaned_rows
        logger.info(f"CSV contains {len(rows)} rows with columns: {list(rows[0].keys())}")
        
        # Convert CSV rows to text documents for vector storage
        documents = []
        metadatas = []
        ids = []
        
        for i, row in enumerate(rows):
            # Create a text representation of the row
            row_text = f"CSV Record {i+1}:\n"
            for key, value in row.items():
                if value:  # Only include non-empty values
                    row_text += f"{key}: {value}\n"
            
            documents.append(row_text.strip())
            
            # Ensure all metadata values are valid types for ChromaDB
            clean_filename = str(csv_file.filename).strip() if csv_file.filename else "unknown.csv"
            column_names = ", ".join(str(key).strip() for key in row.keys() if key)
            
            metadatas.append({
                "source": clean_filename,
                "type": "csv_row",
                "row_number": int(i + 1),
                "columns": column_names
            })
            ids.append(f"csv_{clean_filename.replace('.', '_')}_{i+1}")
        
        # Store in vector database
        if prd_collection and documents:
            try:
                prd_collection.add(
                    documents=documents,
                    metadatas=metadatas,
                    ids=ids
                )
                logger.info(f"Successfully stored {len(documents)} CSV rows in vector database")
                
                # Create a summary of the CSV content
                csv_summary = f"CSV File: {csv_file.filename}\n"
                csv_summary += f"Total Rows: {len(rows)}\n"
                csv_summary += f"Columns: {', '.join(rows[0].keys())}\n\n"
                
                # Add sample data (first few rows)
                csv_summary += "Sample Data:\n"
                for i, row in enumerate(rows[:5]):  # Show first 5 rows
                    csv_summary += f"Row {i+1}: {dict(row)}\n"
                
                if len(rows) > 5:
                    csv_summary += f"... and {len(rows) - 5} more rows\n"
                
                return csv_summary
                
            except Exception as e:
                logger.error(f"Error storing CSV in vector database: {e}")
                return f"Error storing CSV in database: {str(e)}"
        else:
            logger.error("Vector database collection not available")
            return "Vector database not available"
            
    except Exception as e:
        logger.error(f"Error processing CSV file: {e}")
        return f"Error processing CSV: {str(e)}"

def ask_assistant_from_file_optimized(code_filepath, user_prompt):
    """Optimized assistant interaction using direct chat completion instead of deprecated Assistants API."""
    start_time = time.time()
    logger.info(f"Starting optimized assistant interaction: {code_filepath}")
    
    try:
        # Read assistant configuration/instructions with UTF-8 encoding
        with open(f"agents/{code_filepath}", "r", encoding="utf-8") as file:
            assistant_instructions = file.read().strip()
            
        logger.info(f"Using assistant instructions from: {code_filepath}")
        
        # Initialize OpenAI client
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        # Create system message from assistant instructions
        system_message = f"""You are an expert assistant for PRD (Product Requirements Document) processing and epic generation.

{assistant_instructions}

Please follow the instructions above and process the user's request accordingly."""        # Use chat completions API instead of deprecated Assistants API
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,
            max_tokens=4000
        )
        
        assistant_response = response.choices[0].message.content
        
        processing_time = time.time() - start_time
        logger.info(f"Assistant response received in {processing_time:.2f} seconds")
        logger.info(f"Response length: {len(assistant_response)} characters")
        
        return assistant_response
            
    except Exception as e:
        processing_time = time.time() - start_time
        logger.error(f"Error in assistant interaction after {processing_time:.2f}s: {str(e)}")
        return f"Error: {str(e)}"

def create_intelligent_summary_fallback(content, filename, max_summary_length=5000):
    """Create an intelligent summary when vector database is not available."""
    logger.info(f"Creating intelligent fallback summary for {filename}")
    
    if len(content) <= max_summary_length:
        return content
    
    # Extract key sections using keyword-based approach
    important_keywords = [
        'requirement', 'feature', 'user story', 'acceptance criteria',
        'functional', 'non-functional', 'business rule', 'constraint',
        'objective', 'goal', 'scope', 'assumption', 'dependency',
        'architecture', 'design', 'integration', 'security', 'performance'
    ]
    
    # Split content into sections
    sections = content.split('\n\n')
    scored_sections = []
    
    for section in sections:
        score = 0
        section_lower = section.lower()
        
        # Score based on keywords
        for keyword in important_keywords:
            score += section_lower.count(keyword) * 10
        
        # Prefer sections with structured content (lists, numbers)
        score += section.count('•') * 5
        score += section.count('-') * 3
        score += section.count('1.') * 5
        score += section.count('2.') * 4
        score += section.count('3.') * 3
        
        # Prefer longer, more detailed sections
        if len(section) > 200:
            score += 5
        if len(section) > 500:
            score += 5
        
        scored_sections.append((score, section))
    
    # Sort by score and take top sections
    scored_sections.sort(key=lambda x: x[0], reverse=True)
    
    # Combine top sections until we reach max length
    summary_parts = []
    current_length = 0
    
    for score, section in scored_sections:
        if current_length + len(section) <= max_summary_length:
            summary_parts.append(section)
            current_length += len(section)
        else:
            # Try to fit partial section
            remaining_space = max_summary_length - current_length - 50  # Leave space for truncation message
            if remaining_space > 200:  # Only add if meaningful portion fits
                partial_section = section[:remaining_space] + "..."
                summary_parts.append(partial_section)
            break
    
    final_summary = "\n\n".join(summary_parts)
    
    if len(final_summary) > max_summary_length:
        final_summary = final_summary[:max_summary_length] + "... [Content truncated]"
    
    logger.info(f"Intelligent fallback summary created, length: {len(final_summary)} characters")
    
    # Print Fallback Summary for debugging
    logger.info("=" * 80)
    logger.info("INTELLIGENT FALLBACK SUMMARY GENERATED:")
    logger.info("=" * 80)
    logger.info(final_summary[:2000] + "..." if len(final_summary) > 2000 else final_summary)
    logger.info("=" * 80)
    
    return final_summary.strip()

def is_valid_content(content):
    """Check if content is valid for RAG processing."""
    if not content or len(content.strip()) < 50:
        return False
    
    # Check for error messages
    error_indicators = [
        "[Unable to decode",
        "[Unable to read",
        "[Error reading",
        "[File reading failed",
        "[File appears to be empty"
    ]
    
    for indicator in error_indicators:
        if content.startswith(indicator):
            return False
    
    return True

def debug_file_content(content, filename):
    """Debug helper to print file content details."""
    logger.info(f"=== FILE CONTENT DEBUG: {filename} ===")
    logger.info(f"Content length: {len(content) if content else 0}")
    logger.info(f"Content type: {type(content)}")
    logger.info(f"Is valid for RAG: {is_valid_content(content)}")
    
    if content:
        # Show first 200 characters, safely handling non-printable characters
        try:
            preview = content[:200].replace('\n', '\\n').replace('\r', '\\r')
            # Replace any problematic Unicode characters for logging
            safe_preview = preview.encode('ascii', errors='replace').decode('ascii')
            logger.info(f"Content preview: {safe_preview}")
        except Exception as e:
            logger.info(f"Content preview: [Unable to display preview due to encoding: {str(e)}]")
        
        # Count different character types
        if len(content) > 0:
            alpha_count = sum(1 for c in content if c.isalpha())
            digit_count = sum(1 for c in content if c.isdigit())
            space_count = sum(1 for c in content if c.isspace())
            other_count = len(content) - alpha_count - digit_count - space_count
            
            logger.info(f"Character breakdown - Alpha: {alpha_count}, Digits: {digit_count}, Spaces: {space_count}, Other: {other_count}")
    
    logger.info("=== END FILE DEBUG ===")

# Flask App Startup
@app.route("/vector-db-status", methods=["GET"])
def vector_db_status():
    """View vector database status and contents via web interface."""
    logger.info("GET request received for vector database status")
    
    try:
        if not prd_collection:
            return {
                "status": "error",
                "message": "Vector database not initialized",
                "count": 0,
                "documents": []
            }
        
        # Get collection info
        collection_info = {
            "name": prd_collection.name,
            "metadata": prd_collection.metadata,
            "count": prd_collection.count()
        }
        
        # Get all documents (limit to 50 for performance)
        all_docs = prd_collection.get(limit=50)
        
        documents = []
        if all_docs['documents']:
            for i, (doc, metadata, doc_id) in enumerate(zip(
                all_docs['documents'], 
                all_docs['metadatas'], 
                all_docs['ids']
            )):
                documents.append({
                    "id": doc_id,
                    "content_preview": doc[:200] + "..." if len(doc) > 200 else doc,
                    "content_length": len(doc),
                    "metadata": metadata,
                    "index": i
                })
        
        # Create HTML response for better viewing
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Vector Database Status</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .header {{ background-color: #f0f0f0; padding: 15px; border-radius: 5px; }}
                .document {{ border: 1px solid #ddd; margin: 10px 0; padding: 15px; border-radius: 5px; }}
                .metadata {{ background-color: #f9f9f9; padding: 10px; margin: 5px 0; border-radius: 3px; }}
                .content {{ background-color: #fff; padding: 10px; border-left: 3px solid #007bff; }}
                pre {{ white-space: pre-wrap; word-wrap: break-word; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Vector Database Status</h1>
                <p><strong>Collection:</strong> {collection_info['name']}</p>
                <p><strong>Total Documents:</strong> {collection_info['count']}</p>
                <p><strong>Showing:</strong> {len(documents)} documents (limited to 50)</p>
            </div>
        """
        
        if documents:
            for doc in documents:
                html_content += f"""
                <div class="document">
                    <h3>Document #{doc['index'] + 1}</h3>
                    <div class="metadata">
                        <strong>ID:</strong> {doc['id']}<br>
                        <strong>Length:</strong> {doc['content_length']} characters<br>
                        <strong>Filename:</strong> {doc['metadata'].get('filename', 'Unknown')}<br>
                        <strong>Type:</strong> {doc['metadata'].get('doc_type', 'Unknown')}<br>
                        <strong>Chunk:</strong> {doc['metadata'].get('chunk_index', 'N/A')} / {doc['metadata'].get('total_chunks', 'N/A')}<br>
                        <strong>Timestamp:</strong> {doc['metadata'].get('timestamp', 'Unknown')}
                    </div>
                    <div class="content">
                        <strong>Content Preview:</strong>
                        <pre>{doc['content_preview']}</pre>
                    </div>
                </div>
                """
        else:
            html_content += "<p>No documents found in the vector database.</p>"
        
        html_content += """
            <div style="margin-top: 30px; padding: 15px; background-color: #e7f3ff; border-radius: 5px;">
                <h3>How to Query the Vector Database</h3>
                <p>You can search the vector database by making a POST request to <code>/vector-db-search</code> with a JSON body:</p>
                <pre>{"query": "your search query", "top_k": 5}</pre>
                <p>Or use the search interface: <a href="/vector-db-search">Vector DB Search</a></p>
            </div>
        </body>
        </html>
        """
        
        return html_content, 200, {'Content-Type': 'text/html'}
        
    except Exception as e:
        logger.error(f"Error accessing vector database: {str(e)}")
        return {
            "status": "error",
            "message": f"Error accessing vector database: {str(e)}",
            "count": 0,
            "documents": []
        }

@app.route("/vector-db-search", methods=["GET", "POST"])
def vector_db_search():
    """Search the vector database via web interface."""
    logger.info(f"{request.method} request received for vector database search")
    
    if request.method == "GET":
        # Render search form
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Vector Database Search</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .search-form { background-color: #f0f0f0; padding: 20px; border-radius: 5px; margin-bottom: 20px; }
                .result { border: 1px solid #ddd; margin: 10px 0; padding: 15px; border-radius: 5px; }
                .metadata { background-color: #f9f9f9; padding: 10px; margin: 5px 0; border-radius: 3px; }
                .content { background-color: #fff; padding: 10px; border-left: 3px solid #28a745; }
                input[type="text"] { width: 70%; padding: 10px; margin: 5px; }
                input[type="number"] { width: 100px; padding: 10px; margin: 5px; }
                button { padding: 10px 20px; background-color: #007bff; color: white; border: none; border-radius: 3px; cursor: pointer; }
                button:hover { background-color: #0056b3; }
            </style>
        </head>
        <body>
            <h1>Vector Database Search</h1>
            <div class="search-form">
                <form method="POST">
                    <label>Search Query:</label><br>
                    <input type="text" name="query" placeholder="Enter your search query..." required><br>
                    <label>Number of Results:</label><br>
                    <input type="number" name="top_k" value="5" min="1" max="20"><br><br>
                    <button type="submit">Search</button>
                </form>
            </div>
            <p><a href="/vector-db-status">← Back to Vector DB Status</a></p>
        </body>
        </html>
        """
        return html_content, 200, {'Content-Type': 'text/html'}
    
    else:  # POST request
        try:
            # Get search parameters
            if request.is_json:
                data = request.get_json()
                query = data.get("query", "")
                top_k = data.get("top_k", 5)
            else:
                query = request.form.get("query", "")
                top_k = int(request.form.get("top_k", 5))
            
            if not query:
                return {"error": "Query is required"}, 400
            
            # Search the vector database
            results = retrieve_relevant_content(query, top_k=top_k)
            
            # Create HTML response
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Search Results</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 20px; }}
                    .header {{ background-color: #f0f0f0; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
                    .result {{ border: 1px solid #ddd; margin: 10px 0; padding: 15px; border-radius: 5px; }}
                    .metadata {{ background-color: #f9f9f9; padding: 10px; margin: 5px 0; border-radius: 3px; }}
                    .content {{ background-color: #fff; padding: 10px; border-left: 3px solid #28a745; }}
                    .relevance {{ color: #007bff; font-weight: bold; }}
                    pre {{ white-space: pre-wrap; word-wrap: break-word; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>Search Results</h1>
                    <p><strong>Query:</strong> "{query}"</p>
                    <p><strong>Results Found:</strong> {len(results)}</p>
                </div>
            """
            
            if results:
                for i, result in enumerate(results):
                    metadata = result.get('metadata', {})
                    content = result.get('content', '')
                    rank = result.get('relevance_rank', i + 1)
                    
                    html_content += f"""
                    <div class="result">
                        <h3>Result #{rank} <span class="relevance">(Relevance Rank: {rank})</span></h3>
                        <div class="metadata">
                            <strong>Filename:</strong> {metadata.get('filename', 'Unknown')}<br>
                            <strong>Document Type:</strong> {metadata.get('doc_type', 'Unknown')}<br>
                            <strong>Chunk:</strong> {metadata.get('chunk_index', 'N/A')} / {metadata.get('total_chunks', 'N/A')}<br>
                            <strong>Timestamp:</strong> {metadata.get('timestamp', 'Unknown')}
                        </div>
                        <div class="content">
                            <strong>Content:</strong>
                            <pre>{content}</pre>
                        </div>
                    </div>
                    """
            else:
                html_content += "<p>No results found for your query.</p>"
            
            html_content += """
                <div style="margin-top: 30px;">
                    <a href="/vector-db-search">← New Search</a> | 
                    <a href="/vector-db-status">Vector DB Status</a>
                </div>
            </body>
            </html>
            """
            
            return html_content, 200, {'Content-Type': 'text/html'}
            
        except Exception as e:
            logger.error(f"Error searching vector database: {str(e)}")
            return {"error": f"Search failed: {str(e)}"}, 500

@app.route("/vector-db-clear", methods=["POST"])
def vector_db_clear():
    """Clear all documents from the vector database."""
    global prd_collection
    logger.info("POST request received to clear vector database")
    
    try:
        if not prd_collection:
            return {"status": "error", "message": "Vector database not initialized"}
        
        # Get current count
        current_count = prd_collection.count()
        
        # Delete the collection and recreate it
        chroma_client.delete_collection(prd_collection.name)
        
        # Recreate the collection
        prd_collection = chroma_client.get_or_create_collection(
            name="prd_documents",
            metadata={"description": "PRD and documentation storage for RAG"}
        )
        
        logger.info(f"Vector database cleared. Removed {current_count} documents.")
        
        return {
            "status": "success",
            "message": f"Vector database cleared successfully. Removed {current_count} documents.",
            "documents_removed": current_count
        }
        
    except Exception as e:
        logger.error(f"Error clearing vector database: {str(e)}")
        return {"status": "error", "message": f"Error clearing database: {str(e)}"}

@app.route("/document-upload", methods=["POST"])
def document_upload_preview():
    """Step 1: Upload and preview documents with optimized processing for large files."""
    logger.info("=== DOCUMENT UPLOAD ENDPOINT CALLED ===")
    start_time = time.time()
    logger.info("POST request received for document upload and preview")
    
    try:
        logger.info("=== INSIDE TRY BLOCK ===")
        
        # Collect input
        logger.info("Collecting form data and file uploads for preview")
        context = request.form.get("context", "")
        prd_file = request.files.get("prd_file")
        additional_docs = request.files.get("additional_docs")
        csv_file = request.files.get("csv_file")
        
        logger.info(f"Context provided: {'Yes' if context else 'No'}")
        logger.info(f"PRD file uploaded: {'Yes' if prd_file else 'No'}")
        logger.info(f"Additional docs uploaded: {'Yes' if additional_docs else 'No'}")
        logger.info(f"CSV file uploaded: {'Yes' if csv_file else 'No'}")
        
        # Log file sizes for performance monitoring
        if prd_file:
            prd_file.seek(0, 2)  # Seek to end
            prd_size = prd_file.tell()
            prd_file.seek(0)  # Reset to beginning
            logger.info(f"PRD file: {prd_file.filename}, size: {prd_size:,} bytes")
        
        if additional_docs:
            additional_docs.seek(0, 2)
            docs_size = additional_docs.tell()
            additional_docs.seek(0)
            logger.info(f"Additional docs: {additional_docs.filename}, size: {docs_size:,} bytes")

        # Process CSV file with smart optimization based on file size
        csv_summary = ""
        if csv_file and csv_file.filename.lower().endswith('.csv'):
            csv_file.seek(0, 2)  # Seek to end to get size
            csv_size = csv_file.tell()
            csv_file.seek(0)  # Reset to beginning
            
            # Apply fast preview mode only for large CSV files (>10KB)
            if csv_size > 10240:  # 10KB threshold
                logger.info(f"Large CSV file detected ({csv_size:,} bytes), using fast preview mode")
                try:
                    # For large files, read first few rows only for performance
                    content = csv_file.read(2048).decode('utf-8', errors='ignore')  # Read first 2KB only
                    lines = content.split('\n')[:10]  # First 10 lines only
                    csv_summary = f"CSV Preview ({csv_file.filename}):\n" + '\n'.join(lines)
                    if len(lines) >= 10:
                        csv_summary += "\n[... more rows available - full processing during epic generation]"
                    logger.info(f"Large CSV preview completed: {len(lines)} lines shown")
                except Exception as e:
                    logger.warning(f"CSV preview failed: {e}")
                    csv_summary = f"CSV file uploaded: {csv_file.filename} (will be processed during epic generation)"
            else:
                logger.info(f"Small CSV file detected ({csv_size:,} bytes), using full processing")
                try:
                    # For small files, process fully for richer preview
                    csv_summary = process_csv_to_vector_db(csv_file)
                    if not csv_summary:
                        # Fallback to simple preview if processing fails
                        csv_file.seek(0)
                        content = csv_file.read().decode('utf-8', errors='ignore')
                        lines = content.split('\n')[:20]  # Show more lines for small files
                        csv_summary = f"CSV Content ({csv_file.filename}):\n" + '\n'.join(lines)
                except Exception as e:
                    logger.warning(f"CSV processing failed: {e}")
                    csv_summary = f"CSV file uploaded: {csv_file.filename}"
        elif csv_file:
            logger.warning(f"Uploaded file {csv_file.filename} is not a CSV file")
            csv_summary = f"Warning: {csv_file.filename} is not a CSV file"

        # Smart file reading with size-based optimization approach
        logger.info("Reading uploaded files with smart size-based processing")
        
        prd_content = ""
        docs_content = ""
        
        # Read PRD file with smart optimization based on size
        if prd_file:
            logger.info(f"Reading PRD file: {prd_file.filename}")
            prd_content = safe_read(prd_file)
            if prd_content:
                logger.info(f"PRD content length: {len(prd_content):,} characters")
                # Apply aggressive truncation only for very large files (>30KB)
                if len(prd_content) > 30000:  # 30KB text threshold
                    logger.info("Large PRD detected, using aggressive truncation for preview speed")
                    prd_content = prd_content[:30000] + "\n\n[Content truncated for fast preview - full content will be used in epic generation]"
                else:
                    logger.info("Small-medium PRD detected, preserving full content for rich preview")
        
        # Read additional docs with smart optimization based on size
        if additional_docs:
            logger.info(f"Reading additional docs: {additional_docs.filename}")
            docs_content = safe_read(additional_docs)
            if docs_content:
                logger.info(f"Docs content length: {len(docs_content):,} characters")
                # Apply aggressive truncation only for large files (>20KB)
                if len(docs_content) > 20000:  # 20KB text threshold
                    logger.info("Large additional docs detected, using aggressive truncation for preview speed")
                    docs_content = docs_content[:20000] + "\n\n[Content truncated for fast preview - full content will be used in epic generation]"
                else:
                    logger.info("Small-medium additional docs detected, preserving full content for rich preview")

        # Create summaries with smart processing based on file size
        logger.info("Creating summaries with size-adaptive processing")
        
        # Process PRD summary with smart approach
        prd_summary = ""
        if prd_content and len(prd_content.strip()) > 50:
            prd_raw_length = len(prd_content)
            logger.info(f"Processing PRD summary (content length: {prd_raw_length:,} chars)")
            
            # For small files (≤30KB), use full RAG processing for rich summaries
            if prd_raw_length <= 30000 and not prd_content.endswith("[Content truncated for fast preview - full content will be used in epic generation]"):
                logger.info("Small PRD file detected, using full RAG processing for rich preview")
                try:
                    prd_summary = create_rag_summary(prd_content, prd_file.filename if prd_file else "prd_document", max_summary_length=8000)
                except Exception as e:
                    logger.warning(f"RAG processing failed for PRD, using fallback: {e}")
                    prd_summary = create_intelligent_summary_fallback(prd_content, prd_file.filename if prd_file else "prd_document", max_summary_length=5000)
            else:
                # For large files, use fast simple truncation
                logger.info("Large PRD file detected, using fast truncation for preview")
                if len(prd_content) > 3000:
                    prd_summary = prd_content[:3000] + "\n\n[Preview truncated - full content will be processed during epic generation]"
                else:
                    prd_summary = prd_content
        else:
            logger.warning("PRD content is too short or invalid")
            prd_summary = "No valid PRD content available - please check file format and content."

        # Process additional docs summary with smart approach
        docs_summary = ""
        if docs_content and len(docs_content.strip()) > 50:
            docs_raw_length = len(docs_content)
            logger.info(f"Processing additional docs summary (content length: {docs_raw_length:,} chars)")
            
            # For small files (≤20KB), use full RAG processing for rich summaries
            if docs_raw_length <= 20000 and not docs_content.endswith("[Content truncated for fast preview - full content will be used in epic generation]"):
                logger.info("Small additional docs detected, using full RAG processing for rich preview")
                try:
                    docs_summary = create_rag_summary(docs_content, additional_docs.filename if additional_docs else "additional_docs", max_summary_length=5000)
                except Exception as e:
                    logger.warning(f"RAG processing failed for additional docs, using fallback: {e}")
                    docs_summary = create_intelligent_summary_fallback(docs_content, additional_docs.filename if additional_docs else "additional_docs", max_summary_length=3000)
            else:
                # For large files, use fast simple truncation
                logger.info("Large additional docs detected, using fast truncation for preview")
                if len(docs_content) > 2000:
                    docs_summary = docs_content[:2000] + "\n\n[Preview truncated - full content will be processed during epic generation]"
                else:
                    docs_summary = docs_content
        else:
            logger.warning("Additional docs content is too short or invalid")
            docs_summary = "No valid additional documentation available - please check file format and content."
        
        # Log final processing results
        logger.info(f"PRD summary length: {len(prd_summary)} characters")
        logger.info(f"Additional docs summary length: {len(docs_summary)} characters")
        
        # Smart debug output - detailed for small files, minimal for large files
        if (prd_content and len(prd_content) <= 30000) or (docs_content and len(docs_content) <= 20000):
            logger.info("Small files detected - debug output enabled for detailed analysis")
        else:
            logger.info("Large files detected - debug output minimized for performance")
        
        # Create preview data
        preview_data = {
            "user_context": context,
            "prd_summary": prd_summary,
            "docs_summary": docs_summary,
            "csv_summary": csv_summary,
            "prd_filename": prd_file.filename if prd_file else None,
            "docs_filename": additional_docs.filename if additional_docs else None,
            "csv_filename": csv_file.filename if csv_file else None,
            "processing_time": time.time() - start_time
        }
        
        logger.info(f"Document preview generated in {preview_data['processing_time']:.2f} seconds")
        
        # Return JSON response for AJAX handling
        return jsonify({
            "success": True,
            "preview_data": preview_data,
            "message": "Documents processed successfully. Review the summary below and click 'Generate Epics' to proceed."
        })
        
    except Exception as e:
        print(f"Exception in /document-upload: {e}")
        logger.info(f"Error in document_upload_preview: {str(e)}")
        logger.error(f"Exception type: {type(e).__name__}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        return jsonify({
            "success": False,
            "error": f"Error processing documents: {str(e)}"
        }), 500

@app.route("/generate-epics", methods=["POST"])
def generate_epics_from_preview():
    """Step 2: Generate epics using the previewed RAG summary."""
    start_time = time.time()
    logger.info("POST request received for epic generation from preview")
    
    try:
        # Get the preview data from the request
        request_data = request.get_json()
        if not request_data:
            logger.error("No preview data received for epic generation")
            return jsonify({"success": False, "error": "No preview data received"}), 400
        
        user_context = request_data.get("user_context", "")
        prd_summary = request_data.get("prd_summary", "")
        docs_summary = request_data.get("docs_summary", "")
        
        logger.info(f"Epic generation - Context: {'Yes' if user_context else 'No'}")
        logger.info(f"Epic generation - PRD summary: {len(prd_summary)} characters")
        logger.info(f"Epic generation - Docs summary: {len(docs_summary)} characters")

        # Create enhanced context for Epic Generator
        enhanced_context = f"""
        {user_context}

        RAG-Enhanced PRD Analysis:
        {prd_summary}

        Additional Context:
        {docs_summary}

        Instructions: The above content has been intelligently extracted and summarized using RAG (Retrieval Augmented Generation). 
        It contains the most relevant requirements, user stories, and business objectives from the original documents.
        Use this curated information to generate comprehensive epics and user stories.        """
        logger.info(f"Enhanced context length: {len(enhanced_context)} characters")
        
        # Check cache first
        prompt_hash = get_cache_key(enhanced_context)
        cached_result = response_cache.get(prompt_hash)
        if cached_result:
            logger.info("Cache hit! Returning cached response")
            processing_time = time.time() - start_time
            logger.info(f"Total processing time (cached): {processing_time:.2f} seconds")
            return jsonify({
                "success": True,
                "epics": cached_result,
                "processing_time": processing_time,
                "cached": True
            })          # Log token count for the enhanced context
        context_tokens = count_tokens(enhanced_context, "gpt-4o")
        logger.info(f"Enhanced context token count: {context_tokens:,} tokens")
        
        # Check if context is within token limits
        if context_tokens > 120000:
            logger.warning(f"Enhanced context token count ({context_tokens:,}) is approaching the 128k limit!")
            # Further optimize if needed
            enhanced_context = enhanced_context[:100000] + "\n[Content truncated due to token limits]"
            context_tokens = count_tokens(enhanced_context, "gpt-4o")
            logger.info(f"Truncated context token count: {context_tokens:,} tokens")
        else:
            logger.info(f"Enhanced context is within safe token limits ({context_tokens:,}/128,000 tokens)")    # Skip PRD Parser Agent - RAG has already done the parsing and summarization
        logger.info("****************Skipping PRD Parser - Using RAG Summary Directly")
        logger.info("Starting Epic Generator with RAG-enhanced content")
        
        # Print Enhanced Context being sent to Epic Agent
        print("\n" + "=" * 100)
        print("ENHANCED CONTEXT INPUT TO EPIC AGENT:")
        print("=" * 100)
        print(enhanced_context)  # Print the FULL content
        print("=" * 100)
        print(f"Total length: {len(enhanced_context)} characters")
        print(f"Total tokens: {count_tokens(enhanced_context, 'gpt-4o'):,} tokens")
        print("=" * 100 + "\n")
        
        # Also log to file (truncated for file logs)
        logger.info("=" * 80)
        logger.info("ENHANCED CONTEXT INPUT TO EPIC AGENT:")
        logger.info("=" * 80)
        logger.info(enhanced_context[:5000] + "..." if len(enhanced_context) > 5000 else enhanced_context)
        logger.info("=" * 80)
        
        # Directly use Epic Generator with RAG-enhanced content
        start_epic_time = time.time()
        epic_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", enhanced_context)
        epic_processing_time = time.time() - start_epic_time
        
        logger.info("################Epic Generator response received")
        logger.info(f"Epic Generator processing time: {epic_processing_time:.2f} seconds")
        
        # Print Epic Agent Response to console
        print("\n" + "=" * 100)
        print("EPIC AGENT RESPONSE OUTPUT:")
        print("=" * 100)
        print(epic_response)
        print("=" * 100)
        print(f"Response length: {len(epic_response)} characters")
        print("=" * 100 + "\n")                
        # Print Epic Agent Response
        logger.info("=" * 80)
        logger.info("EPIC AGENT RESPONSE OUTPUT:")
        logger.info("=" * 80)
        logger.info(epic_response)
        logger.info("=" * 80)
        
        # Log token usage for Epic Generator interaction
        log_token_usage(enhanced_context, epic_response, model="gpt-4o", context="RAG-Enhanced Epic Generator")

        final_output = epic_response
        logger.info(f"Final output length: {len(final_output)} characters")
        
        # Cache the result for future requests
        response_cache[prompt_hash] = final_output
        logger.info("Response cached for future requests")
        
        # Store the initial epic data in session for navigation consistency
        session['current_epics'] = final_output
        session['current_user_stories'] = ''  # Clear any previous user stories
        logger.info("Stored initial epic data in session")
        
        processing_time = time.time() - start_time
        logger.info(f"Total RAG-optimized processing time: {processing_time:.2f} seconds")
        
        # Calculate time savings
        estimated_traditional_time = processing_time * 2  # Estimate of traditional two-agent approach
        time_saved = estimated_traditional_time - processing_time
        logger.info(f"Estimated time saved by RAG optimization: {time_saved:.2f} seconds")
        
        # Return JSON response for AJAX handling
        return jsonify({
            "success": True,
            "epics": final_output,
            "processing_time": processing_time,
            "time_saved": time_saved,
            "cached": False
        })
        
    except Exception as e:
        logger.error(f"Error in generate_epics_from_preview: {str(e)}")
        logger.error(f"Exception type: {type(e).__name__}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        return jsonify({
            "success": False,
            "error": f"Error generating epics: {str(e)}"
        }), 500

@app.route("/epic-results", methods=["GET", "POST"])
def show_epic_results():
    """Display the epic results page after generation."""
    if request.method == "POST":
        epics = request.form.get("epics", "")
        user_stories = request.form.get("user_stories", "")
        
        # Store the epic and user story data in the session for back navigation
        session['current_epics'] = epics
        session['current_user_stories'] = user_stories
        logger.info("Stored epic and user story data in session for back navigation")
        
        return render_template("poc2_epic_story_screen.html", epics=epics, user_stories=user_stories)
    else:
        # Check if we have stored epic data in the session first
        stored_epics = session.get('current_epics')
        stored_user_stories = session.get('current_user_stories', '')
        
        # Check for navigation context from user story details back button
        navigation_context = session.get('user_story_navigation_context')
        selected_story_ids = None
        
        if navigation_context:
            # User is coming back from user story details - restore selected stories
            selected_story_ids = navigation_context.get('selected_story_ids', [])
            logger.info(f"Restoring selected story context for back navigation: {selected_story_ids}")
            
            # Clear the navigation context after use to prevent stale state
            session.pop('user_story_navigation_context', None)
            session.modified = True
        
        if stored_epics:
            logger.info("Retrieved epic and user story data from session")
            
            # Check if stored user stories are valid/displayable
            user_stories_to_display = stored_user_stories
            
            # If user stories are empty or just raw JSON, we need to provide a placeholder
            if not stored_user_stories or stored_user_stories.strip() == "":
                logger.info("No user stories found in session, user needs to select epics first")
                user_stories_to_display = ""
            elif stored_user_stories.startswith('[') or stored_user_stories.startswith('{'):
                # Stored user stories are raw JSON from agent - provide information message
                logger.info("User stories in session are raw JSON format, showing instruction message")
                user_stories_to_display = """
                <div class="info-message" style="text-align: center; padding: 2rem; background: var(--surface-secondary); border-radius: 8px; border: 1px dashed var(--border-medium);">
                    <h6 style="color: var(--primary-red); margin-bottom: 1rem;">🔄 User Stories Available</h6>
                    <p style="color: var(--text-secondary); margin-bottom: 1.5rem;">
                        User stories have been generated for your selected epics. 
                        <br>Please select the epics above and click "Approve & Generate User Stories" to view them.
                    </p>
                    <div style="background: var(--primary-red-lighter); padding: 1rem; border-radius: 4px; color: var(--text-primary);">
                        <small><strong>💡 Tip:</strong> Your previous work has been preserved. Simply reselect your epics to continue.</small>
                    </div>
                </div>
                """
            
            return render_template("poc2_epic_story_screen.html", 
                                 epics=stored_epics, 
                                 user_stories=user_stories_to_display,
                                 selected_story_ids=selected_story_ids)
        
        # Fallback: Provide sample epics for testing Epic Chat functionality
        logger.info("No stored data found, using sample epics")
        
        # Check for navigation context even for sample data
        navigation_context = session.get('user_story_navigation_context')
        selected_story_ids = None
        
        if navigation_context:
            selected_story_ids = navigation_context.get('selected_story_ids', [])
            logger.info(f"Restoring selected story context for sample data: {selected_story_ids}")
            session.pop('user_story_navigation_context', None)
            session.modified = True
        
        sample_epics = """
        <div class="epic-card" data-epic-id="epic_1">
          <div style="display: flex; align-items: center; gap: 10px;">
            <input type="checkbox" name="epic_ids" value="epic_1" id="epic_1" checked>
            <h5 style="margin: 0;">Epic 1: User Authentication and Security</h5>
          </div>
          <p><strong>Description:</strong> Implement comprehensive user authentication system with multi-factor authentication, secure password policies, and role-based access control to ensure platform security and user data protection.</p>
          <p><strong>Business Value:</strong> Enhanced security reduces risk of data breaches and builds user trust, essential for regulatory compliance and business reputation.</p>
          <p><strong>Acceptance Criteria:</strong></p>
          <ul>
            <li>Users can register with email verification</li>
            <li>Multi-factor authentication is required for admin roles</li>
            <li>Password policies enforce strong passwords</li>
            <li>Session management with automatic timeout</li>
            <li>Role-based permissions control feature access</li>
          </ul>
        </div>
        
        <div class="epic-card" data-epic-id="epic_2">
          <div style="display: flex; align-items: center; gap: 10px;">
            <input type="checkbox" name="epic_ids" value="epic_2" id="epic_2">
            <h5 style="margin: 0;">Epic 2: Data Analytics Dashboard</h5>
          </div>
          <p><strong>Description:</strong> Create a comprehensive analytics dashboard that provides real-time insights into user behavior, system performance, and business metrics with customizable reporting capabilities.</p>
          <p><strong>Business Value:</strong> Data-driven decision making capabilities will improve operational efficiency and enable better strategic planning.</p>
          <p><strong>Acceptance Criteria:</strong></p>
          <ul>
            <li>Real-time data visualization with charts and graphs</li>
            <li>Customizable dashboard layouts</li>
            <li>Export functionality for reports</li>
            <li>Performance metrics tracking</li>
            <li>User activity analytics</li>
          </ul>
        </div>
        
        <div class="epic-card" data-epic-id="epic_3">
          <div style="display: flex; align-items: center; gap: 10px;">
            <input type="checkbox" name="epic_ids" value="epic_3" id="epic_3">
            <h5 style="margin: 0;">Epic 3: Mobile Application Development</h5>
          </div>
          <p><strong>Description:</strong> Develop a responsive mobile application that provides core platform functionality with offline capabilities and push notifications for enhanced user engagement.</p>
          <p><strong>Business Value:</strong> Mobile accessibility increases user engagement and market reach, providing competitive advantage in mobile-first market.</p>
          <p><strong>Acceptance Criteria:</strong></p>
          <ul>
            <li>Native iOS and Android applications</li>
            <li>Offline functionality for core features</li>
            <li>Push notification system</li>
            <li>Responsive design for various screen sizes</li>
            <li>Integration with existing web platform</li>
          </ul>
        </div>
        """
        
        # Provide sample user stories for testing User Story Chat functionality
        sample_user_stories = """[
          {
            "story_id": "US-001",
            "name": "Secure User Registration",
            "priority": "High",
            "systems": "Authentication Service, Email Service",
            "description": "As a new user, I want to register for an account with email verification so that I can securely access the platform"
          },
          {
            "story_id": "US-002", 
            "name": "Multi-Factor Authentication Setup",
            "priority": "High",
            "systems": "Authentication Service, SMS Service",
            "description": "As a user, I want to enable multi-factor authentication so that my account has an additional layer of security"
          },
          {
            "story_id": "US-003",
            "name": "Password Policy Enforcement", 
            "priority": "Medium",
            "systems": "Authentication Service",
            "description": "As a system, I want to enforce strong password policies so that user accounts are protected from common attacks"
          },
          {
            "story_id": "US-004",
            "name": "Real-time Dashboard Creation",
            "priority": "High", 
            "systems": "Analytics Engine, Dashboard Service",
            "description": "As a business user, I want to create customizable real-time dashboards so that I can monitor key metrics and KPIs"
          },
          {
            "story_id": "US-005",
            "name": "Data Visualization Charts",
            "priority": "Medium",
            "systems": "Visualization Service, Chart Library", 
            "description": "As an analyst, I want to create various chart types (bar, line, pie) so that I can visualize data in the most appropriate format"
          },
          {
            "story_id": "US-006",
            "name": "Mobile App User Interface",
            "priority": "High",
            "systems": "Mobile Frontend, API Gateway",
            "description": "As a mobile user, I want an intuitive and responsive interface so that I can easily navigate and use the app on my device"
          },
          {
            "story_id": "US-007",
            "name": "Offline Data Synchronization", 
            "priority": "Medium",
            "systems": "Mobile App, Sync Service",
            "description": "As a mobile user, I want the app to work offline and sync data when reconnected so that I can use core features without internet"
          },
          {
            "story_id": "US-008",
            "name": "Push Notification System",
            "priority": "Low", 
            "systems": "Notification Service, Mobile App",
            "description": "As a user, I want to receive push notifications for important updates so that I stay informed about relevant activities"
          }
        ]"""
        
        return render_template("poc2_epic_story_screen.html", 
                                 epics=sample_epics, 
                                 user_stories=sample_user_stories,
                                 selected_story_ids=selected_story_ids)

@app.route("/approve-epics", methods=["POST"])
def approve_epics():
    """Process approved epics and generate user stories for them."""
    start_time = time.time()
    logger.info("POST request received for epic approval and user story generation")
    user_story_list = []  # <-- new list to collect structured stories
    try:
        # Get the approved epic IDs from the form
        epic_ids_str = request.form.get("epic_ids", "")
        epic_ids = [epic_id.strip() for epic_id in epic_ids_str.split(",") if epic_id.strip()]
        
        logger.info(f"Processing approval for epic IDs: {epic_ids}")
        
        if not epic_ids:
            logger.error("No epic IDs provided for approval")
            return render_template("poc2_epic_story_screen.html", 
                                 epics="No epics selected for approval", 
                                 user_stories="")
          # Get the current epics content to preserve it
        current_epics = request.form.get("current_epics", "")
        logger.info(f"Current epics content length: {len(current_epics)} characters")
        
        # Store current epics in session for Epic Chat context
        session['current_epics'] = current_epics
        session.modified = True
        logger.info("Stored current epics in session for Epic Chat context")
          # Get the selected epic contents (actual epic data)
        selected_epic_contents_str = request.form.get("selected_epic_contents", "{}")
        try:
            selected_epic_contents = json.loads(selected_epic_contents_str)
            logger.info(f"Received epic contents for {len(selected_epic_contents)} epics")
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse selected epic contents: {e}")
            selected_epic_contents = {}
        
        # Generate user stories for each approved epic
        user_stories = []
        logger.info("Starting user story generation for approved epics")
        
        for i, epic_id in enumerate(epic_ids):
            logger.info(f"Generating user stories for epic: {epic_id}")            # Get the actual epic content instead of just the ID
            epic_content = selected_epic_contents.get(epic_id, {})
            epic_title = epic_content.get('title', f'Epic {i+1}')
            epic_description = epic_content.get('description', 'No description available')
            
            logger.info(f"Epic {epic_id} - Title: {epic_title[:50]}...")
            logger.info(f"Epic {epic_id} - Description: {epic_description[:100]}...")
            
            # Create a prompt that aligns with the agent's system instructions
            prompt = f"""
            Epic Title: {epic_title}
            Epic Description: {epic_description}
            
            Generate user stories for this approved epic. Return ONLY the JSON array format with story_id, name, priority, and systems. Do not include descriptions or acceptance criteria.
            """
              # Print what's being sent to User Story Agent
            logger.info("=" * 80)
            logger.info(f"USER STORY AGENT INPUT FOR EPIC {i+1}:")
            logger.info("=" * 80)
            logger.info(f"Epic ID: {epic_id}")
            logger.info(f"Epic Title: {epic_title}")
            logger.info(f"Epic Description: {epic_description}")
            logger.info("-" * 50)
            logger.info("Full Prompt:")
            logger.info(prompt)
            logger.info("=" * 80)
            logger.info(f"Prompt length: {len(prompt)} characters")
            logger.info("=" * 80 + "\n")
            
            try:
                story_response = ask_assistant_from_file_optimized("poc2_agent3_basic_user_story", prompt)
                
                # Print User Story Agent Response
                logger.info("=" * 80)
                logger.info(f"USER STORY AGENT RESPONSE FOR EPIC {i+1}:")
                logger.info("=" * 80)
                logger.info(story_response)
                logger.info("=" * 80)
                logger.info(f"Response length: {len(story_response)} characters")
                logger.info("=" * 80 + "\n")                
                if story_response:
                    # Send the raw JSON response for frontend parsing instead of HTML cards
                    user_stories.append(story_response)
                    user_story_list.append({
                        "epic": epic_title,
                        "stories": story_response  # Raw agent response for JS parsing
                    })
                    logger.info(f"Successfully generated user stories for {epic_id}")
                else:
                    logger.warning(f"No response received for epic {epic_id}")
                    user_stories.append(f"""
                    <div class='user-story-card'>
                    <h6>User Stories for: {epic_title}</h6>
                    <label>
                        <input type="radio" name="user_story_id" value="story_{i+1}">
                        <div style='padding: 10px; background-color: #f8f9fa; border-radius: 4px; margin: 5px 0;'>
                        {story_response}
                        </div>
                    </label>
                    </div>
                    """)
            except Exception as e:
                logger.error(f"Error generating user stories for epic {epic_id}: {str(e)}")
                user_stories.append(f"""
                <div class='user-story-card'>
                    <h6>User Stories for {epic_id}</h6>
                    <p class="text-danger">Error generating user stories: {str(e)}</p>
                </div>
                """)
        
        # Combine all user stories - send raw content for JS parsing
        user_stories_content = "\n".join(user_stories)
        
        # Store the updated user stories in session for back navigation
        # Store both the raw content and a more structured format for better restoration
        session['current_user_stories'] = user_stories_content
        session['current_user_stories_structured'] = user_story_list  # Store structured data too
        logger.info("Updated user stories data in session with both raw and structured formats")
        
        processing_time = time.time() - start_time
        logger.info(f"User story generation completed in {processing_time:.2f} seconds")
        logger.info(f"Generated user stories for {len(epic_ids)} epics")
        
        
        # Render the template with both epics (preserved) and user stories (raw for JS parsing)
        return render_template("poc2_epic_story_screen.html",
                       epics=current_epics,  # This should be epics HTML
                       user_stories=user_stories_content)  # Raw content for JS to parse into table

        
    except Exception as e:
        logger.error(f"Error in approve_epics: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        # Return error page with preserved epics
        error_message = f"""
        <div class='user-story-card'>
            <h6>Error Generating User Stories</h6>
            <p class="text-danger">An error occurred while generating user stories: {str(e)}</p>
            <p class="text-muted">Please try again or contact support if the issue persists.</p>
        </div>
        """
        
        return render_template(
            "poc2_epic_story_screen.html", 
            epics=request.form.get("current_epics", ""), 
            user_stories=error_message
        )

@app.route("/user-story-details", methods=["GET", "POST"])
def user_story_details():
    """Process selected user story and display details with acceptance criteria."""
    logger.info("Request received for user story details")
    
    try:
        # Handle POST request (form submission from epic story screen)
        if request.method == "POST":
            logger.info("POST request received for user story selection")
            
            # Get the selected story IDs and details (now supporting multiple selections)
            if request.is_json:
                data = request.get_json()
                selected_story_id = data.get("selected_story_id", "")  # Comma-separated IDs
                story_name = data.get("selected_story_name", "")
                selected_story_description = data.get("selected_story_description", "")
                selected_stories_data = data.get("selected_stories_data", "")  # JSON string of story objects
                epic_title = data.get("epic_title", "")
                priority = data.get("selected_story_priority", data.get("priority", "High"))
                responsible_systems = data.get("selected_story_systems", "")
                logger.info(f"Processing selected user story IDs (JSON): {selected_story_id}")
                logger.info(f"Processing priority (JSON): {priority}")
                logger.info(f"Processing systems (JSON): {responsible_systems}")
            else:
                selected_story_id = request.form.get("selected_story_id", "")  # Comma-separated IDs
                story_name = request.form.get("selected_story_name", "")
                selected_story_description = request.form.get("selected_story_description", "")
                selected_stories_data = request.form.get("selected_stories_data", "")  # JSON string of story objects
                epic_title = request.form.get("epic_title", "")
                priority = request.form.get("selected_story_priority", request.form.get("priority", "High"))
                responsible_systems = request.form.get("selected_story_systems", "")
                logger.info(f"Processing selected user story IDs (Form): {selected_story_id}")
                logger.info(f"Processing priority (Form): {priority}")
                logger.info(f"Processing systems (Form): {responsible_systems}")
            
            # Parse multiple story IDs
            story_ids = [id.strip() for id in selected_story_id.split(',') if id.strip()] if selected_story_id else []
            logger.info(f"Processing {len(story_ids)} selected user stories: {story_ids}")
            
            # Parse detailed story data if available
            selected_stories = []
            if selected_stories_data:
                try:
                    selected_stories = json.loads(selected_stories_data)
                    logger.info(f"Parsed detailed data for {len(selected_stories)} stories")
                except json.JSONDecodeError as e:
                    logger.warning(f"Failed to parse selected stories data: {e}")
                    selected_stories = []
            
            logger.info(f"Processing selected story names: {story_name}")
            logger.info(f"Processing selected story descriptions: {selected_story_description[:100] if selected_story_description else 'N/A'}...")
            
            # Always proceed - if no stories selected, use defaults to ensure agents are called
            if not story_ids:
                logger.info("No user stories selected, using default values to ensure agents are called")
                story_ids = ['default-story-1']
                story_name = story_name or "Lock critical fields for charged-off accounts"
                selected_story_description = selected_story_description or "As a system architect, I want to identify and classify the necessary data fields for charged-off accounts so that we can ensure only required and relevant data is transmitted securely."
                selected_story_id = 'default-story-1'
            else:
                selected_story_id = ','.join(story_ids)
            
            # Provide default values if story details are missing
            if not story_name or not selected_story_description:
                logger.info("Story name or description missing, using default values")
                default_story_name = "Lock critical fields for charged-off accounts"
                default_story_description = "As a system architect, I want to identify and classify the necessary data fields for charged-off accounts so that we can ensure only required and relevant data is transmitted securely."

                story_name = story_name or default_story_name
                selected_story_description = selected_story_description or default_story_description
                
                logger.info(f"Using default story name: {story_name}")
                logger.info(f"Using default story description: {selected_story_description[:100]}...")
            
            logger.info(f"Successfully processed user story selection: {selected_story_id}")
            
            # Extract data from selected stories for the new agent format
            if selected_stories:
                # Use data from the first selected story as primary, combine others
                primary_story = selected_stories[0]
                
                # Extract systems information from all selected stories
                all_systems = set()
                all_priorities = []
                combined_names = []
                combined_descriptions = []
                
                for story in selected_stories:
                    # Handle new agent response format
                    story_name_field = story.get('name', story.get('title', 'Untitled Story'))
                    story_desc_field = story.get('description', story_name_field)
                    story_priority = story.get('priority', 'Medium')
                    story_systems = story.get('systems', [])
                    
                    combined_names.append(story_name_field)
                    combined_descriptions.append(story_desc_field)
                    all_priorities.append(story_priority)
                    
                    # Handle systems field - can be array or string
                    if isinstance(story_systems, list):
                        all_systems.update(story_systems)
                    elif isinstance(story_systems, str):
                        all_systems.update([s.strip() for s in story_systems.split(',') if s.strip()])
                
                # Override with combined data, but preserve form values as fallback
                story_name = ' | '.join(combined_names) if len(combined_names) > 1 else combined_names[0] if combined_names else story_name
                selected_story_description = ' | '.join(combined_descriptions) if len(combined_descriptions) > 1 else combined_descriptions[0] if combined_descriptions else selected_story_description
                
                # Use extracted systems if available, otherwise keep form value
                extracted_systems = ', '.join(sorted(all_systems)) if all_systems else ""
                if not extracted_systems and responsible_systems:
                    # Keep the form value if no systems were extracted from stories
                    pass  # responsible_systems already has form value
                else:
                    responsible_systems = extracted_systems or responsible_systems or "TBD"
                
                # Use extracted priority if available, otherwise keep form value  
                extracted_priority = all_priorities[0] if all_priorities else ""
                if not extracted_priority and priority:
                    # Keep the form value if no priority was extracted from stories
                    pass  # priority already has form value
                else:
                    priority = extracted_priority or priority or "Medium"
                
                logger.info(f"Extracted from {len(selected_stories)} stories:")
                logger.info(f"  - Combined name: {story_name}")
                logger.info(f"  - Systems: {responsible_systems}")
                logger.info(f"  - Priority: {priority}")
            
            # Generate additional data fields for the new UI
            # Ensure we have values for responsible_systems and priority
            if 'responsible_systems' not in locals() or not responsible_systems:
                responsible_systems = "CAPS, CMS"  # Default value if not set from form or stories
            if 'priority' not in locals() or not priority:
                priority = "High"  # Default value if not set from form or stories
            tagged_requirements = [
                "TBD"
            ]
            
            # Call the acceptance criteria agent - pass only story name
            prompt = story_name
            logger.info(f"Processing selected prompt ****Acceptance Criteria Prompt **** : {prompt}")

            acceptance_response = ask_assistant_from_file_optimized("poc2_agent4_acceptanceCriteria_gen", prompt)

            # Call the new User Description Agent - pass only story name
            description_prompt = story_name
            logger.info(f"Processing User Description Agent prompt: {description_prompt}")
            
            description_response = ask_assistant_from_file_optimized("poc2_agent5_description", description_prompt)
            
            # Call the Traceability Agent to map User Stories to PRD requirements
            logger.info("Processing Traceability Agent to map user stories to PRD requirements")
            
            # Prepare traceability context with user story details and PRD content from vector DB
            traceability_context = f"""
User Story Analysis for Traceability Mapping:

Selected User Stories:
- Story ID: {selected_story_id}
- Story Name: {story_name}
- Story Description: {selected_story_description}
- Epic Title: {epic_title}
- Priority: {priority}
- Systems: {responsible_systems}

Instructions: Map these user stories to the corresponding requirements in the PRD using the vector database. 
Create a traceability matrix showing the relationship between each user story and its source requirements.
"""
            
            try:
                traceability_response = ask_assistant_from_file_optimized("poc2_traceability_agent", traceability_context)
                logger.info(f"Traceability mapping completed, response length: {len(traceability_response)} characters")
            except Exception as e:
                logger.error(f"Traceability agent failed: {e}")
                traceability_response = "Traceability mapping temporarily unavailable. Please try again later."
            
            # Parse the description response
            enhanced_description = selected_story_description  # Default fallback
            try:
                # Clean the response first - remove markdown code blocks if present
                clean_response = description_response.strip()
                
                # Remove markdown JSON code blocks
                if clean_response.startswith('```json') and clean_response.endswith('```'):
                    clean_response = clean_response[7:-3].strip()
                elif clean_response.startswith('```') and clean_response.endswith('```'):
                    clean_response = clean_response[3:-3].strip()
                
                # Try to parse as JSON
                description_data = json.loads(clean_response)
                if isinstance(description_data, dict) and 'description' in description_data:
                    enhanced_description = description_data['description']
                    logger.info(f"Enhanced description generated: {enhanced_description[:100]}...")
                else:
                    logger.warning("User Description Agent did not return expected format")
                    # If it's a dict but no 'description' key, try to get the first value
                    if isinstance(description_data, dict):
                        values = list(description_data.values())
                        if values:
                            enhanced_description = str(values[0])
            except json.JSONDecodeError as e:
                logger.warning(f"Failed to parse User Description Agent response: {e}")
                # Use the raw response if it's not JSON, but try to extract content from markdown
                if description_response and len(description_response.strip()) > 10:
                    clean_response = description_response.strip()
                    # Remove any remaining markdown artifacts
                    if clean_response.startswith('```') and clean_response.endswith('```'):
                        clean_response = clean_response[3:-3].strip()
                    enhanced_description = clean_response

            # Parse criteria (assume response is a bullet list or parse JSON if needed)
            try:
                # First, clean the response to remove markdown code blocks if present
                clean_response = acceptance_response.strip()
                if clean_response.startswith('```json') and clean_response.endswith('```'):
                    clean_response = clean_response[7:-3].strip()
                elif clean_response.startswith('```') and clean_response.endswith('```'):
                    clean_response = clean_response[3:-3].strip()
                
                acceptance_criteria = json.loads(clean_response)
                if isinstance(acceptance_criteria, list):
                    # If it's already a list, use it
                    pass
                elif isinstance(acceptance_criteria, dict) and 'acceptance_criteria' in acceptance_criteria:
                    # If it's a dict with 'acceptance_criteria' key, extract the list
                    acceptance_criteria = acceptance_criteria['acceptance_criteria']
                elif isinstance(acceptance_criteria, dict) and 'criteria' in acceptance_criteria:
                    # If it's a dict with 'criteria' key, extract the list
                    acceptance_criteria = acceptance_criteria['criteria']
                elif isinstance(acceptance_criteria, dict):
                    # If it's a dict but doesn't have expected keys, get the first list value
                    for value in acceptance_criteria.values():
                        if isinstance(value, list):
                            acceptance_criteria = value
                            break
                    else:
                        # If no list found in values, convert all values to list
                        acceptance_criteria = list(acceptance_criteria.values())
                else:
                    # If it's some other format, convert to list
                    acceptance_criteria = [str(acceptance_criteria)]
                
                # Ensure we have a flat list of strings
                if not isinstance(acceptance_criteria, list):
                    acceptance_criteria = [str(acceptance_criteria)]
                    
            except json.JSONDecodeError:
                # If JSON parsing fails, split by newlines and clean up
                acceptance_criteria = [
                    line.strip() 
                    for line in acceptance_response.strip().split("\n") 
                    if line.strip() and not line.strip().startswith('#')
                ]
                # Remove bullet points and numbering
                cleaned_criteria = []
                for criterion in acceptance_criteria:
                    # Remove common bullet point patterns and JSON artifacts
                    criterion = criterion.lstrip('•-*').lstrip('123456789.').strip()
                    # Remove JSON brackets or quotes if they somehow leaked through
                    criterion = criterion.strip('[]{}"\'\n\r\t')
                    if criterion and not criterion.startswith('{') and not criterion.startswith('['):
                        cleaned_criteria.append(criterion)
                acceptance_criteria = cleaned_criteria


            logger.info(f"Processing selected prompt ****Acceptance Criteria **** : {acceptance_criteria}")
            
            # Store selected story details in session for back navigation
            session['selected_story_data'] = {
                'story_id': selected_story_id,
                'story_name': story_name,
                'enhanced_description': enhanced_description,
                'original_description': selected_story_description,
                'epic_title': epic_title,
                'priority': priority,
                'responsible_systems': responsible_systems,
                'acceptance_criteria': acceptance_criteria,
                'tagged_requirements': tagged_requirements,
                'traceability_matrix': traceability_response
            }
            logger.info("Stored selected story details in session with traceability matrix")
            
              # For AJAX requests, return JSON for navigation
            if request.is_json:
                return jsonify({
                    "success": True,
                    "story_id": selected_story_id,
                    "story_name": story_name,
                    "enhanced_description": enhanced_description,
                    "traceability_matrix": traceability_response,
                    "html": render_template(
                        "poc2_user_story_details.html",
                        epic_title=epic_title,
                        user_story_name=story_name,
                        user_story_description=enhanced_description,  # Use enhanced description
                        original_description=selected_story_description,  # Keep original for reference
                        priority=priority,
                        responsible_systems=responsible_systems,
                        acceptance_criteria=acceptance_criteria,
                        tagged_requirements=tagged_requirements,
                        traceability_matrix=traceability_response,
                        story_id=selected_story_id
                    ),
                    "message": "User story processed successfully with enhanced description and traceability mapping"
                })
            
            # Store navigation context in session for back button functionality
            session['user_story_navigation_context'] = {
                'selected_story_ids': story_ids,
                'selected_stories_data': selected_stories_data,
                'epic_title': epic_title,
                'source_epic_data': session.get('current_epics', ''),
                'source_user_stories': session.get('current_user_stories', ''),
                'navigation_timestamp': datetime.now().isoformat()
            }
            session.modified = True
            logger.info(f"Stored navigation context for story IDs: {story_ids}")
            
            # For regular form submission, render the template directly
            return render_template(
                "poc2_user_story_details.html",
                epic_title=epic_title,
                user_story_name=story_name,
                user_story_description=enhanced_description,  # Use enhanced description
                original_description=selected_story_description,  # Keep original for reference
                priority=priority,
                responsible_systems=responsible_systems,
                acceptance_criteria=acceptance_criteria,
                tagged_requirements=tagged_requirements,
                traceability_matrix=traceability_response,
                story_id=selected_story_id
            )
            
        # Handle GET request (direct access or fallback)
        else:
            logger.info("GET request received for user story details")
            
            # Try to get data from session first (backward compatibility)
            session_data = session.get('selected_story_data') or session.get('current_story')
            
            if session_data:
                # Use session data
                logger.info("Using story data from session")
                epic_title = session_data.get('epic_title', '')
                user_story_title = session_data.get('story_name', '')
                user_story_description = session_data.get('enhanced_description') or session_data.get('story_description', '')
                original_description = session_data.get('original_description', user_story_description)
                priority = session_data.get('priority', 'High')
                story_id = session_data.get('story_id', '')
                acceptance_criteria = session_data.get('acceptance_criteria', [])
                responsible_systems = session_data.get('responsible_systems', 'CAPS, CMS')
                tagged_requirements = session_data.get('tagged_requirements', [
                    "TBD"
                ])
                traceability_matrix = session_data.get('traceability_matrix', 'Traceability mapping not available.')
                
                # Clean up session data after use
                session.pop('selected_story_data', None)
                session.pop('current_story', None)
                
            else:
                # Fallback to query parameters
                logger.info("No session data found, checking query parameters")
                epic_title = request.args.get("epic_title", "")
                user_story_title = request.args.get("user_story_title", "")
                user_story_description = request.args.get("user_story_description", "")
                priority = request.args.get("priority", "High")
                story_id = request.args.get("story_id", "")
                
                # Use consistent variable names - with default fallback
                final_story_name = user_story_title or "Lock critical fields for charged-off accounts"
                final_story_description = user_story_description or "No description available"
                
                # Generate acceptance criteria if not in session
                logger.info(f"Processing story: {final_story_name}")
                logger.info(f"Description: {final_story_description[:100]}...")

                # Generate additional data fields for the new UI
                responsible_systems = "CAPS, CMS"  # Default value, could be enhanced later
                tagged_requirements = [
                    "TBD"
                ]

                # Call the acceptance criteria agent - pass only story name
                prompt = final_story_name
                acceptance_response = ask_assistant_from_file_optimized("poc2_agent4_acceptanceCriteria_gen", prompt)
                
                # Call the new User Description Agent - pass only story name
                description_prompt = final_story_name
                logger.info(f"Processing User Description Agent prompt: {description_prompt}")
                
                description_response = ask_assistant_from_file_optimized("poc2_agent5_description", description_prompt)
                
                # Call the Traceability Agent to map User Stories to PRD requirements
                logger.info("Processing Traceability Agent to map user stories to PRD requirements (GET request)")
                
                # Prepare traceability context with user story details and PRD content from vector DB
                traceability_context = f"""
User Story Analysis for Traceability Mapping:

Selected User Stories:
- Story Name: {final_story_name}
- Story Description: {final_story_description}
- Epic Title: {epic_title}
- Priority: {priority}

Instructions: Map these user stories to the corresponding requirements in the PRD using the vector database. 
Create a traceability matrix showing the relationship between each user story and its source requirements.
"""
                
                try:
                    traceability_response = ask_assistant_from_file_optimized("poc2_traceability_agent", traceability_context)
                    logger.info(f"Traceability mapping completed (GET), response length: {len(traceability_response)} characters")
                except Exception as e:
                    logger.error(f"Traceability agent failed (GET): {e}")
                    traceability_response = "Traceability mapping temporarily unavailable. Please try again later."
                
                # Parse the description response
                enhanced_description = final_story_description  # Default fallback
                try:
                    # Clean the response first - remove markdown code blocks if present
                    clean_response = description_response.strip()
                    
                    # Remove markdown JSON code blocks
                    if clean_response.startswith('```json') and clean_response.endswith('```'):
                        clean_response = clean_response[7:-3].strip()
                    elif clean_response.startswith('```') and clean_response.endswith('```'):
                        clean_response = clean_response[3:-3].strip()
                    
                    # Try to parse as JSON
                    description_data = json.loads(clean_response)
                    if isinstance(description_data, dict) and 'description' in description_data:
                        enhanced_description = description_data['description']
                        logger.info(f"Enhanced description generated: {enhanced_description[:100]}...")
                    else:
                        logger.warning("User Description Agent did not return expected format")
                        # If it's a dict but no 'description' key, try to get the first value
                        if isinstance(description_data, dict):
                            values = list(description_data.values())
                            if values:
                                enhanced_description = str(values[0])
                except json.JSONDecodeError as e:
                    logger.warning(f"Failed to parse User Description Agent response: {e}")
                    # Use the raw response if it's not JSON, but try to extract content from markdown
                    if description_response and len(description_response.strip()) > 10:
                        clean_response = description_response.strip()
                        # Remove any remaining markdown artifacts
                        if clean_response.startswith('```') and clean_response.endswith('```'):
                            clean_response = clean_response[3:-3].strip()
                        enhanced_description = clean_response
                  # Parse criteria (assume response is a bullet list or parse JSON if needed)
                try:
                    acceptance_criteria = json.loads(acceptance_response)
                    if isinstance(acceptance_criteria, list):
                        # If it's already a list, use it
                        pass
                    elif isinstance(acceptance_criteria, dict) and 'criteria' in acceptance_criteria:
                        # If it's a dict with criteria key, extract the list
                        acceptance_criteria = acceptance_criteria['criteria']
                    elif isinstance(acceptance_criteria, dict):
                        # If it's a dict but doesn't have 'criteria' key, convert to list of values
                        acceptance_criteria = list(acceptance_criteria.values())
                    else:
                        # If it's some other format, convert to list
                        acceptance_criteria = [str(acceptance_criteria)]
                except:
                    # If JSON parsing fails, split by newlines and clean up
                    acceptance_criteria = [
                        line.strip() 
                        for line in acceptance_response.strip().split("\n") 
                        if line.strip() and not line.strip().startswith('#')
                    ]
                    # Remove bullet points and numbering
                    cleaned_criteria = []
                    for criterion in acceptance_criteria:
                        # Remove common bullet point patterns and JSON artifacts
                        criterion = criterion.lstrip('•-*').lstrip('123456789.').strip()
                        # Remove JSON brackets or quotes if they somehow leaked through
                        criterion = criterion.strip('[]{}"\'\n\r\t')
                        if criterion and not criterion.startswith('{') and not criterion.startswith('['):
                            cleaned_criteria.append(criterion)
                acceptance_criteria = cleaned_criteria
                
                user_story_title = final_story_name
                user_story_description = enhanced_description  # Use enhanced description
                traceability_matrix = traceability_response
        
            return render_template(
                "poc2_user_story_details.html",
                epic_title=epic_title,
                user_story_name=user_story_title,
                user_story_description=user_story_description,
                original_description=original_description if 'original_description' in locals() else (final_story_description if 'final_story_description' in locals() else user_story_description),  # Keep original for reference
                priority=priority,
                responsible_systems=responsible_systems,
                acceptance_criteria=acceptance_criteria,
                tagged_requirements=tagged_requirements,
                traceability_matrix=traceability_matrix if 'traceability_matrix' in locals() else 'Traceability mapping not available.',
                story_id=story_id
            )
        
    except Exception as e:
        logger.error(f"Error in user_story_details: {str(e)}")
        if request.is_json:
            return jsonify({
                "success": False,
                "error": f"Error processing user story details: {str(e)}"
            }), 500
        else:
            return render_template("poc2_user_story_details.html", 
                                 error=f"Error processing user story details: {str(e)}"), 500



@app.route("/epic-chat", methods=["POST"])
def epic_chat():
    """Handle Epic Chat interactions with session persistence and automatic cleanup."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No data provided"}), 400
        
        user_message = data.get("message", "").strip()
        if not user_message:
            return jsonify({"success": False, "error": "Message cannot be empty"}), 400
        
        # Initialize chat history in session if not exists or if expired
        current_time = datetime.now()
        chat_session_timeout = 30 * 60  # 30 minutes in seconds
        
        if 'epic_chat_history' not in session:
            session['epic_chat_history'] = []
            session['epic_chat_last_activity'] = current_time.isoformat()
            logger.info("Initialized new Epic Chat session")
        else:
            # Check if session has expired (30 minutes of inactivity)
            last_activity_str = session.get('epic_chat_last_activity')
            if last_activity_str:
                try:
                    last_activity = datetime.fromisoformat(last_activity_str)
                    time_diff = (current_time - last_activity).total_seconds()
                    
                    if time_diff > chat_session_timeout:
                        # Clear expired session
                        session['epic_chat_history'] = []
                        logger.info(f"Cleared expired Epic Chat session (inactive for {time_diff/60:.1f} minutes)")
                except Exception as e:
                    logger.warning(f"Error parsing last activity time: {e}")
                    session['epic_chat_history'] = []
        
        # Limit chat history to last 50 messages to prevent session bloat
        if len(session['epic_chat_history']) > 50:
            session['epic_chat_history'] = session['epic_chat_history'][-40:]  # Keep last 40 messages
            logger.info("Trimmed Epic Chat history to prevent session bloat")
        
        # Update last activity time
        session['epic_chat_last_activity'] = current_time.isoformat()
        
        # Add user message to chat history
        session['epic_chat_history'].append({
            "role": "user",
            "message": user_message,
            "timestamp": current_time.strftime("%Y-%m-%d %H:%M:%S")
        })
        
        # Get current epics from the frontend request (this is the displayed epics context)
        current_epics = data.get('current_epics_content', '').strip()
        
        # Log the current epics context for debugging
        logger.info(f"Epic Chat received current epics context: {current_epics[:500]}...")
        
        # Format current epics for better context understanding
        epics_context = current_epics
        if current_epics:
            try:
                # Try to parse as JSON first (structured format from frontend)
                import json
                epics_data = json.loads(current_epics)
                if isinstance(epics_data, list):
                    epics_context = "CURRENT EPICS (JSON FORMAT):\n" + json.dumps(epics_data, indent=2)
                else:
                    epics_context = current_epics
            except:
                # If not valid JSON, use as-is (might be HTML content)
                epics_context = "CURRENT EPICS (HTML/TEXT FORMAT):\n" + current_epics
        
        # Prepare context for Epic Agent with focus on updating existing epics
        chat_context = f"""You are an Epic Agent specializing in refining and improving existing epics. Your primary role is to help users update, correct, and enhance their current epics.

{epics_context if epics_context else 'No epics currently defined. Please ask the user to create epics first through the main workflow.'}

CRITICAL INSTRUCTIONS - READ CAREFULLY:
- DEFAULT BEHAVIOR: Always focus on modifying/updating/correcting the EXISTING epics shown above
- NEVER create new epics unless the user EXPLICITLY uses clear phrases like:
  * "add new epic"
  * "create additional epic" 
  * "new epic"
  * "more epics"
  * "another epic"
  * "additional epic"
- When users ask questions, give feedback, or make requests, they want to CHANGE/UPDATE the existing epics, NOT add new ones
- Provide UPDATED versions of the existing epics that incorporate the user's feedback
- Keep the same epic IDs and structure when updating
- If no current epics exist, guide user to the main workflow first

COMMON USER REQUESTS (all should UPDATE existing epics):
- "Make the epics more detailed" → Update existing epics with more detail
- "Add acceptance criteria" → Update existing epics to include acceptance criteria
- "Change the priority" → Update existing epic priorities
- "The epic needs more clarity" → Update existing epic descriptions
- "Fix the requirements" → Update existing epic requirements
- "Improve the user stories" → Update existing epic user story breakdown
- "These need work" → Update existing epics based on feedback

OUTPUT FORMAT REQUIREMENTS:
When providing updated epics, use this EXACT format:

```json
[
  {{
    "epic_id": "epic_1",
    "epic_title": "Updated Epic Title",
    "epic_description": "Updated detailed description of what this epic covers and accomplishes"
  }},
  {{
    "epic_id": "epic_2", 
    "epic_title": "Another Updated Epic Title",
    "epic_description": "Updated detailed description for the second epic"
  }}
]
```

IMPORTANT: 
- Always wrap the JSON in ```json code blocks
- Use "epic_id", "epic_title", and "epic_description" as field names
- Preserve existing epic_id values when updating
- If just providing conversational response without updates, don't include JSON

RESPOND WITH: Updated versions of the EXISTING epics using the JSON format above, or conversational response if no updates are needed.

"""
        if session['epic_chat_history']:
            chat_context += "Previous conversation:\n"
            for msg in session['epic_chat_history'][-10:]:  # Last 10 messages for context
                role = "User" if msg["role"] == "user" else "Assistant"
                chat_context += f"{role}: {msg['message']}\n"
            chat_context += f"\nLatest user question: {user_message}"
        else:
            chat_context += f"User request: {user_message}"
        
        # Call Epic Agent
        logger.info(f"Calling Epic Agent with message: {user_message[:100]}...")
        logger.info(f"Epic Agent context length: {len(chat_context)} characters")
        epic_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", chat_context)
        
        logger.info(f"Epic Agent response length: {len(epic_response)} characters")
        logger.info(f"Epic Agent response preview: {epic_response[:200]}...")
        
        if epic_response.startswith("Error:"):
            logger.error(f"Epic Agent error: {epic_response}")
            return jsonify({
                "success": False,
                "error": "Failed to get response from Epic Agent"
            }), 500
        
        # Check if the response contains structured epic data (JSON format)
        epics_updated = False
        try:
            # Try to detect if response contains JSON with epic data
            import re
            json_match = re.search(r'```json\s*([\s\S]*?)\s*```', epic_response, re.MULTILINE)
            if json_match:
                json_content = json_match.group(1).strip()
                epic_data = json.loads(json_content)
                if isinstance(epic_data, list) and len(epic_data) > 0:
                    # Check if it looks like epic data
                    first_item = epic_data[0]
                    if isinstance(first_item, dict) and any(key in first_item for key in ['title', 'name', 'description', 'epic_id']):
                        epics_updated = True
                        logger.info("Epic Chat: Detected updated epic data in response")
            else:
                # Try parsing the entire response as JSON
                try:
                    epic_data = json.loads(epic_response.strip())
                    if isinstance(epic_data, list) and len(epic_data) > 0:
                        first_item = epic_data[0]
                        if isinstance(first_item, dict) and any(key in first_item for key in ['title', 'name', 'description', 'epic_id']):
                            epics_updated = True
                            logger.info("Epic Chat: Detected epic data in plain JSON response")
                except:
                    pass
        except Exception as e:
            logger.debug(f"Epic detection parsing error (normal): {e}")
        
        # Add assistant response to chat history
        session['epic_chat_history'].append({
            "role": "assistant",
            "message": epic_response,
            "timestamp": current_time.strftime("%Y-%m-%d %H:%M:%S")
        })
        
        # Update last activity time again
        session['epic_chat_last_activity'] = current_time.isoformat()
        
        # Save session changes
        session.modified = True
        
        logger.info(f"Epic Chat response generated successfully, history length: {len(session['epic_chat_history'])}")
        
        return jsonify({
            "success": True,
            "response": epic_response,
            "epics_updated": epics_updated,
            "chat_history": session['epic_chat_history']
        })
        
    except Exception as e:
        logger.error(f"Error in Epic Chat: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/epic-chat-history", methods=["GET"])
def get_epic_chat_history():
    """Get the current Epic Chat history from session."""
    try:
        chat_history = session.get('epic_chat_history', [])
        return jsonify({
            "success": True,
            "chat_history": chat_history
        })
    except Exception as e:
        logger.error(f"Error getting Epic Chat history: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/epic-chat-clear", methods=["POST"])
def clear_epic_chat():
    """Clear the Epic Chat history from session."""
    try:
        if 'epic_chat_history' in session:
            del session['epic_chat_history']
        if 'epic_chat_last_activity' in session:
            del session['epic_chat_last_activity']
        session.modified = True
        logger.info("Epic Chat history and activity timestamp cleared")
        return jsonify({"success": True, "message": "Chat history cleared"})
    except Exception as e:
        logger.error(f"Error clearing Epic Chat history: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/user-story-chat", methods=["POST"])
def user_story_chat():
    """Handle User Story Chat interactions with system mapping functionality."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No data provided"}), 400
        
        user_message = data.get("message", "").strip()
        if not user_message:
            return jsonify({"success": False, "error": "Message cannot be empty"}), 400
        
        # Initialize chat history in session if not exists or if expired
        current_time = datetime.now()
        chat_session_timeout = 30 * 60  # 30 minutes in seconds
        
        if 'user_story_chat_history' not in session:
            session['user_story_chat_history'] = []
            session['user_story_chat_last_activity'] = current_time.isoformat()
            logger.info("Initialized new User Story Chat session")
        else:
            # Check if session has expired (30 minutes of inactivity)
            last_activity_str = session.get('user_story_chat_last_activity')
            if last_activity_str:
                try:
                    last_activity = datetime.fromisoformat(last_activity_str)
                    time_diff = (current_time - last_activity).total_seconds()
                    
                    if time_diff > chat_session_timeout:
                        # Clear expired session
                        session['user_story_chat_history'] = []
                        logger.info(f"Cleared expired User Story Chat session (inactive for {time_diff/60:.1f} minutes)")
                except Exception as e:
                    logger.warning(f"Error parsing last activity time: {e}")
                    session['user_story_chat_history'] = []
        
        # Limit chat history to last 50 messages to prevent session bloat
        if len(session['user_story_chat_history']) > 50:
            session['user_story_chat_history'] = session['user_story_chat_history'][-40:]  # Keep last 40 messages
            logger.info("Trimmed User Story Chat history to prevent session bloat")
        
        # Update last activity time
        session['user_story_chat_last_activity'] = current_time.isoformat()
        
        # Add user message to chat history
        session['user_story_chat_history'].append({
            "role": "user",
            "message": user_message,
            "timestamp": current_time.strftime("%Y-%m-%d %H:%M:%S")
        })
        
        # Get current user stories and epic context from request
        current_user_stories_content = data.get("current_user_stories_content", "")
        epic_context = data.get("epic_context", "")
        
        # Get system information from session (if available)
        system_info = session.get('system_info', '')
        
        # Prepare context for User Story Agent with current context
        chat_context = f"""You are a User Story Agent helping to create, refine, and map user stories to specific systems.

CURRENT USER STORIES CONTEXT:
{current_user_stories_content if current_user_stories_content else 'No current user stories available.'}

RELATED EPICS:
{epic_context if epic_context else 'No epic context available.'}

SYSTEM INFORMATION:
{system_info if system_info else 'No system information provided. User stories will not include system mapping.'}

INSTRUCTIONS:
- When refining or updating user stories, base your changes on the CURRENT USER STORIES CONTEXT provided above
- Maintain the existing user story IDs and structure when making improvements
- If creating new user stories, ensure they align with the related epics
- If system information is available, map each user story to the appropriate system based on the functionality described
- Include a "System Name" for each user story when system mapping is possible
- Format user stories with: Story ID, Title, Description, Acceptance Criteria, Priority, and System Name (if applicable)
- Focus on creating clear, actionable user stories that deliver business value
- When returning updated user stories, provide them in JSON format that matches the existing structure

"""
        if session['user_story_chat_history']:
            chat_context += "Previous conversation:\n"
            for msg in session['user_story_chat_history'][-10:]:  # Last 10 messages for context
                role = "User" if msg["role"] == "user" else "Assistant"
                chat_context += f"{role}: {msg['message']}\n"
            chat_context += f"\nLatest user question: {user_message}"
        else:
            chat_context += f"User request: {user_message}"
        
        # Determine if this is a request for user story updates or just conversation
        is_update_request = any(keyword in user_message.lower() for keyword in [
            'update', 'modify', 'change', 'improve', 'refine', 'edit', 'enhance', 
            'fix', 'adjust', 'revise', 'generate', 'create', 'add'
        ])
        
        if is_update_request and current_user_stories_content and 'No user stories' not in current_user_stories_content:
            # This is a request to update existing user stories - use the structured agent
            chat_context += f"\nUser Request: {user_message}\n\nIMPORTANT: Return the updated user stories as a JSON array using this exact format:\n"
            chat_context += """[
  {
    "story_id": "existing_or_new_id",
    "name": "Story name",
    "description": "Detailed description",
    "priority": "High/Medium/Low",
    "systems": ["System1", "System2"]
  }
]

Include ALL existing user stories with any requested changes, plus any new ones if requested."""
            
            agent_to_use = "poc2_agent3_basic_user_story"
        else:
            # This is a conversational request - use a general assistant approach
            chat_context += f"\nUser Request: {user_message}\n\nProvide helpful conversational guidance about user story best practices, suggestions, or answer their question. Be conversational and helpful."
            agent_to_use = "poc2_agent2_epic_generator"  # Use epic generator as it's more conversational
        
        # Call appropriate agent
        logger.info(f"Calling User Story Agent ({agent_to_use}) with message: {user_message[:100]}...")
        logger.info(f"Is update request: {is_update_request}, Has current stories: {bool(current_user_stories_content and 'No user stories' not in current_user_stories_content)}")
        story_response = ask_assistant_from_file_optimized(agent_to_use, chat_context)
        
        if story_response.startswith("Error:"):
            logger.error(f"User Story Agent error: {story_response}")
            return jsonify({
                "success": False,
                "error": "Failed to get response from User Story Agent"
            }), 500
        
        # Add assistant response to chat history
        session['user_story_chat_history'].append({
            "role": "assistant",
            "message": story_response,
            "timestamp": current_time.strftime("%Y-%m-%d %H:%M:%S")
        })
        
        # Update last activity time again
        session['user_story_chat_last_activity'] = current_time.isoformat()
        
        # Save session changes
        session.modified = True
        
        logger.info(f"User Story Chat response generated successfully, history length: {len(session['user_story_chat_history'])}")
        
        return jsonify({
            "success": True,
            "response": story_response,
            "chat_history": session['user_story_chat_history'],
            "has_system_info": bool(system_info)
        })
        
    except Exception as e:
        logger.error(f"Error in User Story Chat: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/user-story-chat-history", methods=["GET"])
def get_user_story_chat_history():
    """Get the current User Story Chat history from session."""
    try:
        chat_history = session.get('user_story_chat_history', [])
        system_info = session.get('system_info', '')
        return jsonify({
            "success": True,
            "chat_history": chat_history,
            "has_system_info": bool(system_info)
        })
    except Exception as e:
        logger.error(f"Error getting User Story Chat history: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/user-story-chat-clear", methods=["POST"])
def clear_user_story_chat():
    """Clear the User Story Chat history from session."""
    try:
        if 'user_story_chat_history' in session:
            del session['user_story_chat_history']
        if 'user_story_chat_last_activity' in session:
            del session['user_story_chat_last_activity']
        session.modified = True
        logger.info("User Story Chat history and activity timestamp cleared")
        return jsonify({"success": True, "message": "Chat history cleared"})
    except Exception as e:
        logger.error(f"Error clearing User Story Chat history: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/upload-system-info", methods=["POST"])
def upload_system_info():
    """Upload and process system information file for user story mapping."""
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"success": False, "error": "No file selected"}), 400
        
        # Check file size (limit to 10MB)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > 10 * 1024 * 1024:  # 10MB limit
            return jsonify({"success": False, "error": "File size too large (max 10MB)"}), 400
        
        # Read and process file content
        content = ""
        filename = file.filename.lower()
        
        try:
            if filename.endswith('.txt'):
                content = file.read().decode('utf-8')
            elif filename.endswith('.json'):
                json_data = json.load(file)
                content = json.dumps(json_data, indent=2)
            elif filename.endswith('.csv'):
                content = file.read().decode('utf-8')
            else:
                # Try to read as text for other formats
                content = file.read().decode('utf-8')
        except UnicodeDecodeError:
            return jsonify({"success": False, "error": "File encoding not supported. Please use UTF-8 encoded files."}), 400
        except json.JSONDecodeError:
            return jsonify({"success": False, "error": "Invalid JSON file format"}), 400
        
        if not content.strip():
            return jsonify({"success": False, "error": "File appears to be empty"}), 400
        
        # Store system information in session
        session['system_info'] = content
        session['system_info_filename'] = file.filename
        session['system_info_uploaded'] = datetime.now().isoformat()
        session.modified = True
        
        # Parse basic system information for preview
        lines = content.split('\n')[:10]  # First 10 lines for preview
        preview = '\n'.join(lines)
        if len(content.split('\n')) > 10:
            preview += "\n... (file continues)"
        
        logger.info(f"System information uploaded: {file.filename} ({file_size} bytes)")
        
        return jsonify({
            "success": True,
            "message": f"System information uploaded successfully: {file.filename}",
            "filename": file.filename,
            "size": file_size,
            "preview": preview[:500]  # Limit preview to 500 chars
        })
        
    except Exception as e:
        logger.error(f"Error uploading system info: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "Failed to upload system information"}), 500

@app.route("/get-system-info", methods=["GET"])
def get_system_info():
    """Get current system information status."""
    try:
        system_info = session.get('system_info', '')
        filename = session.get('system_info_filename', '')
        uploaded_time = session.get('system_info_uploaded', '')
        
        return jsonify({
            "success": True,
            "has_system_info": bool(system_info),
            "filename": filename,
            "uploaded_time": uploaded_time,
            "preview": system_info[:500] if system_info else ""
        })
    except Exception as e:
        logger.error(f"Error getting system info: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/clear-system-info", methods=["POST"])
def clear_system_info():
    """Clear uploaded system information."""
    try:
        if 'system_info' in session:
            del session['system_info']
        if 'system_info_filename' in session:
            del session['system_info_filename']
        if 'system_info_uploaded' in session:
            del session['system_info_uploaded']
        session.modified = True
        
        logger.info("System information cleared from session")
        return jsonify({"success": True, "message": "System information cleared"})
    except Exception as e:
        logger.error(f"Error clearing system info: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/chat-description", methods=["POST"])
def chat_description():
    """Handle chat requests for description refinement."""
    logger.info("Request received for description chat")
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No JSON data provided"}), 400
            
        user_message = data.get("message", "").strip()
        current_description = data.get("current_description", "").strip()
        
        if not user_message:
            return jsonify({"success": False, "error": "Message is required"}), 400
            
        if not current_description:
            return jsonify({"success": False, "error": "Current description is required"}), 400
        
        logger.info(f"Processing chat request: {user_message[:100]}...")
        
        # Create a prompt for the AI to refine the description
        system_prompt = """You are an expert business analyst helping to refine user story descriptions. 
        Based on the user's request, improve the given description while maintaining its core meaning.
        
        Guidelines:
        - Keep the description clear and actionable
        - Maintain focus on user value and business outcomes
        - Use professional but accessible language
        - Ensure the description supports downstream development activities
        - Be specific about what the user wants to achieve
        
        Return only the improved description without additional commentary."""
        
        user_prompt = f"""Current Description: {current_description}

User Request: {user_message}

Please provide an improved version of the description based on the user's request."""

        # Use OpenAI to generate the improved description
        try:
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=500,
                temperature=0.7
            )
            
            improved_description = response.choices[0].message.content.strip()
            
            logger.info("Successfully generated improved description")
            
            return jsonify({
                "success": True,
                "message": "I've improved the description based on your request.",
                "updated_description": improved_description
            })
            
        except Exception as ai_error:
            logger.error(f"OpenAI API error: {str(ai_error)}")
            
            # Fallback to rule-based processing
            fallback_response = process_description_fallback(user_message, current_description)
            return jsonify({
                "success": True,
                "message": fallback_response["message"],
                "updated_description": fallback_response["updated_description"]
            })
            
    except Exception as e:
        logger.error(f"Error in description chat: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

def process_description_fallback(user_message, current_description):
    """Fallback processing when AI is not available."""
    lowerMessage = user_message.lower()
    
    if any(word in lowerMessage for word in ['shorter', 'concise', 'brief']):
        # Make it more concise
        sentences = current_description.split('.')
        shortened = '. '.join(sentences[:2]).strip()
        if not shortened.endswith('.'):
            shortened += '.'
        return {
            "message": "I've made the description more concise while keeping the key information.",
            "updated_description": shortened
        }
    elif any(word in lowerMessage for word in ['longer', 'detailed', 'elaborate', 'expand']):
        # Make it more detailed
        expanded = current_description + " This functionality ensures data accuracy and supports downstream processes including identity verification, credit checks, and compliance reporting. The system should provide clear feedback for any validation errors and guide users through successful completion."
        return {
            "message": "I've expanded the description with more detail and context.",
            "updated_description": expanded
        }
    elif any(word in lowerMessage for word in ['technical', 'specific', 'implementation']):
        # Add technical details
        technical = current_description + " The system should implement real-time validation using industry-standard APIs, secure data encryption during transmission, and integration with existing customer management systems (CMS) and credit assessment platform (CAPS)."
        return {
            "message": "I've added more technical details and specific requirements.",
            "updated_description": technical
        }
    elif any(word in lowerMessage for word in ['business', 'user focused', 'value']):
        # Add business context
        business = current_description + " This feature directly supports our customer acquisition strategy by reducing onboarding friction while ensuring regulatory compliance. Accurate data capture at this stage prevents costly corrections later and improves overall customer experience."
        return {
            "message": "I've focused more on the business value and user perspective.",
            "updated_description": business
        }
    else:
        return {
            "message": "I understand you want to modify the description. Could you be more specific about what changes you'd like? For example: 'Make it shorter', 'Add more technical details', 'Focus on business value', etc.",
            "updated_description": current_description
        }

@app.route("/chat-acceptance-criteria", methods=["POST"])
def chat_acceptance_criteria():
    """Handle chat requests for acceptance criteria refinement."""
    logger.info("Request received for acceptance criteria chat")
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No JSON data provided"}), 400
            
        user_message = data.get("message", "").strip()
        current_criteria = data.get("current_criteria", [])
        story_context = data.get("story_context", "").strip()
        
        if not user_message:
            return jsonify({"success": False, "error": "Message is required"}), 400
        
        logger.info(f"Processing acceptance criteria chat request: {user_message[:100]}...")
        
        # Format current criteria for AI context
        criteria_text = ""
        if current_criteria and isinstance(current_criteria, list):
            criteria_text = "\n".join([f"• {criterion}" for criterion in current_criteria])
        elif isinstance(current_criteria, str):
            criteria_text = current_criteria
        
        # Create a prompt for the AI to refine acceptance criteria
        system_prompt = """You are an expert business analyst specializing in acceptance criteria for user stories. 
        Your role is to help refine, improve, and enhance acceptance criteria based on user requests.
        
        Guidelines for good acceptance criteria:
        - Use clear, specific, and testable language
        - Follow the Given/When/Then format when appropriate
        - Focus on user behavior and system responses
        - Include both positive and negative test cases
        - Ensure criteria are measurable and verifiable
        - Address security, performance, and accessibility when relevant
        - Keep criteria independent and atomic
        
        Return the improved acceptance criteria as a JSON array of strings.
        Format: ["criterion 1", "criterion 2", "criterion 3"]"""
        
        user_prompt = f"""User Story Context: {story_context}

Current Acceptance Criteria:
{criteria_text}

User Request: {user_message}

Please provide improved acceptance criteria based on the user's request. Return as a JSON array of strings."""

        # Use OpenAI to generate improved acceptance criteria
        try:
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=800,
                temperature=0.7
            )
            
            ai_response = response.choices[0].message.content.strip()
            
            # Try to parse the JSON response
            try:
                import json
                improved_criteria = json.loads(ai_response)
                if not isinstance(improved_criteria, list):
                    # If not a list, try to extract from the response
                    lines = ai_response.split('\n')
                    improved_criteria = [line.strip('• -').strip() for line in lines if line.strip() and not line.startswith('[') and not line.startswith(']')]
            except json.JSONDecodeError:
                # Fallback: parse as bullet points
                lines = ai_response.split('\n')
                improved_criteria = [line.strip('• -').strip() for line in lines if line.strip() and not line.startswith('[') and not line.startswith(']')]
            
            logger.info("Successfully generated improved acceptance criteria")
            
            return jsonify({
                "success": True,
                "message": "I've improved the acceptance criteria based on your request.",
                "updated_criteria": improved_criteria
            })
            
        except Exception as ai_error:
            logger.error(f"OpenAI API error: {str(ai_error)}")
            
            # Fallback to rule-based processing
            fallback_response = process_criteria_fallback(user_message, current_criteria, story_context)
            return jsonify({
                "success": True,
                "message": fallback_response["message"],
                "updated_criteria": fallback_response["updated_criteria"]
            })
            
    except Exception as e:
        logger.error(f"Error in acceptance criteria chat: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

def process_criteria_fallback(user_message, current_criteria, story_context):
    """Fallback processing for acceptance criteria when AI is not available."""
    lowerMessage = user_message.lower()
    
    # Ensure current_criteria is a list
    if isinstance(current_criteria, str):
        current_criteria = [current_criteria] if current_criteria.strip() else []
    elif not isinstance(current_criteria, list):
        current_criteria = []
    
    if any(word in lowerMessage for word in ['more detailed', 'specific', 'elaborate']):
        # Make criteria more detailed
        enhanced_criteria = []
        for criterion in current_criteria:
            enhanced_criteria.append(criterion)
            if 'input' in criterion.lower() or 'enter' in criterion.lower():
                enhanced_criteria.append("System validates input format in real-time")
                enhanced_criteria.append("Clear error messages are displayed for invalid inputs")
        return {
            "message": "I've made the acceptance criteria more detailed and specific.",
            "updated_criteria": enhanced_criteria
        }
    elif any(word in lowerMessage for word in ['security', 'secure', 'validation']):
        # Add security-focused criteria
        security_criteria = current_criteria.copy()
        security_criteria.extend([
            "All data inputs are validated and sanitized",
            "System logs security-relevant events",
            "Unauthorized access attempts are blocked and logged",
            "Data is encrypted during transmission and storage"
        ])
        return {
            "message": "I've added security-focused acceptance criteria.",
            "updated_criteria": security_criteria
        }
    elif any(word in lowerMessage for word in ['performance', 'speed', 'fast']):
        # Add performance criteria
        performance_criteria = current_criteria.copy()
        performance_criteria.extend([
            "System responds within 2 seconds for normal operations",
            "Page load time does not exceed 3 seconds",
            "System handles concurrent users without degradation"
        ])
        return {
            "message": "I've added performance-related acceptance criteria.",
            "updated_criteria": performance_criteria
        }
    elif any(word in lowerMessage for word in ['accessibility', 'a11y', 'screen reader']):
        # Add accessibility criteria
        accessibility_criteria = current_criteria.copy()
        accessibility_criteria.extend([
            "Interface is compatible with screen readers",
            "All interactive elements are keyboard accessible",
            "Color contrast meets WCAG 2.1 AA standards",
            "Form fields have proper labels and descriptions"
        ])
        return {
            "message": "I've added accessibility-focused acceptance criteria.",
            "updated_criteria": accessibility_criteria
        }
    elif any(word in lowerMessage for word in ['given when then', 'gherkin', 'format']):
        # Convert to Given/When/Then format
        formatted_criteria = []
        for criterion in current_criteria:
            if not criterion.startswith(('Given', 'When', 'Then')):
                formatted_criteria.append(f"Given a user wants to {criterion.lower()}")
                formatted_criteria.append(f"When they perform the required action")
                formatted_criteria.append(f"Then the system should {criterion.lower()}")
            else:
                formatted_criteria.append(criterion)
        return {
            "message": "I've formatted the acceptance criteria using Given/When/Then structure.",
            "updated_criteria": formatted_criteria
        }
    elif any(word in lowerMessage for word in ['shorter', 'concise', 'simplify']):
        # Simplify criteria
        simplified_criteria = []
        for criterion in current_criteria:
            # Keep only essential parts
            simplified = criterion.split('.')[0] if '.' in criterion else criterion
            simplified_criteria.append(simplified.strip())
        return {
            "message": "I've simplified the acceptance criteria to be more concise.",
            "updated_criteria": list(set(simplified_criteria))  # Remove duplicates
        }
    else:
        return {
            "message": "I understand you want to modify the acceptance criteria. Could you be more specific? For example: 'Make them more detailed', 'Add security requirements', 'Include performance criteria', 'Use Given/When/Then format', etc.",
            "updated_criteria": current_criteria
        }

@app.route("/chat-requirements", methods=["POST"])
def chat_requirements():
    """Handle chat requests for tagged requirements refinement."""
    logger.info("Request received for tagged requirements chat")
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No JSON data provided"}), 400
            
        user_message = data.get("message", "").strip()
        current_requirements = data.get("current_requirements", [])
        story_context = data.get("story_context", "").strip()
        
        if not user_message:
            return jsonify({"success": False, "error": "Message is required"}), 400
        
        logger.info(f"Processing tagged requirements chat request: {user_message[:100]}...")
        
        # Format current requirements for AI context
        requirements_text = ""
        if current_requirements and isinstance(current_requirements, list):
            requirements_text = "\n".join([f"• {req}" for req in current_requirements])
        elif isinstance(current_requirements, str):
            requirements_text = current_requirements
        
        # Create a prompt for the AI to refine tagged requirements
        system_prompt = """You are an expert business analyst specializing in tagged requirements for user stories. 
        Tagged requirements typically include regulatory, compliance, technical, and business requirements that support the user story.
        
        Guidelines for good tagged requirements:
        - Include relevant regulatory and compliance requirements
        - Add technical constraints and dependencies
        - Specify security and privacy requirements
        - Include performance and scalability requirements
        - Address integration and interoperability needs
        - Consider audit, logging, and monitoring requirements
        - Include accessibility and usability standards
        
        Return the improved tagged requirements as a JSON array of strings.
        Format: ["requirement 1", "requirement 2", "requirement 3"]"""
        
        user_prompt = f"""User Story Context: {story_context}

Current Tagged Requirements:
{requirements_text}

User Request: {user_message}

Please provide improved tagged requirements based on the user's request. Return as a JSON array of strings."""

        # Use OpenAI to generate improved tagged requirements
        try:
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=800,
                temperature=0.7
            )
            
            ai_response = response.choices[0].message.content.strip()
            
            # Try to parse the JSON response
            try:
                import json
                improved_requirements = json.loads(ai_response)
                if not isinstance(improved_requirements, list):
                    # If not a list, try to extract from the response
                    lines = ai_response.split('\n')
                    improved_requirements = [line.strip('• -').strip() for line in lines if line.strip() and not line.startswith('[') and not line.startswith(']')]
            except json.JSONDecodeError:
                # Fallback: parse as bullet points
                lines = ai_response.split('\n')
                improved_requirements = [line.strip('• -').strip() for line in lines if line.strip() and not line.startswith('[') and not line.startswith(']')]
            
            logger.info("Successfully generated improved tagged requirements")
            
            return jsonify({
                "success": True,
                "message": "I've improved the tagged requirements based on your request.",
                "updated_requirements": improved_requirements
            })
            
        except Exception as ai_error:
            logger.error(f"OpenAI API error: {str(ai_error)}")
            
            # Fallback to rule-based processing
            fallback_response = process_requirements_fallback(user_message, current_requirements, story_context)
            return jsonify({
                "success": True,
                "message": fallback_response["message"],
                "updated_requirements": fallback_response["updated_requirements"]
            })
            
    except Exception as e:
        logger.error(f"Error in tagged requirements chat: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

def process_requirements_fallback(user_message, current_requirements, story_context):
    """Fallback processing for tagged requirements when AI is not available."""
    lowerMessage = user_message.lower()
    
    # Ensure current_requirements is a list
    if isinstance(current_requirements, str):
        current_requirements = [current_requirements] if current_requirements.strip() else []
    elif not isinstance(current_requirements, list):
        current_requirements = []
    
    if any(word in lowerMessage for word in ['compliance', 'regulatory', 'regulation']):
        # Add compliance requirements
        compliance_requirements = current_requirements.copy()
        compliance_requirements.extend([
            "Must comply with GDPR data protection requirements",
            "SOX compliance for financial data handling",
            "PCI DSS compliance for payment card data",
            "Audit trail must be maintained for all transactions"
        ])
        return {
            "message": "I've added compliance and regulatory requirements.",
            "updated_requirements": compliance_requirements
        }
    elif any(word in lowerMessage for word in ['security', 'secure', 'encryption']):
        # Add security requirements
        security_requirements = current_requirements.copy()
        security_requirements.extend([
            "All data must be encrypted in transit and at rest",
            "Multi-factor authentication required for sensitive operations",
            "Security audit logging must be implemented",
            "Role-based access control must be enforced"
        ])
        return {
            "message": "I've added security-focused requirements.",
            "updated_requirements": security_requirements
        }
    elif any(word in lowerMessage for word in ['performance', 'scalability', 'load']):
        # Add performance requirements
        performance_requirements = current_requirements.copy()
        performance_requirements.extend([
            "System must support 1000+ concurrent users",
            "Response time must be under 2 seconds",
            "99.9% uptime availability required",
            "Horizontal scaling capability must be supported"
        ])
        return {
            "message": "I've added performance and scalability requirements.",
            "updated_requirements": performance_requirements
        }
    elif any(word in lowerMessage for word in ['integration', 'api', 'interoperability']):
        # Add integration requirements
        integration_requirements = current_requirements.copy()
        integration_requirements.extend([
            "RESTful API endpoints must be provided",
            "Real-time data synchronization with existing systems",
            "Backward compatibility with legacy systems",
            "Standard data formats (JSON/XML) must be supported"
        ])
        return {
            "message": "I've added integration and interoperability requirements.",
            "updated_requirements": integration_requirements
        }
    elif any(word in lowerMessage for word in ['accessibility', 'wcag', 'a11y']):
        # Add accessibility requirements
        accessibility_requirements = current_requirements.copy()
        accessibility_requirements.extend([
            "WCAG 2.1 Level AA compliance required",
            "Screen reader compatibility must be ensured",
            "Keyboard navigation support required",
            "Color contrast standards must be met"
        ])
        return {
            "message": "I've added accessibility requirements.",
            "updated_requirements": accessibility_requirements
        }
    elif any(word in lowerMessage for word in ['audit', 'logging', 'monitoring']):
        # Add audit and monitoring requirements
        audit_requirements = current_requirements.copy()
        audit_requirements.extend([
            "Comprehensive audit logging must be implemented",
            "Real-time monitoring and alerting required",
            "Data retention policies must be enforced",
            "Change tracking for all data modifications"
        ])
        return {
            "message": "I've added audit, logging, and monitoring requirements.",
            "updated_requirements": audit_requirements
        }
    else:
        # Generic enhancement
        if not current_requirements or current_requirements == ["TBD"]:
            generic_requirements = [
                "Data validation and sanitization required",
                "Error handling and user feedback mechanisms",
                "Performance monitoring and logging",
                "Security best practices implementation"
            ]
        else:
            generic_requirements = current_requirements.copy()
            generic_requirements.append("Additional requirements to be defined based on detailed analysis")
        
        return {
            "message": "I understand you want to enhance the tagged requirements. Could you be more specific? For example: 'Add compliance requirements', 'Include security standards', 'Add performance requirements', etc.",
            "updated_requirements": generic_requirements
        }

@app.route("/chat-user-story", methods=["POST"])
def chat_user_story():
    """Handle chat requests for user story refinement and context management."""
    logger.info("Request received for user story chat")
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No JSON data provided"}), 400
            
        user_message = data.get("message", "").strip()
        context = data.get("context", {})
        
        if not user_message:
            return jsonify({"success": False, "error": "Message is required"}), 400
        
        # Extract context information
        user_story_name = context.get("user_story_name", "")
        user_story_description = context.get("user_story_description", "")
        epic_title = context.get("epic_title", "")
        acceptance_criteria = context.get("acceptance_criteria", [])
        tagged_requirements = context.get("tagged_requirements", [])
        
        logger.info(f"User story chat message: {user_message}")
        logger.info(f"Context - Story: {user_story_name}")
        
        # Build context string for AI
        context_text = f"""
Current User Story Context:
- Story Name: {user_story_name or 'Not specified'}
- Description: {user_story_description or 'Not specified'}
- Epic: {epic_title or 'Not specified'}
- Acceptance Criteria: {len(acceptance_criteria)} items
- Tagged Requirements: {len(tagged_requirements)} items
"""
        
        if acceptance_criteria:
            context_text += "\nAcceptance Criteria:\n" + "\n".join([f"• {criterion}" for criterion in acceptance_criteria])
        
        if tagged_requirements:
            context_text += "\nTagged Requirements:\n" + "\n".join([f"• {req}" for req in tagged_requirements])
        
        # System prompt for user story chat
        system_prompt = f"""You are an expert business analyst and product manager specializing in user story development and refinement. You help improve user stories by making them more clear, specific, and valuable.

{context_text}

IMPORTANT CONTEXT AWARENESS:
- This is a single user story details page, not a multi-story management interface
- When users mention removing "US-5" or similar story IDs, they are likely confused about the interface
- Always maintain context of the current user story: "{user_story_name}"
- If asked to remove stories or story IDs, clarify that this page is for refining ONE specific story

Your role is to:
1. Help refine the user story name and description using the "As a... I want... So that..." format
2. Suggest improvements to make the story more specific and actionable
3. Help identify missing acceptance criteria or requirements
4. Ensure the story is properly scoped and testable
5. Handle requests to modify, remove, or restructure story components WITHIN this single story
6. Maintain context of the current user story throughout the conversation
7. If users request to "remove US-X" or similar, explain this is a single story page and offer alternatives

When users ask to remove story identifiers (like US-5), respond with context-aware guidance:
- Explain this is a details page for ONE user story
- Suggest they may want to go back to the epic screen to manage multiple stories
- Offer to help refine THIS story instead
- Keep the conversation focused on improving the current story

For any significant changes, provide the updated content in your response and indicate what specific parts should be updated."""

        user_prompt = f"User Request: {user_message}"

        # Use OpenAI to generate user story guidance
        try:
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=1000,
                temperature=0.7
            )
            
            ai_response = response.choices[0].message.content.strip()
            
            # Check if the response suggests specific updates
            updates = {}
            lower_response = ai_response.lower()
            
            # Look for specific update patterns in the AI response
            if "user story name:" in lower_response or "story name:" in lower_response:
                # Try to extract suggested story name (this is a simple heuristic)
                lines = ai_response.split('\n')
                for line in lines:
                    if "story name:" in line.lower() or "user story name:" in line.lower():
                        # Extract the suggested name after the colon
                        parts = line.split(':', 1)
                        if len(parts) > 1:
                            suggested_name = parts[1].strip().strip('"').strip("'")
                            if suggested_name and len(suggested_name) > 10:  # Reasonable length check
                                updates["user_story_name"] = suggested_name
                        break
            
            if "description:" in lower_response and "as a" in lower_response:
                # Try to extract user story description in proper format
                lines = ai_response.split('\n')
                for i, line in enumerate(lines):
                    if "description:" in line.lower() and i + 1 < len(lines):
                        next_line = lines[i + 1].strip()
                        if next_line.lower().startswith("as a"):
                            updates["user_story_description"] = next_line
                        break
            
            # Look for acceptance criteria suggestions
            if "acceptance criteria:" in lower_response:
                criteria_lines = []
                lines = ai_response.split('\n')
                capturing = False
                for line in lines:
                    if "acceptance criteria:" in line.lower():
                        capturing = True
                        continue
                    elif capturing:
                        if line.strip().startswith('•') or line.strip().startswith('-') or line.strip().startswith('*'):
                            criteria_lines.append(line.strip().lstrip('•-* '))
                        elif line.strip() and not line.strip().startswith('•') and not line.strip().startswith('-'):
                            # Stop capturing if we hit a non-bullet line
                            break
                
                if criteria_lines:
                    updates["acceptance_criteria"] = criteria_lines
            
            logger.info("Successfully generated user story chat response")
            
            response_data = {
                "success": True,
                "message": ai_response
            }
            
            if updates:
                response_data["updates"] = updates
                logger.info(f"Suggested updates: {updates}")
            
            return jsonify(response_data)
            
        except Exception as ai_error:
            logger.error(f"OpenAI API error: {str(ai_error)}")
            
            # Fallback response for user story chat
            fallback_response = process_user_story_fallback(user_message, context)
            return jsonify({
                "success": True,
                "message": fallback_response
            })
            
    except Exception as e:
        logger.error(f"Error in user story chat: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

def process_user_story_fallback(user_message, context):
    """Fallback processing for user story chat when AI is not available."""
    user_story_name = context.get("user_story_name", "")
    user_story_description = context.get("user_story_description", "")
    epic_title = context.get("epic_title", "")
    
    lower_message = user_message.lower()
    
    # Handle removal requests with context awareness
    if any(word in lower_message for word in ['remove', 'delete']) and any(story_id in lower_message for story_id in ['us-', 'story', 'user story']):
        return f"""🎯 **Context Clarification Needed**

I understand you mentioned removing something, but I want to make sure we're on the same page:

**Current Context:** You're viewing the details page for the user story "{user_story_name or 'this specific story'}" which is part of the epic "{epic_title or 'the current epic'}".

This page is designed to refine and improve **this single user story**, not to manage multiple stories.

**If you want to:**
• **Remove this entire story**: Please go back to the epic screen where you can manage multiple stories
• **Modify this story**: I can help you refine the description, acceptance criteria, or requirements right here
• **Split this story**: I can help you identify how to break it into smaller, more manageable stories
• **Remove specific parts**: I can help remove or modify acceptance criteria, requirements, or description elements

**What specifically would you like to do with the current user story: "{user_story_name}"?**"""
    
    elif any(word in lower_message for word in ['improve', 'better', 'enhance', 'refine']):
        suggestions = []
        
        if not user_story_description or len(user_story_description) < 50:
            suggestions.append("• Expand the user story description to follow the 'As a... I want... So that...' format")
        
        if "as a" not in user_story_description.lower():
            suggestions.append("• Rewrite the description to clearly identify the user persona")
        
        if "so that" not in user_story_description.lower():
            suggestions.append("• Add the business value and rationale ('So that...' clause)")
        
        suggestions.append("• Consider adding more specific acceptance criteria")
        suggestions.append("• Review if the story is properly sized for a single sprint")
        
        return f"""🔧 **Improvement Suggestions for "{user_story_name or 'Your User Story'}"**

Here are some suggestions to improve your user story:

{chr(10).join(suggestions)}

**Current Story Context:**
• **Epic:** {epic_title or 'Not specified'}  
• **Story:** {user_story_name or 'Not specified'}
• **Description Length:** {len(user_story_description) if user_story_description else 0} characters

Would you like me to help with any specific aspect? You can ask me to:
- "Rewrite the description in proper format"
- "Suggest additional acceptance criteria"  
- "Help break down a large story into smaller ones"
- "Add missing requirements or constraints"

I'll maintain context of this specific story throughout our conversation."""
    
    elif any(word in lower_message for word in ['format', 'as a', 'structure']):
        return f"""📝 **User Story Format Guidance**

I can help you format the user story properly. A good user story follows this structure:

**As a** [user persona]
**I want** [specific functionality] 
**So that** [business value/benefit]

**Current Story Context:**
• **Epic:** {epic_title or 'Not specified'}
• **Story Name:** {user_story_name or 'Not specified'}
• **Current Description:** "{user_story_description or 'No description provided'}"

Would you like me to help rewrite this in the proper format? Please provide more details about:
- Who is the user? (customer, admin, employee, etc.)
- What specific action do they want to perform?
- What benefit or value will they get from this feature?

I'll keep this focused on improving your current story: "{user_story_name or 'this story'}"."""
    
    else:
        return f"""👋 **Welcome to User Story Chat**

I'm here to help you refine and improve your user story: **"{user_story_name or 'Current Story'}"**

**Current Context:**
• **Epic:** {epic_title or 'Not specified'}
• **Story:** {user_story_name or 'Not specified'}
• **Description:** {user_story_description[:100] + '...' if user_story_description and len(user_story_description) > 100 else user_story_description or 'Not specified'}

**I can help you with:**
• **Story Format**: Ensure it follows "As a... I want... So that..." structure
• **Clarity**: Make the description more specific and actionable
• **Acceptance Criteria**: Add or improve the criteria for testing
• **Requirements**: Identify missing technical or business requirements
• **Scope**: Ensure the story is appropriately sized

**What would you like to work on?** You can say things like:
- "Help me rewrite this in proper format"
- "Add more acceptance criteria"
- "Make this more specific"
- "Is this story too big?"
- "Improve the user story description"

💡 **Note:** This page focuses on refining THIS specific user story. If you need to manage multiple stories or remove stories entirely, please use the epic management screen."""

@app.route("/chat_traceability", methods=["POST"])
def chat_traceability():
    """Handle chat requests for traceability matrix refinement and enhancement."""
    logger.info("Request received for traceability matrix chat")
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No JSON data provided"}), 400
            
        user_message = data.get("userMessage", "").strip()
        context = data.get("sessionContext", "{}")
        
        if not user_message:
            return jsonify({"success": False, "error": "Message is required"}), 400
        
        # Extract context information
        user_story_title = data.get("userStory", {}).get("title", "")
        user_story_description = data.get("userStory", {}).get("description", "")
        acceptance_criteria = data.get("userStory", {}).get("acceptanceCriteria", "")
        current_traceability = data.get("currentTraceability", "")
        
        logger.info(f"Traceability chat message: {user_message}")
        logger.info(f"Context - Story: {user_story_title}")
        
        # Build context string for AI
        context_text = f"""
Current User Story Context:
- Story Title: {user_story_title or 'Not specified'}
- Description: {user_story_description or 'Not specified'}
- Acceptance Criteria: {acceptance_criteria or 'Not specified'}
- Current Traceability Matrix: {current_traceability or 'Not available'}
"""
        
        # System prompt for traceability matrix chat
        system_prompt = f"""You are an expert systems analyst and requirements traceability specialist. You help create, analyze, and improve traceability matrices that map user stories to PRD requirements, design documents, and testing artifacts.

{context_text}

Your role is to:
1. Analyze and improve traceability matrices between user stories and PRD requirements
2. Identify missing traceability links and suggest improvements
3. Help create bidirectional traceability relationships
4. Ensure comprehensive coverage of requirements traceability
5. Suggest impact analysis approaches for requirement changes
6. Help with compliance and audit traceability documentation
7. Create structured, readable traceability matrices in table format

When enhancing traceability matrices, you should:
- Create clear mappings between user stories and specific PRD requirements
- Include requirement IDs, descriptions, and traceability relationships
- Suggest forward and backward traceability links
- Include test case traceability where relevant
- Format output as readable tables or structured text
- Identify gaps in requirement coverage
- Suggest additional traceability dimensions (design documents, test cases, compliance standards)

For traceability matrix outputs, use this format:
| User Story | PRD Requirement ID | Requirement Description | Traceability Type | Test Coverage |
|------------|-------------------|------------------------|-------------------|---------------|
| [Story ID] | [REQ-ID] | [Description] | [Forward/Backward/Bidirectional] | [Test Case IDs] |

Always provide structured, professional traceability information that would be suitable for project documentation and compliance audits."""

        user_prompt = f"User Request: {user_message}\n\nContext: Please analyze and enhance the traceability matrix for the current user story."

        # Use OpenAI to generate traceability guidance
        try:
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            ai_response = response.choices[0].message.content.strip()
            logger.info(f"Generated traceability chat response: {ai_response[:200]}...")
            
            # Check if the response includes an updated traceability matrix
            updated_traceability = None
            if "|" in ai_response and "User Story" in ai_response:
                # Response contains a table format - extract it as updated traceability
                updated_traceability = ai_response
            
            return jsonify({
                "success": True,
                "response": ai_response,
                "updated_traceability": updated_traceability
            })
            
        except Exception as e:
            logger.error(f"OpenAI API error in traceability chat: {str(e)}")
            
            # Provide a fallback response for common traceability questions
            fallback_response = generate_traceability_fallback_response(user_message, user_story_title, current_traceability)
            
            return jsonify({
                "success": True,
                "response": fallback_response,
                "updated_traceability": None
            })
    
    except Exception as e:
        logger.error(f"Error in traceability chat: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

def generate_traceability_fallback_response(user_message, user_story_title, current_traceability):
    """Generate a fallback response for traceability chat when OpenAI is unavailable."""
    lower_message = user_message.lower()
    
    if any(word in lower_message for word in ['improve', 'enhance', 'better']):
        return f"""🔗 **Traceability Matrix Enhancement Suggestions**

For the user story "{user_story_title or 'Current Story'}", here are ways to improve traceability:

**1. Forward Traceability:**
• Map story to specific PRD requirements (REQ-001, REQ-002, etc.)
• Link to design documents and technical specifications
• Connect to test cases and test scenarios

**2. Backward Traceability:**
• Trace back to business objectives and goals
• Link to stakeholder requirements and needs
• Connect to compliance and regulatory requirements

**3. Bidirectional Traceability:**
• Ensure changes can be traced in both directions
• Maintain impact analysis capabilities
• Enable change management tracking

**4. Coverage Analysis:**
• Verify all acceptance criteria map to requirements
• Ensure no orphaned requirements exist
• Validate test coverage completeness

Would you like me to help create a specific traceability matrix table for this user story?"""
    
    elif any(word in lower_message for word in ['table', 'format', 'matrix']):
        return f"""📊 **Traceability Matrix Table Format**

Here's a recommended structure for your traceability matrix:

| User Story | PRD Requirement | Requirement Description | Traceability Type | Test Case |
|------------|----------------|------------------------|-------------------|-----------|
| {user_story_title or 'US-001'} | REQ-001 | User authentication and authorization | Forward | TC-001 |
| {user_story_title or 'US-001'} | REQ-002 | Data validation and input sanitization | Forward | TC-002 |
| {user_story_title or 'US-001'} | REQ-003 | Error handling and user feedback | Forward | TC-003 |

**Traceability Types:**
• **Forward**: User Story → Requirements → Design → Tests
• **Backward**: Tests ← Design ← Requirements ← User Story  
• **Bidirectional**: Maintains links in both directions

This format helps with compliance audits, impact analysis, and requirement coverage verification."""
    
    elif any(word in lower_message for word in ['missing', 'gaps', 'coverage']):
        return f"""🔍 **Traceability Gap Analysis**

Common gaps in traceability matrices:

**Missing Forward Links:**
• User story not mapped to specific PRD requirements
• Requirements not linked to design documents
• Design elements not connected to test cases

**Missing Backward Links:**
• Requirements not traced to business objectives
• User stories not linked to stakeholder needs
• Test cases not mapped back to acceptance criteria

**Coverage Gaps:**
• Orphaned requirements with no implementing stories
• User stories without corresponding test coverage
• Acceptance criteria not validated by specific tests

**For "{user_story_title or 'your story'}":**
Consider mapping to requirements for data validation, user interface design, error handling, security, and performance specifications.

Would you like help identifying specific requirements that should be traced to this user story?"""
    
    else:
        return f"""🔗 **Traceability Matrix Assistant**

I can help you enhance the traceability matrix for "{user_story_title or 'your user story'}".

**Current Traceability Status:**
{current_traceability[:200] + '...' if current_traceability and len(current_traceability) > 200 else current_traceability or 'No traceability matrix currently available'}

**I can help with:**
• Creating structured traceability tables
• Identifying missing requirement mappings
• Analyzing traceability coverage gaps
• Suggesting bidirectional traceability links
• Formatting for compliance documentation

**Common requests:**
- "Create a traceability table for this user story"
- "What requirements should this story map to?"
- "Improve the traceability matrix format"
- "Identify missing traceability links"

What aspect of traceability would you like to work on?"""

@app.route("/submit-jira-ticket", methods=["POST"])
def submit_jira_ticket():
    """Submit user story details to Jira as a ticket."""
    logger.info("Request received for Jira ticket submission")
    
    try:
        # Get form data
        epic_title = request.form.get("epic_title", "")
        user_story_name = request.form.get("user_story_name", "")
        user_story_description = request.form.get("user_story_description", "")
        priority = request.form.get("priority", "High")
        responsible_systems = request.form.get("responsible_systems", "")
        acceptance_criteria_raw = request.form.get("acceptance_criteria", "")
        tagged_requirements_raw = request.form.get("tagged_requirements", "")
        traceability_matrix = request.form.get("traceability_matrix", "")
        
        logger.info(f"Submitting Jira ticket for story: {user_story_name}")
        logger.info(f"Epic: {epic_title}")
        logger.info(f"Priority: {priority}")
        logger.info(f"Systems: {responsible_systems}")
        
        # Parse acceptance criteria (split by |||)
        acceptance_criteria = []
        if acceptance_criteria_raw:
            acceptance_criteria = [criteria.strip() for criteria in acceptance_criteria_raw.split('|||') if criteria.strip()]
        
        # Parse tagged requirements (split by |||)
        tagged_requirements = []
        if tagged_requirements_raw:
            tagged_requirements = [req.strip() for req in tagged_requirements_raw.split('|||') if req.strip()]
        
        # Format the description for Jira
        jira_description = user_story_description
        
        # Add acceptance criteria to description
        if acceptance_criteria:
            jira_description += "\n\n*Acceptance Criteria:*\n"
            for i, criterion in enumerate(acceptance_criteria, 1):
                jira_description += f"• {criterion}\n"
        
        # Add tagged requirements to description
        if tagged_requirements:
            jira_description += "\n\n*Tagged Requirements:*\n"
            for i, requirement in enumerate(tagged_requirements, 1):
                jira_description += f"• {requirement}\n"
        
        # Add traceability matrix to description
        if traceability_matrix and traceability_matrix.strip() != 'Traceability mapping not available.':
            jira_description += "\n\n*Traceability Matrix (User Story → PRD Requirements):*\n"
            jira_description += f"{traceability_matrix}\n"
        
        # Add system information
        if responsible_systems:
            jira_description += f"\n\n*Responsible Systems:* {responsible_systems}"
        
        # Prepare Jira ticket data
        jira_data = {
            'epic_title': epic_title,
            'summary': user_story_name,
            'description': jira_description,
            'priority': priority,
            'responsible_systems': responsible_systems,
            'acceptance_criteria': acceptance_criteria,
            'tagged_requirements': tagged_requirements,
            'traceability_matrix': traceability_matrix
        }
        
        # Try to create Jira ticket using the Jira connector
        try:
            # Import and use Jira connector
            from jira import JIRA
            import os
            from dotenv import load_dotenv
            
            load_dotenv()
            
            # Jira configuration
            JIRA_SERVER = 'https://lalluluke.atlassian.net/'
            EMAIL = 'lalluluke@gmail.com'
            API_TOKEN = os.getenv("JIRA_API_TOKEN")
            
            if not API_TOKEN:
                logger.error("JIRA_API_TOKEN not found in environment variables")
                return render_template('poc2_user_story_details.html', 
                                     error_message="Jira API token not configured. Please contact administrator.",
                                     **jira_data)
            
            # Connect to Jira
            jira = JIRA(server=JIRA_SERVER, basic_auth=(EMAIL, API_TOKEN))
            
            # Create issue dictionary (removed priority field due to Jira screen configuration)
            issue_dict = {
                'project': {'key': 'SCRUM'},  # Default project key
                'summary': user_story_name,
                'description': jira_description,
                'issuetype': {'name': 'Story'}
            }
            
            # Create the Jira issue
            new_issue = jira.create_issue(fields=issue_dict)
            
            logger.info(f"Successfully created Jira ticket: {new_issue.key}")
            
            # Store the ticket information for display
            success_data = {
                'ticket_key': new_issue.key,
                'ticket_url': f"{JIRA_SERVER}browse/{new_issue.key}",
                'success_message': f"Jira ticket {new_issue.key} created successfully!"
            }
            
            # Return success page with ticket details
            return render_template('jira_success.html', **success_data, **jira_data)
            
        except ImportError as e:
            logger.error(f"Jira library not available: {e}")
            error_msg = "Jira integration not available. Please install the required dependencies."
            return render_template('poc2_user_story_details.html', 
                                 error_message=error_msg,
                                 **jira_data)
            
        except Exception as e:
            logger.error(f"Error creating Jira ticket: {e}")
            error_msg = f"Failed to create Jira ticket: {str(e)}"
            return render_template('poc2_user_story_details.html', 
                                 error_message=error_msg,
                                 **jira_data)
    
    except Exception as e:
        logger.error(f"Error in Jira ticket submission: {str(e)}")
        return render_template('poc2_user_story_details.html', 
                             error_message=f"Error submitting to Jira: {str(e)}")

@app.route("/system-mapping", methods=["GET"])
def system_mapping():
    """Serve the system mapping drag-and-drop interface."""
    logger.info("GET request received for system mapping interface")
    try:
        return render_template("system_mapping.html")
    except Exception as e:
        logger.error(f"Error serving system mapping page: {str(e)}")
        return f"Error loading system mapping page: {str(e)}", 500

@app.route("/download-system-mapping", methods=["POST"])
def download_system_mapping():
    """Handle system mapping CSV download."""
    logger.info("POST request received for system mapping CSV download")
    try:
        # Get the selected systems from the request
        selected_systems = request.json.get('systems', [])
        logger.info(f"Received {len(selected_systems)} selected systems for CSV download")
        
        # Create CSV content
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Write header
        writer.writerow(['System Name', 'Description', 'Selected'])
        
        # Write selected systems
        for system in selected_systems:
            writer.writerow([
                system.get('name', ''),
                system.get('description', ''),
                'Yes'
            ])
        
        csv_content = output.getvalue()
        output.close()
        
        return jsonify({
            'success': True,
            'csv_content': csv_content,
            'filename': f'system_mapping_{datetime.now().strftime("%Y%m%d_%H%M%S")}.csv'
        })
        
    except Exception as e:
        logger.error(f"Error generating system mapping CSV: {str(e)}")
        return jsonify({
            'success': False,
            'error': str(e)
        }), 500

@app.route("/system-mapping-chat", methods=["POST"])
def system_mapping_chat():
    """Handle chat requests for system mapping assistance."""
    logger.info("Request received for system mapping chat")
    
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No JSON data provided"}), 400
            
        user_message = data.get("message", "").strip()
        current_systems = data.get("current_systems", [])
        available_systems = data.get("available_systems", [])
        
        if not user_message:
            return jsonify({"success": False, "error": "Message is required"}), 400
        
        logger.info(f"Processing system mapping chat request: {user_message[:100]}...")
        
        # Create a prompt for the AI to help with system mapping
        system_prompt = """You are a technical architect AI assistant specializing in system mapping and technical architecture. 
        You help users select and organize systems for their technical implementations based on their requirements.
        
        Available Systems:
        - Customer acquisition platform: Captures applications from digital channels
        - Credit decision engine: Automated credit policies and risk scoring
        - Application review system: Manual interventions and compliance checks
        - Fraud detection platform: Monitors suspicious patterns and behaviors
        - Document management platform: Document storage and verification
        - Agent support tool: UI for agents to manage applications
        - Customer data repository: Golden source for personal data
        - Account creation engine: Initializes financial product accounts
        - Product configuration database: Product metadata and rules
        - Card issuance manager: Generates cards and configures features
        - Card manufacturing coordinator: External vendor coordination
        - Payment setup module: Payment preferences configuration
        - Financial ledger initializer: Journal entries and GL accounts
        - Settlement configuration module: Fund routing configuration
        - Statement preference manager: Statement delivery preferences
        - Authorization control layer: Real-time transaction validation
        - Notification delivery engine: Customer alerts and updates
        - Credit behavior monitor: Post-booking risk monitoring
        - Credit bureau interface: Bureau data transmission
        - Loyalty program manager: Rewards program integration
        - Customer servicing portal: Customer self-service requests
        
        Guidelines:
        - Suggest relevant systems based on user requirements
        - Explain why certain systems are needed
        - Consider system dependencies and integration points
        - Provide architectural guidance
        - Be specific about system capabilities
        
        Format your response as JSON with:
        {
          "message": "Your helpful response",
          "suggestions": [
            {"system": "System Name", "reason": "Why this system is recommended"}
          ],
          "warnings": ["Any warnings or considerations"]
        }"""
        
        current_systems_text = ", ".join(current_systems) if current_systems else "None selected"
        available_systems_text = ", ".join(available_systems) if available_systems else "All systems available"
        
        user_prompt = f"""Current Selected Systems: {current_systems_text}

Available Systems: {available_systems_text}

User Request: {user_message}

Please provide system mapping recommendations and guidance based on the user's request."""

        # Use OpenAI to generate system mapping assistance
        try:
            client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                max_tokens=800,
                temperature=0.7
            )
            
            ai_response = response.choices[0].message.content.strip()
            
            # Try to parse as JSON, fallback to plain text
            try:
                import json
                parsed_response = json.loads(ai_response)
                return jsonify({
                    "success": True,
                    "message": parsed_response.get("message", ai_response),
                    "suggestions": parsed_response.get("suggestions", []),
                    "warnings": parsed_response.get("warnings", [])
                })
            except json.JSONDecodeError:
                return jsonify({
                    "success": True,
                    "message": ai_response,
                    "suggestions": [],
                    "warnings": []
                })
            
        except Exception as ai_error:
            logger.error(f"OpenAI API error: {str(ai_error)}")
            
            # Fallback to rule-based processing
            fallback_response = process_system_mapping_fallback(user_message, current_systems)
            return jsonify({
                "success": True,
                "message": fallback_response["message"],
                "suggestions": fallback_response["suggestions"],
                "warnings": fallback_response["warnings"]
            })
            
    except Exception as e:
        logger.error(f"Error in system mapping chat: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

def process_system_mapping_fallback(user_message, current_systems):
    """Fallback processing for system mapping when AI is not available."""
    lower_message = user_message.lower()
    
    suggestions = []
    warnings = []
    message = "I can help you with system mapping! "
    
    # Analyze user intent
    if any(word in lower_message for word in ['credit', 'loan', 'approve', 'decision']):
        suggestions.extend([
            {"system": "Credit decision engine", "reason": "Essential for credit assessment and approval workflows"},
            {"system": "Credit bureau interface", "reason": "Required for credit history and reporting"},
            {"system": "Customer data repository", "reason": "Needed for customer information storage"}
        ])
        message += "For credit-related functionality, I recommend these core systems."
        
    elif any(word in lower_message for word in ['fraud', 'security', 'detect', 'risk']):
        suggestions.extend([
            {"system": "Fraud detection platform", "reason": "Primary fraud monitoring and prevention"},
            {"system": "Authorization control layer", "reason": "Real-time transaction security"},
            {"system": "Customer data repository", "reason": "Customer profile for risk assessment"}
        ])
        message += "For fraud detection and security, these systems work together."
        
    elif any(word in lower_message for word in ['card', 'payment', 'transaction']):
        suggestions.extend([
            {"system": "Card issuance manager", "reason": "Manages card creation and configuration"},
            {"system": "Payment setup module", "reason": "Handles payment preferences and setup"},
            {"system": "Authorization control layer", "reason": "Controls transaction approvals"},
            {"system": "Settlement configuration module", "reason": "Manages payment routing and settlement"}
        ])
        message += "For card and payment processing, these systems are essential."
        
    elif any(word in lower_message for word in ['customer', 'service', 'support']):
        suggestions.extend([
            {"system": "Customer servicing portal", "reason": "Self-service customer interface"},
            {"system": "Agent support tool", "reason": "Agent interface for customer support"},
            {"system": "Customer data repository", "reason": "Customer information access"}
        ])
        message += "For customer service functionality, these systems are recommended."
        
    elif any(word in lower_message for word in ['onboard', 'acquisition', 'application']):
        suggestions.extend([
            {"system": "Customer acquisition platform", "reason": "Captures and processes new applications"},
            {"system": "Document management platform", "reason": "Handles required documentation"},
            {"system": "Application review system", "reason": "Manual review and compliance checks"},
            {"system": "Account creation engine", "reason": "Creates accounts after approval"}
        ])
        message += "For customer onboarding, this end-to-end flow is recommended."
        
    else:
        message += "Could you be more specific about your use case? For example: 'I need systems for credit approval', 'What do I need for fraud detection?', 'Help me set up payment processing', etc."
    
    # Add warnings based on current selection
    if len(current_systems) > 10:
        warnings.append("You have selected many systems. Consider if all are necessary for your specific use case.")
    
    if "Customer data repository" not in current_systems and len(current_systems) > 3:
        warnings.append("Consider adding Customer data repository as most systems depend on customer data.")
    
    return {
        "message": message,
        "suggestions": suggestions,
        "warnings": warnings
    }

if __name__ == "__main__":
    try:
        logger.info("Starting Flask application...")
        app.run(
            debug=False,  # Set to False for production
            host="0.0.0.0",  # Allow external connections
            port=5000,
            threaded=True,
            use_reloader=False  # Disable reloader to prevent double initialization
        )
    except Exception as e:
        logger.error(f"Failed to start Flask application: {str(e)}")
        raise
