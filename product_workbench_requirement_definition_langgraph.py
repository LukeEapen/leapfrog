import os
import json
import time
import logging
from datetime import timedelta
from functools import wraps
from concurrent.futures import ThreadPoolExecutor

from flask import Flask, jsonify, request, session, redirect, url_for, render_template, send_file
from dotenv import load_dotenv

# Reuse vector DB blueprint
from vector_db_api import vector_db_api

# LangGraph / LangChain
from langgraph.graph import StateGraph, END
from langgraph.checkpoint.memory import MemorySaver
from langchain_core.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI

# Utilities reused from the existing backend
from product_workbench_requirement_definition import (
    dedupe_legacy_citations,
    extract_references_from_outputs,
    initialize_document_styles,
)
from io import BytesIO
from docx import Document
from bs4 import BeautifulSoup
from markdown2 import markdown
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

load_dotenv()

app = Flask(__name__)
app.register_blueprint(vector_db_api)

app.secret_key = os.getenv('FLASK_SECRET_KEY')
app.config.update(
    SESSION_COOKIE_SECURE=False,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',
    PERMANENT_SESSION_LIFETIME=timedelta(hours=1)
)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

THREAD_POOL_SIZE = 8
executor = ThreadPoolExecutor(max_workers=THREAD_POOL_SIZE)

# ---------- LangGraph Orchestrator ----------
# Minimal state carrying inputs and outputs
class OrchestratorState(dict):
    pass

llm = ChatOpenAI(model=os.getenv('OPENAI_MODEL', 'gpt-4o-mini'), temperature=0.2)

# Prompts similar to existing agents
PROMPT_PRODUCT = ChatPromptTemplate.from_messages([
    ("system", "You synthesize a concise Product Overview for a PRD."),
    ("user", "{context}")
])
PROMPT_FEATURE = ChatPromptTemplate.from_messages([
    ("system", "You generate a Feature Overview from context."),
    ("user", "{context}")
])
PROMPT_REQ = ChatPromptTemplate.from_messages([
    ("system", "Generate the {section} for a PRD. If using legacy cues, include 'Source: Legacy Code' lines where appropriate."),
    ("user", "{context}")
])

# Nodes
async def node_product(state: OrchestratorState):
    try:
        ctx = state.get('combined', '')
        chain = PROMPT_PRODUCT | llm
        resp = await chain.ainvoke({"context": ctx})
        state['product_overview'] = resp.content
    except Exception as e:
        state['product_overview'] = f"Error: {e}"
    return state

async def node_feature(state: OrchestratorState):
    try:
        ctx = state.get('combined', '')
        chain = PROMPT_FEATURE | llm
        resp = await chain.ainvoke({"context": ctx})
        state['feature_overview'] = resp.content
    except Exception as e:
        state['feature_overview'] = f"Error: {e}"
    return state

async def node_requirements(state: OrchestratorState):
    try:
        ctx = state.get('combined', '')
        outputs = {}
        for key, title in [
            ('agent_4_2', 'Functional Requirements'),
            ('agent_4_3', 'Non-Functional Requirements'),
            ('agent_4_4', 'Data Requirements'),
            ('agent_4_5', 'Legal, Regulatory, and Compliance Requirements'),
        ]:
            chain = PROMPT_REQ | llm
            resp = await chain.ainvoke({"context": ctx, "section": title})
            outputs[key] = resp.content
        state['combined_outputs'] = outputs
    except Exception as e:
        state['combined_outputs'] = {"error": str(e)}
    return state

# Graph: product + feature in parallel, then requirements
memory = MemorySaver()
workflow = StateGraph(OrchestratorState)
workflow.add_node("product", node_product)
workflow.add_node("feature", node_feature)
workflow.add_node("requirements", node_requirements)
workflow.set_entry_point("product")
workflow.add_edge("product", "feature")
workflow.add_edge("feature", "requirements")
workflow.add_edge("requirements", END)
app_graph = workflow.compile(checkpointer=memory)

# ---------- Helpers (reuse minimal parts) ----------
def add_hyperlink(paragraph, url, text=None):
    if text is None:
        text = url
    part = paragraph.part
    r_id = part.relate_to(
        url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    new_run.append(rPr)

    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

# ---------- Auth ----------
def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

def verify_credentials(username, password):
    return (username == os.getenv('ADMIN_USERNAME') and password == os.getenv('ADMIN_PASSWORD'))

@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if verify_credentials(request.form.get('username'), request.form.get('password')):
            session['logged_in'] = True
            return redirect(url_for('page1'))
        return render_template('page0_login.html', error=True)
    return render_template('page0_login.html', error=False)

@app.route('/page1', methods=['GET', 'POST'])
@require_auth
def page1():
    if request.method == 'POST':
        inputs = {key: request.form.get(key, '') for key in ['industry', 'sector', 'geography', 'intent', 'features']}
        session['inputs'] = inputs

        # Compose combined context (reuse the simpler version from the existing backend)
        ui = f"""
        # Original User Inputs
        Industry: {inputs.get('industry','')}
        Sector: {inputs.get('sector','')}
        Geography: {inputs.get('geography','')}
        Intent: {inputs.get('intent','')}
        Features: {inputs.get('features','')}
        """
        # Include legacy description if provided
        legacy_desc = request.form.get('legacy_business_description', '').strip()
        combined = ui + "\n\n" + (legacy_desc or '')

        # Run graph: product -> feature -> requirements later (page2)
        # For page1 we'll kick off product + feature only
        state = {"combined": combined}
        # Execute first two steps by running graph to "feature"
        # We can step until requirements by relying on memory checkpoints; for simplicity, invoke synchronous run
        result = app_graph.invoke(state)
        session['data_key'] = 'lg_session'
        # Persist interim outputs in session (simple storage)
        session['product_overview'] = result.get('product_overview', '')
        session['feature_overview'] = result.get('feature_overview', '')
        return redirect('/page2')

    return render_template('page1_input.html')

@app.route('/page2', methods=['GET', 'POST'])
@require_auth
def page2():
    if request.method == 'POST':
        # Build combined context with Page1 results
        inputs = session.get('inputs', {})
        po = session.get('product_overview', '')
        fo = session.get('feature_overview', '')
        legacy = request.form.get('legacy_business_description', '') or ''
        combined = json.dumps({"inputs": inputs, "product_overview": po, "feature_overview": fo, "legacy": legacy})
        result = app_graph.invoke({"combined": combined})
        session['combined_outputs'] = result.get('combined_outputs', {})
        return redirect('/page4')

    return render_template('page2_agents.html',
                           agent11_output=session.get('product_overview',''),
                           agent2_output=session.get('feature_overview',''),
                           legacy_business_description='')

@app.route('/page4')
@require_auth
def page4():
    outputs = {
        'product_overview': dedupe_legacy_citations(session.get('product_overview', '')),
        'feature_overview': dedupe_legacy_citations(session.get('feature_overview', '')),
        'agent_4_2': dedupe_legacy_citations(session.get('combined_outputs', {}).get('agent_4_2', '')),
        'agent_4_3': dedupe_legacy_citations(session.get('combined_outputs', {}).get('agent_4_3', '')),
        'agent_4_4': dedupe_legacy_citations(session.get('combined_outputs', {}).get('agent_4_4', '')),
        'agent_4_5': dedupe_legacy_citations(session.get('combined_outputs', {}).get('agent_4_5', '')),
    }
    return render_template('page4_final_output.html', outputs=outputs)

@app.route('/generate_word_doc', methods=['POST'])
@require_auth
def generate_word_doc():
    data = {
        'product_overview': session.get('product_overview',''),
        'feature_overview': session.get('feature_overview',''),
        'combined_outputs': session.get('combined_outputs', {})
    }

    doc = Document()
    initialize_document_styles(doc)
    doc.add_heading("Product Requirements Document", level=1)
    doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph()

    sections = {
        "Product Overview": dedupe_legacy_citations(data.get("product_overview", "")),
        "Feature Overview": dedupe_legacy_citations(data.get("feature_overview", "")),
        "Functional Requirements": dedupe_legacy_citations(data.get("combined_outputs", {}).get("agent_4_2", "")),
        "Non-Functional Requirements": dedupe_legacy_citations(data.get("combined_outputs", {}).get("agent_4_3", "")),
        "Data Requirements": dedupe_legacy_citations(data.get("combined_outputs", {}).get("agent_4_4", "")),
        "Legal & Compliance Requirements": dedupe_legacy_citations(data.get("combined_outputs", {}).get("agent_4_5", ""))
    }

    for title, content in sections.items():
        if not content:
            continue
        doc.add_heading(title, level=2)
        from bs4 import BeautifulSoup
        html_content = markdown(content, extras=["fenced-code-blocks", "tables", "header-ids", "break-on-newline"])        
        soup = BeautifulSoup(html_content, "html.parser")
        for el in soup.find_all(["p", "li", "h3", "h4", "pre", "code"]):
            p = doc.add_paragraph()
            p.add_run(el.get_text(strip=True))
            p.paragraph_format.space_after = Pt(6)

    refs = extract_references_from_outputs(data.get('combined_outputs', {}))
    if refs:
        doc.add_page_break()
        doc.add_heading("References", level=2)
        for ref in refs:
            para = doc.add_paragraph(style="List Bullet")
            if ref.startswith(("http://","https://")):
                add_hyperlink(para, ref)
            else:
                para.add_run(ref)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name=f"PRD_Draft_{time.strftime('%Y%m%d')}.docx",
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5011))
    app.run(host='0.0.0.0', port=port, debug=True)
