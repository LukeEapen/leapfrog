# --- openai-app-v3.py (updated) ---
import openai
from flask import Flask, request, jsonify, render_template, redirect, url_for, session, send_file
from concurrent.futures import ThreadPoolExecutor
import time
import logging
import os
import sys
import codecs
from dotenv import load_dotenv
from docx import Document
from io import BytesIO

if sys.platform.startswith('win'):
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('agent_responses.log', encoding='utf-8')
    ]
)

app = Flask(__name__)
app.secret_key = os.urandom(24)

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

ASSISTANTS = {
    'agent_1': ("asst_EvIwemZYiG4cCmYc7GnTZoQZ", "Prompt Structuring Agent"),
    'agent_2': ("asst_EkihtJQe9qFiztRdRXPhiy2G", "Requirements Generator"),
    'agent_3': ("asst_Si7JAfL2Ov80wvcly6GKLJcN", "Validator Agent"),
    'agent_4_1': ("asst_vlIf20wJ6Dred3TefHEfeAkp", "PRD Requirement Generator"),
    'agent_4_2': ("asst_iGTOOi1HRZBov6QHhmy126Hx", "PRD NFR Generator"),
    'agent_4_3': ("asst_7t6Lr9E3c6wf18uRUyXdz1DR", "PRD Data Requirement Generator"),
    'agent_4_4': ("asst_Po5dW8wTVHwqyYlZePVQp15E", "PRD LRC Generator")
}

def call_agent(assistant_id: str, message: str, agent_name: str):
    start_time = time.time()
    logging.info(f"[START] {agent_name} (ID: {assistant_id})")

    try:
        thread = openai.beta.threads.create()
        openai.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=message
        )

        run = openai.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=assistant_id
        )

        while True:
            status = openai.beta.threads.runs.retrieve(
                thread_id=thread.id,
                run_id=run.id
            )
            if status.status == "completed":
                break
            elif status.status == "failed":
                raise Exception(f"Run failed for {agent_name}")
            time.sleep(0.25)

        messages = openai.beta.threads.messages.list(thread_id=thread.id)
        response = messages.data[0].content[0].text.value
        elapsed = time.time() - start_time

        logging.info(f"[{agent_name}] completed in {elapsed:.2f}s")
        return response

    except Exception as e:
        logging.error(f"[ERROR] {agent_name}: {str(e)}")
        return f"[ERROR] {str(e)}"

@app.route('/api/agent_1', methods=['POST'])
def run_agent_1():
    question = request.json.get('question', '').strip()
    output = call_agent(ASSISTANTS['agent_1'][0], question, ASSISTANTS['agent_1'][1])
    return jsonify({'output': output})

@app.route('/api/agent_2', methods=['POST'])
def run_agent_2():
    input_text = request.json.get('question', '').strip()
    output = call_agent(ASSISTANTS['agent_2'][0], input_text, ASSISTANTS['agent_2'][1])
    return jsonify({'output': output})

@app.route('/api/agent_3', methods=['POST'])
def run_agent_3():
    input_text = request.json.get('question', '').strip()
    output = call_agent(ASSISTANTS['agent_3'][0], input_text, ASSISTANTS['agent_3'][1])
    return jsonify({'output': output})

@app.route('/api/query_agents', methods=['POST'])
def run_agents_4():
    input_text = request.json.get('question', '').strip()

    def run(agent_key):
        assistant_id, name = ASSISTANTS[agent_key]
        result = call_agent(assistant_id, input_text, name)
        return agent_key, result

    agent_keys = ['agent_4_1', 'agent_4_2', 'agent_4_3', 'agent_4_4']

    results = {}
    with ThreadPoolExecutor(max_workers=4) as executor:
        for key, output in executor.map(run, agent_keys):
            results[key] = output

    return jsonify({
        "agent_4_1": results.get("agent_4_1", "No response"),
        "agent_4_2": results.get("agent_4_2", "No response"),
        "agent_4_3": results.get("agent_4_3", "No response"),
        "agent_4_4": results.get("agent_4_4", "No response")
    })

@app.route('/agenticAI')
def index():
    if not session.get("logged_in"):
        return redirect(url_for('login'))
    return render_template('prd-openai-v1.html')

@app.route('/')
def login():
    if session.get("logged_in"):
        return redirect(url_for('index'))
    return render_template('login.html')

@app.route('/login', methods=['POST'])
def do_login():
    username = request.form.get("username")
    password = request.form.get("password")
    if username == "admin" and password == "secure123":
        session["logged_in"] = True
        return redirect(url_for('intermediate'))
    return "Invalid credentials", 401

@app.route('/intermediate', methods=['GET'])
def intermediate():
    if not session.get("logged_in"):
        return redirect(url_for('login'))
    return render_template('intermediate.html')

@app.route('/submit_intermediate', methods=['POST'])
def submit_intermediate():
    if not session.get("logged_in"):
        return redirect(url_for('login'))

    # Extract form fields
    industry = request.form.get("industry")
    sub_industry = request.form.get("sub_industry")
    intent = request.form.get("intent")
    features = request.form.get("features")

    # Handle file upload
    context_file = request.files.get('context_file')
    if context_file and context_file.filename:
        file_content = context_file.read().decode('utf-8', errors='ignore')
    else:
        file_content = ""

    # Prepare input for Agent 1
    agent_input = f"""Industry: {industry}
                Sub-industry: {sub_industry}
                Intent: {intent}
                Features: {features}
                Supporting Material:
                {file_content}"""

    # Store for use in session
    session['agent_input'] = agent_input

    # Call Agent 1
    output = call_agent(ASSISTANTS['agent_1'][0], agent_input, ASSISTANTS['agent_1'][1])
    session['agent1_output'] = output

    return redirect(url_for('index'))



@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/download_doc', methods=['POST'])
def download_doc():
    data = request.json
    doc = Document()
    doc.add_heading("Agentic AI Output Report", 0)

    for key, title in {
        "agent_4_1": "PRD Requirement Generator",
        "agent_4_2": "PRD NFR Generator",
        "agent_4_3": "PRD Data Requirement Generator",
        "agent_4_4": "PRD LRC Generator"
    }.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(data.get(key, "No response"))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="agent_report.docx")

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5001))
    app.run(host="0.0.0.0", port=port)