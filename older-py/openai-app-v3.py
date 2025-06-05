import openai
from flask import Flask, request, jsonify, render_template
from concurrent.futures import ThreadPoolExecutor
import time
import logging
import os
import sys
import codecs
from dotenv import load_dotenv
from flask import Flask, request, jsonify, render_template, redirect, url_for, session
from docx import Document
from flask import send_file
from io import BytesIO

if sys.platform.startswith('win'):
    sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'strict')
    sys.stderr = codecs.getwriter('utf-8')(sys.stderr.buffer, 'strict')

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        logging.FileHandler('agent_responses.log', encoding='utf-8')
    ]
)

app = Flask(__name__)
app.secret_key = os.urandom(24)

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

ASSISTANTS = {
    'agent_1': ("asst_EvIwemZYiG4cCmYc7GnTZoQZ", "Prompt Structuring Agent"),
    'agent_2': ("asst_EkihtJQe9qFiztRdRXPhiy2G", "Requirements Generator"),
    'agent_3': ("asst_Si7JAfL2Ov80wvcly6GKLJcN", "Validator Agent"),
    'agent_5': ("asst_r29PjUzVwfd6XiiYH3ueV41P", "Legal & Compliance Analyst"),
    'agent_6': ("asst_WG96Jp4VMrLJfUE4RMkGFMjf", "NFR Specialist"),
    'agent_7': ("asst_VWzoRLWbeZJS8I2IwtOcsLMp", "Platform Architect"),
    'agent_8': ("asst_sufR6Spw8EBDDoAzqQJN9iJt", "Requirement Functional - Legacy Parity Agent"),
    'agent_9': ("asst_7BC8vPshg3NSCGnay9iHl6hu", "Data Attribute Generator")
}

def call_agent(assistant_id: str, message: str, agent_name: str):
    start_time = time.time()
    logging.info(f"[START] {agent_name} (ID: {assistant_id})")

    try:
        thread = openai.beta.threads.create()
        openai.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=message
        )

        run = openai.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=assistant_id
        )

        while True:
            status = openai.beta.threads.runs.retrieve(
                thread_id=thread.id,
                run_id=run.id
            )
            if status.status == "completed":
                break
            elif status.status == "failed":
                raise Exception(f"Run failed for {agent_name}")
            time.sleep(0.25)

        messages = openai.beta.threads.messages.list(thread_id=thread.id)
        response = messages.data[0].content[0].text.value
        elapsed = time.time() - start_time

        logging.info(f"[{agent_name}] completed in {elapsed:.2f}s")
        return response

    except Exception as e:
        logging.error(f"[ERROR] {agent_name}: {str(e)}")
        return f"[ERROR] {str(e)}"

@app.route('/api/agent_1', methods=['POST'])
def run_agent_1():
    question = request.json.get('question', '').strip()
    output = call_agent(ASSISTANTS['agent_1'][0], question, ASSISTANTS['agent_1'][1])
    return jsonify({'output': output})

@app.route('/api/agent_2', methods=['POST'])
def run_agent_2():
    input_text = request.json.get('question', '').strip()
    output = call_agent(ASSISTANTS['agent_2'][0], input_text, ASSISTANTS['agent_2'][1])
    return jsonify({'output': output})

@app.route('/api/agent_3', methods=['POST'])
def run_agent_3():
    input_text = request.json.get('question', '').strip()
    output = call_agent(ASSISTANTS['agent_3'][0], input_text, ASSISTANTS['agent_3'][1])
    return jsonify({'output': output})

@app.route('/api/query_agents', methods=['POST'])
def run_agents_5_to_9():
    """Call agents 5–9 in parallel and return each as a separate field."""
    input_text = request.json.get('question', '').strip()

    def run(agent_key):
        assistant_id, name = ASSISTANTS[agent_key]
        result = call_agent(assistant_id, input_text, name)
        return agent_key, result

    agent_keys = ['agent_5', 'agent_6', 'agent_7', 'agent_8', 'agent_9']

    results = {}
    with ThreadPoolExecutor(max_workers=5) as executor:
        for key, output in executor.map(run, agent_keys):
            results[key] = output

    # Return each agent's output as a separate field for table display
    return jsonify({
        "agent_5": results.get("agent_5", "No response"),
        "agent_6": results.get("agent_6", "No response"),
        "agent_7": results.get("agent_7", "No response"),
        "agent_8": results.get("agent_8", "No response"),
        "agent_9": results.get("agent_9", "No response")
    })

@app.route('/agenticAI')
def index():
    if not session.get("logged_in"):
        return redirect(url_for('login'))
    return render_template('index-openai-v2.html')

@app.route('/')
def login():
    if session.get("logged_in"):
        return redirect(url_for('index'))
    return render_template('login.html')

@app.route('/login', methods=['POST'])
def do_login():
    username = request.form.get("username")
    password = request.form.get("password")
    if username == "admin" and password == "secure123":  # Change to your values
        session["logged_in"] = True
        return redirect(url_for('index'))
    return "Invalid credentials", 401

@app.route('/logout')
def logout():
    session.clear()
    return redirect(url_for('login'))

@app.route('/download_doc', methods=['POST'])
def download_doc():
    data = request.json  # expects keys: agent_5, agent_6, agent_7, agent_8, agent_9
    doc = Document()
    doc.add_heading("Agentic AI Output Report", 0)

    for key, title in {
        "agent_5": "Agent 5 – Legal & Compliance",
        "agent_6": "Agent 6 – NFR Specialist",
        "agent_7": "Agent 7 – Platform Architect",
        "agent_8": "Agent 8 – Functional Parity",
        "agent_9": "Agent 9 – Data Attribute Generator"
    }.items():
        doc.add_heading(title, level=1)
        doc.add_paragraph(data.get(key, "No response"))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return send_file(buffer, as_attachment=True, download_name="agent_report.docx")

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5001))
    app.run(host="0.0.0.0", port=port)