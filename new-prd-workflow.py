import os
import re
import json
import uuid
import time
import asyncio
import logging
import tempfile
import logging.config
import random  # Add at top with other imports
from io import BytesIO
from datetime import timedelta
from concurrent.futures import ThreadPoolExecutor

# Security imports
from functools import wraps
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

# Third-party imports
import openai
import redis
from flask import (
    Flask, render_template, request, redirect, 
    url_for, session, send_file, jsonify
)
from dotenv import load_dotenv
from marshmallow import Schema, fields, ValidationError

# Document processing imports
from bs4 import BeautifulSoup
from markdown2 import markdown
from docx.shared import Pt, RGBColor  # Add RGBColor hereE
from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Redis import with fallback
try:
    redis_test = redis.Redis(
        host=os.getenv('REDIS_HOST', 'localhost'),
        port=int(os.getenv('REDIS_PORT', 6379)),
        db=int(os.getenv('REDIS_DB', 0)),
        password=os.getenv('REDIS_PASSWORD'),
        decode_responses=True
    )
    redis_test.ping()
    redis_client = redis_test
    USING_REDIS = True
except redis.ConnectionError as e:
    logging.warning(f"Redis connection failed: {e}. Using file storage.")
    TEMP_DIR = tempfile.gettempdir()
    USING_REDIS = False

load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY')
app.config.update(
    SESSION_COOKIE_SECURE=True,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',
    PERMANENT_SESSION_LIFETIME=timedelta(hours=1)
)

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

executor = ThreadPoolExecutor(max_workers=10)

ASSISTANTS = {
    #'agent_1': 'asst_EvIwemZYiG4cCmYc7GnTZoQZ',
    'agent_1_1': 'asst_sW7IMhE5tQ78Ylx0zQkh6YnZ', # Agent 1.1 - Product Overview Synthesizer – System Instructions
    'agent_2': 'asst_t5hnaKy1wPvD48jTbn8Mx45z',   # Agent 2: Feature Overview Generator – System Instructions
    'agent_3': 'asst_EqkbMBdfOpUoEUaBPxCChVLR',   # Agent 3: Highest-Order Requirements Agent
    'agent_4_1': 'asst_Ed8s7np19IPmjG5aOpMAYcPM', # Agent 4.1: Product Requirements / User Stories Generator - System Instructions
    'agent_4_2': 'asst_CLBdcKGduMvSBM06MC1OJ7bF', # Agent 4.2: Operational Business Requirements Generator – System Instructions
    'agent_4_3': 'asst_61ITzgJTPqkQf4OFnnMnndNb', # Agent 4.3: Capability-Scoped Non-Functional Requirements Generator – System Instructions
    'agent_4_4': 'asst_pPFGsMMqWg04OSHNmyQ5oaAy', # Agent 4.4: Data Attribute Requirement Generator – System Instructions
    'agent_4_5': 'asst_wwgc1Zbl5iknlDtcFLOuTIjd'  # Agent 4.5: LRC: Legal, Regulatory, and Compliance Synthesizer – System Instructions
}

LOGGING_CONFIG = {
    'version': 1,
    'disable_existing_loggers': False,
    'formatters': {
        'standard': {
            'format': '%(asctime)s [%(levelname)s] %(name)s: %(message)s'
        },
    },
    'handlers': {
        'default': {
            'level': 'INFO',
            'formatter': 'standard',
            'class': 'logging.StreamHandler',
        },
        'file': {
            'level': 'INFO',
            'formatter': 'standard',
            'class': 'logging.FileHandler',
            'filename': 'app.log',
            'mode': 'a',
        },
    },
    'loggers': {
        '': {
            'handlers': ['default', 'file'],
            'level': 'INFO',
            'propagate': True
        }
    }
}

logging.config.dictConfig(LOGGING_CONFIG)
logger = logging.getLogger(__name__)
def store_data(data):
    session_id = str(uuid.uuid4())
    if USING_REDIS:
        redis_client.setex(session_id, 3600, json.dumps(data))
    else:
        with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
            json.dump(data, f)
    return session_id

def get_data(session_id):
    if USING_REDIS:
        data = redis_client.get(session_id)
        return json.loads(data) if data else None
    else:
        path = os.path.join(TEMP_DIR, f'prd_session_{session_id}.json')
        if os.path.exists(path):
            with open(path, 'r') as f:
                return json.load(f)
    return None

def monitor_agent_performance(func):
    """Decorator to monitor agent performance."""
    @wraps(func)
    def wrapper(agent_id, input_text, *args, **kwargs):
        start_time = time.time()
        try:
            result = func(agent_id, input_text, *args, **kwargs)
            duration = time.time() - start_time
            logging.info(f"Agent {agent_id} completed in {duration:.2f}s")
            return result
        except Exception as e:
            duration = time.time() - start_time
            logging.error(f"Agent {agent_id} failed after {duration:.2f}s: {str(e)}")
            raise
    return wrapper

async def wait_for_run_completion(thread_id, run_id, timeout=90, poll_interval=0.5):
    start_time = time.time()
    while time.time() - start_time < timeout:
        status = openai.beta.threads.runs.retrieve(run_id=run_id, thread_id=thread_id)
        if status.status == "completed":
            return status
        elif status.status == "failed":
            raise RuntimeError(f"Run {run_id} failed.")
        await asyncio.sleep(poll_interval)
    raise TimeoutError(f"Run {run_id} did not complete in {timeout} seconds.")

@monitor_agent_performance
def call_agent(agent_id, input_text):
    try:
        logging.info(f"[CALL START] Calling agent {agent_id}")
        thread = openai.beta.threads.create()
        
        if not thread or not thread.id:
            raise ValueError("Failed to create thread")
            
        message = openai.beta.threads.messages.create(
            thread_id=thread.id, 
            role="user", 
            content=input_text
        )
        if not message:
            raise ValueError("Failed to create message")
            
        run = openai.beta.threads.runs.create(
            thread_id=thread.id, 
            assistant_id=agent_id
        )
        if not run:
            raise ValueError("Failed to create run")

        start_time = time.time()
        while time.time() - start_time < 60:  # 60 second timeout
            status = openai.beta.threads.runs.retrieve(
                thread_id=thread.id, 
                run_id=run.id
            )
            if status.status == "completed":
                break
            time.sleep(0.25)
            
        messages = openai.beta.threads.messages.list(thread_id=thread.id)
        if not messages.data:
            raise ValueError("No messages returned")
            
        response = messages.data[0].content[0].text.value
        
        # Log the agent's response
        logging.info(f"""
        [AGENT RESPONSE] Agent {agent_id}:
        ----------------------------------------
        {response}
        ----------------------------------------
        """)
            
        return response

    except Exception as e:
        logging.error(f"[ERROR] Agent {agent_id} failed: {e}")
        return f"Error: {str(e)}"

@monitor_agent_performance
async def call_agent_async(agent_id, input_text):
    try:
        logging.info(f"[CALL START] Calling agent {agent_id}")
        thread = openai.beta.threads.create()

        message = openai.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=input_text
        )

        run = openai.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=agent_id
        )

        # Use async polling
        await wait_for_run_completion(thread.id, run.id)

        messages = openai.beta.threads.messages.list(thread_id=thread.id)
        response = messages.data[0].content[0].text.value

        logging.info(f"[AGENT RESPONSE] Agent {agent_id}:\n{response}")
        return response

    except Exception as e:
        logging.error(f"[ERROR] Agent {agent_id} failed: {e}")
        return f"Error: {str(e)}"


# Update call_agents_parallel to use retry logic
async def call_agents_parallel(agent_calls):
    tasks = [
        call_agent_with_cache(agent_id, input_text)
        for agent_id, input_text in agent_calls
    ]
    return await asyncio.gather(*tasks)

agent_semaphore = asyncio.Semaphore(3)  # Limit concurrent calls

async def call_agent_with_limit(agent_id, input_text):
    async with agent_semaphore:
        return await call_agent_with_cache(agent_id, input_text)
def verify_credentials(username, password):
    """Verify user credentials against environment variables."""
    return (username == os.getenv('ADMIN_USERNAME') and 
            password == os.getenv('ADMIN_PASSWORD'))

@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if verify_credentials(request.form.get('username'), request.form.get('password')):
            session['logged_in'] = True
            return redirect(url_for('page1'))
        return render_template('page0_login.html', error=True)
    return render_template('page0_login.html', error=False)


@app.route('/page1', methods=['GET', 'POST'])
def page1():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    if request.method == 'POST':             # Get form inputs
        inputs = {key: request.form[key] for key in ['industry', 'sector', 'geography', 'intent', 'features']}
        
        # Handle file upload
        if 'context_file' in request.files:
            file = request.files['context_file']
            if file and file.filename:
                try:
                    # Read file content
                    file_content = file.read().decode('utf-8')
                    # Add file content to inputs
                    inputs['context_file'] = file_content
                except Exception as e:
                    logging.error(f"File upload error: {str(e)}")
                    return "Error processing file upload", 400

        session.update(inputs)

       # Update context to include file content
        context_parts = [f"{k.replace('_', ' ').title()}: {v}" for k, v in inputs.items() 
                        if k != 'context_file']
        
        if 'context_file' in inputs:
            context_parts.append(f"\nAdditional Context:\n{inputs['context_file']}")
            
        context = "\n".join(context_parts)
        session_id = store_data({
            "inputs": inputs,
            "status": "processing"
        })
        session['data_key'] = session_id

        def run_agents():
            try:
                #a1 = call_agent(ASSISTANTS['agent_1'], context)
                a11, a2, a3 = asyncio.run(call_agents_parallel([
                    (ASSISTANTS['agent_1_1'], context),
                    (ASSISTANTS['agent_2'], context),
                    (ASSISTANTS['agent_3'], context)
                ]))
                logging.info(f"Agent 1.1 Output: {a11}")
                logging.info(f"Agent 2 Output: {a2}")
                final_data = {
                    #"agent_1_output": a1,
                    "product_overview": a11,
                    "feature_overview": a2,
                    "highest_order": a3,
                    "status": "complete"
                }
                if USING_REDIS:
                    redis_client.setex(session_id, 3600, json.dumps(final_data))
                else:
                    with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                        json.dump(final_data, f)
                logging.info(f"[AGENTS DONE] Page 1 background agents complete for session {session_id}")
            except Exception as e:
                logging.error(f"Error during agent execution: {e}")
                fail_data = {"status": "error", "message": str(e)}
                if USING_REDIS:
                    redis_client.setex(session_id, 3600, json.dumps(fail_data))
                else:
                    with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                        json.dump(fail_data, f)

        executor.submit(run_agents)

        return redirect('/page2')

    return render_template('page1_input.html')


@app.route('/page2', methods=['GET', 'POST'])
def page2():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    session_id = session.get('data_key', '')
    data = get_data(session_id) or {}

    # Wait for background job to complete
    start_time = time.time()
    timeout = 120  # max 30 seconds wait
    while True:
        if data.get("status") == "complete":
            break
        if data.get("status") == "error":
            return f"<h3>❌ Agent processing failed:</h3><pre>{data.get('message')}</pre>", 500
        time.sleep(1)
        if time.time() - start_time > timeout:
            return "Processing took too long. Please refresh the page in a few seconds.", 504
        data = get_data(session_id) or {}

    if request.method == 'POST':
        data.update({
            "product_overview": request.form.get("product_overview", ""),
            "feature_overview": request.form.get("feature_overview", "")
        })
        if USING_REDIS:
            redis_client.setex(session_id, 3600, json.dumps(data))
        return redirect('/page3')

    return render_template("page2_agents.html",
        agent11_output=data.get('product_overview', ''),
        agent2_output=data.get('feature_overview', '')
    )

@app.route('/update_content', methods=['POST'])
def update_content():
    if not session.get('logged_in'):
        return jsonify({'error': 'Not authenticated'}), 401
        
    try:
        data = request.get_json()
        content_type = data.get('type')
        new_content = data.get('content')
        
        if not content_type or not new_content:
            return jsonify({'error': 'Missing required fields'}), 400
            
        # Get stored data
        stored_data = get_data(session.get('data_key', '')) or {}
        
        # Update content based on type
        if content_type == 'product':
            # Re-run Agent 1.1 with new content
            new_response = call_agent(ASSISTANTS['agent_1_1'], new_content)
            stored_data['product_overview'] = new_response
        elif content_type == 'feature':
            # Re-run Agent 2 with new content
            new_response = call_agent(ASSISTANTS['agent_2'], new_content)
            stored_data['feature_overview'] = new_response
            
        # Store updated data
        if USING_REDIS:
            redis_client.setex(session['data_key'], 3600, json.dumps(stored_data))
        else:
            with open(os.path.join(TEMP_DIR, f'prd_session_{session["data_key"]}.json'), 'w') as f:
                json.dump(stored_data, f)
                
        return jsonify({'success': True, 'response': new_response})
        
    except Exception as e:
        logging.error(f"Error updating content: {str(e)}")
        return jsonify({'error': str(e)}), 500
    
@app.route('/page3', methods=['GET', 'POST'])
def page3():
    if not session.get('logged_in'): return redirect(url_for('login'))
    data = get_data(session.get('data_key', '')) or {}

    if request.method == 'POST':
        # Get user inputs and agent outputs
        user_inputs = data.get("inputs", {})
        feature_overview = data.get("feature_overview", "")
        # Combine all relevant inputs into structured format
        combined_input = f"""
        # Original User Inputs
        Industry: {user_inputs.get('industry', '')}
        Sector: {user_inputs.get('sector', '')}
        Geography: {user_inputs.get('geography', '')}
        Intent: {user_inputs.get('intent', '')}
        Features: {user_inputs.get('features', '')}

        # Feature Overview (Agent 2 Analysis)
        {feature_overview}
        """

        logging.info(f"[PAGE3] Combined input for agents: {combined_input}")

        keys = [k for k in ASSISTANTS if k.startswith("agent_4_")]
        # Use asyncio to call all agents in parallel
        results = asyncio.run(call_agents_parallel([
            (ASSISTANTS[k], combined_input) for k in keys
        ]))
        outputs = dict(zip(keys, results))

        data['combined_outputs'] = outputs
        session['combined_outputs'] = outputs

        if USING_REDIS:
            redis_client.setex(session['data_key'], 3600, json.dumps(data))
        else:
            with open(os.path.join(TEMP_DIR, f'prd_session_{session["data_key"]}.json'), 'w') as f:
                json.dump(data, f)

        return redirect('/page4')

    return render_template('page3_prompt_picker.html', highest_order=data.get('highest_order', ''))

_agent_response_cache = {}

async def call_agent_with_cache(agent_id, input_text):
    cache_key = f"{agent_id}_{hash(input_text)}"
    if cache_key in _agent_response_cache:
        logging.info(f"[CACHE HIT] Returning cached response for {agent_id}")
        return _agent_response_cache[cache_key]
    
    result = await call_agent_with_retry(agent_id, input_text)
    _agent_response_cache[cache_key] = result
    return result

def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

@app.route('/page4')
@require_auth
def page4():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    try:
        data = get_data(session.get('data_key', ''))
        if not data:
            logger.error("No data found for session key")
            return "No data found", 404  # Or render a friendly template

        outputs = {
            'product_overview': data.get('product_overview', ''),
            'feature_overview': data.get('feature_overview', ''),
            'agent_4_1': data.get('combined_outputs', {}).get('agent_4_1', ''),
            'agent_4_2': data.get('combined_outputs', {}).get('agent_4_2', ''),
            'agent_4_3': data.get('combined_outputs', {}).get('agent_4_3', ''),
            'agent_4_4': data.get('combined_outputs', {}).get('agent_4_4', ''),
            'agent_4_5': data.get('combined_outputs', {}).get('agent_4_5', '')
        }

        logger.info("[PAGE4] Rendering with outputs:")
        for key, value in outputs.items():
            logger.info(f"{key}: {value[:100]}...")

        return render_template('page4_final_output.html', outputs=outputs)

    except Exception as e:
        logger.error(f"Error in page4: {str(e)}")
        return str(e), 500




def add_hyperlink(paragraph, url, text=None):
    if text is None:
        text = url
    part = paragraph.part
    r_id = part.relate_to(
        url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    new_run.append(rPr)

    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def validate_reference(ref: str) -> bool:
    """Validate extracted reference."""
    if not ref or len(ref.strip()) < 5:
        return False
        
    # Check for common invalid patterns
    invalid_patterns = [
        r'^[\W_]+$',                    # Only special characters
        r'^(?:http)?[:/]+$',            # Incomplete URLs
        r'^Reference:\s*$',             # Empty reference labels
        r'^\d+$'                        # Just numbers
    ]
    
    return not any(re.match(pattern, ref.strip()) for pattern in invalid_patterns)

def extract_references_from_outputs(outputs):
    """Extract references from combined outputs."""
    all_refs = set()
    ref_patterns = [
        r'\[([^\]]+)\]\(([^\)]+)\)',                     # Markdown links
        r'Reference:\s+(.+?)(?=\n|$)',                   # Reference lines
        r'Source:\s+(.+?)(?=\n|$)',                      # Source lines
        r'(?:https?://\S+)',                             # Raw URLs
        r'(?<=Reference:).*?[\d]{4}.*?(?=\n|$)',        # Citations with years
        r'\((?:[^()]*?\d{4}[^()]*?)\)',                 # Parenthetical citations
        r'(?<=DOI:)\s*[\d.]+\/[\w.-]+',                 # DOI references
        r'(?:ISBN(?:-1[03])?:?\s*(?:[\d-]+))'          # ISBN references
    ]
    
    def extract_from_text(text):
        refs = set()
        for pattern in ref_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE)
            for match in matches:
                try:
                    # Handle markdown links differently
                    if '[' in pattern:
                        # Check if we have enough groups before accessing
                        if len(match.groups()) >= 2:
                            refs.add(match.group(2))  # Add URL from markdown link
                        else:
                            refs.add(match.group(0))  # Add entire match if not enough groups
                    else:
                        # For other patterns, add the full match
                        refs.add(match.group().strip())
                except (IndexError, AttributeError) as e:
                    logging.warning(f"Error extracting reference: {e}")
                    continue
        return refs

    # Extract from all outputs
    for output in outputs.values():
        if isinstance(output, str):
            try:
                refs = extract_from_text(output)
                all_refs.update(refs)
            except Exception as e:
                logging.error(f"Error processing output text: {e}")
                continue
    
    # Clean and sort references
    cleaned_refs = sorted(list({
        ref.strip(' .,[]()\"\'')  # Clean up punctuation
        for ref in all_refs
        if ref and len(ref.strip()) > 5  # Minimum length to filter noise
    }))
    
    logging.info(f"[REFERENCES] Extracted {len(cleaned_refs)} references")
    return cleaned_refs

limiter = Limiter(app=app, key_func=get_remote_address)

# Replace with the new route
@app.route('/generate_word_doc', methods=['POST'])
@limiter.limit("5 per minute")
def generate_word_doc():
    """Generate Word document from PRD content."""
    if not session.get('logged_in'):
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        # Get data and validate
        data = get_data(session.get("data_key", ""))
        if not data:
            logging.error("No data found in session")
            return jsonify({'error': 'No session data found'}), 404

        # Initialize document with styles
        doc = Document()
        initialize_document_styles(doc)

        # Map all sections
        sections = {
            "Product Overview": data.get("product_overview", ""),
            "Feature Overview": data.get("feature_overview", ""),
            "Product Requirements": data.get("combined_outputs", {}).get("agent_4_1", ""),
            "Functional Requirements": data.get("combined_outputs", {}).get("agent_4_2", ""),
            "Non-Functional Requirements": data.get("combined_outputs", {}).get("agent_4_3", ""),
            "Data Requirements": data.get("combined_outputs", {}).get("agent_4_4", ""),
            "Legal & Compliance Requirements": data.get("combined_outputs", {}).get("agent_4_5", "")
        }

        # Add title and metadata
        doc.add_heading("Product Requirements Document", level=1)
        doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()

        # Process each section
        for title, content in sections.items():
            if content and content.strip():
                try:
                    process_section(doc, title, content)
                except Exception as e:
                    logging.error(f"Error processing section {title}: {str(e)}")
                    doc.add_paragraph(f"Error in section {title}: {str(e)}", style="Error")

        # Add references
        add_references_section(doc, data.get("combined_outputs", {}))

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logging.info("Document generated successfully")
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"PRD_Draft_{time.strftime('%Y%m%d')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logging.error(f"Document generation failed: {str(e)}")
        return jsonify({
            'error': 'Document generation failed',
            'message': str(e)
        }), 500
    

def add_references_section(doc, outputs):
    """Add references section to document."""
    try:
        refs = extract_references_from_outputs(outputs)
        if not refs:
            return

        doc.add_page_break()
        doc.add_heading("References", level=2)
        
        for ref in refs:
            p = doc.add_paragraph(style="List Bullet")
            if ref.startswith(("http://", "https://")):
                add_hyperlink(p, ref)
            else:
                p.add_run(ref)
                
    except Exception as e:
        logging.error(f"Error adding references: {str(e)}")

def process_section(doc, title, raw_markdown):
    """Process a single section of the document."""
    try:
        doc.add_heading(title, level=2)
        
        # Convert markdown to HTML
        html_content = markdown(
            raw_markdown,
            extras=[
                "fenced-code-blocks",
                "tables",
                "header-ids",
                "break-on-newline"
            ]
        )
        
        # Parse HTML
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Process elements
        for el in soup.find_all(["p", "li", "h3", "h4", "pre", "code"]):
            process_element(doc, el)
            
    except Exception as e:
        logging.error(f"Error processing section {title}: {str(e)}")
        # Add error note in document
        doc.add_paragraph(f"Error processing section: {str(e)}", style="Error")

def process_element(doc, el):
    """Process a single HTML element."""
    text = el.get_text(strip=True)
    if not text:
        return

    para = doc.add_paragraph()
    
    # Handle requirement labels
    if match := re.match(r"^(F\d+|NFR\d+|DR\d+|LR\d+)\. (.+)", text):
        label, desc = match.groups()
        run = para.add_run(f"{label}. ")
        run.bold = True
        para.add_run(desc)
    # Handle labeled lines
    elif match := re.match(r"^([A-Z][\w\s]+?): (.+)", text):
        label, desc = match.groups()
        run = para.add_run(f"{label}: ")
        run.bold = True
        para.add_run(desc)
    # Regular text
    else:
        para.add_run(text)
        if el.name == "li":
            para.style = "List Bullet"
    
    para.paragraph_format.space_after = Pt(6)

@app.errorhandler(Exception)
def handle_exception(e):
    """Handle all unhandled exceptions."""
    logging.error(f"Unhandled exception: {str(e)}")
    return jsonify({
        'error': 'An unexpected error occurred',
        'message': str(e)
    }), 500

@app.errorhandler(404)
def not_found_error(e):
    """Handle 404 errors."""
    return jsonify({'error': 'Resource not found'}), 404

@app.errorhandler(401)
def unauthorized_error(e):
    """Handle 401 errors."""
    return redirect(url_for('login'))

def process_section(doc, title, raw_markdown):
    """Process a single section of the document."""
    try:
        doc.add_heading(title, level=2)
        
        # Convert markdown to HTML
        html_content = markdown(
            raw_markdown,
            extras=[
                "fenced-code-blocks",
                "tables",
                "header-ids",
                "break-on-newline"
            ]
        )
        
        # Parse HTML
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Process elements
        for el in soup.find_all(["p", "li", "h3", "h4", "pre", "code"]):
            process_element(doc, el)
            
    except Exception as e:
        logging.error(f"Error processing section {title}: {str(e)}")
        # Add error note in document
        doc.add_paragraph(f"Error processing section: {str(e)}", style="Error")

def initialize_document_styles(doc):
    """Initialize document styles."""
    styles = doc.styles
    
    # Error style
    error_style = styles.add_style('Error', WD_STYLE_TYPE.PARAGRAPH)
    error_style.font.color.rgb = RGBColor(255, 0, 0)  # Now RGBColor will be defined
    error_style.font.italic = True
    
    # Requirement style
    req_style = styles.add_style('Requirement', WD_STYLE_TYPE.PARAGRAPH)
    req_style.font.name = 'Calibri'
    req_style.font.size = Pt(11)
    
    return doc

def cleanup_expired_sessions():
    """Clean up expired session data with better error handling."""
    if USING_REDIS:
        return  # Redis handles expiration
        
    try:
        current_time = time.time()
        cleaned = 0
        failed = 0
        
        for filename in os.listdir(TEMP_DIR):
            if not filename.startswith('prd_session_'):
                continue
                
            filepath = os.path.join(TEMP_DIR, filename)
            try:
                if current_time - os.path.getmtime(filepath) > 3600:
                    os.remove(filepath)
                    cleaned += 1
            except OSError as e:
                failed += 1
                logging.error(f"Failed to remove {filepath}: {e}")
                
        logging.info(f"Session cleanup: {cleaned} removed, {failed} failed")
        
    except Exception as e:
        logging.error(f"Session cleanup failed: {e}")

class UserInputSchema(Schema):
    industry = fields.Str(required=True)
    sector = fields.Str(required=True)
    geography = fields.Str(required=True)
    intent = fields.Str(required=True)
    features = fields.Str(required=True)

def validate_form_input(form_data):
    """Validate form input against schema with improved sanitization."""
    try:
        # Sanitize input before validation
        sanitized_data = {
            key: value.strip() for key, value in form_data.items()
        }
        
        schema = UserInputSchema()
        validated_data = schema.load(sanitized_data)
        
        # Enhanced validation rules
        for key, value in validated_data.items():
            if len(value) < 3:
                return False, {key: ["Input too short (minimum 3 characters)"]}
            if len(value) > 1000:
                return False, {key: ["Input too long (maximum 1000 characters)"]}
            if any(char in value for char in '<>{}[]'):
                return False, {key: ["Invalid characters detected"]}
                
        return True, validated_data
        
    except ValidationError as e:
        logging.warning(f"Validation error: {e.messages}")
        return False, e.messages
    except Exception as e:
        logging.error(f"Unexpected validation error: {str(e)}")
        return False, {"error": "Internal validation error"}
    
def extend_session():
    """Extend session lifetime if user is active."""
    if session.get('logged_in'):
        session.permanent = True
        app.permanent_session_lifetime = timedelta(hours=1)
        session.modified = True

@app.before_request
def before_request():
    extend_session()
    # Add session cleanup check every 100 requests
    if random.random() < 0.01:  # 1% chance
        cleanup_expired_sessions()

@monitor_agent_performance
async def call_agent_with_retry(agent_id, input_text, max_retries=3):
    for attempt in range(max_retries):
        try:
            thread = openai.beta.threads.create(messages=[{"role": "user", "content": input_text}])
            run = openai.beta.threads.runs.create(thread_id=thread.id, assistant_id=agent_id)

            # Use the async event loop to wait non-blockingly
            status = await wait_for_run_completion(thread.id, run.id)

            messages = openai.beta.threads.messages.list(thread_id=thread.id)
            response = messages.data[0].content[0].text.value if messages.data else ""
            return response

        except Exception as e:
            print(f"[Retry {attempt+1}/{max_retries}] Error with agent {agent_id}: {e}")
            if attempt == max_retries - 1:
                return f"Error: {str(e)}"
            await asyncio.sleep(1.5)
            
def store_session_data(data):
    """Store large session data in Redis/files and keep only key in session."""
    try:
        if not data:
            raise ValueError("Cannot store empty data")
            
        session_id = str(uuid.uuid4())
        storage_key = f"prd_data_{session_id}"
        
        # Convert data to JSON
        try:
            json_data = json.dumps(data)
        except (TypeError, ValueError) as e:
            logging.error(f"JSON serialization failed: {e}")
            raise ValueError("Invalid data format")
        
        # Store main data with error handling
        if USING_REDIS:
            try:
                redis_client.setex(storage_key, 3600, json_data)
            except redis.RedisError as e:
                logging.error(f"Redis storage failed: {e}")
                # Fallback to file storage
                with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                    f.write(json_data)
        else:
            with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                f.write(json_data)
                
        # Store minimal data in session
        session['data_key'] = session_id
        logging.info(f"Session data stored with ID: {session_id}")
        return session_id
        
    except Exception as e:
        logging.error(f"Failed to store session data: {e}")
        raise

def get_session_data():
    """Retrieve session data from storage with validation."""
    try:
        session_id = session.get('data_key')
        if not session_id:
            logging.warning("No session key found")
            return None
            
        storage_key = f"prd_data_{session_id}"
        
        # Try Redis first
        if USING_REDIS:
            try:
                data = redis_client.get(storage_key)
                if data:
                    return json.loads(data)
            except (redis.RedisError, json.JSONDecodeError) as e:
                logging.error(f"Redis retrieval failed: {e}")
        
        # Fallback to file storage
        path = os.path.join(TEMP_DIR, f'prd_session_{session_id}.json')
        if os.path.exists(path):
            try:
                with open(path, 'r') as f:
                    return json.load(f)
            except (IOError, json.JSONDecodeError) as e:
                logging.error(f"File retrieval failed: {e}")
                
        return None
        
    except Exception as e:
        logging.error(f"Session data retrieval failed: {e}")
        return None
        
    return get_data(session_id)
# Add to main
if __name__ == '__main__':
    cleanup_expired_sessions()
    port = int(os.environ.get("PORT", 7001))
    app.run(host="0.0.0.0", port=port, debug=True)

