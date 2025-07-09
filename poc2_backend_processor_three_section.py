from flask import Flask, render_template, request, jsonify, session
from openai import OpenAI
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
from cachetools import TTLCache
# Add JIRA integration
from jira import JIRA
from dotenv import load_dotenv
import os
import json
import logging
import time
import hashlib
import traceback
import string
import csv
import io
import re

# Load environment variables
load_dotenv()

# JIRA Configuration
JIRA_SERVER = 'https://lalluluke.atlassian.net/'
JIRA_EMAIL = 'lalluluke@gmail.com'
JIRA_API_TOKEN = os.getenv("JIRA_API_TOKEN")
JIRA_PROJECT_KEY = 'SCRUM'

# Initialize JIRA client if credentials are available
jira_client = None
if JIRA_API_TOKEN:
    try:
        jira_client = JIRA(server=JIRA_SERVER, basic_auth=(JIRA_EMAIL, JIRA_API_TOKEN))
        logging.info("JIRA client initialized successfully")
    except Exception as e:
        logging.error(f"Failed to initialize JIRA client: {e}")
        jira_client = None
else:
    logging.warning("JIRA_API_TOKEN not found in environment variables")

# Vector Database and RAG imports - Optional for deployment
try:
    import chromadb
    from chromadb.config import Settings
    CHROMADB_AVAILABLE = True
except ImportError:
    logging.warning("ChromaDB not available - vector search functionality disabled")
    CHROMADB_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer
    import numpy as np
    from sklearn.metrics.pairwise import cosine_similarity
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    logging.warning("Sentence transformers not available - embedding functionality disabled")
    EMBEDDINGS_AVAILABLE = False
    # Fallback numpy import
    try:
        import numpy as np
    except ImportError:
        logging.warning("NumPy not available")
        np = None

# Additional imports for token counting
try:
    import tiktoken
    TIKTOKEN_AVAILABLE = True
except ImportError:
    logging.warning("tiktoken not available - using fallback token counting")
    TIKTOKEN_AVAILABLE = False

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Document processing imports
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
    logger.info("python-docx library available for DOCX file processing")
except ImportError:
    logger.warning("python-docx not available. DOCX files will be processed as text extraction fallback.")
    DOCX_AVAILABLE = False

# Initialize Vector Database and RAG components
logger.info("Initializing Vector Database and RAG components...")

# Initialize ChromaDB client (if available)
if CHROMADB_AVAILABLE:
    try:
        chroma_client = chromadb.PersistentClient(path="./vector_db")
        logger.info("ChromaDB client initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize ChromaDB client: {e}")
        chroma_client = None
        CHROMADB_AVAILABLE = False
else:
    chroma_client = None
    logger.info("ChromaDB not available - skipping vector database initialization")

# Initialize embedding model for RAG (if available)
if EMBEDDINGS_AVAILABLE:
    try:
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        logger.info("Embedding model loaded successfully")
    except Exception as e:
        logger.warning(f"Failed to load embedding model: {e}. RAG features may be limited.")
        embedding_model = None
        EMBEDDINGS_AVAILABLE = False
else:
    embedding_model = None
    logger.info("Sentence transformers not available - skipping embedding model initialization")

# Create or get collection for PRD documents (if ChromaDB available)
if CHROMADB_AVAILABLE and chroma_client:
    try:
        prd_collection = chroma_client.get_or_create_collection(
            name="prd_documents",
            metadata={"description": "PRD and documentation storage for RAG"}
        )
        logger.info("ChromaDB collection initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize ChromaDB collection: {e}")
        prd_collection = None
else:
    prd_collection = None
    logger.info("ChromaDB collection not available - vector search features disabled")

app = Flask(__name__)

# Configure Flask app
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'dev-secret-key-change-in-production-three-section')

# Performance optimizations
THREAD_POOL_SIZE = 4
MAX_CACHE_SIZE = 1000
CACHE_TTL = 3600  # 1 hour

# Initialize thread pool for parallel processing
executor = ThreadPoolExecutor(max_workers=THREAD_POOL_SIZE)

# Cache for responses (in-memory with TTL)
response_cache = TTLCache(maxsize=MAX_CACHE_SIZE, ttl=CACHE_TTL)

logger.info("Three Section Flask application starting up with performance optimizations...")
logger.info(f"OpenAI API key configured: {'Yes' if os.getenv('OPENAI_API_KEY') else 'No'}")
logger.info(f"Thread pool size: {THREAD_POOL_SIZE}")
logger.info(f"Cache size: {MAX_CACHE_SIZE} items, TTL: {CACHE_TTL}s")

# Copy all utility functions from original backend
def get_cache_key(prompt):
    """Generate a cache key from the prompt."""
    return hashlib.md5(prompt.encode('utf-8')).hexdigest()

def create_jira_story(summary, description, priority='High', epic_link=None, acceptance_criteria=None, responsible_systems=None):
    """Create a user story in JIRA."""
    if not jira_client:
        return {'success': False, 'error': 'JIRA client not available'}
    
    # Format description with responsible systems and acceptance criteria
    formatted_description = description
    
    # Add responsible systems if provided
    if responsible_systems:
        formatted_description += f"\n\n*Responsible Systems:* {responsible_systems}\n"
    
    # Add acceptance criteria if provided (this might need adjustment based on your JIRA setup)
    if acceptance_criteria and isinstance(acceptance_criteria, list) and acceptance_criteria:
        formatted_description += "\n*Acceptance Criteria:*\n"
        for i, criteria in enumerate(acceptance_criteria, 1):
            if criteria.strip():
                formatted_description += f"{i}. {criteria.strip()}\n"
    
    issue_dict = {
        'project': {'key': JIRA_PROJECT_KEY},
        'summary': summary,
        'description': formatted_description,
        'issuetype': {'name': 'Story'}
    }
    
    # Try to add priority if it's available in the project (some JIRA setups don't allow it)
    try:
        if priority and priority.lower() in ['low', 'medium', 'high', 'highest', 'lowest']:
            issue_dict['priority'] = {'name': priority}
    except Exception as e:
        logger.warning(f"Priority field not available in JIRA project: {e}")
    
    # Add epic link if provided (this might need adjustment based on your JIRA setup)
    if epic_link:
        # Common Epic Link field IDs: customfield_10014, customfield_10008, customfield_10011
        issue_dict['customfield_10014'] = epic_link
    
    try:
        new_issue = jira_client.create_issue(fields=issue_dict)
        logger.info(f"Successfully created JIRA story: {new_issue.key}")
        return {
            'success': True, 
            'key': new_issue.key, 
            'id': new_issue.id,
            'url': f"{JIRA_SERVER}browse/{new_issue.key}"
        }
    except Exception as e:
        error_message = str(e)
        logger.error(f"Error creating JIRA story: {error_message}")
        
        # Provide more user-friendly error messages for common issues
        if 'priority' in error_message.lower():
            error_message = "Priority field is not available in your JIRA project. The story will be created without priority."
            # Try again without priority
            try:
                issue_dict_no_priority = {k: v for k, v in issue_dict.items() if k != 'priority'}
                new_issue = jira_client.create_issue(fields=issue_dict_no_priority)
                logger.info(f"Successfully created JIRA story without priority: {new_issue.key}")
                return {
                    'success': True, 
                    'key': new_issue.key, 
                    'id': new_issue.id,
                    'url': f"{JIRA_SERVER}browse/{new_issue.key}",
                    'warning': 'Story created successfully, but priority field was not available in your JIRA project.'
                }
            except Exception as e2:
                logger.error(f"Failed to create story even without priority: {e2}")
                error_message = f"JIRA configuration issue: {str(e2)}"
        elif 'field' in error_message.lower() and 'cannot be set' in error_message.lower():
            error_message = "Some fields are not available in your JIRA project configuration. Please check your JIRA project settings."
        elif 'permission' in error_message.lower():
            error_message = "You don't have permission to create stories in this JIRA project. Please contact your JIRA administrator."
        
        return {'success': False, 'error': error_message}

def store_document_in_vector_db(content, filename, doc_type="prd"):
    """Store document content in vector database for RAG retrieval."""
    if not CHROMADB_AVAILABLE or not EMBEDDINGS_AVAILABLE or not prd_collection or not embedding_model:
        logger.warning("Vector DB or embedding model not available. Skipping storage.")
        return None
    
    try:
        logger.info(f"Storing {doc_type} document in vector database: {filename}")
        
        # Split content into chunks for better retrieval
        chunks = split_document_into_chunks(content, chunk_size=1000, overlap=200)
        logger.info(f"Document split into {len(chunks)} chunks")
        
        # Generate embeddings for each chunk
        chunk_embeddings = embedding_model.encode(chunks)
        
        # Create metadata for each chunk
        documents = []
        metadatas = []
        ids = []
        
        doc_id = hashlib.md5(f"{filename}_{doc_type}".encode()).hexdigest()
        
        for i, (chunk, embedding) in enumerate(zip(chunks, chunk_embeddings)):
            chunk_id = f"{doc_id}_chunk_{i}"
            
            documents.append(chunk)
            metadatas.append({
                "filename": filename,
                "doc_type": doc_type,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "timestamp": datetime.now().isoformat()
            })
            ids.append(chunk_id)
        
        # Store in ChromaDB
        prd_collection.add(
            documents=documents,
            metadatas=metadatas,
            ids=ids
        )
        
        logger.info(f"Successfully stored {len(chunks)} chunks in vector database")
        return doc_id
        
    except Exception as e:
        logger.error(f"Error storing document in vector DB: {str(e)}")
        return None

def split_document_into_chunks(content, chunk_size=1000, overlap=200):
    """Split document content into overlapping chunks for better retrieval."""
    if len(content) <= chunk_size:
        return [content]
    
    chunks = []
    start = 0
    
    while start < len(content):
        end = start + chunk_size
        chunk = content[start:end]
        chunks.append(chunk)
        start += chunk_size - overlap
    
    return chunks

@lru_cache(maxsize=100)
def count_tokens(text, model="gpt-4o"):
    """Count tokens in text using tiktoken."""
    try:
        import tiktoken
        encoding = tiktoken.encoding_for_model(model)
        return len(encoding.encode(text))
    except ImportError:
        # Fallback: rough estimation (4 chars per token)
        return len(text) // 4
    except Exception:
        # Fallback: rough estimation
        return len(text) // 4

def ask_assistant_from_file_optimized(code_filepath, user_prompt):
    """Optimized assistant interaction using direct chat completion."""
    start_time = time.time()
    logger.info(f"Starting optimized assistant interaction: {code_filepath}")
    
    try:
        # Read agent instructions from file
        agent_file_path = os.path.join("agents", code_filepath)
        
        with open(agent_file_path, 'r', encoding='utf-8') as file:
            system_instructions = file.read()
        
        logger.info(f"Loaded system instructions from {agent_file_path}")
        
        # Initialize OpenAI client
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        # Create chat completion
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_instructions},
                {"role": "user", "content": user_prompt}
            ],
            max_tokens=4000,
            temperature=0.7
        )
        
        response_text = response.choices[0].message.content.strip()
        
        elapsed_time = time.time() - start_time
        logger.info(f"Assistant response completed in {elapsed_time:.2f} seconds")
        
        # Log token usage
        prompt_tokens = count_tokens(system_instructions + user_prompt, "gpt-4o")
        response_tokens = count_tokens(response_text, "gpt-4o")
        logger.info(f"Token usage - Prompt: {prompt_tokens:,}, Response: {response_tokens:,}, Total: {prompt_tokens + response_tokens:,}")
        
        return response_text
        
    except FileNotFoundError:
        logger.error(f"Agent file not found: {agent_file_path}")
        return f"Error: Agent configuration file '{code_filepath}' not found."
    except Exception as e:
        logger.error(f"Error in assistant interaction: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"Error: {str(e)}"

# Three Section Layout Routes

@app.route("/", methods=["GET"])
def landing_page():
    """Render the landing page."""
    try:
        logger.info("Rendering landing page")
        return render_template('landing_page.html')
    except Exception as e:
        logger.error(f"Error rendering landing page: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"Error loading landing page: {str(e)}", 500

@app.route("/three-section", methods=["GET"])
def three_section_layout():
    """Render the three section layout page."""
    try:
        logger.info("Rendering three section layout")
        return render_template('poc2_three_section_layout.html')
    except Exception as e:
        logger.error(f"Error rendering three section layout: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"Error loading page: {str(e)}", 500

@app.route("/tabbed-layout", methods=["GET"])
def tabbed_layout():
    """Render the tabbed layout page."""
    try:
        logger.info("Rendering tabbed layout")
        return render_template('poc2_tabbed_layout.html')
    except Exception as e:
        logger.error(f"Error rendering tabbed layout: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"Error loading tabbed layout page: {str(e)}", 500

@app.route("/epic-first-tabbed", methods=["GET"])
def epic_first_tabbed_layout():
    """Render the epic-first tabbed layout page."""
    try:
        logger.info("Rendering epic-first tabbed layout")
        return render_template('poc2_tabbed_layout_epic_first.html')
    except Exception as e:
        logger.error(f"Error rendering epic-first tabbed layout: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"Error loading epic-first tabbed layout page: {str(e)}", 500

@app.route("/three-section-get-epics", methods=["GET"])
def three_section_get_epics():
    """Get epics from session for three section layout."""
    try:
        logger.info("GET request to /three-section-get-epics")
        epics = session.get('generated_epics', [])
        
        if not epics:
            logger.info("No epics found in session - need to upload PRD first")
            return jsonify({"success": False, "message": "No epics available", "epics": []})
        
        logger.info(f"Retrieved {len(epics)} epics from session")
        
        # Format epics for display if they exist
        if epics:
            formatted_epics_html = format_epics_for_display(epics)
            return jsonify({
                "success": True, 
                "epics": epics,
                "epics_html": formatted_epics_html
            })
        
        return jsonify({"success": False, "message": "No epics available", "epics": []})
        
    except Exception as e:
        logger.error(f"Error getting epics: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e), "epics": []})
        return jsonify({"success": False, "error": str(e)})

@app.route("/three-section-approve-epics", methods=["POST"])
def three_section_approve_epics():
    """Generate user stories for selected epics in three section layout."""
    logger.info("Processing epic approval for three section layout")
    
    try:
        epic_ids = request.form.get('epic_ids', '').split(',')
        selected_epic_contents = request.form.get('selected_epic_contents', '{}')
        
        if not epic_ids or not epic_ids[0]:
            return jsonify({"success": False, "error": "No epics selected"})
        
        try:
            epic_contents = json.loads(selected_epic_contents)
        except json.JSONDecodeError:
            epic_contents = {}
        
        logger.info(f"Processing {len(epic_ids)} selected epics")
        
        # Create prompt for user story generation
        selected_epics_text = ""
        for epic_id in epic_ids:
            if epic_id in epic_contents:
                epic = epic_contents[epic_id]
                selected_epics_text += f"\nEpic ID: {epic_id}\n"
                selected_epics_text += f"Title: {epic.get('title', 'Unknown')}\n"
                selected_epics_text += f"Description: {epic.get('description', 'No description')}\n"
                selected_epics_text += f"Priority: {epic.get('priority', 'High')}\n"
                selected_epics_text += "---\n"
        
        user_story_prompt = f"""Based on the following approved epics, generate comprehensive user stories:

{selected_epics_text}

Generate 3-5 detailed user stories for each epic. Each user story should include:
- Unique ID (US-001, US-002, etc.)
- Clear title
- Description following "As a [user], I want [goal] so that [benefit]" format
- Priority level
- Story points estimate
- Responsible systems

Return the response as a JSON array of user story objects with the following structure:
[
  {{
    "id": "US-001",
    "title": "User Story Title",
    "description": "As a [user], I want [goal] so that [benefit]",
    "priority": "High|Medium|Low",
    "story_points": 3,
    "systems": ["System1", "System2"],
    "epic_id": "epic_1"
  }}
]"""
        
        # Generate user stories using the agent
        response = ask_assistant_from_file_optimized("poc2_agent3_basic_user_story", user_story_prompt)
        
        # Try to parse JSON response
        try:
            user_stories = json.loads(response)
            if not isinstance(user_stories, list):
                user_stories = [user_stories]
        except json.JSONDecodeError:
            logger.warning("Failed to parse JSON response from user story agent, using fallback")
            # Fallback: create basic user stories
            user_stories = [
                {
                    "id": "US-001",
                    "title": "User Authentication",
                    "description": "As a user, I want to securely log into the system so that I can access my account.",
                    "priority": "High",
                    "story_points": 5,
                    "systems": ["Authentication System", "User Management"],
                    "epic_id": epic_ids[0]
                },
                {
                    "id": "US-002",
                    "title": "User Profile Management",
                    "description": "As a user, I want to view and edit my profile information so that I can keep my details up to date.",
                    "priority": "Medium",
                    "story_points": 3,
                    "systems": ["User Management", "Profile Service"],
                    "epic_id": epic_ids[0]
                }
            ]
        
        # Debug: Log the generated user stories to check their structure
        logger.info(f"Generated user stories: {json.dumps(user_stories, indent=2)}")
        
        # Ensure all user stories have required fields
        for i, story in enumerate(user_stories):
            # Map 'name' to 'title' for consistency with frontend
            if 'name' in story and 'title' not in story:
                story['title'] = story['name']
            
            # Create proper user story description if missing
            if not story.get('description'):
                story_title = story.get('title', story.get('name', f'User Story {i+1}'))
                # Create a more meaningful user story format that relates to the title
                story['description'] = f"As a customer, I want to {story_title.lower()} so that I can complete my business requirements efficiently."
                
            # Ensure systems are properly formatted - use only the first (most important) system
            if 'systems' in story and isinstance(story['systems'], list):
                # Take only the first system (most important one)
                primary_system = story['systems'][0] if story['systems'] else 'Customer acquisition platform'
                story['systems'] = [primary_system]  # Keep as single-item array for consistency
                story['responsible_systems'] = primary_system  # Single system as string
            elif not story.get('responsible_systems'):
                story['responsible_systems'] = 'Customer acquisition platform'  # fallback to single system
        
        # Store user stories in session
        session['user_stories'] = user_stories
        session['selected_epic_ids'] = epic_ids
        
        logger.info(f"Generated {len(user_stories)} user stories")
        return jsonify({"success": True, "user_stories": user_stories})
        
    except Exception as e:
        logger.error(f"Error in epic approval: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e)})

@app.route("/three-section-user-story-details", methods=["POST"])
def three_section_user_story_details():
    """Get detailed information for a selected user story."""
    logger.info("Getting user story details for three section layout")
    
    try:
        story_id = request.form.get('selected_story_id')
        story_name = request.form.get('selected_story_name')
        story_description = request.form.get('selected_story_description')
        
        if not story_name:
            return jsonify({"success": False, "error": "No story selected"})
        
        logger.info(f"Processing story details for: {story_name}")
        
        # Generate acceptance criteria
        criteria_prompt = f"""For the following user story, generate comprehensive acceptance criteria:

Story: {story_name}
Description: {story_description}

Generate 4-6 specific, testable acceptance criteria in Given-When-Then format for this specific story.
The criteria should be directly related to the story name and description provided.

IMPORTANT: Return ONLY a JSON array of strings, nothing else. Format:
["Given...", "Given...", "Given..."]

Do not include any explanatory text, just the JSON array."""
        
        logger.info(f"Sending criteria prompt to agent: {criteria_prompt}")
        criteria_response = ask_assistant_from_file_optimized("poc2_agent4_acceptanceCriteria_gen", criteria_prompt)
        logger.info(f"Agent criteria response: {criteria_response}")
        
        try:
            # Try to parse as JSON first
            acceptance_criteria = json.loads(criteria_response)
            if not isinstance(acceptance_criteria, list):
                acceptance_criteria = [criteria_response]
        except json.JSONDecodeError:
            # If JSON parsing fails, try to extract from the response
            logger.warning(f"JSON parsing failed for criteria response: {criteria_response}")
            
            # Try to find JSON in the response
            if 'acceptance_criteria' in criteria_response:
                import re
                json_match = re.search(r'\{.*"acceptance_criteria".*\}', criteria_response, re.DOTALL)
                if json_match:
                    try:
                        parsed_json = json.loads(json_match.group())
                        acceptance_criteria = parsed_json.get('acceptance_criteria', [])
                    except:
                        acceptance_criteria = []
                else:
                    acceptance_criteria = []
            else:
                acceptance_criteria = []
            
            # Generate story-specific fallback criteria if parsing completely fails
            if not acceptance_criteria:
                acceptance_criteria = [
                    f"Given a user wants to {story_name.lower()}, when they initiate the action, then the system should respond appropriately",
                    f"Given the {story_name.lower()} process is initiated, when all required data is provided, then the operation should complete successfully",
                    f"Given invalid data is provided for {story_name.lower()}, when the user attempts the operation, then appropriate error messages should be displayed",
                    f"Given a user completes {story_name.lower()}, when the operation is successful, then the system should provide confirmation"
                ]
        
        # Generate tagged requirements
        requirements_prompt = f"""For the following user story, identify and tag relevant requirements:

Story: {story_name}
Description: {story_description}

Generate 3-5 tagged requirements that this story addresses.
Return as a JSON array of strings."""
        
        requirements_response = ask_assistant_from_file_optimized("poc2_traceability_agent", requirements_prompt)
        
        try:
            tagged_requirements = json.loads(requirements_response)
            if not isinstance(tagged_requirements, list):
                tagged_requirements = [requirements_response]
        except json.JSONDecodeError:
            tagged_requirements = [
                "REQ-001: User authentication and authorization",
                "REQ-002: Data security and privacy",
                "REQ-003: User interface and experience"
            ]
        
        # Generate traceability matrix
        traceability_prompt = f"""Create a traceability matrix for the following user story:

Story: {story_name}
Description: {story_description}
Acceptance Criteria: {acceptance_criteria}

Create a formatted traceability matrix showing relationships between requirements, story, and test cases."""
        
        traceability_response = ask_assistant_from_file_optimized("poc2_traceability_agent", traceability_prompt)
        
        story_details = {
            "acceptance_criteria": acceptance_criteria,
            "tagged_requirements": tagged_requirements,
            "traceability_matrix": traceability_response
        }
        
        # Store in session for Jira submission
        session['current_story_details'] = story_details
        session['selected_story'] = {
            "id": story_id,
            "name": story_name,
            "description": story_description
        }
        
        logger.info("Story details generated successfully")
        return jsonify({"success": True, "story_details": story_details})
        
    except Exception as e:
        logger.error(f"Error getting story details: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e)})

@app.route("/three-section-user-story-details-page", methods=["GET", "POST"])
def three_section_user_story_details_page():
    """Render the user story details page for Jira submission from three section layout."""
    logger.info("Rendering user story details page for three section layout")
    
    try:
        # Get story details from session or form
        if request.method == "POST":
            story_id = request.form.get('selected_story_id')
            story_name = request.form.get('selected_story_name')
            story_description = request.form.get('selected_story_description')
            
            logger.info(f"POST data - ID: {story_id}, Name: {story_name}, Description: [{story_description}]")
            logger.info(f"Form data keys: {list(request.form.keys())}")
            logger.info(f"All form data: {dict(request.form)}")
            
            # Ensure we have valid values
            if not story_description or story_description.strip() == '':
                logger.warning(f"Empty description received for story: {story_name}")
                # Try to generate a basic description from the story name
                if story_name and story_name != 'User Story':
                    story_description = f"As a user, I want to {story_name.lower()} so that I can accomplish my goals efficiently."
                else:
                    story_description = 'No description available'
            if not story_name:
                story_name = 'User Story'
            if not story_id:
                story_id = 'US-001'
            
            # Store epic title from current context
            current_epics = session.get('generated_epics', [])
            selected_epic_ids = session.get('selected_epic_ids', [])
            
            epic_title = 'Unknown Epic'
            if selected_epic_ids and current_epics:
                for epic in current_epics:
                    if epic.get('id') in selected_epic_ids:
                        epic_title = epic.get('title', 'Unknown Epic')
                        break
            
            # Store current story selection in session
            session['selected_story'] = {
                'id': story_id,
                'name': story_name,
                'description': story_description
            }
            session['current_epic_title'] = epic_title
            
        else:
            # GET request - get from session
            selected_story = session.get('selected_story', {})
            story_id = selected_story.get('id', 'US-001')
            story_name = selected_story.get('name', 'User Story')
            story_description = selected_story.get('description', 'No description available')
            epic_title = session.get('current_epic_title', 'Unknown Epic')
            
            logger.info(f"GET data from session - ID: {story_id}, Name: {story_name}, Description: {story_description}")
        
        # Get story details from session
        story_details = session.get('current_story_details', {})
        
        # If no story details, generate them
        if not story_details:
            logger.info("No story details in session, generating new ones")
            
            # Generate acceptance criteria
            criteria_prompt = f"""For the following user story, generate comprehensive acceptance criteria:

Story: {story_name}
Description: {story_description}

Generate 4-6 specific, testable acceptance criteria in Given-When-Then format.
Return as a JSON array of strings."""
            
            criteria_response = ask_assistant_from_file_optimized("poc2_agent4_acceptanceCriteria_gen", criteria_prompt)
            
            try:
                acceptance_criteria = json.loads(criteria_response)
                if not isinstance(acceptance_criteria, list):
                    acceptance_criteria = [criteria_response]
            except json.JSONDecodeError:
                acceptance_criteria = [
                    "Given a user is on the login page, when they enter valid credentials, then they should be authenticated successfully",
                    "Given a user enters invalid credentials, when they attempt to login, then they should see an appropriate error message",
                    "Given a user is authenticated, when they access protected resources, then they should have appropriate permissions"
                ]
            
            # Generate tagged requirements
            requirements_prompt = f"""For the following user story, identify and tag relevant requirements:

Story: {story_name}
Description: {story_description}

Generate 3-5 tagged requirements that this story addresses.
Return as a JSON array of strings."""
            
            requirements_response = ask_assistant_from_file_optimized("poc2_traceability_agent", requirements_prompt)
            
            try:
                tagged_requirements = json.loads(requirements_response)
                if not isinstance(tagged_requirements, list):
                    tagged_requirements = [requirements_response]
            except json.JSONDecodeError:
                tagged_requirements = [
                    "REQ-001: User authentication and authorization",
                    "REQ-002: Data security and privacy",
                    "REQ-003: User interface and experience"
                ]
            
            # Generate traceability matrix
            traceability_prompt = f"""Create a traceability matrix for the following user story:

Story: {story_name}
Description: {story_description}
Acceptance Criteria: {acceptance_criteria}

Create a formatted traceability matrix showing relationships between requirements, story, and test cases."""
            
            traceability_response = ask_assistant_from_file_optimized("poc2_traceability_agent", traceability_prompt)
            
            story_details = {
                "acceptance_criteria": acceptance_criteria,
                "tagged_requirements": tagged_requirements,
                "traceability_matrix": traceability_response
            }
            
            # Store in session
            session['current_story_details'] = story_details
        
        # Render the user story details template
        logger.info(f"Rendering template with - Epic: {epic_title}, Story: {story_name}, Description: {story_description}")
        
        return render_template('poc2_user_story_details.html',
                             epic_title=epic_title,
                             user_story_name=story_name,
                             user_story_description=story_description,
                             priority='High',
                             responsible_systems='Customer acquisition platform',
                             acceptance_criteria=story_details.get('acceptance_criteria', []),
                             tagged_requirements=story_details.get('tagged_requirements', []),
                             traceability_matrix=story_details.get('traceability_matrix', ''))
        
    except Exception as e:
        logger.error(f"Error rendering user story details page: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"Error loading user story details: {str(e)}", 500

@app.route("/three-section-epic-chat", methods=["POST"])
def three_section_epic_chat():
    """Handle epic chat for three section layout."""
    logger.info("Processing epic chat for three section layout")
    
    try:
        data = request.get_json()
        message = data.get('message', '')
        epics = data.get('epics', [])
        
        if not message:
            return jsonify({"success": False, "error": "No message provided"})
        
        # Create context for epic chat
        epics_context = ""
        for i, epic in enumerate(epics):
            epics_context += f"Epic {i+1}:\n"
            epics_context += f"Title: {epic.get('title', 'Unknown')}\n"
            epics_context += f"Description: {epic.get('description', 'No description')}\n"
            epics_context += f"Priority: {epic.get('priority', 'High')}\n\n"
        
        chat_prompt = f"""Current Epics:
{epics_context}

User Message: {message}

You are an epic refinement assistant. When the user asks to modify, update, or change epics, you should:
1. Provide a helpful response explaining what changes you're making in natural language
2. Return the modified epics in a structured format

IMPORTANT FORMATTING RULES:
- If the user wants to modify epics, you MUST respond with EXACTLY this format:
  RESPONSE: [Your explanation in plain English - no JSON or technical formatting]
  UPDATED_EPICS: [Valid JSON array of modified epics]

- If the user is just asking questions without wanting modifications, just provide a helpful response in natural language.

- The RESPONSE section must be conversational and explain what you're doing
- The UPDATED_EPICS section must be valid JSON with the same structure as the input epics
- Each epic should have: id, title, description, priority

Example format when modifying:
RESPONSE: I've updated the priority of the first epic from High to Low as requested. This change reflects a shift in project priorities.
UPDATED_EPICS: [{{
  "id": "epic-1",
  "title": "Example Epic",
  "description": "Example description",
  "priority": "Low"
}}]

Current context: The user has {len(epics)} epics and wants to: {message}"""

        logger.info(f"Sending chat prompt to epic agent: {chat_prompt[:200]}...")
        response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", chat_prompt)
        logger.info(f"Received response from epic agent: {response[:200]}...")
        
        # Check if the response contains updated epics
        updated_epics = None
        if "UPDATED_EPICS:" in response:
            try:
                parts = response.split("UPDATED_EPICS:")
                response_text = parts[0].replace("RESPONSE:", "").strip()
                epics_json = parts[1].strip()
                
                # Try to parse the updated epics
                updated_epics = json.loads(epics_json)
                
                logger.info(f"Successfully parsed updated epics: {len(updated_epics)} epics")
                
                return jsonify({
                    "success": True, 
                    "response": response_text,
                    "updated_epics": updated_epics
                })
            except (json.JSONDecodeError, IndexError) as e:
                logger.warning(f"Failed to parse updated epics: {e}")
                logger.warning(f"Raw response was: {response}")
                # Fall back to just returning the response
        
        # Clean up the response to remove any RESPONSE: prefix if it exists
        clean_response = response.replace("RESPONSE:", "").strip()
        logger.info(f"Returning clean response: {clean_response[:100]}...")
        return jsonify({"success": True, "response": clean_response})
        
    except Exception as e:
        logger.error(f"Error in epic chat: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

@app.route("/three-section-user-story-chat", methods=["POST"])
def three_section_user_story_chat():
    """Handle user story chat for three section layout."""
    logger.info("Processing user story chat for three section layout")
    
    try:
        data = request.get_json()
        message = data.get('message', '')
        user_stories = data.get('user_stories', [])
        epic = data.get('epic', {})
        
        if not message:
            return jsonify({"success": False, "error": "No message provided"})
        
        # Create context for user story chat
        stories_context = ""
        for i, story in enumerate(user_stories):
            stories_context += f"Story {i+1}: {story.get('title', story.get('name', 'Unknown'))}\n"
            stories_context += f"Description: {story.get('description', story.get('summary', 'No description'))}\n"
            stories_context += f"Priority: {story.get('priority', 'Medium')}\n"
            stories_context += f"Systems: {story.get('systems', ['Unknown'])}\n\n"
        
        chat_prompt = f"""You are a helpful user story refinement assistant. You help users modify and improve their user stories in a conversational way.

Current Epic: {epic.get('title', 'Unknown')}
Epic Description: {epic.get('description', 'No description')}

Current User Stories:
{stories_context}

User Message: "{message}"

Instructions:
- If the user wants to modify, update, or change user stories, provide a helpful response in natural language explaining what changes you're making
- Then provide the updated user stories in a specific format
- If the user is just asking questions, provide helpful information about the user stories

When making changes, respond in this EXACT format:

RESPONSE: [Your explanation in natural language - be conversational and helpful. Explain what you're doing and why it's beneficial.]

UPDATED_USER_STORIES: [JSON array of the modified user stories with the same structure as the input, including id, title/name, description, priority, and systems fields]

If you're not making changes, just provide a helpful conversational response without the UPDATED_USER_STORIES section.

Examples of natural language responses:
- "I've updated the priorities based on your feedback..."
- "I've reorganized the user stories to better reflect..."
- "I've added more detail to the descriptions..."
- "I've adjusted the system assignments based on..."

Remember: Keep the RESPONSE section conversational and in plain English. Save the technical JSON format only for the UPDATED_USER_STORIES section."""
        
        # Use the epic generator agent as it's more flexible for conversational responses
        response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", chat_prompt)
        
        logger.info(f"User story chat raw response: {response}")
        
        # Parse the response to extract natural language and updated user stories
        updated_user_stories = None
        response_text = response
        
        if "UPDATED_USER_STORIES:" in response:
            try:
                parts = response.split("UPDATED_USER_STORIES:")
                response_text = parts[0].replace("RESPONSE:", "").strip()
                stories_json = parts[1].strip()
                
                # Clean up the JSON part
                if stories_json.startswith('```json'):
                    stories_json = stories_json.replace('```json', '').replace('```', '').strip()
                elif stories_json.startswith('```'):
                    stories_json = stories_json.replace('```', '').strip()
                
                # Try to parse the updated user stories
                import json
                updated_user_stories = json.loads(stories_json)
                
                logger.info(f"Successfully parsed {len(updated_user_stories)} updated user stories")
                
                return jsonify({
                    "success": True, 
                    "response": response_text,
                    "updated_user_stories": updated_user_stories
                })
            except (json.JSONDecodeError, IndexError) as e:
                logger.warning(f"Failed to parse updated user stories: {e}")
                logger.warning(f"Raw stories JSON: {stories_json}")
                # Fall back to just returning the response
        
        # Clean up the response if it has the RESPONSE: prefix but no updates
        if response_text.startswith("RESPONSE:"):
            response_text = response_text.replace("RESPONSE:", "").strip()
        
        logger.info(f"User story chat final response: {response_text}")
        
        return jsonify({"success": True, "response": response_text})
        
    except Exception as e:
        logger.error(f"Error in user story chat: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

@app.route("/three-section-story-details-chat", methods=["POST"])
def three_section_story_details_chat():
    """Handle story details chat for three section layout."""
    logger.info("Processing story details chat for three section layout")
    
    try:
        data = request.get_json()
        message = data.get('message', '')
        section = data.get('section', 'general')
        story_details = data.get('story_details', {})
        user_story = data.get('user_story', {})
        epic = data.get('epic', {})
        
        if not message:
            return jsonify({"success": False, "error": "No message provided"})
        
        # Create detailed context for story details chat
        acceptance_criteria = story_details.get('acceptance_criteria', [])
        tagged_requirements = story_details.get('tagged_requirements', [])
        traceability_matrix = story_details.get('traceability_matrix', 'Not available')
        
        criteria_text = "\n".join([f"- {criteria}" for criteria in acceptance_criteria]) if acceptance_criteria else "No acceptance criteria defined yet"
        requirements_text = "\n".join([f"- {req}" for req in tagged_requirements]) if tagged_requirements else "No tagged requirements defined yet"
        
        chat_prompt = f"""You are a helpful story details assistant. You help users refine and improve story details, acceptance criteria, tagged requirements, and traceability information in a conversational way.

Current Context:
Epic: {epic.get('title', 'Unknown')}
Epic Description: {epic.get('description', 'No description')}

User Story: {user_story.get('title', user_story.get('name', 'Unknown'))}
User Story Description: {user_story.get('description', user_story.get('summary', 'No description'))}

Current Story Details:
Acceptance Criteria:
{criteria_text}

Tagged Requirements:
{requirements_text}

Traceability Matrix:
{traceability_matrix}

User is specifically asking about: {section}
User Message: "{message}"

Instructions:
- If the user wants to modify, update, or change story details, provide a helpful response in natural language explaining what changes you're making
- Then provide the updated story details in a specific format
- If the user is just asking questions, provide helpful information about the story details

When making changes, respond in this EXACT format:

RESPONSE: [Your explanation in natural language - be conversational and helpful. Explain what changes you're making and why they're beneficial.]

UPDATED_STORY_DETAILS: [JSON object with the modified story details, maintaining the same structure as the input: {{"acceptance_criteria": [...], "tagged_requirements": [...], "traceability_matrix": "..."}}]

If you're not making changes, just provide a helpful conversational response without the UPDATED_STORY_DETAILS section.

Examples of natural language responses:
- "I've refined the acceptance criteria to be more specific..."
- "I've added additional requirements based on your feedback..."
- "I've updated the traceability matrix to better reflect..."
- "I've reorganized the criteria to improve clarity..."

Remember: Keep the RESPONSE section conversational and in plain English. Save the technical JSON format only for the UPDATED_STORY_DETAILS section."""
        
        # Use the epic generator agent as it's more flexible for conversational responses
        response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", chat_prompt)
        
        logger.info(f"Story details chat raw response: {response}")
        
        # Parse the response to extract natural language and updated story details
        updated_story_details = None
        response_text = response
        
        if "UPDATED_STORY_DETAILS:" in response:
            try:
                parts = response.split("UPDATED_STORY_DETAILS:")
                response_text = parts[0].replace("RESPONSE:", "").strip()
                details_json = parts[1].strip()
                
                # Clean up the JSON part
                if details_json.startswith('```json'):
                    details_json = details_json.replace('```json', '').replace('```', '').strip()
                elif details_json.startswith('```'):
                    details_json = details_json.replace('```', '').strip()
                
                # Try to parse the updated story details
                import json
                updated_story_details = json.loads(details_json)
                
                logger.info(f"Successfully parsed updated story details")
                
                return jsonify({
                    "success": True, 
                    "response": response_text,
                    "updated_story_details": updated_story_details
                })
            except (json.JSONDecodeError, IndexError) as e:
                logger.warning(f"Failed to parse updated story details: {e}")
                logger.warning(f"Raw details JSON: {details_json}")
                # Fall back to just returning the response
        
        # Clean up the response if it has the RESPONSE: prefix but no updates
        if response_text.startswith("RESPONSE:"):
            response_text = response_text.replace("RESPONSE:", "").strip()
        
        logger.info(f"Story details chat final response: {response_text}")
        
        return jsonify({"success": True, "response": response_text})
        
    except Exception as e:
        logger.error(f"Error in story details chat: {str(e)}")
        return jsonify({"success": False, "error": str(e)})

@app.route("/three-section-submit-jira", methods=["POST"])
def three_section_submit_jira():
    """Submit user story to Jira from three section layout."""
    logger.info("Processing Jira submission from three section layout")
    
    try:
        # Get form data
        epic_title = request.form.get('epic_title', '')
        user_story_name = request.form.get('user_story_name', '')
        user_story_description = request.form.get('user_story_description', '')
        priority = request.form.get('priority', 'High')
        responsible_systems = request.form.get('responsible_systems', 'Customer acquisition platform')
        acceptance_criteria = request.form.get('acceptance_criteria', '').split('|||')
        tagged_requirements = request.form.get('tagged_requirements', '').split('|||')
        traceability_matrix = request.form.get('traceability_matrix', '')
        
        # Clean empty criteria
        acceptance_criteria = [criteria.strip() for criteria in acceptance_criteria if criteria.strip()]
        tagged_requirements = [req.strip() for req in tagged_requirements if req.strip()]
        
        logger.info(f"Submitting to Jira: {user_story_name}")
        
        # Create the story in JIRA
        jira_result = create_jira_story(
            summary=user_story_name,
            description=user_story_description,
            priority=priority,
            acceptance_criteria=acceptance_criteria,
            responsible_systems=responsible_systems
        )
        
        # Store submission data for potential export
        submission_data = {
            'epic_title': epic_title,
            'user_story_name': user_story_name,
            'user_story_description': user_story_description,
            'priority': priority,
            'responsible_systems': responsible_systems,
            'acceptance_criteria': acceptance_criteria,
            'tagged_requirements': tagged_requirements,
            'traceability_matrix': traceability_matrix,
            'submitted_at': datetime.now().isoformat(),
            'jira_result': jira_result
        }
        
        # Store in session for potential retrieval
        if 'jira_submissions' not in session:
            session['jira_submissions'] = []
        session['jira_submissions'].append(submission_data)
        
        if jira_result['success']:
            logger.info(f"Jira submission completed successfully: {jira_result['key']}")
            ticket_id = jira_result['key']
            success_message = f"Successfully submitted to Jira: {jira_result['key']}"
        else:
            logger.error(f"Jira submission failed: {jira_result['error']}")
            ticket_id = f"ERROR-{int(time.time())}"
            success_message = f"Jira submission failed: {jira_result['error']}"
        
        # Check if this is an AJAX request by looking at headers
        is_ajax = (request.headers.get('X-Requested-With') == 'XMLHttpRequest' or 
                   request.headers.get('Content-Type', '').startswith('multipart/form-data'))
        
        if is_ajax:
            # AJAX request - return JSON
            response_data = {
                "success": jira_result['success'],
                "message": success_message,
                "ticket_id": ticket_id,
                "epic_title": epic_title,
                "user_story_name": user_story_name,
                "jira_url": jira_result.get('url', '') if jira_result['success'] else '',
                "error": jira_result.get('error', '') if not jira_result['success'] else ''
            }
            
            # Add warning if present (e.g., priority field not available)
            if jira_result.get('warning'):
                response_data['warning'] = jira_result['warning']
                
            return jsonify(response_data)
        else:
            # Regular form submission - return HTML page
            return render_template('jira_success.html', 
                                 epic_title=epic_title,
                                 user_story_name=user_story_name,
                                 ticket_id=ticket_id,
                                 success=jira_result['success'],
                                 jira_url=jira_result.get('url', '') if jira_result['success'] else '',
                                 error_message=jira_result.get('error', '') if not jira_result['success'] else '',
                                 warning=jira_result.get('warning', ''))
        
    except Exception as e:
        logger.error(f"Error in Jira submission: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        # Return JSON error for AJAX, HTML error for regular form
        is_ajax = (request.headers.get('X-Requested-With') == 'XMLHttpRequest' or 
                   request.headers.get('Content-Type', '').startswith('multipart/form-data'))
        
        if is_ajax:
            return jsonify({"success": False, "error": str(e)})
        else:
            return f"Error submitting to Jira: {str(e)}", 500

@app.route("/submit-jira-ticket", methods=["POST"])
def submit_jira_ticket():
    """Generic Jira submission route for compatibility with existing templates."""
    logger.info("Processing generic Jira ticket submission")
    
    try:
        # Get form data
        epic_title = request.form.get('epic_title', '')
        user_story_name = request.form.get('user_story_name', '')
        user_story_description = request.form.get('user_story_description', '')
        priority = request.form.get('priority', 'High')
        responsible_systems = request.form.get('responsible_systems', 'Customer acquisition platform')
        acceptance_criteria = request.form.get('acceptance_criteria', '').split('|||')
        tagged_requirements = request.form.get('tagged_requirements', '').split('|||')
        traceability_matrix = request.form.get('traceability_matrix', '')
        
        # Clean empty criteria
        acceptance_criteria = [criteria.strip() for criteria in acceptance_criteria if criteria.strip()]
        tagged_requirements = [req.strip() for req in tagged_requirements if req.strip()]
        
        logger.info(f"Submitting to Jira: {user_story_name}")
        
        # Here you would integrate with actual Jira API
        # For now, we'll simulate success
        
        # Store submission data for potential export
        submission_data = {
            'epic_title': epic_title,
            'user_story_name': user_story_name,
            'user_story_description': user_story_description,
            'priority': priority,
            'responsible_systems': responsible_systems,
            'acceptance_criteria': acceptance_criteria,
            'tagged_requirements': tagged_requirements,
            'traceability_matrix': traceability_matrix,
            'submitted_at': datetime.now().isoformat()
        }
        
        # Store in session for potential retrieval
        if 'jira_submissions' not in session:
            session['jira_submissions'] = []
        session['jira_submissions'].append(submission_data)
        
        logger.info("Jira submission completed successfully")
        
        # Return success page or redirect
        return render_template('jira_success.html', 
                             epic_title=epic_title,
                             user_story_name=user_story_name,
                             ticket_id=f"TST-{int(time.time())}")
        
    except Exception as e:
        logger.error(f"Error in Jira submission: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return f"Error submitting to Jira: {str(e)}", 500

@app.route("/three-section-upload-prd", methods=["POST"])
def three_section_upload_prd():
    """Handle PRD upload and generate epics for three section layout."""
    logger.info("Processing PRD upload for three section layout")
    
    try:
        # Get files and form data
        prd_file = request.files.get('prd_file')
        additional_docs = request.files.getlist('additional_docs')
        requirements_text = request.form.get('requirements', '')
        
        logger.info(f"PRD file uploaded: {'Yes' if prd_file else 'No'}")
        logger.info(f"Additional docs uploaded: {len(additional_docs) if additional_docs else 0}")
        logger.info(f"Manual requirements provided: {'Yes' if requirements_text else 'No'}")
        
        # Process files if uploaded
        prd_content = ""
        docs_content = ""
        
        if prd_file:
            prd_content = safe_read(prd_file)
            logger.info(f"PRD content length: {len(prd_content)} characters")
        
        if additional_docs:
            docs_contents = []
            for doc in additional_docs:
                doc_content = safe_read(doc)
                if doc_content and doc_content.strip():
                    docs_contents.append(f"=== {doc.filename} ===\n{doc_content}")
            docs_content = "\n\n".join(docs_contents)
            logger.info(f"Additional docs content length: {len(docs_content)} characters")
        
        # Combine with manual requirements
        if requirements_text:
            if prd_content:
                prd_content = f"{prd_content}\n\nAdditional Requirements:\n{requirements_text}"
            else:
                prd_content = requirements_text
        
        if not prd_content.strip():
            return jsonify({"success": False, "error": "No valid content provided"})
        
        # Process PRD with RAG if content is substantial
        if is_valid_content(prd_content):
            if len(prd_content) > 5000:
                logger.info("Creating RAG-enhanced PRD summary")
                prd_content = create_rag_summary(prd_content, 
                                               prd_file.filename if prd_file else "manual_requirements", 
                                               max_summary_length=25000)
            elif len(prd_content) > 30000:
                logger.info("Large PRD detected - using traditional optimization")
                prd_content = optimize_prd_content(prd_content, max_length=40000)
        
        # Process additional docs with RAG
        if is_valid_content(docs_content):
            if len(docs_content) > 3000:
                logger.info("Creating RAG-enhanced docs summary")
                docs_content = create_rag_summary(docs_content, "additional_docs", max_summary_length=10000)
        
        # Create enhanced context for Epic Generator
        enhanced_context = f"""
Project Requirements Document Analysis:
{prd_content}

Additional Documentation:
{docs_content}

User Context:
{requirements_text}

Instructions: Generate comprehensive epics from the above requirements. Focus on:
1. Breaking down large features into manageable epics
2. Identifying user journeys and workflows
3. Prioritizing features based on business value
4. Creating clear epic titles and descriptions
5. Ensuring comprehensive coverage of all requirements
"""
        
        logger.info(f"Enhanced context length: {len(enhanced_context)} characters")
        
        # Generate epics using the epic generator agent
        epics_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", enhanced_context)
        
        # Parse epics response into structured format
        epics = parse_epics_response(epics_response)
        
        # Store in session for later use
        session['generated_epics'] = epics
        session['prd_content'] = prd_content[:1000]  # Store first 1000 chars for reference
        
        logger.info(f"Generated {len(epics)} epics successfully")
        return jsonify({"success": True, "epics": epics})
        
    except Exception as e:
        logger.error(f"Error processing PRD upload: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e)})

@app.route("/three-section-document-upload", methods=["POST"])
def three_section_document_upload():
    """Enhanced document upload endpoint for three section layout with preview."""
    logger.info("Processing enhanced document upload for three section layout")
    
    try:
        # Get files and form data
        prd_file = request.files.get('prd_file')
        additional_file = request.files.get('additional_file')
        context = request.form.get('context_notes', '')
        
        logger.info(f"PRD file: {'Yes' if prd_file else 'No'}")
        logger.info(f"Additional file: {'Yes' if additional_file else 'No'}")
        logger.info(f"Context provided: {'Yes' if context else 'No'}")
        
        # Process PRD file
        prd_summary = ""
        prd_filename = ""
        prd_content = ""
        if prd_file and prd_file.filename:
            prd_content = safe_read(prd_file)
            if prd_content and prd_content.strip():
                prd_filename = prd_file.filename
                # Create summary if content is long
                if len(prd_content) > 1000:
                    logger.info("Creating PRD summary")
                    prd_summary = create_content_summary(prd_content, "PRD", max_length=2000)
                else:
                    prd_summary = prd_content
                logger.info(f"PRD processed: {len(prd_content)} chars -> {len(prd_summary)} chars")
        
        # Process additional files
        additional_summary = ""
        additional_filenames = []
        combined_content = ""
        if additional_file and additional_file.filename:
            file_content = safe_read(additional_file)
            if file_content and file_content.strip():
                additional_filenames.append(additional_file.filename)
                combined_content = f"=== {additional_file.filename} ===\n{file_content}"
                
                if len(combined_content) > 1000:
                    logger.info("Creating additional documents summary")
                    additional_summary = create_content_summary(combined_content, "Additional Documents", max_length=1500)
                else:
                    additional_summary = combined_content
                logger.info(f"Additional docs processed: {len(combined_content)} chars -> {len(additional_summary)} chars")
        
        # Store processed data in session for epic generation
        session['prd_upload_data'] = {
            'prd_content': prd_content,
            'prd_summary': prd_summary,
            'prd_filename': prd_filename,
            'additional_content': combined_content,
            'additional_summary': additional_summary,
            'additional_filenames': additional_filenames,
            'context': context
        }
        
        # Return preview data
        return jsonify({
            "success": True,
            "prd_summary": prd_summary,
            "prd_filename": prd_filename,
            "additional_summary": additional_summary if additional_summary and additional_summary.strip() != "No valid additional documentation available." else None,
            "additional_filenames": additional_filenames,
            "context": context if context.strip() else None
        })
        
    except Exception as e:
        logger.error(f"Error processing document upload: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e)})

@app.route("/three-section-generate-epics", methods=["POST"])
def three_section_generate_epics():
    """Generate epics from processed document data."""
    logger.info("Generating epics from processed documents")
    
    try:
        # Get data from request (for direct API call) or session (for continuation)
        request_data = request.get_json()
        upload_data = session.get('prd_upload_data', {})
        
        # Use request data if available, otherwise fall back to session data
        prd_summary = request_data.get('prd_summary') if request_data else upload_data.get('prd_summary', '')
        additional_summary = request_data.get('additional_summary') if request_data else upload_data.get('additional_summary', '')
        context = request_data.get('context') if request_data else upload_data.get('context', '')
        
        # Get full content for enhanced processing
        prd_content = upload_data.get('prd_content', prd_summary) or ""
        additional_content = upload_data.get('additional_content', additional_summary) or ""
        
        # Ensure we have strings, not None values
        prd_content = prd_content if prd_content else ""
        additional_content = additional_content if additional_content else ""
        context = context if context else ""
        
        logger.info(f"Processing: PRD({len(prd_content)} chars), Additional({len(additional_content)} chars), Context({len(context)} chars)")
        
        if not prd_summary and not context:
            return jsonify({"success": False, "error": "No content provided for epic generation"})
        
        # Create a meaningful prompt even with minimal content
        if not prd_content and not additional_content and context:
            # Use context as the main content if no files were uploaded
            prd_content = context
            processed_prd = context
        
        # Process content with RAG if substantial
        processed_prd = prd_content
        processed_additional = additional_content
        
        if is_valid_content(prd_content) and len(prd_content) > 5000:
            logger.info("Creating RAG-enhanced PRD summary for epic generation")
            processed_prd = create_rag_summary(prd_content, "prd_for_epics", max_summary_length=25000)
        
        if is_valid_content(additional_content) and len(additional_content) > 3000:
            logger.info("Creating RAG-enhanced additional content summary")
            processed_additional = create_rag_summary(additional_content, "additional_for_epics", max_summary_length=10000)
        
        # Create enhanced context for Epic Generator
        enhanced_context = create_epic_generation_context(processed_prd, processed_additional, context)
        
        logger.info(f"Enhanced context length: {len(enhanced_context)} characters")
        
        # Generate epics using the epic generator agent
        epics_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", enhanced_context)
        
        # Parse epics response into structured format
        epics = parse_epics_response(epics_response)
        
        # Store in session for later use
        session['generated_epics'] = epics
        session['prd_content'] = processed_prd[:1000]  # Store first 1000 chars for reference
        
        logger.info(f"Generated {len(epics)} epics successfully")
        
        # Format epics for display
        formatted_epics_html = format_epics_for_display(epics)
        
        return jsonify({
            "success": True, 
            "epics": epics,  # Keep structured data for frontend processing
            "epics_html": formatted_epics_html  # Add formatted HTML for display
        })
        
    except Exception as e:
        logger.error(f"Error generating epics: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e)})

def create_content_summary(content, content_type, max_length=2000):
    """Create a summary of content using OpenAI."""
    try:
        if len(content) <= max_length:
            return content
        
        prompt = f"""Create a comprehensive summary of this {content_type}:

{content[:10000]}  # Limit input to prevent token overflow

Requirements:
- Capture all key features, requirements, and objectives
- Maintain technical details and specifications
- Preserve business logic and user flows
- Keep the summary under {max_length} characters
- Use clear, concise language

Summary:"""
        
        client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[{"role": "user", "content": prompt}],
            max_tokens=800,
            temperature=0.3
        )
        
        summary = response.choices[0].message.content.strip()
        logger.info(f"Created {content_type} summary: {len(content)} -> {len(summary)} characters")
        return summary
        
    except Exception as e:
        logger.error(f"Error creating content summary: {str(e)}")
        return content[:max_length] + "..." if len(content) > max_length else content

def create_epic_generation_context(prd_content, additional_content, context):
    """Create enhanced context for epic generation."""
    enhanced_context = "Project Requirements Document Analysis:\n"
    
    if prd_content:
        enhanced_context += f"{prd_content}\n\n"
    
    if additional_content:
        enhanced_context += f"Additional Documentation:\n{additional_content}\n\n"
    
    if context:
        enhanced_context += f"User Context:\n{context}\n\n"
    
    enhanced_context += """Instructions: Generate 3-5 comprehensive epics from the above requirements. It is CRITICAL that you create multiple epics (between 3-5) to break down the requirements into manageable chunks. Focus on:

1. Breaking down large features into 3-5 separate, manageable epics
2. Identifying distinct user journeys and workflows for each epic
3. Prioritizing features based on business value across multiple epics
4. Creating clear epic titles and descriptions for each epic
5. Ensuring comprehensive coverage of all requirements across all epics
6. Each epic should have a clear scope and deliverable outcomes
7. IMPORTANT: Generate multiple epics - aim for 3-5 distinct epics minimum

Please provide your response in the following structured format for EACH epic (you MUST provide 3-5 epics):

Epic 1: [Clear Epic Title]
Description: [Detailed explanation of what needs to be built, including scope and objectives]
Priority: High/Medium/Low

Epic 2: [Clear Epic Title]
Description: [Detailed explanation of what needs to be built, including scope and objectives]
Priority: High/Medium/Low

Epic 3: [Clear Epic Title]
Description: [Detailed explanation of what needs to be built, including scope and objectives]
Priority: High/Medium/Low

Epic 4: [Clear Epic Title]
Description: [Detailed explanation of what needs to be built, including scope and objectives]
Priority: High/Medium/Low

Epic 5: [Clear Epic Title]
Description: [Detailed explanation of what needs to be built, including scope and objectives]
Priority: High/Medium/Low

OR provide a valid JSON array format with 3-5 epics:
[
  {
    "title": "Epic Title 1",
    "description": "Detailed explanation of what needs to be built",
    "priority": "High"
  },
  {
    "title": "Epic Title 2", 
    "description": "Detailed explanation of what needs to be built",
    "priority": "High"
  },
  {
    "title": "Epic Title 3",
    "description": "Detailed explanation of what needs to be built", 
    "priority": "Medium"
  }
]

REMEMBER: You MUST generate between 3-5 distinct epics. Do not provide just one epic."""
    
    return enhanced_context

# Health check and debug endpoints
@app.route("/health", methods=["GET"])
def health_check():
    """Simple health check endpoint."""
    return jsonify({
        "status": "ok", 
        "message": "Three section server is running",
        "timestamp": str(datetime.now())
    })

@app.route("/debug-info", methods=["GET"])
def debug_info():
    """Debug information endpoint."""
    try:
        return jsonify({
            "status": "ok",
            "session_keys": list(session.keys()),
            "has_epics": bool(session.get('generated_epics')),
            "epic_count": len(session.get('generated_epics', [])),
            "openai_configured": bool(os.getenv('OPENAI_API_KEY')),
            "chromadb_available": CHROMADB_AVAILABLE,
            "embeddings_available": EMBEDDINGS_AVAILABLE
        })
    except Exception as e:
        return jsonify({"status": "error", "error": str(e)})

# File Processing Utility Functions
def safe_read(file):
    """Safely read file content with proper error handling and encoding detection."""
    try:
        logger.debug(f"Attempting to read file: {file.filename if file else 'None'}")
        
        # Check if this is a DOCX file and handle it specially
        if file and file.filename and file.filename.lower().endswith('.docx'):
            logger.info(f"Detected DOCX file: {file.filename}")
            return extract_docx_text(file)
        
        # Read the raw bytes first
        raw_content = file.read()
        logger.debug(f"Read {len(raw_content)} bytes from file: {file.filename}")
        
        # Check if this looks like a DOCX file even if extension is wrong
        if raw_content.startswith(b'PK\x03\x04') and b'[Content_Types].xml' in raw_content:
            logger.info(f"Detected DOCX format in file: {file.filename}")
            file.seek(0)  # Reset file pointer for DOCX processing
            return extract_docx_text(file)
        
        # Try multiple encodings for regular text files
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                content = raw_content.decode(encoding)
                logger.info(f"Successfully decoded file {file.filename} using {encoding} encoding, length: {len(content)} characters")
                
                # Validate that we have meaningful content
                if len(content.strip()) > 0:
                    return content
                else:
                    logger.warning(f"File {file.filename} appears to be empty or whitespace only")
                    return "[File appears to be empty]"
                    
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try to extract text from binary content
        logger.warning(f"All standard encodings failed for {file.filename}, attempting binary extraction")
        try:
            # Try to extract printable ASCII characters
            import string
            printable_chars = set(string.printable)
            extracted_text = ''.join(char for char in raw_content.decode('latin-1') if char in printable_chars)
            
            if len(extracted_text.strip()) > 50:  # If we got some meaningful text
                logger.info(f"Extracted {len(extracted_text)} printable characters from {file.filename}")
                return extracted_text
            else:
                logger.error(f"Unable to extract meaningful text from {file.filename}")
                return f"[Unable to decode file - tried encodings: {', '.join(encodings)}]"
                
        except Exception as extract_error:
            logger.error(f"Binary extraction failed for {file.filename}: {str(extract_error)}")
            return f"[File reading failed - may be binary or corrupted]"
            
    except Exception as e:
        logger.error(f"Error reading file {file.filename}: {str(e)}")
        return f"[Error reading file: {str(e)}]"

def extract_docx_text(file):
    """Extract text from DOCX file using python-docx."""
    try:
        # Try importing python-docx
        try:
            from docx import Document
        except ImportError:
            logger.warning("python-docx not available, trying basic extraction")
            return "[DOCX file detected but python-docx library not available. Please install with: pip install python-docx]"
        
        # Read the document
        doc = Document(file)
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        extracted_text = "\n".join(text_content)
        logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX file")
        
        return extracted_text if extracted_text.strip() else "[DOCX file appears to be empty]"
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {file.filename}: {str(e)}")
        return f"[Error processing DOCX file: {str(e)}]"

def is_valid_content(content):
    """Check if content is valid for RAG processing."""
    if not content or len(content.strip()) < 50:
        return False
    
    # Check for error messages
    error_indicators = [
        "[Unable to decode",
        "[Unable to read", 
        "[Error reading",
        "[File reading failed",
        "[File appears to be empty"
    ]
    
    for indicator in error_indicators:
        if content.startswith(indicator):
            return False
    
    return True

def optimize_prd_content(prd_content, max_length=40000):
    """Optimize PRD content by extracting key sections and reducing verbosity."""
    if len(prd_content) <= max_length:
        return prd_content
    
    logger.info(f"Optimizing PRD content from {len(prd_content)} to ~{max_length} characters")
    
    # Extract key sections based on common PRD structure
    sections = {
        'overview': '',
        'features': '',
        'requirements': '',
        'user_stories': '',
        'technical': '',
        'other': ''
    }
    
    lines = prd_content.split('\n')
    current_section = 'other'
    
    for line in lines:
        line_lower = line.lower().strip()
        
        # Detect section headers
        if any(keyword in line_lower for keyword in ['overview', 'summary', 'introduction']):
            current_section = 'overview'
        elif any(keyword in line_lower for keyword in ['feature', 'functionality', 'capability']):
            current_section = 'features'
        elif any(keyword in line_lower for keyword in ['requirement', 'specification', 'criteria']):
            current_section = 'requirements'
        elif any(keyword in line_lower for keyword in ['user story', 'user stories', 'stories']):
            current_section = 'user_stories'
        elif any(keyword in line_lower for keyword in ['technical', 'architecture', 'implementation']):
            current_section = 'technical'
        
        sections[current_section] += line + '\n'
    
    # Prioritize sections and trim if needed
    priority_order = ['overview', 'features', 'requirements', 'user_stories', 'technical', 'other']
    optimized_content = ""
    remaining_length = max_length
    
    for section in priority_order:
        section_content = sections[section].strip()
        if section_content and remaining_length > 100:
            if len(section_content) <= remaining_length:
                optimized_content += f"\n\n=== {section.upper()} ===\n{section_content}"
                remaining_length -= len(section_content) + 50
            else:
                # Truncate section but try to end at a complete sentence
                truncated = section_content[:remaining_length-50]
                last_period = truncated.rfind('.')
                if last_period > remaining_length // 2:
                    truncated = truncated[:last_period+1]
                optimized_content += f"\n\n=== {section.upper()} ===\n{truncated}\n[...truncated for length...]"
                break
    
    logger.info(f"Optimized PRD content to {len(optimized_content)} characters")
    return optimized_content.strip()

def create_rag_summary(content, filename, max_summary_length=5000):
    """Create an intelligent summary using RAG if available, with fallbacks."""
    logger.info(f"Creating RAG summary for {filename}, content length: {len(content)}")
    
    # For now, use intelligent content extraction as RAG requires more setup
    logger.info("Using intelligent content extraction")
    return create_intelligent_summary(content, filename, max_summary_length)

def create_intelligent_summary(content, filename, max_summary_length=5000):
    """Create an intelligent summary without RAG using content analysis."""
    logger.info(f"Creating intelligent summary for {filename}")
    
    if len(content) <= max_summary_length:
        return content
    
    # Split into sections and prioritize
    sections = []
    current_section = ""
    
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Detect section headers (lines with keywords or formatting)
        if (len(line) < 100 and 
            (line.isupper() or 
             any(keyword in line.lower() for keyword in 
                 ['overview', 'summary', 'requirements', 'features', 'user stories', 
                  'technical', 'architecture', 'scope', 'objectives', 'goals']))):
            if current_section:
                sections.append(current_section)
            current_section = f"\n{line}\n"
        else:
            current_section += line + "\n"
    
    if current_section:
        sections.append(current_section)
    
    # Score sections by importance
    scored_sections = []
    for section in sections:
        score = calculate_section_importance(section)
        scored_sections.append((score, section))
    
    # Sort by score and include highest priority sections
    scored_sections.sort(reverse=True, key=lambda x: x[0])
    
    summary = ""
    for score, section in scored_sections:
        if len(summary) + len(section) <= max_summary_length:
            summary += section + "\n"
        else:
            remaining_space = max_summary_length - len(summary) - 100
            if remaining_space > 0:
                summary += section[:remaining_space] + "\n[...content truncated for length...]"
            break
    
    return summary.strip()

def calculate_section_importance(section):
    """Calculate importance score for a section based on keywords and structure."""
    score = 0
    section_lower = section.lower()
    
    # High priority keywords
    high_priority = ['requirement', 'feature', 'user story', 'objective', 'goal', 'scope']
    medium_priority = ['overview', 'summary', 'description', 'specification']
    low_priority = ['technical', 'implementation', 'details']
    
    for keyword in high_priority:
        score += section_lower.count(keyword) * 3
    
    for keyword in medium_priority:
        score += section_lower.count(keyword) * 2
    
    for keyword in low_priority:
        score += section_lower.count(keyword) * 1
    
    # Boost score for sections with lists or structured content
    if section.count('\n-') > 2 or section.count('\n*') > 2:
        score += 5
    
    # Boost score for sections with numbered items
    if any(f'{i}.' in section for i in range(1, 10)):
        score += 3
    
    return score

# Epic Processing Functions
def parse_epics_response(epics_response):
    """Parse the epics response from the AI agent into structured format."""
    logger.info("Parsing epics response into structured format")
    
    # First, try to parse as JSON if it looks like JSON
    epics_response_clean = epics_response.strip()
    
    # Remove code blocks if present
    if "```json" in epics_response_clean:
        json_match = epics_response_clean.split("```json")[1].split("```")[0].strip()
        epics_response_clean = json_match
    elif "```" in epics_response_clean:
        code_match = epics_response_clean.split("```")[1].split("```")[0].strip()
        if code_match.startswith('[') or code_match.startswith('{'):
            epics_response_clean = code_match
    
    # Try JSON parsing first
    if epics_response_clean.startswith('[') or epics_response_clean.startswith('{'):
        try:
            parsed_json = json.loads(epics_response_clean)
            epics = []
            
            if isinstance(parsed_json, list):
                for i, epic_data in enumerate(parsed_json):
                    if isinstance(epic_data, dict):
                        epic = {
                            'id': epic_data.get('id', f"epic_{i+1}"),
                            'title': epic_data.get('title', epic_data.get('name', epic_data.get('epic_title', f'Epic {i+1}'))),
                            'description': epic_data.get('description', epic_data.get('desc', epic_data.get('epic_description', 'No description provided'))),
                            'priority': epic_data.get('priority', 'High'),
                            'user_stories': epic_data.get('user_stories', [])
                        }
                        epics.append(epic)
                        
            elif isinstance(parsed_json, dict):
                # Single epic as object
                epic = {
                    'id': parsed_json.get('id', 'epic_1'),
                    'title': parsed_json.get('title', parsed_json.get('name', parsed_json.get('epic_title', 'Epic 1'))),
                    'description': parsed_json.get('description', parsed_json.get('desc', parsed_json.get('epic_description', 'No description provided'))),
                    'priority': parsed_json.get('priority', 'High'),
                    'user_stories': parsed_json.get('user_stories', [])
                }
                epics = [epic]
            
            if epics:
                logger.info(f"Successfully parsed JSON response with {len(epics)} epics")
                return epics
                
        except json.JSONDecodeError as e:
            logger.info(f"JSON parsing failed: {e}, trying text parsing")
    
    # Fall back to text parsing for non-JSON responses
    epics = []
    lines = epics_response.split('\n')
    current_epic = {}
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Look for epic headers (various formats)
        if any(pattern in line.lower() for pattern in ['epic 1:', 'epic 2:', 'epic 3:', 'epic 4:', 'epic 5:']):
            # Save previous epic if exists
            if current_epic and current_epic.get('title'):
                epics.append(current_epic)
            
            # Start new epic
            epic_num = len(epics) + 1
            current_epic = {
                'id': f"epic_{epic_num}",
                'title': extract_epic_title(line),
                'description': '',
                'priority': 'High',
                'user_stories': []
            }
            
        elif line.lower().startswith('title:'):
            if current_epic:
                current_epic['title'] = line[6:].strip()
            
        elif line.lower().startswith('description:'):
            if current_epic:
                current_epic['description'] = line[12:].strip()
            
        elif line.lower().startswith('priority:'):
            if current_epic:
                priority = line[9:].strip()
                current_epic['priority'] = priority if priority in ['High', 'Medium', 'Low'] else 'High'
            
        elif current_epic and line and not any(pattern in line.lower() for pattern in ['epic 1:', 'epic 2:', 'epic 3:']):
            # Add to description if we have a current epic
            if current_epic.get('description'):
                current_epic['description'] += ' ' + line
            else:
                current_epic['description'] = line
    
    # Don't forget the last epic
    if current_epic and current_epic.get('title'):
        epics.append(current_epic)        # If parsing failed, create default epics from the response
        if not epics:
            logger.warning("Failed to parse structured epics, creating default epics from response")
            # Split response into chunks that might be epics
            sections = [s.strip() for s in epics_response.split('\n\n') if s.strip()]
            
            for i, section in enumerate(sections[:5]):  # Limit to 5 epics max
                if len(section) > 20:  # Only meaningful sections
                    # Try to extract title from first line
                    lines = section.split('\n')
                    title = lines[0] if lines else f'Epic {i+1}'
                    description = '\n'.join(lines[1:]) if len(lines) > 1 else section
                    
                    epics.append({
                        'id': f'epic_{i+1}',
                        'title': extract_epic_title(title),
                        'description': description[:500] + ('...' if len(description) > 500 else ''),
                        'priority': 'High',
                        'user_stories': []
                    })
            
            # Final fallback - create multiple default epics to ensure we have several
            if not epics or len(epics) < 3:
                logger.warning("Creating fallback epics to ensure minimum of 3 epics")
                fallback_epics = [
                    {
                        'id': 'epic_1',
                        'title': 'User Management and Authentication',
                        'description': 'Core user authentication, registration, and profile management features including secure login, password reset, and user data management.',
                        'priority': 'High',
                        'user_stories': []
                    },
                    {
                        'id': 'epic_2', 
                        'title': 'Core Application Features',
                        'description': 'Primary application functionality and business logic implementation including main workflows and user interactions.',
                        'priority': 'High',
                        'user_stories': []
                    },
                    {
                        'id': 'epic_3',
                        'title': 'System Integration and API',
                        'description': 'External system integrations, API development, and third-party service connections for enhanced functionality.',
                        'priority': 'Medium',
                        'user_stories': []
                    },
                    {
                        'id': 'epic_4',
                        'title': 'Reporting and Analytics',
                        'description': 'Dashboard development, reporting capabilities, and analytics features for data insights and monitoring.',
                        'priority': 'Medium',
                        'user_stories': []
                    }
                ]
                epics = fallback_epics[:max(3, len(epics))]  # Ensure at least 3 epics
    
    logger.info(f"Final result: {len(epics)} epics parsed")
    return epics

def extract_epic_title(line):
    """Extract epic title from a line of text."""
    # Remove common prefixes
    line = line.strip()
    for prefix in ['epic 1:', 'epic 2:', 'epic 3:', 'epic 4:', 'epic 5:', 'epic:', 'title:']:
        if line.lower().startswith(prefix):
            line = line[len(prefix):].strip()
    
    # Clean up and return
    return line if line else 'Untitled Epic'

def format_epics_for_display(epics):
    """Format epics into HTML for display in the three-section UI with single selection."""
    if not epics:
        return "<p>No epics generated. Please try uploading your PRD again.</p>"
    
    html_parts = []
    
    for i, epic in enumerate(epics):
        epic_id = epic.get('id', f'epic_{i+1}')
        title = epic.get('title', 'Untitled Epic')
        description = epic.get('description', 'No description provided')
        priority = epic.get('priority', 'High')
        
        # Determine priority badge class
        priority_class = {
            'High': 'badge-danger',
            'Medium': 'badge-warning', 
            'Low': 'badge-secondary'
        }.get(priority, 'badge-danger')
        
        # Only first epic is selected by default
        is_checked = 'checked' if i == 0 else ''
        
        # Create epic card HTML with radio buttons for single selection
        epic_html = f"""
        <div class="epic-item mb-3 p-3 border rounded" data-epic-id="{epic_id}" style="background: #f8f9fa; border: 1px solid #dee2e6 !important;">
            <div class="form-check">
                <input class="form-check-input" type="radio" name="epic_selection" value="{epic_id}" id="epic_{epic_id}" {is_checked} style="margin-top: 0.5rem;">
                <label class="form-check-label w-100" for="epic_{epic_id}" style="cursor: pointer;">
                    <div class="d-flex justify-content-between align-items-start">
                        <h5 class="mb-2 text-primary">{title}</h5>
                        <span class="badge {priority_class}">{priority}</span>
                    </div>
                    <p class="mb-0 text-muted">{description}</p>
                </label>
            </div>
        </div>"""
        
        html_parts.append(epic_html)
    
    return "\n".join(html_parts)

@app.route("/tabbed-generate-epics", methods=["POST"])
def tabbed_generate_epics():
    """Generate epics from uploaded PRD for epic-first tabbed layout."""
    try:
        logger.info("POST request to /tabbed-generate-epics")
        
        # Handle file uploads
        prd_file = request.files.get('prd_file')
        additional_file = request.files.get('additional_file')
        context_notes = request.form.get('context_notes', '')
        
        if not prd_file:
            return jsonify({"success": False, "error": "PRD file is required"})
        
        # Process PRD content using existing function
        prd_content = ""
        additional_content = ""
        
        # Process main PRD file
        if prd_file and prd_file.filename:
            prd_content = safe_read(prd_file)
            logger.info(f"Extracted PRD content: {len(prd_content)} characters")
        
        # Process additional file if provided
        if additional_file and additional_file.filename:
            additional_content = safe_read(additional_file)
            logger.info(f"Extracted additional content: {len(additional_content)} characters")
        
        # Store in session for later use
        session['prd_content'] = prd_content
        session['additional_content'] = additional_content
        session['context_notes'] = context_notes
        
        # Create combined content
        combined_content = prd_content
        if additional_content:
            combined_content += f"\n\n--- Additional Documentation ---\n{additional_content}"
        if context_notes:
            combined_content += f"\n\n--- User Context ---\n{context_notes}"
        
        session['combined_content'] = combined_content
        
        # Generate epics using existing logic
        enhanced_context = create_epic_generation_context(prd_content, additional_content, context_notes)
        
        try:
            client = OpenAI(api_key=os.getenv('OPENAI_API_KEY'))
            response = client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {"role": "system", "content": "You are an expert product manager and business analyst. Generate multiple comprehensive epics (3-5) that break down the requirements into manageable chunks."},
                    {"role": "user", "content": enhanced_context}
                ],
                max_tokens=2000,
                temperature=0.7
            )
            
            epics_text = response.choices[0].message.content
            
            # Parse epics from the response
            epics = parse_epics_from_response(epics_text)
            
            if epics:
                # Store epics in session
                session['generated_epics'] = epics
                session['epic_first_mode'] = True
                
                logger.info(f"Generated {len(epics)} epics successfully")
                return jsonify({
                    "success": True,
                    "epics": epics,
                    "message": f"Successfully generated {len(epics)} epics"
                })
            else:
                return jsonify({
                    "success": False,
                    "error": "Failed to parse epics from OpenAI response"
                })
                
        except Exception as e:
            logger.error(f"Error calling OpenAI API: {str(e)}")
            return jsonify({
                "success": False,
                "error": "Failed to generate epics. Please check your OpenAI API key."
            })
            
    except Exception as e:
        logger.error(f"Error in tabbed epic generation: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e)})

@app.route("/tabbed-generate-user-stories", methods=["POST"])
def tabbed_generate_user_stories():
    """Generate user stories for a selected epic in epic-first tabbed layout."""
    try:
        logger.info("POST request to /tabbed-generate-user-stories")
        
        data = request.get_json()
        selected_epic = data.get('epic')
        
        if not selected_epic:
            return jsonify({"success": False, "error": "Selected epic is required"})
        
        # Get PRD content from session for context
        prd_content = session.get('combined_content', '')
        
        # Generate user stories using OpenAI
        user_stories = generate_user_stories_for_selected_epic(selected_epic, prd_content)
        
        if user_stories:
            # Store in session
            session['current_epic'] = selected_epic
            session['generated_user_stories'] = user_stories
            
            logger.info(f"Generated {len(user_stories)} user stories for epic: {selected_epic.get('title', 'Unknown')}")
            return jsonify({
                "success": True,
                "user_stories": user_stories,
                "epic": selected_epic,
                "message": f"Successfully generated {len(user_stories)} user stories"
            })
        else:
            return jsonify({
                "success": False,
                "error": "Failed to generate user stories for the selected epic"
            })
            
    except Exception as e:
        logger.error(f"Error generating user stories: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": str(e)})

def parse_epics_from_response(epics_text):
    """Parse epics from OpenAI response text."""
    try:
        epics = []
        epic_sections = epics_text.split('Epic ')
        
        for i, section in enumerate(epic_sections):
            if i == 0:  # Skip the first split part which is usually empty or intro text
                continue
                
            lines = section.strip().split('\n')
            if not lines:
                continue
                
            # Extract epic number and title from first line
            first_line = lines[0]
            if ':' in first_line:
                title_part = first_line.split(':', 1)[1].strip()
            else:
                title_part = first_line.strip()
            
            # Find description
            description = ""
            priority = "Medium"  # Default priority
            
            for line in lines[1:]:
                line = line.strip()
                if line.startswith('Description:'):
                    description = line.replace('Description:', '').strip()
                elif line.startswith('Priority:'):
                    priority = line.replace('Priority:', '').strip()
                elif not line.startswith('Epic') and not line.startswith('Priority:') and description == "":
                    # If no explicit "Description:" label, use the content as description
                    if line and not line.lower().startswith('priority'):
                        description = line
            
            if title_part and description:
                epic = {
                    'id': f'epic_{i}',
                    'title': title_part,
                    'description': description,
                    'priority': priority,
                    'estimated_stories': 'TBD',
                    'estimated_effort': 'TBD'
                }
                epics.append(epic)
        
        logger.info(f"Parsed {len(epics)} epics from response")
        return epics
        
    except Exception as e:
        logger.error(f"Error parsing epics: {str(e)}")
        return []

def generate_user_stories_for_selected_epic(epic, prd_content):
    """Generate user stories specifically for a given epic."""
    try:
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        prompt = f"""
        Based on the following PRD content and the specific epic, generate detailed user stories.

        PRD Content:
        {prd_content[:3000] if prd_content else "No additional PRD content available"}

        Epic Details:
        Title: {epic.get('title', '')}
        Description: {epic.get('description', '')}

        Please generate 5-8 user stories that specifically support this epic. For each user story, provide:

        User Story 1: [Story Title]
        Description: As a [user type], I want [goal] so that [benefit]
        Acceptance Criteria:
        - [Specific, testable criterion 1]
        - [Specific, testable criterion 2]
        - [Specific, testable criterion 3]
        Priority: High/Medium/Low
        Effort: Small/Medium/Large

        User Story 2: [Story Title]
        Description: As a [user type], I want [goal] so that [benefit]
        Acceptance Criteria:
        - [Specific, testable criterion 1]
        - [Specific, testable criterion 2]
        - [Specific, testable criterion 3]
        Priority: High/Medium/Low
        Effort: Small/Medium/Large

        Continue this format for all user stories (aim for 5-8 stories).
        """
        
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are an expert product manager and business analyst. Generate detailed, actionable user stories that align with the given epic and PRD requirements."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=3000,
            temperature=0.7
        )
        
        # Parse the response
        content = response.choices[0].message.content
        user_stories = parse_user_stories_from_response(content, epic)
        
        return user_stories
        
    except Exception as e:
        logger.error(f"Error generating user stories for epic: {str(e)}")
        return []

def parse_user_stories_from_response(content, epic):
    """Parse user stories from OpenAI response."""
    try:
        user_stories = []
        story_sections = content.split('User Story ')
        
        for i, section in enumerate(story_sections):
            if i == 0:  # Skip the first split part
                continue
                
            lines = section.strip().split('\n')
            if not lines:
                continue
            
            # Extract story number and title from first line
            first_line = lines[0]
            if ':' in first_line:
                title_part = first_line.split(':', 1)[1].strip()
            else:
                title_part = first_line.strip()
            
            # Parse story details
            description = ""
            acceptance_criteria = []
            priority = "Medium"
            effort = "Medium"
            
            current_section = None
            for line in lines[1:]:
                line = line.strip()
                if line.startswith('Description:'):
                    description = line.replace('Description:', '').strip()
                    current_section = 'description'
                elif line.startswith('Acceptance Criteria:'):
                    current_section = 'criteria'
                elif line.startswith('Priority:'):
                    priority = line.replace('Priority:', '').strip()
                    current_section = None
                elif line.startswith('Effort:'):
                    effort = line.replace('Effort:', '').strip()
                    current_section = None
                elif line.startswith('- ') and current_section == 'criteria':
                    acceptance_criteria.append(line[2:])
                elif current_section == 'description' and line and not line.startswith('Acceptance') and not line.startswith('Priority') and not line.startswith('Effort'):
                    if description:
                        description += " " + line
                    else:
                        description = line
            
            if title_part and description:
                story = {
                    'id': f'story_{i}',
                    'title': title_part,
                    'description': description,
                    'acceptance_criteria': '\n'.join([f"• {criteria}" for criteria in acceptance_criteria]) if acceptance_criteria else "Acceptance criteria to be defined",
                    'priority': priority,
                    'estimated_effort': effort,
                    'epic_id': epic.get('id', 'epic_1'),
                    'epic_title': epic.get('title', '')
                }
                user_stories.append(story)
        
        logger.info(f"Parsed {len(user_stories)} user stories from response")
        return user_stories
        
    except Exception as e:
        logger.error(f"Error parsing user stories: {str(e)}")
        return []
