from flask import Flask, jsonify, request, session
app = Flask(__name__)

# Ensure route is registered at top level
@app.route('/api/upload-prd-to-vector-db', methods=['POST'])
def upload_prd_to_vector_db():
    try:
        session_id = session.get('data_key', None)
        if not session_id:
            return jsonify({'success': False, 'error': 'No session key found'}), 400

        prd_data = get_data(session_id)
        if not prd_data:
            return jsonify({'success': False, 'error': 'No PRD data found'}), 404

        from docx import Document
        from io import BytesIO
        doc = Document()
        doc.add_heading("Product Requirements Document", level=1)
        doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()
        doc.add_heading("Product Overview", level=2)
        doc.add_paragraph(prd_data.get("product_overview", ""))
        doc.add_heading("Feature Overview", level=2)
        doc.add_paragraph(prd_data.get("feature_overview", ""))
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        import requests
        files = {'prd_file': ('PRD_Draft.docx', buffer, 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')}
        logger.info("Uploading PRD to Vector DB at /api/upload-to-vector-db...")
        try:
            response = requests.post('http://localhost:5001/api/upload-to-vector-db', files=files)
            logger.info(f"Upload response status: {response.status_code}")
            logger.info(f"Upload response text: {response.text}")
            result = response.json()
            if result.get('success'):
                prd_link = result.get('url')
                logger.info(f"PRD uploaded to Vector DB: {prd_link}")
                return jsonify({'success': True, 'link': prd_link})
            else:
                logger.error(f"Vector DB upload failed: {result.get('error')}")
                return jsonify({'success': False, 'error': result.get('error')}), 500
        except Exception as e:
            logger.error(f"Vector DB upload failed: {str(e)}")
            return jsonify({'success': False, 'error': str(e)}), 500
    except Exception as e:
        logger.error(f"PRD upload failed: {str(e)}")
        return jsonify({'success': False, 'error': str(e)}), 500


# IMPORTS AND SETUP
########################

# Standard library imports
import os
import re
import json
import uuid
import time
import asyncio
import logging
import tempfile
import logging.config
import random  # Add at top with other imports
from io import BytesIO
from datetime import timedelta
from concurrent.futures import ThreadPoolExecutor

# Security and authentication imports
from functools import wraps
from flask_limiter import Limiter
from flask_limiter.util import get_remote_address

# Third-party framework and API imports
import openai
import redis
from flask import (
    render_template, request, redirect, 
    url_for, session, send_file, jsonify
)
from dotenv import load_dotenv
from marshmallow import Schema, fields, ValidationError

# Document processing imports
from bs4 import BeautifulSoup
from markdown2 import markdown
from docx.shared import Pt, RGBColor  # Add RGBColor hereE
from docx import Document
import docx  # <-- Add this import for docx.Document usage
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from vector_db_api import vector_db_api

########################
# REDIS CONFIGURATION
########################

# Redis import with fallback
try:
    redis_test = redis.Redis(
        host=os.getenv('REDIS_HOST', 'localhost'),
        port=int(os.getenv('REDIS_PORT', 6379)),
        db=int(os.getenv('REDIS_DB', 0)),
        password=os.getenv('REDIS_PASSWORD'),
        decode_responses=True
    )
    redis_test.ping()
    redis_client = redis_test
    USING_REDIS = True
except redis.ConnectionError as e:
    logging.warning(f"Redis connection failed: {e}. Using file storage.")
    TEMP_DIR = tempfile.gettempdir()
    USING_REDIS = False

########################
load_dotenv()
openai.api_key = os.getenv("OPENAI_API_KEY")

app.secret_key = os.getenv('FLASK_SECRET_KEY')

# Configure secure session settings
app.config.update(
    SESSION_COOKIE_SECURE=True,
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',
    PERMANENT_SESSION_LIFETIME=timedelta(hours=1)
)
# Development override: disable Secure flag when not running under HTTPS to ensure cookies are sent.
if os.getenv('FLASK_ENV', 'development') != 'production':
    app.config['SESSION_COOKIE_SECURE'] = False
    # Optional: make explicit in logs
    logging.info('SESSION_COOKIE_SECURE disabled for non-production environment')

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

THREAD_POOL_SIZE = 12  # Increased for more parallelism (tune as needed for your hardware)
executor = ThreadPoolExecutor(max_workers=THREAD_POOL_SIZE)

# Hard ceilings to control input size to agents (rough char -> token approximation)
MAX_AGENT_INPUT_CHARS = int(os.getenv('MAX_AGENT_INPUT_CHARS', '12000'))
BUDGETS = {
    'user_inputs': 2000,     # industry/sector/geography/intent/features
    'product': 3500,         # product_overview
    'feature': 3500,         # feature_overview
    'legacy': 2000,          # legacy_business_description
    'guidance': 500          # small guidance appended per agent
}

def _clip(text: str, max_chars: int) -> str:
    try:
        t = text or ""
        if len(t) <= max_chars:
            return t
        return t[:max_chars]
    except Exception:
        return (text or "")[:max_chars]

def _compose_user_inputs_block(user_inputs: dict) -> str:
    block = f"""
    # Original User Inputs
    Industry: {user_inputs.get('industry', '')}
    Sector: {user_inputs.get('sector', '')}
    Geography: {user_inputs.get('geography', '')}
    Intent: {user_inputs.get('intent', '')}
    Features: {user_inputs.get('features', '')}
    """
    return _clip(block, BUDGETS['user_inputs'])

def build_capped_combined_input(user_inputs, product_overview, feature_overview, legacy_desc) -> str:
    """Compose the combined input with per-section budgets and a total cap."""
    ui = _compose_user_inputs_block(user_inputs)
    po = _clip(product_overview or "", BUDGETS['product'])
    fo = _clip(feature_overview or "", BUDGETS['feature'])
    ld = _clip(legacy_desc or "", BUDGETS['legacy'])

    body = f"""
    {ui}

    # Product Overview (Agent 1.1 Output)
    {po}

    # Feature Overview (Agent 2 Analysis)
    {fo}

    # Legacy Business Description (from Legacy Code)
    {ld}

    # Incorporation Guidelines
    - Use insights from the Legacy Business Description to inform and refine Functional, Non-Functional, Data, and Legal/Compliance requirements when relevant.
    - When a requirement or statement is directly derived from legacy artifacts, add a trailing citation line: "Source: Legacy Code" or "Reference: Legacy Code" (one line per item where applicable).
    - Do not over-cite; include a citation only when the point is supported by the legacy artifacts.
    """
    # Enforce global cap as a final guard
    return _clip(body, MAX_AGENT_INPUT_CHARS)

# Utility for parallel agent calls
def run_agents_in_parallel(agent_tasks):
    import time
    from concurrent.futures import as_completed
    results = {}
    start = time.time()
    with ThreadPoolExecutor(max_workers=THREAD_POOL_SIZE) as executor:
        future_to_key = {executor.submit(func, *args): key for key, (func, args) in agent_tasks.items()}
        for future in as_completed(future_to_key):
            key = future_to_key[future]
            try:
                results[key] = future.result()
            except Exception as exc:
                results[key] = f"Error: {exc}"
    logging.info(f"[PERF] Parallel agent batch ({list(agent_tasks.keys())}) took {time.time() - start:.2f}s")
    return results
########################
# OPENAI ASSISTANT IDs
########################

# Map of assistant IDs to their roles
ASSISTANTS = {
    #'agent_1': 'asst_EvIwemZYiG4cCmYc7GnTZoQZ',
    'agent_1_1': 'asst_htCAXHgeveZkjJj84Ldpnv6L', # Agent 1.1 - Product Overview Synthesizer – System Instructions
    'agent_2'  : 'asst_t5hnaKy1wPvD48jTbn8Mx45z',   # Agent 2: Feature Overview Generator – System Instructions
    'agent_4_1': 'asst_Ed8s7np19IPmjG5aOpMAYcPM', # Agent 4.1: Product Requirements / User Stories Generator - System Instructions
    'agent_4_2': 'asst_CLBdcKGduMvSBM06MC1OJ7bF', # Agent 4.2: Operational Business Requirements Generator – System Instructions
    'agent_4_3': 'asst_61ITzgJTPqkQf4OFnnMnndNb', # Agent 4.3: Capability-Scoped Non-Functional Requirements Generator – System Instructions
    'agent_4_4': 'asst_pPFGsMMqWg04OSHNmyQ5oaAy', # Agent 4.4: Data Attribute Requirement Generator – System Instructions
    'agent_4_5': 'asst_wwgc1Zbl5iknlDtcFLOuTIjd',  # Agent 4.5: LRC: Legal, Regulatory, and Compliance Synthesizer – System Instructions
    'agent_4_6': 'asst_JOtY81FnKEkrhgcJmuJSDyip'
}

# Dedicated Legacy Business Description agent (optional). Set LEGACY_AGENT_ID env var.
LEGACY_AGENT_ID = os.getenv('LEGACY_AGENT_ID')

########################
# LOGGING CONFIGURATION
########################

# Define logging configuration
LOGGING_CONFIG = {
    'version': 1,
    'disable_existing_loggers': False,
    'formatters': {
        'standard': {
            'format': '%(asctime)s [%(levelname)s] %(name)s: %(message)s'
        },
    },
    'handlers': {
        'default': {
            'level': 'INFO',
            'formatter': 'standard',
            'class': 'logging.StreamHandler',
        },
        'file': {
            'level': 'INFO',
            'formatter': 'standard',
            'class': 'logging.FileHandler',
            'filename': 'app.log',
            'mode': 'a',
        },
    },
    'loggers': {
        '': {
            'handlers': ['default', 'file'],
            'level': 'INFO',
            'propagate': True
        }
    }
}
# Apply logging configuration
logging.config.dictConfig(LOGGING_CONFIG)
logger = logging.getLogger(__name__)

########################
# DATA STORAGE FUNCTIONS
########################

def store_data(data):
    # Generate unique session ID
    session_id = str(uuid.uuid4())
    if USING_REDIS:
        redis_client.setex(session_id, 3600, json.dumps(data))
    else:
        with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
            json.dump(data, f)
    return session_id

# -------------------------------
# Legacy description extraction
# -------------------------------
def _extract_text_from_file_storage(fs):
    """Read uploaded file (COBOL/JCL/TXT/DOCX) into text."""
    name = (fs.filename or '').lower()
    try:
        if name.endswith('.docx'):
            d = docx.Document(fs)
            return "\n".join(p.text for p in d.paragraphs)
        else:
            # Try utf-8, fallback latin-1
            data = fs.read()
            try:
                return data.decode('utf-8', errors='ignore')
            except Exception:
                return data.decode('latin-1', errors='ignore')
    except Exception as e:
        logging.warning(f"Failed to read file {fs.filename}: {e}")
        return ""

def extract_legacy_business_description(file_texts):
    """
    Heuristic extractor: pull comments and descriptive lines from COBOL/JCL/TXT.
    file_texts: list of tuples (filename, text)
    Returns a concise markdown description.
    """
    highlights = []
    for fname, text in file_texts:
        if not text:
            continue
        lines = text.splitlines()
        picked = []
        for ln in lines[:2000]:  # scan first 2k lines per file
            s = ln.strip()
            if not s:
                continue
            # Capture typical comment lines and key markers
            if s.startswith(('*', '*>', '//*', '/*', '*/', '//')):
                picked.append(s.lstrip('*/ '))
                continue
            upper = s.upper()
            if any(k in upper for k in (
                'IDENTIFICATION DIVISION', 'PROGRAM-ID', 'AUTHOR', 'REMARK', 'REMARKS', 'PURPOSE', 'DESC', 'DESCRIPTION',
                'JOB', 'EXEC PGM=', 'PROC', 'STEP', 'FUNCTION', 'BUSINESS'
            )):
                picked.append(s)
        picked = [p for p in picked if len(p) > 4]
        if picked:
            fn_disp = os.path.basename(fname)
            highlights.append(f"### {fn_disp}\n" + "\n".join(f"- {p}" for p in picked[:25]))
    if not highlights:
        return ""
    header = "## Legacy Business Description\nThe following summary was extracted from uploaded legacy artifacts (COBOL/JCL):\n"
    return header + "\n\n" + "\n\n".join(highlights)

def generate_legacy_business_requirements(file_texts):
    """Convert legacy artifacts into structured Business Requirements Breakdown (heuristic, POC 3b-style).
    - Filters license headers and boilerplate (Apache, copyright, identifiers, divisions).
    - Promotes lines like "Function : ..." to requirement names and derives brief logic steps from business verbs.
    - Emits sections: Requirement N: <Name>, Description, Logic Description (bullets), Dependencies.
    """
    if not file_texts:
        return ""

    noise_substrings = [
        'apache license', 'all rights reserved', 'copyright', 'you may not use this file',
        'http://www.apache.org/licenses', 'https://www.apache.org/licenses', 'licensed under',
        'identification division', 'environment division', 'data division', 'working-storage section',
        'linkage section', 'procedure division', 'program-id', 'author', 'installation', 'date-written',
    ]
    skip_prefixes = ('*>', '//*', '/*', '//', '*', '*/')
    business_verbs = (
        'update', 'post', 'calculate', 'determine', 'validate', 'apply', 'compute', 'assess',
        'flag', 'set', 'assign', 'add', 'remove', 'charge', 'waive', 'approve', 'reject'
    )

    def is_noise(s: str) -> bool:
        u = s.lower()
        if any(tok in u for tok in noise_substrings):
            return True
        if re.match(r"^\s*(program\s*:|application\s*:|type\s*:|author\s*:)", u):
            return True
        # Generic technical starters we want to skip
        if re.match(r"^\s*(display\b|move\b|perform\b|if\b|end-if\b|add\b|subtract\b|select\b|assign\b|open\b|close\b|read\b|write\b|fetch\b|execute\b|evaluate\b|into\b|using\b|call\b)\b", u):
            # Codey statements rarely translate directly; we mine business verbs separately
            return True
        # Paragraph labels like 1050-UPDATE-ACCOUNT
        if re.match(r"^\s*\d{2,5}-[A-Z0-9\-]+\.?$", s):
            return True
        if len(s.strip()) < 4:
            return True
        return False

    # Collect candidates per file
    files_info = []
    for fname, text in file_texts:
        if not text:
            continue
        function_names = []
        logic_lines = []
        deps_tokens = set()
        for raw in text.splitlines()[:2000]:
            s = (raw or '').strip().strip('*').strip()
            if not s or s.startswith(skip_prefixes):
                # Still allow parsing of function lines under comments
                cleaned = s.lstrip('*/ -').strip()
            else:
                cleaned = s

            if not cleaned:
                continue

            low = cleaned.lower()
            if 'function' in low and re.search(r"\bfunction\b\s*:\s*", low):
                # Extract function name after 'Function :'
                name = re.split(r"(?i)function\s*:\s*", cleaned, maxsplit=1)[-1]
                name = name.strip().strip('.').strip()
                if name and name.lower() not in ('n/a', 'tbd'):
                    function_names.append(name)
                continue

            if is_noise(cleaned):
                continue

            # Business verb based logic extraction
            if any(v in low for v in business_verbs):
                # Normalize and end with period
                line = cleaned
                # Skip overtly technical content in logic
                if re.search(r"\b(FILE|RECORD|SELECT|ASSIGN|OPEN|CLOSE|DISPLAY|MOVE|PERFORM|EVALUATE|FETCH|EXEC|SQL)\b", line):
                    continue
                line = re.sub(r"\s{2,}", " ", line)
                if not line.endswith(('.','!','?')):
                    line += '.'
                # Token replacements to business nouns
                repl = [
                    (r"\bTRAN\b", "transaction"),
                    (r"\btrans\b", "transaction"),
                    (r"\bACCT\b", "account"),
                    (r"\bBAL\b", "balance"),
                    (r"\bAMT\b", "amount"),
                ]
                for pat, rep in repl:
                    line = re.sub(pat, rep, line, flags=re.IGNORECASE)
                logic_lines.append(line)

            # Dependencies: uppercase tokens with hyphens or underscores (e.g., ACCT-BAL, TRAN-AMT)
            for tok in re.findall(r"\b[A-Z][A-Z0-9_-]{3,}\b", cleaned):
                # Avoid generic words
                if tok not in ("IDENTIFICATION", "DIVISION", "PROGRAM", "AUTHOR", "SECTION", "DISPLAY", "MOVE"):
                    deps_tokens.add(tok)

        files_info.append({
            'file': os.path.basename(fname or 'artifact'),
            'functions': list(dict.fromkeys(function_names))[:20],  # preserve order, higher cap
            'logic': list(dict.fromkeys(logic_lines))[:15],         # dedupe, higher cap
            'deps': sorted(list(deps_tokens))[:20]
        })

    # Helpers for output prettification
    def prettify_token(tok: str) -> str:
        # Convert ACCT-CREDIT-LIMIT -> account credit limit
        parts = re.split(r"[-_]+", tok)
        words = []
        for p in parts:
            pl = p.lower()
            mapping = {
                'acct': 'account', 'tran': 'transaction', 'trans': 'transaction', 'amt': 'amount', 'bal': 'balance',
                'desc': 'description', 'num': 'number', 'id': 'ID', 'ts': 'timestamp'
            }
            words.append(mapping.get(pl, pl))
        # Capitalize ID if alone
        return ' '.join(words).replace(' id', ' ID')

    tech_deps_blacklist = (
        'file', 'record', 'section', 'division', 'program', 'stmt', 'open', 'close', 'display', 'move', 'perform',
        'select', 'assign', 'exec', 'sql', 'cics', 'db2', 'vsam', 'comp', 'pic'
    )

    def clean_deps(deps: set) -> list:
        cleaned = []
        for t in deps:
            tl = t.lower()
            if any(x in tl for x in tech_deps_blacklist):
                continue
            cleaned.append(prettify_token(t))
        # De-dupe while preserving order
        seen = set()
        out = []
        for c in cleaned:
            if c not in seen:
                out.append(c)
                seen.add(c)
        return out[:10]

    def mk_description(name: str) -> str:
        nl = (name or '').lower()
        if 'interest' in nl:
            return 'Calculates interest based on configured rates, balances, and transaction timing.'
        if 'post' in nl or 'posting' in nl:
            return 'Posts daily transactions and updates affected account balances.'
        if 'account update' in nl or 'update account' in nl:
            return 'Accepts and applies account updates with validation.'
        if 'late fee' in nl:
            return 'Determines eligibility and applies late fees per policy.'
        return f"Provides {name.strip().lower()}." if name else 'Provides business functionality.'

    def mk_business_description(name: str, logic: list[str]) -> str:
        base = mk_description(name)
        if logic:
            # Include up to two key activities as part of the sentence
            steps = []
            for s in logic[:2]:
                c = re.sub(r"\s*[-•]\s*", "", s).strip()
                if c.endswith('.'):
                    c = c[:-1]
                if c:
                    steps.append(c)
            if steps:
                return f"{base} Key activities include: " + "; ".join(steps) + "."
        return base

    # Build Business Description (top-level) + Requirements
    req_idx = 1
    out = ["## Business Description"]
    # Aggregate capability names from functions
    all_funcs = []
    for i in files_info:
        all_funcs.extend(i['functions'])
    if all_funcs:
        out.append("Capabilities identified:")
        # Dedup preserve order
        seen = set()
        caps = []
        for f in all_funcs:
            if f not in seen:
                caps.append(f)
                seen.add(f)
        for f in caps:
            out.append(f"- {f}")
    # List files processed
    if files_info:
        out.append("")
        out.append("Files processed:")
        for i in files_info:
            out.append(f"- {i['file']}")

    out.append("")
    out.append("## Business Requirements Breakdown")
    for info in files_info:
        funcs = info['functions'] or []
        logic = info['logic'] or []
        deps = clean_deps(set(info['deps'] or []))

        # Per-file Business Description
        file_funcs = funcs[:3]
        if file_funcs:
            out.append(f"\n#### {info['file']} — Business Description")
            # Simple 1-2 sentence summary based on functions
            if len(file_funcs) == 1:
                out.append(f"This component supports: {file_funcs[0]}.")
            else:
                out.append("This component supports: " + ", ".join(file_funcs[:-1]) + f" and {file_funcs[-1]}.")

        # If we have explicit functions, create a requirement per function; otherwise, one per file
        targets = funcs if funcs else [f"Business process from {info['file']}"]
        for name in targets:
            title = name[:1].upper() + name[1:]
            out.append(f"\n### Requirement {req_idx}: {title}")
            # Business Description (business-facing summary)
            out.append(f"Business Description: {mk_business_description(title, logic)}")
            # Description prefers the function text, lightly normalized
            out.append(f"Description: {mk_description(title)}")

            if logic:
                out.append("\nLogic Description:")
                for step in logic[:4]:
                    # Make step more businessy
                    cleaned = step
                    cleaned = re.sub(r"\bTRAN\b", "transaction", cleaned, flags=re.IGNORECASE)
                    cleaned = re.sub(r"\bACCT\b", "account", cleaned, flags=re.IGNORECASE)
                    # Remove trailing technical fragments like INTO/TO variable
                    cleaned = re.sub(r"\b(INTO|TO)\b.+", "", cleaned).strip()
                    if cleaned:
                        out.append(f"- {cleaned}")

            if deps:
                out.append("\nDependencies:")
                out.append("- Data elements: " + ", ".join(deps))

            req_idx += 1
            # Cap within inner loop (high ceiling to include most relevant requirements)
            if req_idx > 200:
                break

    return "\n".join(out) if len(out) > 1 else ""

def _summarize_with_agent(file_texts):
    """Use an agent to craft a concise Legacy Business Description from legacy code artifacts."""
    try:
        # If no LEGACY_AGENT_ID is configured, skip agent call and let caller fallback to heuristic
        if not LEGACY_AGENT_ID or not isinstance(LEGACY_AGENT_ID, str) or not LEGACY_AGENT_ID.strip():
            logging.info("[LEGACY] LEGACY_AGENT_ID is not set; skipping agent summarization and using heuristic fallback")
            return ""
        parts = []
        for fname, txt in file_texts[:5]:
            snippet = (txt or '')[:6000]
            parts.append(f"File: {os.path.basename(fname)}\n---\n{snippet}")
        corpus = "\n\n".join(parts)
        prompt = (
            "You are a Legacy Code Business Analyst. Read the following COBOL/JCL/text artifacts and produce a concise, "
            "business-facing 'Legacy Business Description' in markdown. Focus on purpose, key business rules, inputs, outputs, "
            "and major processing steps, avoiding code syntax. Use short paragraphs and bullet points; keep it under 300 words.\n\n"
            f"Artifacts:\n{corpus}"
        )
        resp = call_agent(LEGACY_AGENT_ID, prompt)
        return resp or ""
    except Exception as e:
        logging.warning(f"Agent summarization failed: {e}")
        return ""

@app.route('/api/legacy-business-description', methods=['POST'])
def api_legacy_business_description():
    if not session.get('logged_in'):
        return jsonify({'error': 'Not authenticated'}), 401
    try:
        files = request.files.getlist('legacy_files[]') or []
        if not files:
            return jsonify({'error': 'No files provided'}), 400
        texts = []
        for fs in files:
            fs.stream.seek(0)
            txt = _extract_text_from_file_storage(fs)
            texts.append((fs.filename, txt))
        # First try the agent summarizer
        desc = _summarize_with_agent(texts)
        if not desc:
            # Fallback to heuristic requirements generator; then description if still empty
            desc = generate_legacy_business_requirements(texts) or extract_legacy_business_description(texts)
        # Ensure explicit source citation is present
        if desc and "Legacy Code" not in desc:
            desc = desc + "\n\nSource: Legacy Code"
        return jsonify({'description': desc})
    except Exception as e:
        logging.error(f"Legacy extraction failed: {e}")
        return jsonify({'error': str(e)}), 500

def get_data(session_id):
    if USING_REDIS:
        # Try Redis first
        data = redis_client.get(session_id)
        return json.loads(data) if data else None
    else:
        # Check file storage
        path = os.path.join(TEMP_DIR, f'prd_session_{session_id}.json')
        if os.path.exists(path):
            with open(path, 'r') as f:
                return json.load(f)
    return None

########################
# MONITORING & PERFORMANCE
########################
def monitor_agent_performance(func):
    @wraps(func)
    def wrapper(agent_id, input_text, *args, **kwargs):
        start_time = time.time()
        try:
            result = func(agent_id, input_text, *args, **kwargs)
            duration = time.time() - start_time
            logging.info(f"Agent {agent_id} completed in {duration:.2f}s")
            return result
        except Exception as e:
            duration = time.time() - start_time
            logging.error(f"Agent {agent_id} failed after {duration:.2f}s: {str(e)}")
            raise
    return wrapper

async def wait_for_run_completion(thread_id, run_id, timeout=90, poll_interval=0.5):
    start_time = time.time()
    while time.time() - start_time < timeout:
        status = openai.beta.threads.runs.retrieve(run_id=run_id, thread_id=thread_id)
        if status.status == "completed":
            return status
        elif status.status == "failed":
            raise RuntimeError(f"Run {run_id} failed.")
        await asyncio.sleep(poll_interval)
    raise TimeoutError(f"Run {run_id} did not complete in {timeout} seconds.")


@monitor_agent_performance
def call_agent(agent_id, input_text):
    try:
        # Validate assistant id
        if not agent_id or not isinstance(agent_id, str) or not agent_id.strip():
            logging.error("[ERROR] assistant_id is missing or invalid; aborting call_agent")
            return "Error: assistant_id not configured"
        logging.info(f"[CALL START] Calling agent {agent_id}")
        thread_key = f"thread_{agent_id}"
        if thread_key not in session:
            thread = openai.beta.threads.create()
            session[thread_key] = thread.id
        else:
            thread = openai.beta.threads.retrieve(session[thread_key])
        
        if not thread or not thread.id:
            raise ValueError("Failed to create thread")
            
        message = openai.beta.threads.messages.create(
            thread_id=thread.id, 
            role="user", 
            content=input_text
        )
        if not message:
            raise ValueError("Failed to create message")
            
        run = openai.beta.threads.runs.create(
            thread_id=thread.id, 
            assistant_id=agent_id
        )
        if not run:
            raise ValueError("Failed to create run")

        start_time = time.time()
        while time.time() - start_time < 60:  # 60 second timeout
            status = openai.beta.threads.runs.retrieve(
                thread_id=thread.id, 
                run_id=run.id
            )
            if status.status == "completed":
                break
            time.sleep(0.25)
            
        messages = openai.beta.threads.messages.list(thread_id=thread.id)
        if not messages.data:
            raise ValueError("No messages returned")
            
        response = messages.data[0].content[0].text.value
        
        # Log the agent's response
        logging.info(f"""
        [AGENT RESPONSE] Agent {agent_id}:
        ----------------------------------------
        {response}
        ----------------------------------------
        """)
            
        return response

    except Exception as e:
        logging.error(f"[ERROR] Agent {agent_id} failed: {e}")
        return f"Error: {str(e)}"

@monitor_agent_performance
async def call_agent_async(agent_id, input_text):
    try:
        # Validate assistant id
        if not agent_id or not isinstance(agent_id, str) or not agent_id.strip():
            logging.error("[ERROR] assistant_id is missing or invalid; aborting call_agent_async")
            return "Error: assistant_id not configured"
        logging.info(f"[CALL START] Calling agent {agent_id}")
        thread_key = f"thread_{agent_id}"
        if thread_key not in session:
            thread = openai.beta.threads.create()
            session[thread_key] = thread.id
        else:
            thread = openai.beta.threads.retrieve(session[thread_key])

        message = openai.beta.threads.messages.create(
            thread_id=thread.id,
            role="user",
            content=input_text
        )

        run = openai.beta.threads.runs.create(
            thread_id=thread.id,
            assistant_id=agent_id
        )

        # Use async polling
        await wait_for_run_completion(thread.id, run.id)

        messages = openai.beta.threads.messages.list(thread_id=thread.id)
        response = messages.data[0].content[0].text.value

        logging.info(f"[AGENT RESPONSE] Agent {agent_id}:\n{response}")
        return response

    except Exception as e:
        logging.error(f"[ERROR] Agent {agent_id} failed: {e}")
        return f"Error: {str(e)}"


# Update call_agents_parallel to use retry logic
async def call_agents_parallel(agent_calls):
    tasks = [
        call_agent_with_cache(agent_id, input_text)
        for agent_id, input_text in agent_calls
    ]
    return await asyncio.gather(*tasks)

agent_semaphore = asyncio.Semaphore(3)  # Limit concurrent calls

async def call_agent_with_limit(agent_id, input_text):
    async with agent_semaphore:
        return await call_agent_with_cache(agent_id, input_text)
def verify_credentials(username, password):
    """Verify user credentials against environment variables."""
    return (username == os.getenv('ADMIN_USERNAME') and 
            password == os.getenv('ADMIN_PASSWORD'))

@app.route('/', methods=['GET', 'POST'])
def login():
    if request.method == 'POST':
        if verify_credentials(request.form.get('username'), request.form.get('password')):
            session['logged_in'] = True
            return redirect(url_for('page1'))
        return render_template('page0_login.html', error=True)
    return render_template('page0_login.html', error=False)


@app.route('/page1', methods=['GET', 'POST'])
def page1():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    if request.method == 'POST':             # Get form inputs
        inputs = {key: request.form[key] for key in ['industry', 'sector', 'geography', 'intent', 'features']}
        MAX_FILE_SIZE_KB = 100
        # Handle file upload
        if 'context_file' in request.files:
                file = request.files['context_file']
                if file and file.filename:
                    file.seek(0, os.SEEK_END)
                    file_length = file.tell()
                    file.seek(0)  # Reset stream pointer

                    if file_length > MAX_FILE_SIZE_KB * 1024:
                        logging.warning("Uploaded file too large")
                        return "Uploaded file exceeds the 100KB limit.", 400
                    try:
                        filename = file.filename.lower()
                        if filename.endswith('.txt'):
                            # Read as UTF-8 text
                            file_content = file.read().decode('utf-8')
                        elif filename.endswith('.docx'):
                            # Read as docx (do not decode)
                            doc = docx.Document(file)
                            file_content = "\n".join([para.text for para in doc.paragraphs])
                        else:
                            file_content = "Unsupported file format."

                        # Add file content to inputs
                        inputs['context_file'] = file_content
                    except Exception as e:
                        logging.error(f"File upload error: {str(e)}")
                        return "Error processing file upload", 400

        # Legacy business description via hidden field or derive from uploaded files
        legacy_desc = request.form.get('legacy_business_description', '').strip()
        if not legacy_desc:
            legacy_files = request.files.getlist('legacy_files[]') or []
            if legacy_files:
                texts = []
                for fs in legacy_files:
                    fs.stream.seek(0)
                    txt = _extract_text_from_file_storage(fs)
                    if txt:
                        texts.append((fs.filename, txt))
                if texts:
                    try:
                        # Prefer agent summarization, fallback to heuristic Business Requirements then description
                        legacy_desc = _summarize_with_agent(texts) or generate_legacy_business_requirements(texts) or extract_legacy_business_description(texts)
                    except Exception as e:
                        logging.warning(f"Legacy extraction failed: {e}")
                        # no-op fallback; legacy_desc remains empty
        if legacy_desc:
            # Ensure explicit source citation is present
            if "Legacy Code" not in legacy_desc:
                legacy_desc = legacy_desc + "\n\nSource: Legacy Code"
            inputs['legacy_business_description'] = legacy_desc

        session.update(inputs)

       # Update context to include file content
        context_parts = [f"{k.replace('_', ' ').title()}: {v}" for k, v in inputs.items() 
                        if k != 'context_file']
        
        if 'context_file' in inputs:
            context_parts.append(f"\nAdditional Context:\n{inputs['context_file']}")
            
        context = "\n".join(context_parts)
        session_id = store_data({
            "inputs": inputs,
            "legacy_business_description": inputs.get('legacy_business_description', ''),
            "status": "processing"
        })
        session['data_key'] = session_id

        # Ensure old data is cleared
        if USING_REDIS:
            redis_client.delete(session_id)
        else:
            path = os.path.join(TEMP_DIR, f'prd_session_{session_id}.json')
            if os.path.exists(path):
                os.remove(path)


        # Persistent event loop for agent orchestration
        def run_agents():
            import threading, gc
            t0 = time.time()
            try:
                loop = None
                try:
                    loop = asyncio.get_event_loop()
                except RuntimeError:
                    loop = asyncio.new_event_loop()
                    asyncio.set_event_loop(loop)
                # Run background agent calls (Page 1): Product Overview and Feature Overview only
                a11, a2 = loop.run_until_complete(call_agents_parallel([
                    (ASSISTANTS['agent_1_1'], context),
                    (ASSISTANTS['agent_2'], context)
                ]))
                logging.info(f"[PERF] Page1 agent batch: {time.time() - t0:.2f}s")
                logging.info(f"Agent 1.1 Output: {a11}")
                logging.info(f"Agent 2 Output: {a2}")

                final_data = {
                    "inputs": inputs,
                    "product_overview": a11,
                    "feature_overview": a2,
                    "legacy_business_description": inputs.get('legacy_business_description', ''),
                    "status": "complete"
                }
                if USING_REDIS:
                    redis_client.setex(session_id, 3600, json.dumps(final_data))
                else:
                    with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                        json.dump(final_data, f)
                logging.info(f"[AGENTS DONE] Page 1 background agents complete for session {session_id}")
            except Exception as e:
                logging.error(f"Error during agent execution: {e}")
                fail_data = {"status": "error", "message": str(e)}
                if USING_REDIS:
                    redis_client.setex(session_id, 3600, json.dumps(fail_data))
                else:
                    with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                        json.dump(fail_data, f)
                        json.dump(fail_data, f)
            # Resource usage logging and cleanup
            try:
                logging.info(f"[RESOURCE] Thread count: {threading.active_count()}")
                gc.collect()
            except Exception as e:
                logging.warning(f"Resource logging failed: {e}")

        executor.submit(run_agents)
        return redirect('/page2')

    return render_template('page1_input.html')


@app.route('/page2', methods=['GET', 'POST'])
def page2():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    session_id = session.get('data_key', '')
    data = get_data(session_id) or {}

    # Wait for background job to complete
    start_time = time.time()
    timeout = 120  # seconds

    while time.time() - start_time < timeout:
        current_data = get_data(session_id) or {}
        if current_data.get("status") == "complete":
            data = current_data
            break
        elif current_data.get("status") == "error":
            return f"<h3>❌ Agent processing failed:</h3><pre>{current_data.get('message')}</pre>", 500
        time.sleep(1)
    else:
        return "⏳ Processing took too long. Please refresh the page shortly.", 504

    if request.method == 'POST':
        # Debug: log the incoming form data
        logging.info(f"[PAGE2][POST] request.form: {dict(request.form)}")
        # Update product_overview and feature_overview from form only if not empty
        po = request.form.get("product_overview", "").strip()
        fo = request.form.get("feature_overview", "").strip()
        if po:
            data["product_overview"] = po
        if fo:
            data["feature_overview"] = fo

        # Ensure user inputs are preserved and merged
        if "inputs" not in data or not data["inputs"]:
            # Try to get from session as fallback
            data["inputs"] = session.get("inputs", {})

        # Log the updated data for debugging
        logging.info(f"[PAGE2][POST] Updated data: {json.dumps(data, indent=2)}")

        # Save updated data prior to agent calls
        if USING_REDIS:
            redis_client.setex(session_id, 3600, json.dumps(data))
        else:
            with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                json.dump(data, f)

        # Build combined input from product, feature, and legacy descriptions
        user_inputs = data.get("inputs", {})
        feature_overview = data.get("feature_overview", "")
        product_overview = data.get("product_overview", "")
        legacy_desc_full = data.get("legacy_business_description", "") or user_inputs.get("legacy_business_description", "")
        combined_input = build_capped_combined_input(user_inputs, product_overview, feature_overview, legacy_desc_full)

        # Call agent_4_x in parallel and persist outputs, then go straight to Page 4
        import threading, gc
        loop = None
        try:
            loop = asyncio.get_event_loop()
        except RuntimeError:
            loop = asyncio.new_event_loop()
            asyncio.set_event_loop(loop)

        keys = [k for k in ASSISTANTS if k.startswith("agent_4_")]
        agent_inputs = []
        for k in keys:
            guidance = _clip(
                "When applicable, incorporate insights from 'Legacy Business Description' into this section. "
                "Add 'Source: Legacy Code' on any requirement derived from legacy artifacts.",
                BUDGETS['guidance']
            )
            agent_inputs.append((ASSISTANTS[k], combined_input + "\n\n[Agent Guidance]\n" + guidance))

        results = loop.run_until_complete(call_agents_parallel(agent_inputs))
        outputs = dict(zip(keys, results))
        data['combined_outputs'] = outputs
        session['combined_outputs'] = outputs

        # Persist combined outputs
        if USING_REDIS:
            redis_client.setex(session_id, 3600, json.dumps(data))
        else:
            with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                json.dump(data, f)

        try:
            logging.info(f"[RESOURCE] Thread count: {threading.active_count()}")
            gc.collect()
        except Exception as e:
            logging.warning(f"Resource logging failed: {e}")

        return redirect('/page4')

    return render_template("page2_agents.html",
        agent11_output=data.get('product_overview', ''),
        agent2_output=data.get('feature_overview', ''),
        legacy_business_description=data.get('legacy_business_description', '') or data.get('inputs', {}).get('legacy_business_description', '')
    )

@app.route('/update_content', methods=['POST'])
def update_content():
    if not session.get('logged_in'):
        return jsonify({'error': 'Not authenticated'}), 401
        
    try:
        data = request.get_json()
        content_type = data.get('type')
        new_content = data.get('content')
        
        if not content_type or not new_content:
            return jsonify({'error': 'Missing required fields'}), 400
            
        # Get stored data
        stored_data = get_data(session.get('data_key', '')) or {}
        
        # Update content based on type
        if content_type == 'product':
            # Re-run Agent 1.1 with new content
            existing_content = stored_data.get('product_overview', '')
            full_prompt = f"""You are editing the Product Overview section. Keep existing relevant content unless changes are requested.

            Current content:
            {existing_content}

            User instruction:
            {new_content}
            """
            new_response = call_agent(ASSISTANTS['agent_1_1'], full_prompt)
            stored_data['product_overview'] = new_response
        elif content_type == 'feature':
            # Re-run Agent 2 with new content
            existing_content = stored_data.get('feature_overview', '')
            full_prompt = f"""You are editing the Feature Overview section. Maintain useful content unless instructed otherwise.

            Current content:
            {existing_content}

            User instruction:
            {new_content}
            """
            new_response = call_agent(ASSISTANTS['agent_2'], full_prompt)
            stored_data['feature_overview'] = new_response
        elif content_type == 'highest_order':
            return jsonify({'error': 'Highest-order requirements editing is disabled in this flow.'}), 400
        elif content_type == 'legacy':
            # Update legacy business description directly, enforce citation
            legacy_new = new_content or ''
            if legacy_new and 'Legacy Code' not in legacy_new:
                legacy_new = legacy_new.rstrip() + "\n\nSource: Legacy Code"
            legacy_new = dedupe_legacy_citations(legacy_new)
            stored_data['legacy_business_description'] = legacy_new
            new_response = legacy_new
        elif content_type.startswith('agent_4_'):
                agent_id = ASSISTANTS.get(content_type)
                if not agent_id:
                    return jsonify({'error': f'Unknown agent for {content_type}'}), 400

                # Gather full context for Page 4 agents: user inputs, product, feature, legacy
                user_inputs = stored_data.get('inputs', {})
                feature_overview = stored_data.get('feature_overview', '')
                product_overview = stored_data.get('product_overview', '')
                legacy_desc_full = stored_data.get('legacy_business_description', '') or user_inputs.get('legacy_business_description', '')
                combined_input = build_capped_combined_input(user_inputs, product_overview, feature_overview, legacy_desc_full)

                # Existing content for the targeted section
                existing_content = stored_data.get('combined_outputs', {}).get(content_type, '')

                # Guidance to enforce citation when legacy-derived statements are included
                guidance = _clip(
                    "When applicable, incorporate insights from 'Legacy Business Description' into this section. "
                    "Add 'Source: Legacy Code' on lines that are directly supported by legacy artifacts.",
                    BUDGETS['guidance']
                )

                full_prompt = (
                    "You are updating a section of a Product Requirements Document. Keep all useful information intact, "
                    "and only modify based on the user request. Be conservative with deletions.\n\n"
                    "[Current Section Content]\n" + existing_content + "\n\n"
                    "[User Instruction]\n" + new_content + "\n\n"
                    "[Combined Context: User Inputs, Product, Feature, Legacy]\n" + combined_input + "\n\n"
                    "[Agent Guidance]\n" + guidance
                )

                new_response = call_agent(agent_id, full_prompt)
                stored_data.setdefault('combined_outputs', {})[content_type] = new_response   

        # Store updated data
        if USING_REDIS:
            redis_client.setex(session['data_key'], 3600, json.dumps(stored_data))
        else:
            with open(os.path.join(TEMP_DIR, f'prd_session_{session["data_key"]}.json'), 'w') as f:
                json.dump(stored_data, f)
                
        return jsonify({'success': True, 'response': new_response})
        
    except Exception as e:
        logging.error(f"Error updating content: {str(e)}")
        return jsonify({'error': str(e)}), 500
    
_agent_response_cache = {}

async def call_agent_with_cache(agent_id, input_text):
    cache_key = f"{agent_id}_{hash(input_text)}"
    if cache_key in _agent_response_cache:
        logging.info(f"[CACHE HIT] Returning cached response for {agent_id}")
        return _agent_response_cache[cache_key]
    
    result = await call_agent_with_retry(agent_id, input_text)
    _agent_response_cache[cache_key] = result
    return result

def require_auth(f):
    @wraps(f)
    def decorated(*args, **kwargs):
        if not session.get('logged_in'):
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated

@app.route('/page4')
@require_auth
def page4():
    if not session.get('logged_in'):
        return redirect(url_for('login'))

    try:
        data = get_data(session.get('data_key', ''))
        if not data:
            logger.error("No data found for session key")
            return "No data found", 404  # Or render a friendly template

        outputs = {
            'product_overview': dedupe_legacy_citations(data.get('product_overview', '')),
            'feature_overview': dedupe_legacy_citations(data.get('feature_overview', '')),
            'agent_4_1': dedupe_legacy_citations(data.get('combined_outputs', {}).get('agent_4_1', '')),
            'agent_4_2': dedupe_legacy_citations(data.get('combined_outputs', {}).get('agent_4_2', '')),
            'agent_4_3': dedupe_legacy_citations(data.get('combined_outputs', {}).get('agent_4_3', '')),
            'agent_4_4': dedupe_legacy_citations(data.get('combined_outputs', {}).get('agent_4_4', '')),
            'agent_4_5': dedupe_legacy_citations(data.get('combined_outputs', {}).get('agent_4_5', ''))
        }
        logger.info(json.dumps(outputs, indent=2))

        logger.info("[PAGE4] Rendering with outputs:")
        for key, value in outputs.items():
            logger.info(f"{key}: {value[:100]}...")

        return render_template('page4_final_output.html', outputs=outputs)

    except Exception as e:
        logger.error(f"Error in page4: {str(e)}")
        return str(e), 500




def add_hyperlink(paragraph, url, text=None):
    if text is None:
        text = url
    part = paragraph.part
    r_id = part.relate_to(
        url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')
    rPr.append(color)
    underline = OxmlElement('w:u')
    underline.set(qn('w:val'), 'single')
    rPr.append(underline)
    new_run.append(rPr)

    new_run_text = OxmlElement('w:t')
    new_run_text.text = text
    new_run.append(new_run_text)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return paragraph

def validate_reference(ref: str) -> bool:
    """Validate extracted reference."""
    if not ref or len(ref.strip()) < 5:
        return False
        
    # Check for common invalid patterns
    invalid_patterns = [
        r'^[\W_]+$',                    # Only special characters
        r'^(?:http)?[:/]+$',            # Incomplete URLs
        r'^[Rr]eference:\s*$',             # Empty reference labels
        r'^\d+$'                        # Just numbers
    ]
    
    return not any(re.match(pattern, ref.strip()) for pattern in invalid_patterns)

def extract_references_from_outputs(outputs):
    """Extract references from combined outputs."""
    all_refs = set()
    ref_patterns = [
        r'\[([^\]]+)\]\(([^\)]+)\)',                     # Markdown links
        r'Reference:\s+(.+?)(?=\n|$)',                   # Reference lines
        r'Source:\s+(.+?)(?=\n|$)',                      # Source lines
        r'(?:https?://\S+)',                             # Raw URLs
        r'(?<=Reference:).*?[\d]{4}.*?(?=\n|$)',        # Citations with years
        r'\((?:[^()]*?\d{4}[^()]*?)\)',                 # Parenthetical citations
        r'(?<=DOI:)\s*[\d.]+\/[\w.-]+',                 # DOI references
        r'(?:ISBN(?:-1[03])?:?\s*(?:[\d-]+))'          # ISBN references
    ]
    
    def extract_from_text(text):
        refs = set()
        for pattern in ref_patterns:
            matches = re.finditer(pattern, text, re.MULTILINE | re.IGNORECASE)
            for match in matches:
                try:
                    # Handle markdown links differently
                    if '[' in pattern:
                        # Check if we have enough groups before accessing
                        if len(match.groups()) >= 2:
                            refs.add(match.group(2))  # Add URL from markdown link
                        else:
                            refs.add(match.group(0))  # Add entire match if not enough groups
                    else:
                        # For other patterns, add the full match
                        refs.add(match.group().strip())
                except (IndexError, AttributeError) as e:
                    logging.warning(f"Error extracting reference: {e}")
                    continue
        return refs

    # Extract from all outputs
    for output in outputs.values():
        if isinstance(output, str):
            try:
                refs = extract_from_text(output)
                all_refs.update(refs)
            except Exception as e:
                logging.error(f"Error processing output text: {e}")
                continue
    
    # Clean and sort references
    cleaned_refs = sorted(list({
        ref.strip(' .,[]()\"\'')  # Clean up punctuation
        for ref in all_refs
        if ref and len(ref.strip()) > 5  # Minimum length to filter noise
    }))
    
    logging.info(f"[REFERENCES] Extracted {len(cleaned_refs)} references")
    return cleaned_refs

def dedupe_legacy_citations(text: str) -> str:
    """Remove duplicate trailing 'Source: Legacy Code' lines and collapse consecutive duplicates.
    - If multiple consecutive lines end with the exact citation, keep one.
    - If block ends with duplicate citation lines, keep a single trailing citation.
    """
    if not isinstance(text, str) or not text:
        return text
    lines = text.splitlines()
    deduped = []
    citation = "Source: Legacy Code"
    prev_was_citation = False
    for ln in lines:
        if ln.strip() == citation:
            if not prev_was_citation:
                deduped.append(ln)
                prev_was_citation = True
            else:
                # skip duplicate citation line
                continue
        else:
            deduped.append(ln)
            prev_was_citation = False
    # Also ensure we don't end with multiple blank lines before a citation
    result = "\n".join(deduped)
    result = re.sub(r"(\n\s*){3,}" , "\n\n", result)
    return result

limiter = Limiter(app=app, key_func=get_remote_address)

# Replace with the new route
@app.route('/generate_word_doc', methods=['POST'])
@limiter.limit("5 per minute")
def generate_word_doc():
    """Generate Word document from PRD content."""
    if not session.get('logged_in'):
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        # Get data and validate
        data = get_data(session.get("data_key", ""))
        if not data:
            logging.error("No data found in session")
            return jsonify({'error': 'No session data found'}), 404

        # Initialize document with styles
        doc = Document()
        initialize_document_styles(doc)

        # Prepare legacy section with explicit source note
        legacy_section = data.get("legacy_business_description", "")
        if legacy_section and "Legacy Code" not in legacy_section:
            legacy_section = legacy_section + "\n\nSource: Legacy Code"
        legacy_section = dedupe_legacy_citations(legacy_section)

        # Map all sections
        sections = {
            "Legacy Business Description (from Legacy Code)": legacy_section,
            "Product Overview": dedupe_legacy_citations(data.get("product_overview", "")),
            "Feature Overview": dedupe_legacy_citations(data.get("feature_overview", "")),
         #   "Product Requirements": data.get("combined_outputs", {}).get("agent_4_1", ""),
            "Functional Requirements": dedupe_legacy_citations(data.get("combined_outputs", {}).get("agent_4_2", "")),
            "Non-Functional Requirements": dedupe_legacy_citations(data.get("combined_outputs", {}).get("agent_4_3", "")),
            "Data Requirements": dedupe_legacy_citations(data.get("combined_outputs", {}).get("agent_4_4", "")),
            "Legal & Compliance Requirements": dedupe_legacy_citations(data.get("combined_outputs", {}).get("agent_4_5", ""))
        }

        # Add title and metadata
        doc.add_heading("Product Requirements Document", level=1)
        doc.add_paragraph(f"Generated: {time.strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph()

        # Process each section
        for title, content in sections.items():
            if content and content.strip():
                try:
                    process_section(doc, title, content)
                except Exception as e:
                    logging.error(f"Error processing section {title}: {str(e)}")
                    doc.add_paragraph(f"Error in section {title}: {str(e)}", style="Error")

        # Add references
        add_references_section(doc, data.get("combined_outputs", {}))

        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        logging.info("Document generated successfully")
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"PRD_Draft_{time.strftime('%Y%m%d')}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )

    except Exception as e:
        logging.error(f"Document generation failed: {str(e)}")
        return jsonify({
            'error': 'Document generation failed',
            'message': str(e)
        }), 500
    

def add_references_section(doc, outputs):
    """Add references section to document."""
    try:
        refs = extract_references_from_outputs(outputs)
        if not refs:
            return

        doc.add_page_break()
        doc.add_heading("References", level=2)
        
        for ref in refs:
            p = doc.add_paragraph(style="List Bullet")
            if ref.startswith(("http://", "https://")):
                add_hyperlink(p, ref)
            else:
                p.add_run(ref)
                
    except Exception as e:
        logging.error(f"Error adding references: {str(e)}")

def process_section(doc, title, raw_markdown):
    """Process a single section of the document."""
    try:
        doc.add_heading(title, level=2)
        
        # Convert markdown to HTML
        html_content = markdown(
            raw_markdown,
            extras=[
                "fenced-code-blocks",
                "tables",
                "header-ids",
                "break-on-newline"
            ]
        )
        
        # Parse HTML
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Process elements
        for el in soup.find_all(["p", "li", "h3", "h4", "pre", "code"]):
            process_element(doc, el)
            
    except Exception as e:
        logging.error(f"Error processing section {title}: {str(e)}")
        # Add error note in document
        doc.add_paragraph(f"Error processing section: {str(e)}", style="Error")

def process_element(doc, el):
    """Process a single HTML element."""
    text = el.get_text(strip=True)
    if not text:
        return

    para = doc.add_paragraph()
    
    # Handle requirement labels
    if match := re.match(r"^(F\d+|NFR\d+|DR\d+|LR\d+)\. (.+)", text):
        label, desc = match.groups()
        run = para.add_run(f"{label}. ")
        run.bold = True
        para.add_run(desc)
    # Handle labeled lines
    elif match := re.match(r"^([A-Z][\w\s]+?): (.+)", text):
        label, desc = match.groups()
        run = para.add_run(f"{label}: ")
        run.bold = True
        para.add_run(desc)
    # Regular text
    else:
        para.add_run(text)
        if el.name == "li":
            para.style = "List Bullet"
    
    para.paragraph_format.space_after = Pt(6)

@app.errorhandler(Exception)
def handle_exception(e):
    """Handle all unhandled exceptions."""
    logging.error(f"Unhandled exception: {str(e)}")
    return jsonify({
        'error': 'An unexpected error occurred',
        'message': str(e)
    }), 500

@app.errorhandler(404)
def not_found_error(e):
    """Handle 404 errors."""
    return jsonify({'error': 'Resource not found'}), 404

@app.errorhandler(401)
def unauthorized_error(e):
    """Handle 401 errors."""
    return redirect(url_for('login'))

def process_section(doc, title, raw_markdown):
    """Process a single section of the document."""
    try:
        doc.add_heading(title, level=2)
        
        # Convert markdown to HTML
        html_content = markdown(
            raw_markdown,
            extras=[
                "fenced-code-blocks",
                "tables",
                "header-ids",
                "break-on-newline"
            ]
        )
        
        # Parse HTML
        soup = BeautifulSoup(html_content, "html.parser")
        
        # Process elements
        for el in soup.find_all(["p", "li", "h3", "h4", "pre", "code"]):
            process_element(doc, el)
            
    except Exception as e:
        logging.error(f"Error processing section {title}: {str(e)}")
        # Add error note in document
        doc.add_paragraph(f"Error processing section: {str(e)}", style="Error")

def initialize_document_styles(doc):
    """Initialize document styles."""
    styles = doc.styles
    
    # Error style
    error_style = styles.add_style('Error', WD_STYLE_TYPE.PARAGRAPH)
    error_style.font.color.rgb = RGBColor(255, 0, 0)  # Now RGBColor will be defined
    error_style.font.italic = True
    
    # Requirement style
    req_style = styles.add_style('Requirement', WD_STYLE_TYPE.PARAGRAPH)
    req_style.font.name = 'Calibri'
    req_style.font.size = Pt(11)
    
    return doc

def cleanup_expired_sessions():
    """Clean up expired session data with better error handling."""
    if USING_REDIS:
        return  # Redis handles expiration
    # TODO: Implement file cleanup logic here
        
    try:
        current_time = time.time()
        cleaned = 0
        failed = 0
        
        for filename in os.listdir(TEMP_DIR):
            if not filename.startswith('prd_session_'):
                continue
                
            filepath = os.path.join(TEMP_DIR, filename)
            try:
                if current_time - os.path.getmtime(filepath) > 3600:
                    os.remove(filepath)
                    cleaned += 1
            except OSError as e:
                failed += 1
                logging.error(f"Failed to remove {filepath}: {e}")
                
        logging.info(f"Session cleanup: {cleaned} removed, {failed} failed")
        
    except Exception as e:
        logging.error(f"Session cleanup failed: {e}")

class UserInputSchema(Schema):
    industry = fields.Str(required=True)
    sector = fields.Str(required=True)
    geography = fields.Str(required=True)
    intent = fields.Str(required=True)
    features = fields.Str(required=True)

def validate_form_input(form_data):
    """Validate form input against schema with improved sanitization."""
    try:
        # Sanitize input before validation
        sanitized_data = {
            key: value.strip() for key, value in form_data.items()
        }
        
        schema = UserInputSchema()
        validated_data = schema.load(sanitized_data)
        
        # Enhanced validation rules
        for key, value in validated_data.items():
            if len(value) < 3:
                return False, {key: ["Input too short (minimum 3 characters)"]}
            if len(value) > 1000:
                return False, {key: ["Input too long (maximum 1000 characters)"]}
            if any(char in value for char in '<>{}[]'):
                return False, {key: ["Invalid characters detected"]}
                
        return True, validated_data
        
    except ValidationError as e:
        logging.warning(f"Validation error: {e.messages}")
        return False, e.messages
    except Exception as e:
        logging.error(f"Unexpected validation error: {str(e)}")
        return False, {"error": "Internal validation error"}
    
def extend_session():
    """Extend session lifetime if user is active."""
    if session.get('logged_in'):
        session.permanent = True
        app.permanent_session_lifetime = timedelta(hours=1)
        session.modified = True

@app.before_request
def before_request():
    """
    Performs pre-request operations such as extending the user session and occasionally cleaning up expired sessions.

    This function should be called before handling each request. It ensures the session is kept active and, with a small probability (1% per call), triggers cleanup of expired sessions to manage resources efficiently.
    """
    extend_session()
    # Add session cleanup check every 100 requests
    if random.random() < 0.01:  # 1% chance
        cleanup_expired_sessions()

########################
# OPENAI AGENT INTERACTION
########################
@monitor_agent_performance
async def call_agent_with_retry(agent_id, input_text, max_retries=3):
    # Validate assistant id before attempting retries
    if not agent_id or not isinstance(agent_id, str) or not agent_id.strip():
        logging.error("[ERROR] assistant_id is missing or invalid; aborting call_agent_with_retry")
        return "Error: assistant_id not configured"
    for attempt in range(max_retries):
        try:
            # Create thread with initial message
            thread = openai.beta.threads.create(messages=[{"role": "user", "content": input_text}])
            if not thread or not thread.id:
                raise ValueError("Thread creation failed")

            # Start agent run
            run = openai.beta.threads.runs.create(thread_id=thread.id, assistant_id=agent_id)

            # Use the async event loop to wait non-blockingly
            status = await wait_for_run_completion(thread.id, run.id)

            # Get response messages
            messages = openai.beta.threads.messages.list(thread_id=thread.id)
            response = messages.data[0].content[0].text.value if messages.data else ""
            return response

        except Exception as e:
            print(f"[Retry {attempt+1}/{max_retries}] Error with agent {agent_id}: {e}")
            if attempt == max_retries - 1:
                return f"Error: {str(e)}"
            await asyncio.sleep(1.5)
            
def store_session_data(data):
    """Store large session data in Redis/files and keep only key in session."""
    try:
        if not data:
            raise ValueError("Cannot store empty data")
            
        session_id = str(uuid.uuid4())
        storage_key = f"prd_data_{session_id}"
        
        # Convert data to JSON
        try:
            json_data = json.dumps(data)
        except (TypeError, ValueError) as e:
            logging.error(f"JSON serialization failed: {e}")
            raise ValueError("Invalid data format")
        
        # Store main data with error handling
        if USING_REDIS:
            try:
                redis_client.setex(storage_key, 3600, json_data)
            except redis.RedisError as e:
                logging.error(f"Redis storage failed: {e}")
                # Fallback to file storage
                with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                    f.write(json_data)
        else:
            with open(os.path.join(TEMP_DIR, f'prd_session_{session_id}.json'), 'w') as f:
                f.write(json_data)
                
        # Store minimal data in session
        session['data_key'] = session_id
        logging.info(f"Session data stored with ID: {session_id}")
        return session_id
        
    except Exception as e:
        logging.error(f"Failed to store session data: {e}")
        raise

def get_session_data():
    """Retrieve session data from storage with validation."""
    try:
        session_id = session.get('data_key')
        if not session_id:
            logging.warning("No session key found")
            return None
            
        storage_key = f"prd_data_{session_id}"
        
        # Try Redis first
        if USING_REDIS:
            try:
                data = redis_client.get(storage_key)
                if data:
                    return json.loads(data)
            except (redis.RedisError, json.JSONDecodeError) as e:
                logging.error(f"Redis retrieval failed: {e}")
        
        # Fallback to file storage
        path = os.path.join(TEMP_DIR, f'prd_session_{session_id}.json')
        if os.path.exists(path):
            try:
                with open(path, 'r') as f:
                    return json.load(f)
            except (IOError, json.JSONDecodeError) as e:
                logging.error(f"File retrieval failed: {e}")
                
        return None
        
    except Exception as e:
        logging.error(f"Session data retrieval failed: {e}")
        return None
        
    return get_data(session_id)

@app.route('/chat_with_agent', methods=['POST'])
def chat_with_agent():
    if not session.get('logged_in'):
        return jsonify({'error': 'Not authenticated'}), 401

    try:
        data = request.get_json()
        message = data.get('message', '')
        if not message:
            return jsonify({'error': 'Message is required'}), 400

        # 👇 THIS is where it talks to Agent 1.1
        reply = call_agent(ASSISTANTS['agent_1_1'], message)

        return jsonify({'reply': reply})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

    
     

@app.route("/review_prd", methods=["POST"])
def review_prd():
    data = request.get_json()
    prd_sections = data.get("sections", {})

    full_prd = "\n\n".join(
        f"{key.replace('_', ' ').title()}:\n{value}"
        for key, value in prd_sections.items()
        if value.strip()
    )

  # Construct review prompt for Agent 4.6
    review_prompt = (
        "As the PRD Review and Quality Assurance Agent, analyze this PRD content. "
        "For each section, evaluate completeness, clarity, and identify any gaps. "
        "Provide specific, actionable improvements. "
        "Format response as JSON array with fields: section, issue, suggestion.\n\n"
        f"{full_prd}"
    )

    # Call Agent 4.6 (PRD Review Agent)
    response = call_agent(ASSISTANTS['agent_4_6'], review_prompt)
   # response = call_openai_agent("agent_4_6", review_prompt)  # Assumes this helper exists
    try:
        parsed = json.loads(response)
    except json.JSONDecodeError:
        parsed = [{
            "section": "ReviewAgent", 
            "issue": "Invalid JSON response format", 
            "suggestion": "Please try the review again"
        }]

    return jsonify({"issues": parsed})
    
# Add to main
if __name__ == '__main__':
    cleanup_expired_sessions()
    app.register_blueprint(vector_db_api)
    host = os.environ.get("HOST", "0.0.0.0")
    port = int(os.environ.get("PORT", 5001))
    app.run(host=host, port=port, debug=True)

