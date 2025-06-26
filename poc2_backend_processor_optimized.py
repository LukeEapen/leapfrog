from flask import Flask, render_template, request, jsonify, session
import openai
from openai import OpenAI
import os
import time
import logging
import tiktoken
import asyncio
import aiohttp
import threading
import string
import traceback
from datetime import datetime
from concurrent.futures import ThreadPoolExecutor, as_completed
from functools import lru_cache
import hashlib
import json
from cachetools import TTLCache
import tenacity
import csv
import io

# Vector Database and RAG imports
import chromadb
from chromadb.config import Settings
from sentence_transformers import SentenceTransformer
import numpy as np
from sklearn.metrics.pairwise import cosine_similarity

# Configure logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('app.log'),
        logging.StreamHandler()
    ]
)
logger = logging.getLogger(__name__)

# Document processing imports
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
    logger.info("python-docx library available for DOCX file processing")
except ImportError:
    logger.warning("python-docx not available. DOCX files will be processed as text extraction fallback.")
    DOCX_AVAILABLE = False

# Initialize Vector Database and RAG components
logger.info("Initializing Vector Database and RAG components...")

# Initialize ChromaDB client
chroma_client = chromadb.PersistentClient(path="./vector_db")

# Initialize embedding model for RAG
try:
    embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
    logger.info("Embedding model loaded successfully")
except Exception as e:
    logger.warning(f"Failed to load embedding model: {e}. RAG features may be limited.")
    embedding_model = None

# Create or get collection for PRD documents
try:
    prd_collection = chroma_client.get_or_create_collection(
        name="prd_documents",
        metadata={"description": "PRD and documentation storage for RAG"}
    )
    logger.info("ChromaDB collection initialized successfully")
except Exception as e:
    logger.error(f"Failed to initialize ChromaDB collection: {e}")
    prd_collection = None

app = Flask(__name__)

# Configure Flask app
app.secret_key = os.getenv('FLASK_SECRET_KEY', 'dev-secret-key-change-in-production')
# In production, make sure to set FLASK_SECRET_KEY environment variable

# Performance optimizations
THREAD_POOL_SIZE = 4
MAX_CACHE_SIZE = 1000
CACHE_TTL = 3600  # 1 hour

# Initialize thread pool for parallel processing
executor = ThreadPoolExecutor(max_workers=THREAD_POOL_SIZE)

# Cache for responses (in-memory with TTL)
response_cache = TTLCache(maxsize=MAX_CACHE_SIZE, ttl=CACHE_TTL)

logger.info("Flask application starting up with performance optimizations...")
logger.info(f"OpenAI API key configured: {'Yes' if os.getenv('OPENAI_API_KEY') else 'No'}")
logger.info(f"Thread pool size: {THREAD_POOL_SIZE}")
logger.info(f"Cache size: {MAX_CACHE_SIZE} items, TTL: {CACHE_TTL}s")

@app.route("/prd_parser", methods=["GET"])
def user_story_input():
    logger.info("GET request received for user story input page")
    try:
        logger.info("Rendering poc2_user_story_input.html template")
        return render_template("poc2_user_story_input.html")
    except Exception as e:
        logger.error(f"Error rendering user story input template: {str(e)}")
        raise

@app.route("/", methods=["GET"])
def home():
    """Home page with navigation to all features."""
    logger.info("GET request received for home page")
    
    html_content = """
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>RAG-Enhanced PRD to Epic Generator</title>
        <link rel="stylesheet" href="https://stackpath.bootstrapcdn.com/bootstrap/4.5.2/css/bootstrap.min.css"/>
        <style>
            body { 
                font-family: Arial, sans-serif; 
                margin: 20px; 
                background-color: #2f323a; 
                color: #fff;
            }
            .container { 
                max-width: 1200px; 
                margin: 0 auto; 
                background-color: #2f323a; 
                padding: 30px; 
                border-radius: 10px; 
            }
            .header-bar {
                background-color: #d0021b;
                color: white;
                text-align: center;
                font-weight: bold;
                padding: 15px;
                border-radius: 6px;
                margin-bottom: 20px;
            }
            .header { 
                text-align: center; 
                margin-bottom: 30px; 
                color: #fff;
            }
            .header h1 {
                color: #fff;
                margin-bottom: 15px;
            }
            .header p {
                color: #ccc;
                font-size: 1.1rem;
            }
            .feature-grid { 
                display: grid; 
                grid-template-columns: 1fr 1fr; 
                gap: 20px; 
                margin-top: 30px; 
            }
            @media (max-width: 768px) {
                .feature-grid {
                    grid-template-columns: 1fr;
                }
            }
            .feature-card { 
                background-color: #f8f9fa; 
                color: #000;
                padding: 20px; 
                border-radius: 8px; 
                border: 1px solid #dee2e6;
                transition: transform 0.2s ease;
            }
            .feature-card:hover {
                transform: translateY(-2px);
                box-shadow: 0 4px 8px rgba(0,0,0,0.2);
            }
            .feature-card h3 { 
                color: #d0021b; 
                margin-top: 0; 
                font-weight: bold;
            }
            .feature-card p, .feature-card ul {
                color: #333;
            }
            .btn { 
                display: inline-block; 
                padding: 12px 24px; 
                background-color: #d0021b; 
                color: white; 
                text-decoration: none; 
                border-radius: 5px; 
                margin: 10px 0; 
                font-weight: bold;
                transition: background-color 0.2s ease;
            }
            .btn:hover { 
                background-color: #b0011a; 
                color: white;
                text-decoration: none;
            }
            .btn-secondary { 
                background-color: #444; 
            }
            .btn-secondary:hover { 
                background-color: #333; 
            }
            .status { 
                background-color: #444; 
                border: 1px solid #555; 
                color: #fff; 
                padding: 15px; 
                border-radius: 8px; 
                margin-bottom: 20px; 
                text-align: center;
                font-weight: bold;
            }
            .footer-text {
                margin-top: 30px; 
                text-align: center; 
                color: #ccc;
                font-style: italic;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <div class="header-bar">
                RAG-Enhanced PRD to Epic Generator
            </div>
            
            <div class="header">
                <h1>Transform PRDs into Comprehensive Epics & User Stories</h1>
                <p>Advanced RAG technology for intelligent document processing and epic generation</p>
                <div class="status">
                    Vector Database Active | RAG Processing Ready | Epic Generator Online
                </div>
            </div>
            
            <div class="feature-grid">
                <div class="feature-card">
                    <h3>Generate Epics & User Stories</h3>
                    <p>Upload your PRD and additional documentation to generate comprehensive epics and user stories using RAG-enhanced processing.</p>
                    <a href="/prd_parser" class="btn">Start Processing</a>
                </div>
                
                <div class="feature-card">
                    <h3>Vector Database Status</h3>
                    <p>View all documents stored in the vector database, including chunks, metadata, and content previews.</p>
                    <a href="/vector-db-status" class="btn btn-secondary">View Database</a>
                </div>
                
                <div class="feature-card">
                    <h3>Search Vector Database</h3>
                    <p>Search through stored documents using semantic similarity to find relevant content for any query.</p>
                    <a href="/vector-db-search" class="btn btn-secondary">Search Database</a>
                </div>
                
                <div class="feature-card">
                    <h3>RAG Features</h3>
                    <p>This system uses Retrieval Augmented Generation with:</p>
                    <ul>
                        <li>ChromaDB vector storage</li>
                        <li>SentenceTransformers embeddings</li>
                        <li>Intelligent document chunking</li>
                        <li>Semantic content retrieval</li>
                        <li>Single-agent optimization</li>
                    </ul>
                </div>
            </div>
            
            <div class="footer-text">
                <p>PRD Parser Agent replaced with RAG summaries | 50%+ faster processing | Reduced token costs</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return html_content, 200, {'Content-Type': 'text/html'}

@app.route("/user-story-upload", methods=["POST"])
def process_user_story():
    start_time = time.time()
    logger.info("POST request received for user story processing")
    
    # Collect input
    logger.info("Collecting form data and file uploads")
    context = request.form.get("context", "")
    prd_file = request.files.get("prd_file")
    additional_docs = request.files.get("additional_docs")
    
    logger.info(f"Context provided: {'Yes' if context else 'No'}")
    logger.info(f"PRD file uploaded: {'Yes' if prd_file else 'No'}")
    logger.info(f"Additional docs uploaded: {'Yes' if additional_docs else 'No'}")
    
    if prd_file:
        logger.info(f"PRD file name: {prd_file.filename}")
    if additional_docs:
        logger.info(f"Additional docs file name: {additional_docs.filename}")

    # Parallel file reading for better performance
    logger.info("Reading uploaded files in parallel")
    
    def read_files_parallel():
        with ThreadPoolExecutor(max_workers=2) as file_executor:
            prd_future = file_executor.submit(safe_read, prd_file) if prd_file else None
            docs_future = file_executor.submit(safe_read, additional_docs) if additional_docs else None
            
            prd_content = prd_future.result() if prd_future else ""
            docs_content = docs_future.result() if docs_future else ""
            
            return prd_content, docs_content
    
    prd_content, docs_content = read_files_parallel()
    
    # Debug file content
    if prd_content:
        debug_file_content(prd_content, prd_file.filename if prd_file else "prd_content")
    if docs_content:
        debug_file_content(docs_content, additional_docs.filename if additional_docs else "additional_docs")    # Store documents in vector database and create RAG summaries
    logger.info("Processing documents with RAG enhancement")
    
    # Process PRD with RAG if content is substantial and valid
    if is_valid_content(prd_content):
        if len(prd_content) > 5000:
            logger.info("Creating RAG-enhanced PRD summary")
            prd_content = create_rag_summary(prd_content, prd_file.filename if prd_file else "prd_content", max_summary_length=25000)
        elif len(prd_content) > 30000:
            logger.info("Large PRD detected - using traditional optimization")
            prd_content = optimize_prd_content(prd_content, max_length=40000)
        else:
            logger.info("PRD content is valid but small, using as-is")
    else:
        logger.warning(f"PRD content is invalid for RAG processing: {prd_content[:100] if prd_content else 'None'}")
        prd_content = "No valid PRD content available - please check file format and encoding."
    
    # Process additional docs with RAG
    if is_valid_content(docs_content):
        if len(docs_content) > 3000:
            logger.info("Creating RAG-enhanced docs summary")
            docs_content = create_rag_summary(docs_content, additional_docs.filename if additional_docs else "additional_docs", max_summary_length=10000)
        elif len(docs_content) > 15000:
            logger.info("Large additional docs detected - truncating for faster processing")
            docs_content = docs_content[:15000] + "...\n[Content truncated for performance]"
        else:
            logger.info("Additional docs content is valid but small, using as-is")
    else:
        logger.warning(f"Additional docs content is invalid: {docs_content[:100] if docs_content else 'None'}")
        docs_content = "No valid additional documentation available."
    
    logger.info(f"PRD content length (post-RAG): {len(prd_content)} characters")
    logger.info(f"Additional docs content length (post-RAG): {len(docs_content)} characters")    # Combine prompt with RAG-enhanced content
    logger.info("Combining context and RAG-enhanced content")
    
    # Create enhanced context for Epic Generator
    enhanced_context = f"""
{context}

RAG-Enhanced PRD Analysis:
{prd_content}

Additional Context:
{docs_content}

Instructions: The above content has been intelligently extracted and summarized using RAG (Retrieval Augmented Generation). 
It contains the most relevant requirements, user stories, and business objectives from the original documents.
Use this curated information to generate comprehensive epics and user stories.
"""
    logger.info(f"Enhanced context length: {len(enhanced_context)} characters")
    
    # Check cache first with more granular caching
    prompt_hash = get_cache_key(enhanced_context)
    cached_result = response_cache.get(prompt_hash)
    if cached_result:
        logger.info("Cache hit! Returning cached response")
        processing_time = time.time() - start_time
        logger.info(f"Total processing time (cached): {processing_time:.2f} seconds")
        return render_template("poc2_epic_story_screen.html", epics=cached_result)
      # Log token count for the enhanced context
    context_tokens = count_tokens(enhanced_context, "gpt-4o")
    logger.info(f"Enhanced context token count: {context_tokens:,} tokens")
    
    # Check if context is within token limits
    if context_tokens > 120000:
        logger.warning(f"Enhanced context token count ({context_tokens:,}) is approaching the 128k limit!")
        # Further optimize if needed
        enhanced_context = enhanced_context[:100000] + "\n[Content truncated due to token limits]"
        context_tokens = count_tokens(enhanced_context, "gpt-4o")
        logger.info(f"Truncated context token count: {context_tokens:,} tokens")
    else:
        logger.info(f"Enhanced context is within safe token limits ({context_tokens:,}/128,000 tokens)")    # Skip PRD Parser Agent - RAG has already done the parsing and summarization
    logger.info("****************Skipping PRD Parser - Using RAG Summary Directly")
    logger.info("Starting Epic Generator with RAG-enhanced content")
    
    # Print Enhanced Context being sent to Epic Agent
    logger.info("=" * 80)
    logger.info("ENHANCED CONTEXT INPUT TO EPIC AGENT:")
    logger.info("=" * 80)
    logger.info(enhanced_context[:3000] + "..." if len(enhanced_context) > 3000 else enhanced_context)
    logger.info("=" * 80)
    
    try:
        # Directly use Epic Generator with RAG-enhanced content
        start_epic_time = time.time()
        epic_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", enhanced_context)
        epic_processing_time = time.time() - start_epic_time
        
        logger.info("################Epic Generator response received")
        logger.info(f"Epic Generator processing time: {epic_processing_time:.2f} seconds")
        
        # Print Epic Agent Response
        logger.info("=" * 80)
        logger.info("EPIC AGENT RESPONSE OUTPUT:")
        logger.info("=" * 80)
        logger.info(epic_response)
        logger.info("=" * 80)
        
        # Log token usage for Epic Generator interaction
        log_token_usage(enhanced_context, epic_response, model="gpt-4o", context="RAG-Enhanced Epic Generator")

        final_output = epic_response
        logger.info(f"Final output length: {len(final_output)} characters")
        
        # Cache the result for future requests
        response_cache[prompt_hash] = final_output
        logger.info("Response cached for future requests")
        
        processing_time = time.time() - start_time
        logger.info(f"Total RAG-optimized processing time: {processing_time:.2f} seconds")
        
        # Calculate time savings
        estimated_traditional_time = processing_time * 2  # Estimate of traditional two-agent approach
        time_saved = estimated_traditional_time - processing_time
        logger.info(f"Estimated time saved by RAG optimization: {time_saved:.2f} seconds")        # Render page 2 with response
        logger.info("Rendering epic story screen with RAG-optimized output")
        return render_template("poc2_epic_story_screen.html", epics=final_output)
        
    except Exception as e:
        logger.error(f"Error in RAG-optimized processing: {str(e)}")
        raise

def process_large_prompt_chunked(prompt, start_time):
    """Process large prompts using RAG-enhanced chunking for better performance."""
    logger.info("Processing large prompt using RAG-enhanced chunked approach")
    
    # Check cache first
    prompt_hash = get_cache_key(prompt)
    cached_result = response_cache.get(prompt_hash)
    if cached_result:
        logger.info("Cache hit! Returning cached response")
        processing_time = time.time() - start_time
        logger.info(f"Total processing time (cached): {processing_time:.2f} seconds")
        return render_template("poc2_epic_story_screen.html", epics=cached_result)
    
    # Split prompt into sections and extract content
    sections = prompt.split('\n\n')
    context = sections[0] if sections else ""
    prd_section = ""
    docs_section = ""
    
    # Extract PRD and Docs sections
    for i, section in enumerate(sections):
        if 'PRD:' in section or 'RAG-Enhanced PRD Analysis:' in section:
            prd_section = '\n\n'.join(sections[i:i+15])  # Take next 15 sections for PRD
            break
    
    for i, section in enumerate(sections):
        if 'Additional Docs:' in section or 'Additional Context:' in section:
            docs_section = '\n\n'.join(sections[i:])  # Take remaining for docs
            break
    
    # Create RAG-optimized prompt for direct Epic Generator use
    rag_optimized_prompt = f"""
{context}

RAG-Enhanced PRD Analysis (Chunked Processing):
{prd_section[:35000]}

Additional Context:
{docs_section[:12000]}

Instructions: The above content has been intelligently processed and chunked for optimal performance. 
It contains the most relevant requirements, user stories, and business objectives from the original documents.
Generate comprehensive epics and user stories directly from this curated information.
"""
    
    logger.info(f"RAG-optimized chunked prompt length: {len(rag_optimized_prompt)} characters (reduced from {len(prompt)})")      # Log token count for the optimized prompt
    optimized_tokens = count_tokens(rag_optimized_prompt, "gpt-4o")
    logger.info(f"RAG-optimized chunked prompt token count: {optimized_tokens:,} tokens")
    
    try:
        # Skip PRD Parser - Use Epic Generator directly with RAG-enhanced content
        logger.info("****************Skipping PRD Parser in chunked processing - Using RAG-enhanced content directly")
        
        start_epic_time = time.time()
        epic_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", rag_optimized_prompt)
        epic_processing_time = time.time() - start_epic_time
        
        logger.info("Epic Generator (chunked) response received")
        logger.info(f"Epic Generator processing time: {epic_processing_time:.2f} seconds")
          # Log token usage for Epic Generator interaction
        log_token_usage(rag_optimized_prompt, epic_response, model="gpt-4o", context="RAG-Enhanced Epic Generator (Chunked)")
        
        final_output = epic_response
        
        # Cache the result
        response_cache[prompt_hash] = final_output
        logger.info("Chunked response cached for future requests")
        
        processing_time = time.time() - start_time
        logger.info(f"Total RAG-enhanced chunked processing time: {processing_time:.2f} seconds")
        
        # Calculate time savings compared to traditional two-agent approach
        estimated_traditional_time = processing_time * 2.5  # Estimate of traditional approach with chunking
        time_saved = estimated_traditional_time - processing_time
        logger.info(f"Estimated time saved by RAG chunked optimization: {time_saved:.2f} seconds")
        
        return render_template("poc2_epic_story_screen.html", epics=final_output)
        
    except Exception as e:
        logger.error(f"Error in RAG-enhanced chunked processing: {str(e)}")
        raise

def get_cache_key(prompt):
    """Generate a cache key from the prompt."""
    return hashlib.md5(prompt.encode('utf-8')).hexdigest()

def store_document_in_vector_db(content, filename, doc_type="prd"):
    """Store document content in vector database for RAG retrieval."""
    if not prd_collection or not embedding_model:
        logger.warning("Vector DB or embedding model not available. Skipping storage.")
        return None
    
    try:
        logger.info(f"Storing {doc_type} document in vector database: {filename}")
        
        # Split content into chunks for better retrieval
        chunks = split_document_into_chunks(content, chunk_size=1000, overlap=200)
        logger.info(f"Document split into {len(chunks)} chunks")
        
        # Generate embeddings for each chunk
        chunk_embeddings = embedding_model.encode(chunks)
        
        # Create metadata for each chunk
        documents = []
        metadatas = []
        ids = []
        
        doc_id = hashlib.md5(f"{filename}_{doc_type}".encode()).hexdigest()
        
        for i, (chunk, embedding) in enumerate(zip(chunks, chunk_embeddings)):
            chunk_id = f"{doc_id}_chunk_{i}"
            
            documents.append(chunk)
            metadatas.append({
                "filename": filename,
                "doc_type": doc_type,
                "chunk_index": i,
                "total_chunks": len(chunks),
                "timestamp": datetime.now().isoformat()
            })
            ids.append(chunk_id)
        
        # Store in ChromaDB
        prd_collection.add(
            documents=documents,
            metadatas=metadatas,
            ids=ids
        )
        
        logger.info(f"Successfully stored {len(chunks)} chunks in vector database")
        return doc_id
        
    except Exception as e:
        logger.error(f"Error storing document in vector DB: {str(e)}")
        return None

def split_document_into_chunks(content, chunk_size=1000, overlap=200):
    """Split document content into overlapping chunks for better retrieval."""
    if len(content) <= chunk_size:
        return [content]
    
    chunks = []
    start = 0
    
    while start < len(content):
        end = start + chunk_size
        
        # Try to break at sentence boundaries
        if end < len(content):
            # Look for sentence endings within overlap range
            for punct in ['. ', '.\n', '! ', '!\n', '? ', '?\n']:
                last_punct = content.rfind(punct, start, end)
                if last_punct > start:
                    end = last_punct + len(punct)
                    break
        
        chunk = content[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        start = end - overlap
        
        # Avoid infinite loop
        if start >= end:
            start = end
    
    return chunks

def retrieve_relevant_content(query, top_k=5, doc_type=None):
    """Retrieve relevant content from vector database using RAG."""
    if not prd_collection or not embedding_model:
        logger.warning("Vector DB or embedding model not available. Skipping RAG retrieval.")
        return []
    
    try:
        logger.info(f"Retrieving relevant content for query (top {top_k} results)")
        
        # Create query filter if doc_type is specified
        where_filter = {"doc_type": doc_type} if doc_type else None
        
        # Query the collection
        results = prd_collection.query(
            query_texts=[query],
            n_results=top_k,
            where=where_filter
        )
        
        relevant_chunks = []
        if results['documents'] and results['documents'][0]:
            for i, (doc, metadata) in enumerate(zip(results['documents'][0], results['metadatas'][0])):
                relevant_chunks.append({
                    'content': doc,
                    'metadata': metadata,
                    'relevance_rank': i + 1
                })
                
        logger.info(f"Retrieved {len(relevant_chunks)} relevant chunks")
        return relevant_chunks
        
    except Exception as e:
        logger.error(f"Error retrieving from vector DB: {str(e)}")
        return []

def create_rag_summary(content, filename, max_summary_length=5000):
    """Create a RAG-enhanced summary of the document content."""
    logger.info(f"Creating RAG-enhanced summary for {filename}")
      # Store the document in vector DB
    doc_id = store_document_in_vector_db(content, filename)
    
    if not doc_id:
        # Fallback to intelligent content extraction if vector DB fails
        logger.warning("Vector DB storage failed, using intelligent content extraction")
        return create_intelligent_summary_fallback(content, filename, max_summary_length)
    
    # Create comprehensive summary queries to extract key information
    summary_queries = [
        "requirements and functional specifications",
        "user stories and acceptance criteria", 
        "business objectives and goals",
        "system constraints and dependencies",
        "key features and capabilities",
        "technical architecture and design",
        "data flow and integration requirements",
        "security and compliance requirements",
        "performance and scalability requirements",
        "user interface and user experience requirements"
    ]
    
    # Retrieve relevant content for each query
    summary_parts = []
    for query in summary_queries:
        relevant_chunks = retrieve_relevant_content(query, top_k=2, doc_type="prd")
        
        if relevant_chunks:
            # Take the most relevant chunk for each query
            best_chunk = relevant_chunks[0]['content']
            summary_parts.append(f"[{query.title()}]\n{best_chunk}\n")
    
    # Combine summary parts with intelligent ordering
    ordered_summary = order_summary_parts(summary_parts)
    final_summary = "\n\n".join(ordered_summary)
    
    # Truncate to max length
    if len(final_summary) > max_summary_length:
        logger.info(f"Final summary length ({len(final_summary)}) exceeds max length ({max_summary_length}), truncating")
        final_summary = final_summary[:max_summary_length] + "... [Content truncated]"
    
    logger.info(f"RAG summary created, length: {len(final_summary)} characters")
    
    # Print RAG Summary for debugging
    logger.info("=" * 80)
    logger.info("RAG SUMMARY GENERATED:")
    logger.info("=" * 80)
    logger.info(final_summary[:2000] + "..." if len(final_summary) > 2000 else final_summary)
    logger.info("=" * 80)
    
    return final_summary.strip()

def order_summary_parts(summary_parts):
    """Order summary parts based on importance and relevance."""
    logger.info("Ordering summary parts by importance")
    
    # Simple heuristic: prioritize sections with more details and specific requirements
    def section_score(section):
        score = 0
        # More weight to user stories and requirements
        if "user story" in section.lower():
            score += 10
        if "requirement" in section.lower():
            score += 8
        # Add score for each bullet point or numbered item
        score += section.count('•') * 2
        score += section.count('-') * 2
        score += section.count('1.') * 2
        score += section.count('2.') * 2
        return score
    
    # Sort sections by computed score
    ordered_sections = sorted(summary_parts, key=section_score, reverse=True)
    logger.info(f"Ordered {len(ordered_sections)} summary parts by importance")
    return ordered_sections

def extract_docx_text(file):
    """Extract text content from a DOCX file."""
    try:
        if not DOCX_AVAILABLE:
            logger.warning(f"Cannot process DOCX file {file.filename}: python-docx not available")
            return "[DOCX file detected but python-docx library not available. Please install python-docx to process Word documents.]"
        
        logger.info(f"Extracting text from DOCX file: {file.filename}")
        
        # Create a temporary BytesIO object from the file
        from io import BytesIO
        file_content = file.read()
        file.seek(0)  # Reset file pointer in case it's needed again
        
        # Load the document using python-docx
        doc = DocxDocument(BytesIO(file_content))
        
        # Extract text from all paragraphs
        text_content = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text.strip())
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        final_text = "\n\n".join(text_content)
        
        if final_text.strip():
            logger.info(f"Successfully extracted {len(final_text)} characters from DOCX file {file.filename}")
            return final_text
        else:
            logger.warning(f"No text content found in DOCX file {file.filename}")
            return "[DOCX file processed but no readable text content found]"
            
    except Exception as e:
        logger.error(f"Error extracting text from DOCX file {file.filename}: {str(e)}")
        return f"[Error processing DOCX file: {str(e)}]"

def safe_read(file):
    """Safely read file content with proper error handling and encoding detection."""
    try:
        logger.debug(f"Attempting to read file: {file.filename if file else 'None'}")
        
        # Check if this is a DOCX file and handle it specially
        if file and file.filename and file.filename.lower().endswith('.docx'):
            logger.info(f"Detected DOCX file: {file.filename}")
            return extract_docx_text(file)
        
        # Read the raw bytes first
        raw_content = file.read()
        logger.debug(f"Read {len(raw_content)} bytes from file: {file.filename}")
        
        # Check if this looks like a DOCX file even if extension is wrong
        if raw_content.startswith(b'PK\x03\x04') and b'[Content_Types].xml' in raw_content:
            logger.info(f"Detected DOCX format in file: {file.filename}")
            file.seek(0)  # Reset file pointer for DOCX processing
            return extract_docx_text(file)
        
        # Try multiple encodings for regular text files
        encodings = ['utf-8', 'utf-8-sig', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                content = raw_content.decode(encoding)
                logger.info(f"Successfully decoded file {file.filename} using {encoding} encoding, length: {len(content)} characters")
                
                # Validate that we have meaningful content
                if len(content.strip()) > 0:
                    return content
                else:
                    logger.warning(f"File {file.filename} appears to be empty or whitespace only")
                    return "[File appears to be empty]"
                    
            except UnicodeDecodeError:
                continue
        
        # If all encodings fail, try to extract text from binary content
        logger.warning(f"All standard encodings failed for {file.filename}, attempting binary extraction")
        try:
            # Try to extract printable ASCII characters
            import string
            printable_chars = set(string.printable)
            extracted_text = ''.join(char for char in raw_content.decode('latin-1') if char in printable_chars)
            
            if len(extracted_text.strip()) > 50:  # If we got some meaningful text
                logger.info(f"Extracted {len(extracted_text)} printable characters from {file.filename}")
                return extracted_text
            else:
                logger.error(f"Unable to extract meaningful text from {file.filename}")
                return f"[Unable to decode file - tried encodings: {', '.join(encodings)}]"
                
        except Exception as extract_error:
            logger.error(f"Binary extraction failed for {file.filename}: {str(extract_error)}")
            return f"[File reading failed - may be binary or corrupted]"
            
    except Exception as e:
        logger.error(f"Error reading file {file.filename}: {str(e)}")
        return f"[Error reading file: {str(e)}]"

def optimize_prd_content(prd_content, max_length=40000):
    """Optimize PRD content by extracting key sections and reducing verbosity."""
    if len(prd_content) <= max_length:
        return prd_content
    
    logger.info(f"Optimizing PRD content from {len(prd_content)} to ~{max_length} characters")
    
    # Split into sections and prioritize important ones
    sections = prd_content.split('\n\n')
    important_keywords = [
        'requirement', 'feature', 'user story', 'acceptance criteria', 
        'functional', 'non-functional', 'business rule', 'constraint',
        'objective', 'goal', 'scope', 'assumption', 'dependency'
    ]
    
    # Score sections based on importance
    scored_sections = []
    for section in sections:
        score = 0
        section_lower = section.lower()
        
        # Score based on keywords
        for keyword in important_keywords:
            score += section_lower.count(keyword) * 10
        
        # Prefer sections with structured content
        score += section.count('•') * 5
        score += section.count('-') * 3
        score += section.count('1.') * 5
        
        scored_sections.append((score, section))
    
    # Sort by score and take top sections
    scored_sections.sort(key=lambda x: x[0], reverse=True)
    
    optimized_content = ""
    for score, section in scored_sections:
        if len(optimized_content) + len(section) <= max_length:
            optimized_content += section + "\n\n"
        else:
            break
    
    logger.info(f"Optimized PRD content to {len(optimized_content)} characters")
    return optimized_content.strip()

@lru_cache(maxsize=100)
def count_tokens(text, model="gpt-4o"):
    """Count tokens in text using tiktoken. Cached for performance."""
    try:
        # Map gpt-4o and gpt-4o to gpt-4 for tiktoken compatibility
        tiktoken_model = "gpt-4" if model in ["gpt-4-turbo", "gpt-4o"] else model
        encoding = tiktoken.encoding_for_model(tiktoken_model)
        return len(encoding.encode(text))
    except Exception as e:
        logger.warning(f"Error counting tokens: {e}, using fallback estimate")
        # Fallback: rough estimate of 4 characters per token
        return len(text) // 4

def log_token_usage(prompt_text, response_text, model="gpt-4o", context=""):
    """Log detailed token usage and cost estimation."""
    prompt_tokens = count_tokens(prompt_text, model)
    response_tokens = count_tokens(response_text, model)
    total_tokens = prompt_tokens + response_tokens
    
    # Cost estimation for GPT-4 Turbo (approximate rates)
    input_cost_per_1k = 0.01   # $0.01 per 1K input tokens
    output_cost_per_1k = 0.03  # $0.03 per 1K output tokens
    
    input_cost = (prompt_tokens / 1000) * input_cost_per_1k
    output_cost = (response_tokens / 1000) * output_cost_per_1k
    total_cost = input_cost + output_cost
    
    logger.info(f"Token Usage - {context}")
    logger.info(f"  Input tokens: {prompt_tokens:,}")
    logger.info(f"  Output tokens: {response_tokens:,}")
    logger.info(f"  Total tokens: {total_tokens:,}")
    logger.info(f"  Estimated cost: ${total_cost:.4f}")

def process_csv_to_vector_db(csv_file):
    """Process CSV file and store it in the vector database."""
    try:
        logger.info(f"Processing CSV file: {csv_file.filename}")
        
        # Read CSV content with proper encoding handling
        csv_content = csv_file.read().decode('utf-8-sig')  # utf-8-sig handles BOM
        csv_file.seek(0)  # Reset file pointer
        
        # Clean up the content - remove non-breaking spaces and other problematic characters
        csv_content = csv_content.replace('\xa0', ' ').replace('\ufeff', '')
        
        # Parse CSV
        csv_reader = csv.DictReader(io.StringIO(csv_content))
        rows = list(csv_reader)
        
        if not rows:
            logger.warning("CSV file is empty or has no data rows")
            return "CSV file is empty"
        
        # Clean column names
        cleaned_rows = []
        for row in rows:
            cleaned_row = {}
            for key, value in row.items():
                # Clean column names and values
                clean_key = key.strip().replace('\xa0', ' ').replace('\ufeff', '') if key else 'Unknown'
                clean_value = value.strip().replace('\xa0', ' ').replace('\ufeff', '') if value else ''
                cleaned_row[clean_key] = clean_value
            cleaned_rows.append(cleaned_row)
        
        rows = cleaned_rows
        logger.info(f"CSV contains {len(rows)} rows with columns: {list(rows[0].keys())}")
        
        # Convert CSV rows to text documents for vector storage
        documents = []
        metadatas = []
        ids = []
        
        for i, row in enumerate(rows):
            # Create a text representation of the row
            row_text = f"CSV Record {i+1}:\n"
            for key, value in row.items():
                if value:  # Only include non-empty values
                    row_text += f"{key}: {value}\n"
            
            documents.append(row_text.strip())
            
            # Ensure all metadata values are valid types for ChromaDB
            clean_filename = str(csv_file.filename).strip() if csv_file.filename else "unknown.csv"
            column_names = ", ".join(str(key).strip() for key in row.keys() if key)
            
            metadatas.append({
                "source": clean_filename,
                "type": "csv_row",
                "row_number": int(i + 1),
                "columns": column_names
            })
            ids.append(f"csv_{clean_filename.replace('.', '_')}_{i+1}")
        
        # Store in vector database
        if prd_collection and documents:
            try:
                prd_collection.add(
                    documents=documents,
                    metadatas=metadatas,
                    ids=ids
                )
                logger.info(f"Successfully stored {len(documents)} CSV rows in vector database")
                
                # Create a summary of the CSV content
                csv_summary = f"CSV File: {csv_file.filename}\n"
                csv_summary += f"Total Rows: {len(rows)}\n"
                csv_summary += f"Columns: {', '.join(rows[0].keys())}\n\n"
                
                # Add sample data (first few rows)
                csv_summary += "Sample Data:\n"
                for i, row in enumerate(rows[:5]):  # Show first 5 rows
                    csv_summary += f"Row {i+1}: {dict(row)}\n"
                
                if len(rows) > 5:
                    csv_summary += f"... and {len(rows) - 5} more rows\n"
                
                return csv_summary
                
            except Exception as e:
                logger.error(f"Error storing CSV in vector database: {e}")
                return f"Error storing CSV in database: {str(e)}"
        else:
            logger.error("Vector database collection not available")
            return "Vector database not available"
            
    except Exception as e:
        logger.error(f"Error processing CSV file: {e}")
        return f"Error processing CSV: {str(e)}"

def ask_assistant_from_file_optimized(code_filepath, user_prompt):
    """Optimized assistant interaction using direct chat completion instead of deprecated Assistants API."""
    start_time = time.time()
    logger.info(f"Starting optimized assistant interaction: {code_filepath}")
    
    try:
        # Read assistant configuration/instructions with UTF-8 encoding
        with open(f"agents/{code_filepath}", "r", encoding="utf-8") as file:
            assistant_instructions = file.read().strip()
            
        logger.info(f"Using assistant instructions from: {code_filepath}")
        
        # Initialize OpenAI client
        client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
        
        # Create system message from assistant instructions
        system_message = f"""You are an expert assistant for PRD (Product Requirements Document) processing and epic generation.

{assistant_instructions}

Please follow the instructions above and process the user's request accordingly."""        # Use chat completions API instead of deprecated Assistants API
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": system_message},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.7,
            max_tokens=4000
        )
        
        assistant_response = response.choices[0].message.content
        
        processing_time = time.time() - start_time
        logger.info(f"Assistant response received in {processing_time:.2f} seconds")
        logger.info(f"Response length: {len(assistant_response)} characters")
        
        return assistant_response
            
    except Exception as e:
        processing_time = time.time() - start_time
        logger.error(f"Error in assistant interaction after {processing_time:.2f}s: {str(e)}")
        return f"Error: {str(e)}"

def create_intelligent_summary_fallback(content, filename, max_summary_length=5000):
    """Create an intelligent summary when vector database is not available."""
    logger.info(f"Creating intelligent fallback summary for {filename}")
    
    if len(content) <= max_summary_length:
        return content
    
    # Extract key sections using keyword-based approach
    important_keywords = [
        'requirement', 'feature', 'user story', 'acceptance criteria',
        'functional', 'non-functional', 'business rule', 'constraint',
        'objective', 'goal', 'scope', 'assumption', 'dependency',
        'architecture', 'design', 'integration', 'security', 'performance'
    ]
    
    # Split content into sections
    sections = content.split('\n\n')
    scored_sections = []
    
    for section in sections:
        score = 0
        section_lower = section.lower()
        
        # Score based on keywords
        for keyword in important_keywords:
            score += section_lower.count(keyword) * 10
        
        # Prefer sections with structured content (lists, numbers)
        score += section.count('•') * 5
        score += section.count('-') * 3
        score += section.count('1.') * 5
        score += section.count('2.') * 4
        score += section.count('3.') * 3
        
        # Prefer longer, more detailed sections
        if len(section) > 200:
            score += 5
        if len(section) > 500:
            score += 5
        
        scored_sections.append((score, section))
    
    # Sort by score and take top sections
    scored_sections.sort(key=lambda x: x[0], reverse=True)
    
    # Combine top sections until we reach max length
    summary_parts = []
    current_length = 0
    
    for score, section in scored_sections:
        if current_length + len(section) <= max_summary_length:
            summary_parts.append(section)
            current_length += len(section)
        else:
            # Try to fit partial section
            remaining_space = max_summary_length - current_length - 50  # Leave space for truncation message
            if remaining_space > 200:  # Only add if meaningful portion fits
                partial_section = section[:remaining_space] + "..."
                summary_parts.append(partial_section)
            break
    
    final_summary = "\n\n".join(summary_parts)
    
    if len(final_summary) > max_summary_length:
        final_summary = final_summary[:max_summary_length] + "... [Content truncated]"
    
    logger.info(f"Intelligent fallback summary created, length: {len(final_summary)} characters")
    
    # Print Fallback Summary for debugging
    logger.info("=" * 80)
    logger.info("INTELLIGENT FALLBACK SUMMARY GENERATED:")
    logger.info("=" * 80)
    logger.info(final_summary[:2000] + "..." if len(final_summary) > 2000 else final_summary)
    logger.info("=" * 80)
    
    return final_summary.strip()

def is_valid_content(content):
    """Check if content is valid for RAG processing."""
    if not content or len(content.strip()) < 50:
        return False
    
    # Check for error messages
    error_indicators = [
        "[Unable to decode",
        "[Unable to read",
        "[Error reading",
        "[File reading failed",
        "[File appears to be empty"
    ]
    
    for indicator in error_indicators:
        if content.startswith(indicator):
            return False
    
    return True

def debug_file_content(content, filename):
    """Debug helper to print file content details."""
    logger.info(f"=== FILE CONTENT DEBUG: {filename} ===")
    logger.info(f"Content length: {len(content) if content else 0}")
    logger.info(f"Content type: {type(content)}")
    logger.info(f"Is valid for RAG: {is_valid_content(content)}")
    
    if content:
        # Show first 200 characters, safely handling non-printable characters
        try:
            preview = content[:200].replace('\n', '\\n').replace('\r', '\\r')
            # Replace any problematic Unicode characters for logging
            safe_preview = preview.encode('ascii', errors='replace').decode('ascii')
            logger.info(f"Content preview: {safe_preview}")
        except Exception as e:
            logger.info(f"Content preview: [Unable to display preview due to encoding: {str(e)}]")
        
        # Count different character types
        if len(content) > 0:
            alpha_count = sum(1 for c in content if c.isalpha())
            digit_count = sum(1 for c in content if c.isdigit())
            space_count = sum(1 for c in content if c.isspace())
            other_count = len(content) - alpha_count - digit_count - space_count
            
            logger.info(f"Character breakdown - Alpha: {alpha_count}, Digits: {digit_count}, Spaces: {space_count}, Other: {other_count}")
    
    logger.info("=== END FILE DEBUG ===")

# Flask App Startup
@app.route("/vector-db-status", methods=["GET"])
def vector_db_status():
    """View vector database status and contents via web interface."""
    logger.info("GET request received for vector database status")
    
    try:
        if not prd_collection:
            return {
                "status": "error",
                "message": "Vector database not initialized",
                "count": 0,
                "documents": []
            }
        
        # Get collection info
        collection_info = {
            "name": prd_collection.name,
            "metadata": prd_collection.metadata,
            "count": prd_collection.count()
        }
        
        # Get all documents (limit to 50 for performance)
        all_docs = prd_collection.get(limit=50)
        
        documents = []
        if all_docs['documents']:
            for i, (doc, metadata, doc_id) in enumerate(zip(
                all_docs['documents'], 
                all_docs['metadatas'], 
                all_docs['ids']
            )):
                documents.append({
                    "id": doc_id,
                    "content_preview": doc[:200] + "..." if len(doc) > 200 else doc,
                    "content_length": len(doc),
                    "metadata": metadata,
                    "index": i
                })
        
        # Create HTML response for better viewing
        html_content = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <title>Vector Database Status</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                .header {{ background-color: #f0f0f0; padding: 15px; border-radius: 5px; }}
                .document {{ border: 1px solid #ddd; margin: 10px 0; padding: 15px; border-radius: 5px; }}
                .metadata {{ background-color: #f9f9f9; padding: 10px; margin: 5px 0; border-radius: 3px; }}
                .content {{ background-color: #fff; padding: 10px; border-left: 3px solid #007bff; }}
                pre {{ white-space: pre-wrap; word-wrap: break-word; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Vector Database Status</h1>
                <p><strong>Collection:</strong> {collection_info['name']}</p>
                <p><strong>Total Documents:</strong> {collection_info['count']}</p>
                <p><strong>Showing:</strong> {len(documents)} documents (limited to 50)</p>
            </div>
        """
        
        if documents:
            for doc in documents:
                html_content += f"""
                <div class="document">
                    <h3>Document #{doc['index'] + 1}</h3>
                    <div class="metadata">
                        <strong>ID:</strong> {doc['id']}<br>
                        <strong>Length:</strong> {doc['content_length']} characters<br>
                        <strong>Filename:</strong> {doc['metadata'].get('filename', 'Unknown')}<br>
                        <strong>Type:</strong> {doc['metadata'].get('doc_type', 'Unknown')}<br>
                        <strong>Chunk:</strong> {doc['metadata'].get('chunk_index', 'N/A')} / {doc['metadata'].get('total_chunks', 'N/A')}<br>
                        <strong>Timestamp:</strong> {doc['metadata'].get('timestamp', 'Unknown')}
                    </div>
                    <div class="content">
                        <strong>Content Preview:</strong>
                        <pre>{doc['content_preview']}</pre>
                    </div>
                </div>
                """
        else:
            html_content += "<p>No documents found in the vector database.</p>"
        
        html_content += """
            <div style="margin-top: 30px; padding: 15px; background-color: #e7f3ff; border-radius: 5px;">
                <h3>How to Query the Vector Database</h3>
                <p>You can search the vector database by making a POST request to <code>/vector-db-search</code> with a JSON body:</p>
                <pre>{"query": "your search query", "top_k": 5}</pre>
                <p>Or use the search interface: <a href="/vector-db-search">Vector DB Search</a></p>
            </div>
        </body>
        </html>
        """
        
        return html_content, 200, {'Content-Type': 'text/html'}
        
    except Exception as e:
        logger.error(f"Error accessing vector database: {str(e)}")
        return {
            "status": "error",
            "message": f"Error accessing vector database: {str(e)}",
            "count": 0,
            "documents": []
        }

@app.route("/vector-db-search", methods=["GET", "POST"])
def vector_db_search():
    """Search the vector database via web interface."""
    logger.info(f"{request.method} request received for vector database search")
    
    if request.method == "GET":
        # Render search form
        html_content = """
        <!DOCTYPE html>
        <html>
        <head>
            <title>Vector Database Search</title>
            <style>
                body { font-family: Arial, sans-serif; margin: 20px; }
                .search-form { background-color: #f0f0f0; padding: 20px; border-radius: 5px; margin-bottom: 20px; }
                .result { border: 1px solid #ddd; margin: 10px 0; padding: 15px; border-radius: 5px; }
                .metadata { background-color: #f9f9f9; padding: 10px; margin: 5px 0; border-radius: 3px; }
                .content { background-color: #fff; padding: 10px; border-left: 3px solid #28a745; }
                input[type="text"] { width: 70%; padding: 10px; margin: 5px; }
                input[type="number"] { width: 100px; padding: 10px; margin: 5px; }
                button { padding: 10px 20px; background-color: #007bff; color: white; border: none; border-radius: 3px; cursor: pointer; }
                button:hover { background-color: #0056b3; }
            </style>
        </head>
        <body>
            <h1>Vector Database Search</h1>
            <div class="search-form">
                <form method="POST">
                    <label>Search Query:</label><br>
                    <input type="text" name="query" placeholder="Enter your search query..." required><br>
                    <label>Number of Results:</label><br>
                    <input type="number" name="top_k" value="5" min="1" max="20"><br><br>
                    <button type="submit">Search</button>
                </form>
            </div>
            <p><a href="/vector-db-status">← Back to Vector DB Status</a></p>
        </body>
        </html>
        """
        return html_content, 200, {'Content-Type': 'text/html'}
    
    else:  # POST request
        try:
            # Get search parameters
            if request.is_json:
                data = request.get_json()
                query = data.get("query", "")
                top_k = data.get("top_k", 5)
            else:
                query = request.form.get("query", "")
                top_k = int(request.form.get("top_k", 5))
            
            if not query:
                return {"error": "Query is required"}, 400
            
            # Search the vector database
            results = retrieve_relevant_content(query, top_k=top_k)
            
            # Create HTML response
            html_content = f"""
            <!DOCTYPE html>
            <html>
            <head>
                <title>Search Results</title>
                <style>
                    body {{ font-family: Arial, sans-serif; margin: 20px; }}
                    .header {{ background-color: #f0f0f0; padding: 15px; border-radius: 5px; margin-bottom: 20px; }}
                    .result {{ border: 1px solid #ddd; margin: 10px 0; padding: 15px; border-radius: 5px; }}
                    .metadata {{ background-color: #f9f9f9; padding: 10px; margin: 5px 0; border-radius: 3px; }}
                    .content {{ background-color: #fff; padding: 10px; border-left: 3px solid #28a745; }}
                    .relevance {{ color: #007bff; font-weight: bold; }}
                    pre {{ white-space: pre-wrap; word-wrap: break-word; }}
                </style>
            </head>
            <body>
                <div class="header">
                    <h1>Search Results</h1>
                    <p><strong>Query:</strong> "{query}"</p>
                    <p><strong>Results Found:</strong> {len(results)}</p>
                </div>
            """
            
            if results:
                for i, result in enumerate(results):
                    metadata = result.get('metadata', {})
                    content = result.get('content', '')
                    rank = result.get('relevance_rank', i + 1)
                    
                    html_content += f"""
                    <div class="result">
                        <h3>Result #{rank} <span class="relevance">(Relevance Rank: {rank})</span></h3>
                        <div class="metadata">
                            <strong>Filename:</strong> {metadata.get('filename', 'Unknown')}<br>
                            <strong>Document Type:</strong> {metadata.get('doc_type', 'Unknown')}<br>
                            <strong>Chunk:</strong> {metadata.get('chunk_index', 'N/A')} / {metadata.get('total_chunks', 'N/A')}<br>
                            <strong>Timestamp:</strong> {metadata.get('timestamp', 'Unknown')}
                        </div>
                        <div class="content">
                            <strong>Content:</strong>
                            <pre>{content}</pre>
                        </div>
                    </div>
                    """
            else:
                html_content += "<p>No results found for your query.</p>"
            
            html_content += """
                <div style="margin-top: 30px;">
                    <a href="/vector-db-search">← New Search</a> | 
                    <a href="/vector-db-status">Vector DB Status</a>
                </div>
            </body>
            </html>
            """
            
            return html_content, 200, {'Content-Type': 'text/html'}
            
        except Exception as e:
            logger.error(f"Error searching vector database: {str(e)}")
            return {"error": f"Search failed: {str(e)}"}, 500

@app.route("/vector-db-clear", methods=["POST"])
def vector_db_clear():
    """Clear all documents from the vector database."""
    global prd_collection
    logger.info("POST request received to clear vector database")
    
    try:
        if not prd_collection:
            return {"status": "error", "message": "Vector database not initialized"}
        
        # Get current count
        current_count = prd_collection.count()
        
        # Delete the collection and recreate it
        chroma_client.delete_collection(prd_collection.name)
        
        # Recreate the collection
        prd_collection = chroma_client.get_or_create_collection(
            name="prd_documents",
            metadata={"description": "PRD and documentation storage for RAG"}
        )
        
        logger.info(f"Vector database cleared. Removed {current_count} documents.")
        
        return {
            "status": "success",
            "message": f"Vector database cleared successfully. Removed {current_count} documents.",
            "documents_removed": current_count
        }
        
    except Exception as e:
        logger.error(f"Error clearing vector database: {str(e)}")
        return {"status": "error", "message": f"Error clearing database: {str(e)}"}

@app.route("/document-upload", methods=["POST"])
def document_upload_preview():
    """Step 1: Upload and preview documents with RAG summary before epic generation."""
    logger.info("=== DOCUMENT UPLOAD ENDPOINT CALLED ===")
    start_time = time.time()
    logger.info("POST request received for document upload and preview")
    
    try:
        logger.info("=== INSIDE TRY BLOCK ===")
        
        # Test logging to ensure we reach this point
        logger.info("About to process document upload...")
        
        # Collect input
        logger.info("Collecting form data and file uploads for preview")
        context = request.form.get("context", "")
        prd_file = request.files.get("prd_file")
        additional_docs = request.files.get("additional_docs")
        csv_file = request.files.get("csv_file")
        
        logger.info(f"Context provided: {'Yes' if context else 'No'}")
        logger.info(f"PRD file uploaded: {'Yes' if prd_file else 'No'}")
        logger.info(f"Additional docs uploaded: {'Yes' if additional_docs else 'No'}")
        logger.info(f"CSV file uploaded: {'Yes' if csv_file else 'No'}")
        
        if prd_file:
            logger.info(f"PRD file name: {prd_file.filename}")
        if additional_docs:
            logger.info(f"Additional docs file name: {additional_docs.filename}")
        if csv_file:
            logger.info(f"CSV file name: {csv_file.filename}")

        # Process CSV file first if provided (store in vector DB)
        csv_summary = ""
        if csv_file and csv_file.filename.lower().endswith('.csv'):
            logger.info("Processing CSV file for vector database storage")
            csv_summary = process_csv_to_vector_db(csv_file)
            logger.info(f"CSV processing completed: {csv_summary[:100]}...")
        elif csv_file:
            logger.warning(f"Uploaded file {csv_file.filename} is not a CSV file")
            csv_summary = f"Warning: {csv_file.filename} is not a CSV file"

        # Parallel file reading for better performance
        logger.info("Reading uploaded files in parallel for preview")
        
        def read_files_parallel():
            with ThreadPoolExecutor(max_workers=2) as file_executor:
                prd_future = file_executor.submit(safe_read, prd_file) if prd_file else None
                docs_future = file_executor.submit(safe_read, additional_docs) if additional_docs else None
                
                prd_content = prd_future.result() if prd_future else ""
                docs_content = docs_future.result() if docs_future else ""
                
                return prd_content, docs_content
        
        prd_content, docs_content = read_files_parallel()
        
        # Debug file content
        if prd_content:
            debug_file_content(prd_content, prd_file.filename if prd_file else "prd_content")
        if docs_content:
            debug_file_content(docs_content, additional_docs.filename if additional_docs else "additional_docs")

        # Store documents in vector database and create RAG summaries for preview
        logger.info("Processing documents with RAG enhancement for preview")
        
        # Process PRD with RAG if content is substantial and valid
        prd_summary = ""
        if is_valid_content(prd_content):
            if len(prd_content) > 1000:  # Lower threshold for preview
                logger.info("Creating RAG-enhanced PRD summary for preview")
                prd_summary = create_rag_summary(prd_content, prd_file.filename if prd_file else "prd_content", max_summary_length=15000)
            else:
                logger.info("PRD content is valid but small, using as-is for preview")
                prd_summary = prd_content
        else:
            logger.warning(f"PRD content is invalid for RAG processing: {prd_content[:100] if prd_content else 'None'}")
            prd_summary = "No valid PRD content available - please check file format and encoding."
          # Process additional docs with RAG
        docs_summary = ""
        if is_valid_content(docs_content):
            if len(docs_content) > 1000:  # Lower threshold for preview
                logger.info("Creating RAG-enhanced docs summary for preview")
                docs_summary = create_rag_summary(docs_content, additional_docs.filename if additional_docs else "additional_docs", max_summary_length=8000)
            else:
                logger.info("Additional docs content is valid but small, using as-is for preview")
                docs_summary = docs_content
        else:
            logger.warning(f"Additional docs content is invalid: {docs_content[:100] if docs_content else 'None'}")
            docs_summary = "No valid additional documentation available."
        
        logger.info(f"PRD summary length: {len(prd_summary)} characters")
        logger.info(f"Additional docs summary length: {len(docs_summary)} characters")
        
        # Print RAG summaries to console for debugging
        print("\n" + "=" * 100)
        print("RAG-PROCESSED DOCUMENT SUMMARIES:")
        print("=" * 100)
        
        if context:
            print("USER CONTEXT:")
            print("-" * 50)
            print(context)
            print("-" * 50)
        
        if prd_summary:
            print("PRD SUMMARY:")
            print("-" * 50)
            print(prd_summary)
            print("-" * 50)
        
        if docs_summary:
            print("ADDITIONAL DOCS SUMMARY:")
            print("-" * 50)
            print(docs_summary)
            print("-" * 50)
        
        print("=" * 100 + "\n")
        
        # Create preview data
        preview_data = {
            "user_context": context,
            "prd_summary": prd_summary,
            "docs_summary": docs_summary,
            "csv_summary": csv_summary,
            "prd_filename": prd_file.filename if prd_file else None,
            "docs_filename": additional_docs.filename if additional_docs else None,
            "csv_filename": csv_file.filename if csv_file else None,
            "processing_time": time.time() - start_time
        }
        
        logger.info(f"Document preview generated in {preview_data['processing_time']:.2f} seconds")
        
        # Return JSON response for AJAX handling
        return jsonify({
            "success": True,
            "preview_data": preview_data,
            "message": "Documents processed successfully. Review the summary below and click 'Generate Epics' to proceed."
        })
        
    except Exception as e:
        print(f"Exception in /document-upload: {e}")
        logger.info(f"Error in document_upload_preview: {str(e)}")
        logger.error(f"Exception type: {type(e).__name__}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        return jsonify({
            "success": False,
            "error": f"Error processing documents: {str(e)}"
        }), 500

@app.route("/generate-epics", methods=["POST"])
def generate_epics_from_preview():
    """Step 2: Generate epics using the previewed RAG summary."""
    start_time = time.time()
    logger.info("POST request received for epic generation from preview")
    
    try:
        # Get the preview data from the request
        request_data = request.get_json()
        if not request_data:
            logger.error("No preview data received for epic generation")
            return jsonify({"success": False, "error": "No preview data received"}), 400
        
        user_context = request_data.get("user_context", "")
        prd_summary = request_data.get("prd_summary", "")
        docs_summary = request_data.get("docs_summary", "")
        
        logger.info(f"Epic generation - Context: {'Yes' if user_context else 'No'}")
        logger.info(f"Epic generation - PRD summary: {len(prd_summary)} characters")
        logger.info(f"Epic generation - Docs summary: {len(docs_summary)} characters")

        # Create enhanced context for Epic Generator
        enhanced_context = f"""
        {user_context}

        RAG-Enhanced PRD Analysis:
        {prd_summary}

        Additional Context:
        {docs_summary}

        Instructions: The above content has been intelligently extracted and summarized using RAG (Retrieval Augmented Generation). 
        It contains the most relevant requirements, user stories, and business objectives from the original documents.
        Use this curated information to generate comprehensive epics and user stories.        """
        logger.info(f"Enhanced context length: {len(enhanced_context)} characters")
        
        # Check cache first
        prompt_hash = get_cache_key(enhanced_context)
        cached_result = response_cache.get(prompt_hash)
        if cached_result:
            logger.info("Cache hit! Returning cached response")
            processing_time = time.time() - start_time
            logger.info(f"Total processing time (cached): {processing_time:.2f} seconds")
            return jsonify({
                "success": True,
                "epics": cached_result,
                "processing_time": processing_time,
                "cached": True
            })          # Log token count for the enhanced context
        context_tokens = count_tokens(enhanced_context, "gpt-4o")
        logger.info(f"Enhanced context token count: {context_tokens:,} tokens")
        
        # Check if context is within token limits
        if context_tokens > 120000:
            logger.warning(f"Enhanced context token count ({context_tokens:,}) is approaching the 128k limit!")
            # Further optimize if needed
            enhanced_context = enhanced_context[:100000] + "\n[Content truncated due to token limits]"
            context_tokens = count_tokens(enhanced_context, "gpt-4o")
            logger.info(f"Truncated context token count: {context_tokens:,} tokens")
        else:
            logger.info(f"Enhanced context is within safe token limits ({context_tokens:,}/128,000 tokens)")    # Skip PRD Parser Agent - RAG has already done the parsing and summarization
        logger.info("****************Skipping PRD Parser - Using RAG Summary Directly")
        logger.info("Starting Epic Generator with RAG-enhanced content")
        
        # Print Enhanced Context being sent to Epic Agent
        print("\n" + "=" * 100)
        print("ENHANCED CONTEXT INPUT TO EPIC AGENT:")
        print("=" * 100)
        print(enhanced_context)  # Print the FULL content
        print("=" * 100)
        print(f"Total length: {len(enhanced_context)} characters")
        print(f"Total tokens: {count_tokens(enhanced_context, 'gpt-4o'):,} tokens")
        print("=" * 100 + "\n")
        
        # Also log to file (truncated for file logs)
        logger.info("=" * 80)
        logger.info("ENHANCED CONTEXT INPUT TO EPIC AGENT:")
        logger.info("=" * 80)
        logger.info(enhanced_context[:5000] + "..." if len(enhanced_context) > 5000 else enhanced_context)
        logger.info("=" * 80)
        
        # Directly use Epic Generator with RAG-enhanced content
        start_epic_time = time.time()
        epic_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", enhanced_context)
        epic_processing_time = time.time() - start_epic_time
        
        logger.info("################Epic Generator response received")
        logger.info(f"Epic Generator processing time: {epic_processing_time:.2f} seconds")
          # Print Epic Agent Response to console
        print("\n" + "=" * 100)
        print("EPIC AGENT RESPONSE OUTPUT:")
        print("=" * 100)
        print(epic_response)
        print("=" * 100)
        print(f"Response length: {len(epic_response)} characters")
        print("=" * 100 + "\n")                
        # Print Epic Agent Response
        logger.info("=" * 80)
        logger.info("EPIC AGENT RESPONSE OUTPUT:")
        logger.info("=" * 80)
        logger.info(epic_response)
        logger.info("=" * 80)
        
        # Log token usage for Epic Generator interaction
        log_token_usage(enhanced_context, epic_response, model="gpt-4o", context="RAG-Enhanced Epic Generator")

        final_output = epic_response
        logger.info(f"Final output length: {len(final_output)} characters")
        
        # Cache the result for future requests
        response_cache[prompt_hash] = final_output
        logger.info("Response cached for future requests")
        
        processing_time = time.time() - start_time
        logger.info(f"Total RAG-optimized processing time: {processing_time:.2f} seconds")
        
        # Calculate time savings
        estimated_traditional_time = processing_time * 2  # Estimate of traditional two-agent approach
        time_saved = estimated_traditional_time - processing_time
        logger.info(f"Estimated time saved by RAG optimization: {time_saved:.2f} seconds")
        
        # Return JSON response for AJAX handling
        return jsonify({
            "success": True,
            "epics": final_output,
            "processing_time": processing_time,
            "time_saved": time_saved,
            "cached": False
        })
        
    except Exception as e:
        logger.error(f"Error in generate_epics_from_preview: {str(e)}")
        logger.error(f"Exception type: {type(e).__name__}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        return jsonify({
            "success": False,
            "error": f"Error generating epics: {str(e)}"
        }), 500

@app.route("/epic-results", methods=["GET", "POST"])
def show_epic_results():
    """Display the epic results page after generation."""
    if request.method == "POST":
        epics = request.form.get("epics", "")
        return render_template("poc2_epic_story_screen.html", epics=epics)
    else:
        return render_template("poc2_epic_story_screen.html", epics="")

@app.route("/approve-epics", methods=["POST"])
def approve_epics():
    """Process approved epics and generate user stories for them."""
    start_time = time.time()
    logger.info("POST request received for epic approval and user story generation")
    user_story_list = []  # <-- new list to collect structured stories
    try:
        # Get the approved epic IDs from the form
        epic_ids_str = request.form.get("epic_ids", "")
        epic_ids = [epic_id.strip() for epic_id in epic_ids_str.split(",") if epic_id.strip()]
        
        logger.info(f"Processing approval for epic IDs: {epic_ids}")
        
        if not epic_ids:
            logger.error("No epic IDs provided for approval")
            return render_template("poc2_epic_story_screen.html", 
                                 epics="No epics selected for approval", 
                                 user_stories="")
          # Get the current epics content to preserve it
        current_epics = request.form.get("current_epics", "")
        logger.info(f"Current epics content length: {len(current_epics)} characters")
        
        # Store current epics in session for Epic Chat context
        session['current_epics'] = current_epics
        session.modified = True
        logger.info("Stored current epics in session for Epic Chat context")
          # Get the selected epic contents (actual epic data)
        selected_epic_contents_str = request.form.get("selected_epic_contents", "{}")
        try:
            selected_epic_contents = json.loads(selected_epic_contents_str)
            logger.info(f"Received epic contents for {len(selected_epic_contents)} epics")
        except json.JSONDecodeError as e:
            logger.error(f"Failed to parse selected epic contents: {e}")
            selected_epic_contents = {}
        
        # Generate user stories for each approved epic
        user_stories = []
        logger.info("Starting user story generation for approved epics")
        
        for i, epic_id in enumerate(epic_ids):
            logger.info(f"Generating user stories for epic: {epic_id}")            # Get the actual epic content instead of just the ID
            epic_content = selected_epic_contents.get(epic_id, {})
            epic_title = epic_content.get('title', f'Epic {i+1}')
            epic_description = epic_content.get('description', 'No description available')
            
            logger.info(f"Epic {epic_id} - Title: {epic_title[:50]}...")
            logger.info(f"Epic {epic_id} - Description: {epic_description[:100]}...")
            
            # Create a prompt that aligns with the agent's system instructions
            prompt = f"""
            Epic Title: {epic_title}
            Epic Description: {epic_description}
            
            Generate user stories for this approved epic. Return ONLY the JSON array format with story_id, name, priority, and systems. Do not include descriptions or acceptance criteria.
            """
              # Print what's being sent to User Story Agent
            logger.info("=" * 80)
            logger.info(f"USER STORY AGENT INPUT FOR EPIC {i+1}:")
            logger.info("=" * 80)
            logger.info(f"Epic ID: {epic_id}")
            logger.info(f"Epic Title: {epic_title}")
            logger.info(f"Epic Description: {epic_description}")
            logger.info("-" * 50)
            logger.info("Full Prompt:")
            logger.info(prompt)
            logger.info("=" * 80)
            logger.info(f"Prompt length: {len(prompt)} characters")
            logger.info("=" * 80 + "\n")
            
            try:
                story_response = ask_assistant_from_file_optimized("poc2_agent3_basic_user_story", prompt)
                
                # Print User Story Agent Response
                logger.info("=" * 80)
                logger.info(f"USER STORY AGENT RESPONSE FOR EPIC {i+1}:")
                logger.info("=" * 80)
                logger.info(story_response)
                logger.info("=" * 80)
                logger.info(f"Response length: {len(story_response)} characters")
                logger.info("=" * 80 + "\n")                
                if story_response:
                    # Send the raw JSON response for frontend parsing instead of HTML cards
                    user_stories.append(story_response)
                    user_story_list.append({
                        "epic": epic_title,
                        "stories": story_response  # Raw agent response for JS parsing
                    })
                    logger.info(f"Successfully generated user stories for {epic_id}")
                else:
                    logger.warning(f"No response received for epic {epic_id}")
                    user_stories.append(f"""
                    <div class='user-story-card'>
                    <h6>User Stories for: {epic_title}</h6>
                    <label>
                        <input type="radio" name="user_story_id" value="story_{i+1}">
                        <div style='padding: 10px; background-color: #f8f9fa; border-radius: 4px; margin: 5px 0;'>
                        {story_response}
                        </div>
                    </label>
                    </div>
                    """)
            except Exception as e:
                logger.error(f"Error generating user stories for epic {epic_id}: {str(e)}")
                user_stories.append(f"""
                <div class='user-story-card'>
                    <h6>User Stories for {epic_id}</h6>
                    <p class="text-danger">Error generating user stories: {str(e)}</p>
                </div>
                """)
        
        # Combine all user stories - send raw content for JS parsing
        user_stories_content = "\n".join(user_stories)
        
        processing_time = time.time() - start_time
        logger.info(f"User story generation completed in {processing_time:.2f} seconds")
        logger.info(f"Generated user stories for {len(epic_ids)} epics")
        
        
        # Render the template with both epics (preserved) and user stories (raw for JS parsing)
        return render_template("poc2_epic_story_screen.html",
                       epics=current_epics,  # This should be epics HTML
                       user_stories=user_stories_content)  # Raw content for JS to parse into table

        
    except Exception as e:
        logger.error(f"Error in approve_epics: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        
        # Return error page with preserved epics
        error_message = f"""
        <div class='user-story-card'>
            <h6>Error Generating User Stories</h6>
            <p class="text-danger">An error occurred while generating user stories: {str(e)}</p>
            <p class="text-muted">Please try again or contact support if the issue persists.</p>
        </div>
        """
        
        return render_template(
            "poc2_epic_story_screen.html", 
            epics=request.form.get("current_epics", ""), 
            user_stories=error_message
        )

@app.route("/user-story-details", methods=["GET", "POST"])
def user_story_details():
    """Process selected user story and display details with acceptance criteria."""
    logger.info("Request received for user story details")
    
    try:
        # Handle POST request (form submission from epic story screen)
        if request.method == "POST":
            logger.info("POST request received for user story selection")
            
            # Get the selected story IDs and details (now supporting multiple selections)
            if request.is_json:
                data = request.get_json()
                selected_story_id = data.get("selected_story_id", "")  # Comma-separated IDs
                story_name = data.get("selected_story_name", "")
                selected_story_description = data.get("selected_story_description", "")
                selected_stories_data = data.get("selected_stories_data", "")  # JSON string of story objects
                epic_title = data.get("epic_title", "")
                priority = data.get("priority", "High")
                logger.info(f"Processing selected user story IDs (JSON): {selected_story_id}")
            else:
                selected_story_id = request.form.get("selected_story_id", "")  # Comma-separated IDs
                story_name = request.form.get("selected_story_name", "")
                selected_story_description = request.form.get("selected_story_description", "")
                selected_stories_data = request.form.get("selected_stories_data", "")  # JSON string of story objects
                epic_title = request.form.get("epic_title", "")
                priority = request.form.get("priority", "High")
                logger.info(f"Processing selected user story IDs (Form): {selected_story_id}")
            
            # Parse multiple story IDs
            story_ids = [id.strip() for id in selected_story_id.split(',') if id.strip()] if selected_story_id else []
            logger.info(f"Processing {len(story_ids)} selected user stories: {story_ids}")
            
            # Parse detailed story data if available
            selected_stories = []
            if selected_stories_data:
                try:
                    selected_stories = json.loads(selected_stories_data)
                    logger.info(f"Parsed detailed data for {len(selected_stories)} stories")
                except json.JSONDecodeError as e:
                    logger.warning(f"Failed to parse selected stories data: {e}")
                    selected_stories = []
            
            logger.info(f"Processing selected story names: {story_name}")
            logger.info(f"Processing selected story descriptions: {selected_story_description[:100] if selected_story_description else 'N/A'}...")
            
            # Always proceed - if no stories selected, use defaults to ensure agents are called
            if not story_ids:
                logger.info("No user stories selected, using default values to ensure agents are called")
                story_ids = ['default-story-1']
                story_name = story_name or "Lock critical fields for charged-off accounts"
                selected_story_description = selected_story_description or "As a system architect, I want to identify and classify the necessary data fields for charged-off accounts so that we can ensure only required and relevant data is transmitted securely."
                selected_story_id = 'default-story-1'
            else:
                selected_story_id = ','.join(story_ids)
            
            # Provide default values if story details are missing
            if not story_name or not selected_story_description:
                logger.info("Story name or description missing, using default values")
                default_story_name = "Lock critical fields for charged-off accounts"
                default_story_description = "As a system architect, I want to identify and classify the necessary data fields for charged-off accounts so that we can ensure only required and relevant data is transmitted securely."

                story_name = story_name or default_story_name
                selected_story_description = selected_story_description or default_story_description
                
                logger.info(f"Using default story name: {story_name}")
                logger.info(f"Using default story description: {selected_story_description[:100]}...")
            
            logger.info(f"Successfully processed user story selection: {selected_story_id}")
            
            # Extract data from selected stories for the new agent format
            if selected_stories:
                # Use data from the first selected story as primary, combine others
                primary_story = selected_stories[0]
                
                # Extract systems information from all selected stories
                all_systems = set()
                all_priorities = []
                combined_names = []
                combined_descriptions = []
                
                for story in selected_stories:
                    # Handle new agent response format
                    story_name_field = story.get('name', story.get('title', 'Untitled Story'))
                    story_desc_field = story.get('description', story_name_field)
                    story_priority = story.get('priority', 'Medium')
                    story_systems = story.get('systems', [])
                    
                    combined_names.append(story_name_field)
                    combined_descriptions.append(story_desc_field)
                    all_priorities.append(story_priority)
                    
                    # Handle systems field - can be array or string
                    if isinstance(story_systems, list):
                        all_systems.update(story_systems)
                    elif isinstance(story_systems, str):
                        all_systems.update([s.strip() for s in story_systems.split(',') if s.strip()])
                
                # Override with combined data
                story_name = ' | '.join(combined_names) if len(combined_names) > 1 else combined_names[0] if combined_names else story_name
                selected_story_description = ' | '.join(combined_descriptions) if len(combined_descriptions) > 1 else combined_descriptions[0] if combined_descriptions else selected_story_description
                responsible_systems = ', '.join(sorted(all_systems)) if all_systems else "TBD"
                priority = all_priorities[0] if all_priorities else "Medium"  # Use first story's priority
                
                logger.info(f"Extracted from {len(selected_stories)} stories:")
                logger.info(f"  - Combined name: {story_name}")
                logger.info(f"  - Systems: {responsible_systems}")
                logger.info(f"  - Priority: {priority}")
            
            # Generate additional data fields for the new UI
            if 'responsible_systems' not in locals():
                responsible_systems = "CAPS, CMS"  # Default value if not extracted from stories
            tagged_requirements = [
                "F1. Capture Personal Information for New Customer Account",
                "KDA2. Customer full name",
                "KDA3. Customer date of birth", 
                "KDA4. Customer SSN or equivalent"
            ]
            
            # Call the acceptance criteria agent - pass only story name
            prompt = story_name
            logger.info(f"Processing selected prompt ****Acceptance Criteria Prompt **** : {prompt}")

            acceptance_response = ask_assistant_from_file_optimized("poc2_agent4_acceptanceCriteria_gen", prompt)

            # Call the new User Description Agent - pass only story name
            description_prompt = story_name
            logger.info(f"Processing User Description Agent prompt: {description_prompt}")
            
            description_response = ask_assistant_from_file_optimized("poc2_agent5_description", description_prompt)
            
            # Parse the description response
            enhanced_description = selected_story_description  # Default fallback
            try:
                # Clean the response first - remove markdown code blocks if present
                clean_response = description_response.strip()
                
                # Remove markdown JSON code blocks
                if clean_response.startswith('```json') and clean_response.endswith('```'):
                    clean_response = clean_response[7:-3].strip()
                elif clean_response.startswith('```') and clean_response.endswith('```'):
                    clean_response = clean_response[3:-3].strip()
                
                # Try to parse as JSON
                description_data = json.loads(clean_response)
                if isinstance(description_data, dict) and 'description' in description_data:
                    enhanced_description = description_data['description']
                    logger.info(f"Enhanced description generated: {enhanced_description[:100]}...")
                else:
                    logger.warning("User Description Agent did not return expected format")
                    # If it's a dict but no 'description' key, try to get the first value
                    if isinstance(description_data, dict):
                        values = list(description_data.values())
                        if values:
                            enhanced_description = str(values[0])
            except json.JSONDecodeError as e:
                logger.warning(f"Failed to parse User Description Agent response: {e}")
                # Use the raw response if it's not JSON, but try to extract content from markdown
                if description_response and len(description_response.strip()) > 10:
                    clean_response = description_response.strip()
                    # Remove any remaining markdown artifacts
                    if clean_response.startswith('```') and clean_response.endswith('```'):
                        clean_response = clean_response[3:-3].strip()
                    enhanced_description = clean_response

            # Parse criteria (assume response is a bullet list or parse JSON if needed)
            try:
                # First, clean the response to remove markdown code blocks if present
                clean_response = acceptance_response.strip()
                if clean_response.startswith('```json') and clean_response.endswith('```'):
                    clean_response = clean_response[7:-3].strip()
                elif clean_response.startswith('```') and clean_response.endswith('```'):
                    clean_response = clean_response[3:-3].strip()
                
                acceptance_criteria = json.loads(clean_response)
                if isinstance(acceptance_criteria, list):
                    # If it's already a list, use it
                    pass
                elif isinstance(acceptance_criteria, dict) and 'acceptance_criteria' in acceptance_criteria:
                    # If it's a dict with 'acceptance_criteria' key, extract the list
                    acceptance_criteria = acceptance_criteria['acceptance_criteria']
                elif isinstance(acceptance_criteria, dict) and 'criteria' in acceptance_criteria:
                    # If it's a dict with 'criteria' key, extract the list
                    acceptance_criteria = acceptance_criteria['criteria']
                elif isinstance(acceptance_criteria, dict):
                    # If it's a dict but doesn't have expected keys, get the first list value
                    for value in acceptance_criteria.values():
                        if isinstance(value, list):
                            acceptance_criteria = value
                            break
                    else:
                        # If no list found in values, convert all values to list
                        acceptance_criteria = list(acceptance_criteria.values())
                else:
                    # If it's some other format, convert to list
                    acceptance_criteria = [str(acceptance_criteria)]
                
                # Ensure we have a flat list of strings
                if not isinstance(acceptance_criteria, list):
                    acceptance_criteria = [str(acceptance_criteria)]
                    
            except json.JSONDecodeError:
                # If JSON parsing fails, split by newlines and clean up
                acceptance_criteria = [
                    line.strip() 
                    for line in acceptance_response.strip().split("\n") 
                    if line.strip() and not line.strip().startswith('#')
                ]
                # Remove bullet points and numbering
                cleaned_criteria = []
                for criterion in acceptance_criteria:
                    # Remove common bullet point patterns and JSON artifacts
                    criterion = criterion.lstrip('•-*').lstrip('123456789.').strip()
                    # Remove JSON brackets or quotes if they somehow leaked through
                    criterion = criterion.strip('[]{}"\'\n\r\t')
                    if criterion and not criterion.startswith('{') and not criterion.startswith('['):
                        cleaned_criteria.append(criterion)
                acceptance_criteria = cleaned_criteria


            logger.info(f"Processing selected prompt ****Acceptance Criteria **** : {acceptance_criteria}")
              # For AJAX requests, return JSON for navigation
            if request.is_json:
                return jsonify({
                    "success": True,
                    "story_id": selected_story_id,
                    "story_name": story_name,
                    "enhanced_description": enhanced_description,
                    "html": render_template(
                        "poc2_user_story_details.html",
                        epic_title=epic_title,
                        user_story_name=story_name,
                        user_story_description=enhanced_description,  # Use enhanced description
                        original_description=selected_story_description,  # Keep original for reference
                        priority=priority,
                        responsible_systems=responsible_systems,
                        acceptance_criteria=acceptance_criteria,
                        tagged_requirements=tagged_requirements,
                        story_id=selected_story_id
                    ),
                    "message": "User story processed successfully with enhanced description"
                })
            
            # For regular form submission, render the template directly
            return render_template(
                "poc2_user_story_details.html",
                epic_title=epic_title,
                user_story_name=story_name,
                user_story_description=enhanced_description,  # Use enhanced description
                original_description=selected_story_description,  # Keep original for reference
                priority=priority,
                responsible_systems=responsible_systems,
                acceptance_criteria=acceptance_criteria,
                tagged_requirements=tagged_requirements,
                story_id=selected_story_id
            )
            
        # Handle GET request (direct access or fallback)
        else:
            logger.info("GET request received for user story details")
            
            # Try to get data from session first (backward compatibility)
            session_data = session.get('current_story')
            
            if session_data:
                # Use session data
                logger.info("Using story data from session")
                epic_title = session_data.get('epic_title', '')
                user_story_title = session_data.get('story_name', '')
                user_story_description = session_data.get('story_description', '')
                priority = session_data.get('priority', 'High')
                story_id = session_data.get('story_id', '')
                acceptance_criteria = session_data.get('acceptance_criteria', [])
                responsible_systems = session_data.get('responsible_systems', 'CAPS, CMS')
                tagged_requirements = session_data.get('tagged_requirements', [
                    "F1. Capture Personal Information for New Customer Account",
                    "KDA2. Customer full name",
                    "KDA3. Customer date of birth", 
                    "KDA4. Customer SSN or equivalent"
                ])
                
                # Clean up session data after use
                session.pop('current_story', None)
                
            else:
                # Fallback to query parameters
                logger.info("No session data found, checking query parameters")
                epic_title = request.args.get("epic_title", "")
                user_story_title = request.args.get("user_story_title", "")
                user_story_description = request.args.get("user_story_description", "")
                priority = request.args.get("priority", "High")
                story_id = request.args.get("story_id", "")
                
                # Use consistent variable names - with default fallback
                final_story_name = user_story_title or "Lock critical fields for charged-off accounts"
                final_story_description = user_story_description or "No description available"
                
                # Generate acceptance criteria if not in session
                logger.info(f"Processing story: {final_story_name}")
                logger.info(f"Description: {final_story_description[:100]}...")

                # Generate additional data fields for the new UI
                responsible_systems = "CAPS, CMS"  # Default value, could be enhanced later
                tagged_requirements = [
                    "F1. Capture Personal Information for New Customer Account",
                    "KDA2. Customer full name",
                    "KDA3. Customer date of birth", 
                    "KDA4. Customer SSN or equivalent"
                ]

                # Call the acceptance criteria agent - pass only story name
                prompt = final_story_name
                acceptance_response = ask_assistant_from_file_optimized("poc2_agent4_acceptanceCriteria_gen", prompt)
                
                # Call the new User Description Agent - pass only story name
                description_prompt = final_story_name
                logger.info(f"Processing User Description Agent prompt: {description_prompt}")
                
                description_response = ask_assistant_from_file_optimized("poc2_agent5_description", description_prompt)
                
                # Parse the description response
                enhanced_description = final_story_description  # Default fallback
                try:
                    # Clean the response first - remove markdown code blocks if present
                    clean_response = description_response.strip()
                    
                    # Remove markdown JSON code blocks
                    if clean_response.startswith('```json') and clean_response.endswith('```'):
                        clean_response = clean_response[7:-3].strip()
                    elif clean_response.startswith('```') and clean_response.endswith('```'):
                        clean_response = clean_response[3:-3].strip()
                    
                    # Try to parse as JSON
                    description_data = json.loads(clean_response)
                    if isinstance(description_data, dict) and 'description' in description_data:
                        enhanced_description = description_data['description']
                        logger.info(f"Enhanced description generated: {enhanced_description[:100]}...")
                    else:
                        logger.warning("User Description Agent did not return expected format")
                        # If it's a dict but no 'description' key, try to get the first value
                        if isinstance(description_data, dict):
                            values = list(description_data.values())
                            if values:
                                enhanced_description = str(values[0])
                except json.JSONDecodeError as e:
                    logger.warning(f"Failed to parse User Description Agent response: {e}")
                    # Use the raw response if it's not JSON, but try to extract content from markdown
                    if description_response and len(description_response.strip()) > 10:
                        clean_response = description_response.strip()
                        # Remove any remaining markdown artifacts
                        if clean_response.startswith('```') and clean_response.endswith('```'):
                            clean_response = clean_response[3:-3].strip()
                        enhanced_description = clean_response
                  # Parse criteria (assume response is a bullet list or parse JSON if needed)
                try:
                    acceptance_criteria = json.loads(acceptance_response)
                    if isinstance(acceptance_criteria, list):
                        # If it's already a list, use it
                        pass
                    elif isinstance(acceptance_criteria, dict) and 'criteria' in acceptance_criteria:
                        # If it's a dict with criteria key, extract the list
                        acceptance_criteria = acceptance_criteria['criteria']
                    elif isinstance(acceptance_criteria, dict):
                        # If it's a dict but doesn't have 'criteria' key, convert to list of values
                        acceptance_criteria = list(acceptance_criteria.values())
                    else:
                        # If it's some other format, convert to list
                        acceptance_criteria = [str(acceptance_criteria)]
                except:
                    # If JSON parsing fails, split by newlines and clean up
                    acceptance_criteria = [
                        line.strip() 
                        for line in acceptance_response.strip().split("\n") 
                        if line.strip() and not line.strip().startswith('#')
                    ]
                    # Remove bullet points and numbering
                    cleaned_criteria = []
                    for criterion in acceptance_criteria:
                        # Remove common bullet point patterns and JSON artifacts
                        criterion = criterion.lstrip('•-*').lstrip('123456789.').strip()
                        # Remove JSON brackets or quotes if they somehow leaked through
                        criterion = criterion.strip('[]{}"\'\n\r\t')
                        if criterion and not criterion.startswith('{') and not criterion.startswith('['):
                            cleaned_criteria.append(criterion)
                    acceptance_criteria = cleaned_criteria
                
                user_story_title = final_story_name
                user_story_description = enhanced_description  # Use enhanced description
        
            return render_template(
                "poc2_user_story_details.html",
                epic_title=epic_title,
                user_story_name=user_story_title,
                user_story_description=user_story_description,
                original_description=final_story_description if 'final_story_description' in locals() else user_story_description,  # Keep original for reference
                priority=priority,
                responsible_systems=responsible_systems,
                acceptance_criteria=acceptance_criteria,
                tagged_requirements=tagged_requirements,
                story_id=story_id
            )
        
    except Exception as e:
        logger.error(f"Error in user_story_details: {str(e)}")
        if request.is_json:
            return jsonify({
                "success": False,
                "error": f"Error processing user story details: {str(e)}"
            }), 500
        else:
            return render_template("poc2_user_story_details.html", 
                                 error=f"Error processing user story details: {str(e)}"), 500



@app.route("/epic-chat", methods=["POST"])
def epic_chat():
    """Handle Epic Chat interactions with session persistence and automatic cleanup."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No data provided"}), 400
        
        user_message = data.get("message", "").strip()
        if not user_message:
            return jsonify({"success": False, "error": "Message cannot be empty"}), 400
        
        # Initialize chat history in session if not exists or if expired
        current_time = datetime.now()
        chat_session_timeout = 30 * 60  # 30 minutes in seconds
        
        if 'epic_chat_history' not in session:
            session['epic_chat_history'] = []
            session['epic_chat_last_activity'] = current_time.isoformat()
            logger.info("Initialized new Epic Chat session")
        else:
            # Check if session has expired (30 minutes of inactivity)
            last_activity_str = session.get('epic_chat_last_activity')
            if last_activity_str:
                try:
                    last_activity = datetime.fromisoformat(last_activity_str)
                    time_diff = (current_time - last_activity).total_seconds()
                    
                    if time_diff > chat_session_timeout:
                        # Clear expired session
                        session['epic_chat_history'] = []
                        logger.info(f"Cleared expired Epic Chat session (inactive for {time_diff/60:.1f} minutes)")
                except Exception as e:
                    logger.warning(f"Error parsing last activity time: {e}")
                    session['epic_chat_history'] = []
        
        # Limit chat history to last 50 messages to prevent session bloat
        if len(session['epic_chat_history']) > 50:
            session['epic_chat_history'] = session['epic_chat_history'][-40:]  # Keep last 40 messages
            logger.info("Trimmed Epic Chat history to prevent session bloat")
        
        # Update last activity time
        session['epic_chat_last_activity'] = current_time.isoformat()
        
        # Add user message to chat history
        session['epic_chat_history'].append({
            "role": "user",
            "message": user_message,
            "timestamp": current_time.strftime("%Y-%m-%d %H:%M:%S")
        })
        
        # Get current epics from the frontend request (this is the displayed epics context)
        current_epics = data.get('current_epics_content', '').strip()
        
        # Log the current epics context for debugging
        logger.info(f"Epic Chat received current epics context: {current_epics[:500]}...")
        
        # Format current epics for better context understanding
        epics_context = current_epics
        if current_epics:
            try:
                # Try to parse as JSON first (structured format from frontend)
                import json
                epics_data = json.loads(current_epics)
                if isinstance(epics_data, list):
                    epics_context = "CURRENT EPICS (JSON FORMAT):\n" + json.dumps(epics_data, indent=2)
                else:
                    epics_context = current_epics
            except:
                # If not valid JSON, use as-is (might be HTML content)
                epics_context = "CURRENT EPICS (HTML/TEXT FORMAT):\n" + current_epics
        
        # Prepare context for Epic Agent with focus on updating existing epics
        chat_context = f"""You are an Epic Agent specializing in refining and improving existing epics. Your primary role is to help users update, correct, and enhance their current epics.

{epics_context if epics_context else 'No epics currently defined. Please ask the user to create epics first through the main workflow.'}

CRITICAL INSTRUCTIONS - READ CAREFULLY:
- DEFAULT BEHAVIOR: Always focus on modifying/updating/correcting the EXISTING epics shown above
- NEVER create new epics unless the user EXPLICITLY uses clear phrases like:
  * "add new epic"
  * "create additional epic" 
  * "new epic"
  * "more epics"
  * "another epic"
  * "additional epic"
- When users ask questions, give feedback, or make requests, they want to CHANGE/UPDATE the existing epics, NOT add new ones
- Provide UPDATED versions of the existing epics that incorporate the user's feedback
- Keep the same epic IDs and structure when updating
- If no current epics exist, guide user to the main workflow first

COMMON USER REQUESTS (all should UPDATE existing epics):
- "Make the epics more detailed" → Update existing epics with more detail
- "Add acceptance criteria" → Update existing epics to include acceptance criteria
- "Change the priority" → Update existing epic priorities
- "The epic needs more clarity" → Update existing epic descriptions
- "Fix the requirements" → Update existing epic requirements
- "Improve the user stories" → Update existing epic user story breakdown
- "These need work" → Update existing epics based on feedback

RESPOND WITH: Updated versions of the EXISTING epics, not new ones.

"""
        if session['epic_chat_history']:
            chat_context += "Previous conversation:\n"
            for msg in session['epic_chat_history'][-10:]:  # Last 10 messages for context
                role = "User" if msg["role"] == "user" else "Assistant"
                chat_context += f"{role}: {msg['message']}\n"
            chat_context += f"\nLatest user question: {user_message}"
        else:
            chat_context += f"User request: {user_message}"
        
        # Call Epic Agent
        logger.info(f"Calling Epic Agent with message: {user_message[:100]}...")
        logger.info(f"Epic Agent context length: {len(chat_context)} characters")
        epic_response = ask_assistant_from_file_optimized("poc2_agent2_epic_generator", chat_context)
        
        logger.info(f"Epic Agent response length: {len(epic_response)} characters")
        logger.info(f"Epic Agent response preview: {epic_response[:200]}...")
        
        if epic_response.startswith("Error:"):
            logger.error(f"Epic Agent error: {epic_response}")
            return jsonify({
                "success": False,
                "error": "Failed to get response from Epic Agent"
            }), 500
        
        # Add assistant response to chat history
        session['epic_chat_history'].append({
            "role": "assistant",
            "message": epic_response,
            "timestamp": current_time.strftime("%Y-%m-%d %H:%M:%S")
        })
        
        # Update last activity time again
        session['epic_chat_last_activity'] = current_time.isoformat()
        
        # Save session changes
        session.modified = True
        
        logger.info(f"Epic Chat response generated successfully, history length: {len(session['epic_chat_history'])}")
        
        return jsonify({
            "success": True,
            "response": epic_response,
            "chat_history": session['epic_chat_history']
        })
        
    except Exception as e:
        logger.error(f"Error in Epic Chat: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/epic-chat-history", methods=["GET"])
def get_epic_chat_history():
    """Get the current Epic Chat history from session."""
    try:
        chat_history = session.get('epic_chat_history', [])
        return jsonify({
            "success": True,
            "chat_history": chat_history
        })
    except Exception as e:
        logger.error(f"Error getting Epic Chat history: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/epic-chat-clear", methods=["POST"])
def clear_epic_chat():
    """Clear the Epic Chat history from session."""
    try:
        if 'epic_chat_history' in session:
            del session['epic_chat_history']
        if 'epic_chat_last_activity' in session:
            del session['epic_chat_last_activity']
        session.modified = True
        logger.info("Epic Chat history and activity timestamp cleared")
        return jsonify({"success": True, "message": "Chat history cleared"})
    except Exception as e:
        logger.error(f"Error clearing Epic Chat history: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/user-story-chat", methods=["POST"])
def user_story_chat():
    """Handle User Story Chat interactions with system mapping functionality."""
    try:
        data = request.get_json()
        if not data:
            return jsonify({"success": False, "error": "No data provided"}), 400
        
        user_message = data.get("message", "").strip()
        if not user_message:
            return jsonify({"success": False, "error": "Message cannot be empty"}), 400
        
        # Initialize chat history in session if not exists or if expired
        current_time = datetime.now()
        chat_session_timeout = 30 * 60  # 30 minutes in seconds
        
        if 'user_story_chat_history' not in session:
            session['user_story_chat_history'] = []
            session['user_story_chat_last_activity'] = current_time.isoformat()
            logger.info("Initialized new User Story Chat session")
        else:
            # Check if session has expired (30 minutes of inactivity)
            last_activity_str = session.get('user_story_chat_last_activity')
            if last_activity_str:
                try:
                    last_activity = datetime.fromisoformat(last_activity_str)
                    time_diff = (current_time - last_activity).total_seconds()
                    
                    if time_diff > chat_session_timeout:
                        # Clear expired session
                        session['user_story_chat_history'] = []
                        logger.info(f"Cleared expired User Story Chat session (inactive for {time_diff/60:.1f} minutes)")
                except Exception as e:
                    logger.warning(f"Error parsing last activity time: {e}")
                    session['user_story_chat_history'] = []
        
        # Limit chat history to last 50 messages to prevent session bloat
        if len(session['user_story_chat_history']) > 50:
            session['user_story_chat_history'] = session['user_story_chat_history'][-40:]  # Keep last 40 messages
            logger.info("Trimmed User Story Chat history to prevent session bloat")
        
        # Update last activity time
        session['user_story_chat_last_activity'] = current_time.isoformat()
        
        # Add user message to chat history
        session['user_story_chat_history'].append({
            "role": "user",
            "message": user_message,
            "timestamp": current_time.strftime("%Y-%m-%d %H:%M:%S")
        })
        
        # Get system information from session (if available)
        system_info = session.get('system_info', '')
        
        # Prepare context for User Story Agent with system mapping
        chat_context = f"""You are a User Story Agent helping to create, refine, and map user stories to specific systems.

SYSTEM INFORMATION:
{system_info if system_info else 'No system information provided. User stories will not include system mapping.'}

INSTRUCTIONS:
- Create detailed user stories following best practices
- If system information is available, map each user story to the appropriate system based on the functionality described
- Include a "System Name" for each user story when system mapping is possible
- Format user stories with: Title, Description, Acceptance Criteria, and System Name (if applicable)
- Focus on creating clear, actionable user stories that deliver business value

"""
        if session['user_story_chat_history']:
            chat_context += "Previous conversation:\n"
            for msg in session['user_story_chat_history'][-10:]:  # Last 10 messages for context
                role = "User" if msg["role"] == "user" else "Assistant"
                chat_context += f"{role}: {msg['message']}\n"
            chat_context += f"\nLatest user question: {user_message}"
        else:
            chat_context += f"User request: {user_message}"
        
        # Call User Story Agent (using same agent as Epic for consistency)
        logger.info(f"Calling User Story Agent with message: {user_message[:100]}...")
        story_response = ask_assistant_from_file_optimized("poc2_agent3_basic_user_story", chat_context)
        
        if story_response.startswith("Error:"):
            logger.error(f"User Story Agent error: {story_response}")
            return jsonify({
                "success": False,
                "error": "Failed to get response from User Story Agent"
            }), 500
        
        # Add assistant response to chat history
        session['user_story_chat_history'].append({
            "role": "assistant",
            "message": story_response,
            "timestamp": current_time.strftime("%Y-%m-%d %H:%M:%S")
        })
        
        # Update last activity time again
        session['user_story_chat_last_activity'] = current_time.isoformat()
        
        # Save session changes
        session.modified = True
        
        logger.info(f"User Story Chat response generated successfully, history length: {len(session['user_story_chat_history'])}")
        
        return jsonify({
            "success": True,
            "response": story_response,
            "chat_history": session['user_story_chat_history'],
            "has_system_info": bool(system_info)
        })
        
    except Exception as e:
        logger.error(f"Error in User Story Chat: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/user-story-chat-history", methods=["GET"])
def get_user_story_chat_history():
    """Get the current User Story Chat history from session."""
    try:
        chat_history = session.get('user_story_chat_history', [])
        system_info = session.get('system_info', '')
        return jsonify({
            "success": True,
            "chat_history": chat_history,
            "has_system_info": bool(system_info)
        })
    except Exception as e:
        logger.error(f"Error getting User Story Chat history: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/user-story-chat-clear", methods=["POST"])
def clear_user_story_chat():
    """Clear the User Story Chat history from session."""
    try:
        if 'user_story_chat_history' in session:
            del session['user_story_chat_history']
        if 'user_story_chat_last_activity' in session:
            del session['user_story_chat_last_activity']
        session.modified = True
        logger.info("User Story Chat history and activity timestamp cleared")
        return jsonify({"success": True, "message": "Chat history cleared"})
    except Exception as e:
        logger.error(f"Error clearing User Story Chat history: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/upload-system-info", methods=["POST"])
def upload_system_info():
    """Upload and process system information file for user story mapping."""
    try:
        if 'file' not in request.files:
            return jsonify({"success": False, "error": "No file provided"}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({"success": False, "error": "No file selected"}), 400
        
        # Check file size (limit to 10MB)
        file.seek(0, 2)  # Seek to end
        file_size = file.tell()
        file.seek(0)  # Reset to beginning
        
        if file_size > 10 * 1024 * 1024:  # 10MB limit
            return jsonify({"success": False, "error": "File size too large (max 10MB)"}), 400
        
        # Read and process file content
        content = ""
        filename = file.filename.lower()
        
        try:
            if filename.endswith('.txt'):
                content = file.read().decode('utf-8')
            elif filename.endswith('.json'):
                json_data = json.load(file)
                content = json.dumps(json_data, indent=2)
            elif filename.endswith('.csv'):
                content = file.read().decode('utf-8')
            else:
                # Try to read as text for other formats
                content = file.read().decode('utf-8')
        except UnicodeDecodeError:
            return jsonify({"success": False, "error": "File encoding not supported. Please use UTF-8 encoded files."}), 400
        except json.JSONDecodeError:
            return jsonify({"success": False, "error": "Invalid JSON file format"}), 400
        
        if not content.strip():
            return jsonify({"success": False, "error": "File appears to be empty"}), 400
        
        # Store system information in session
        session['system_info'] = content
        session['system_info_filename'] = file.filename
        session['system_info_uploaded'] = datetime.now().isoformat()
        session.modified = True
        
        # Parse basic system information for preview
        lines = content.split('\n')[:10]  # First 10 lines for preview
        preview = '\n'.join(lines)
        if len(content.split('\n')) > 10:
            preview += "\n... (file continues)"
        
        logger.info(f"System information uploaded: {file.filename} ({file_size} bytes)")
        
        return jsonify({
            "success": True,
            "message": f"System information uploaded successfully: {file.filename}",
            "filename": file.filename,
            "size": file_size,
            "preview": preview[:500]  # Limit preview to 500 chars
        })
        
    except Exception as e:
        logger.error(f"Error uploading system info: {str(e)}")
        logger.error(f"Traceback: {traceback.format_exc()}")
        return jsonify({"success": False, "error": "Failed to upload system information"}), 500

@app.route("/get-system-info", methods=["GET"])
def get_system_info():
    """Get current system information status."""
    try:
        system_info = session.get('system_info', '')
        filename = session.get('system_info_filename', '')
        uploaded_time = session.get('system_info_uploaded', '')
        
        return jsonify({
            "success": True,
            "has_system_info": bool(system_info),
            "filename": filename,
            "uploaded_time": uploaded_time,
            "preview": system_info[:500] if system_info else ""
        })
    except Exception as e:
        logger.error(f"Error getting system info: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.route("/clear-system-info", methods=["POST"])
def clear_system_info():
    """Clear uploaded system information."""
    try:
        if 'system_info' in session:
            del session['system_info']
        if 'system_info_filename' in session:
            del session['system_info_filename']
        if 'system_info_uploaded' in session:
            del session['system_info_uploaded']
        session.modified = True
        
        logger.info("System information cleared from session")
        return jsonify({"success": True, "message": "System information cleared"})
    except Exception as e:
        logger.error(f"Error clearing system info: {str(e)}")
        return jsonify({"success": False, "error": "Internal server error"}), 500

@app.errorhandler(404)
def not_found(error):
    logger.warning(f"404 error for URL: {request.url}")
    return jsonify({"success": False, "error": "Endpoint not found"}), 404

@app.errorhandler(500)
def internal_error(error):
    logger.error(f"500 error: {str(error)}")
    return jsonify({"success": False, "error": "Internal server error"}), 500

@app.errorhandler(Exception)
def handle_exception(e):
    logger.error(f"Unhandled exception: {str(e)}")
    logger.error(f"Exception type: {type(e).__name__}")
    logger.error(f"Request URL: {request.url}")
    logger.error(f"Request method: {request.method}")
    logger.error(f"Traceback: {traceback.format_exc()}")
    return jsonify({"success": False, "error": f"Unexpected error: {str(e)}"}), 500

if __name__ == "__main__":
    logger.info("Starting Flask application...")
    logger.info("RAG-Enhanced PRD to Epic/User Story Generator")
    logger.info("=" * 60)
    logger.info("Features:")
    logger.info("- RAG-based document processing with ChromaDB")
    logger.info("- SentenceTransformers embeddings for semantic search")
    logger.info("- PRD Parser Agent replaced with intelligent RAG summaries")
    logger.info("- Single-agent workflow using only Epic Generator")
    logger.info("- Advanced caching and performance optimizations")
    logger.info("=" * 60)
    
    try:
        # Run Flask app in debug mode for development
        app.run(
            host="0.0.0.0",
            port=5000,
            debug=True,
            threaded=True,
            use_reloader=False  # Disable reloader to prevent double initialization
        )
    except Exception as e:
        logger.error(f"Failed to start Flask application: {str(e)}")
        raise
