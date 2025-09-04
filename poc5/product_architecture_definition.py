from flask import Flask, render_template, request, session, redirect, url_for, make_response, send_file, jsonify
from flask_session import Session
import os
import json
import csv
import re
from jinja2 import ChoiceLoader, FileSystemLoader
import time
from io import BytesIO
import subprocess
import tempfile
import shutil
import os
import base64
import urllib.request
import urllib.parse
try:
    from docx import Document
    try:
        from docx.shared import Inches
    except Exception:
        Inches = None
except Exception:
    Document = None

# Serve shared root-level static assets (../static) and templates (../templates)
ROOT_DIR = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
STATIC_DIR = os.path.join(ROOT_DIR, 'static')
TEMPLATES_DIR = os.path.join(ROOT_DIR, 'templates')

app = Flask(__name__, static_folder=STATIC_DIR, static_url_path='/static', template_folder=TEMPLATES_DIR)
app.secret_key = "your_secret_key"
app.config.update(
    SESSION_TYPE='filesystem',
    SESSION_PERMANENT=False,
    SESSION_FILE_DIR=os.path.join(os.path.dirname(__file__), '.flask_session'),
)
Session(app)

# Prefer POC5 templates, but fall back to root-level templates
POC5_TEMPLATES = os.path.join(os.path.dirname(__file__), 'templates')
app.jinja_loader = ChoiceLoader([
    FileSystemLoader(POC5_TEMPLATES),
    FileSystemLoader(TEMPLATES_DIR),
])

AGENT_PROMPTS = {
    "agent_5_1": "agents/agent_5_1.txt",
    "agent_5_2": "agents/agent_5_2.txt",
    "agent_5_3": "agents/agent_5_3.txt",
    "agent_5_4": "agents/agent_5_4.txt",
    "agent_5_5": "agents/agent_5_5.txt",
    "agent_mermaid": "agents/agent_mermaid.txt",
    "agent_5_6": "agents/agent_5_6.txt",
}

# Tab index mapping for chat/approval navigation (nav order):
# Default flow now (Context Doc merged into Tab 0):
# 0 (Context + Context Doc) → 1 (Initial) → 2 (System Interaction) → 5 (DB) → 6 (Service)
# → 3 (Decisions+Integration+Pros) → 4 (Docs)
TAB_NAV_ORDER = [0, 1, 2, 5, 6, 3, 4]
AGENT_BY_TAB = {
    0: 'agent_5_1',       # Context & Requirements (includes merged Context Doc)
    1: 'agent_mermaid',   # Initial Blueprint (diagram-focused)
    2: 'agent_mermaid',   # System Interaction (sequenceDiagram)
    3: 'agent_5_2',       # Combined Decisions/Integration/Pros (text)
    4: 'agent_5_5',       # Documentation & Review
    5: 'agent_mermaid',   # Database Solution Design (ER/data schemas)
    6: 'agent_mermaid',   # Service Solution Design (service decomposition/interfaces)
}

# Allowed fields per tab for chat-driven updates
TAB_ALLOWED_FIELDS: dict[int, list[str]] = {
    0: [
        'business_goals','legacy_system','constraints','additional_requirements',
        'architecture_style','desired_architecture_style','cloud_provider','region_strategy','deployment_model',
        'database_family','messaging_backbone','identity_provider','api_style','caching_strategy','observability_stack',
        'data_residency','compliance_standards',
        'availability_sla','rpo_rto','performance_targets','throughput','peak_concurrency','data_volume_retention',
        'dr_strategy','encryption_at_rest','encryption_in_transit','kms','api_gateway_type','service_mesh','edge_controls',
        'secrets_manager','iac_tool','ci_cd','environments','release_strategy','deployment_topology','tenancy_model',
        'observability_requirements','security_posture','cost_constraints','capacity_estimates','workloads',
        'high_level_decision_points','one_way_door_decisions','other_relevant_questions',
        # Merged Context Doc fields (formerly Tab 6)
        'architecture_context','assumptions','stakeholders','in_scope','out_of_scope',
        'migration_strategy','risks','open_questions'
    ],
    1: ['architecture_diagram','major_components','interface_definitions','reusable_patterns'],
    2: [
        # System Interaction
        'system_interaction_diagram','system_interaction_notes'
    ],
    3: [
        # Combined: Decisions + Communication & Integration + Pros & Cons
        'architectural_decisions','rationale_tradeoffs','blueprint_references',
        'high_level_decision_points','one_way_door_decisions','other_relevant_questions',
        'communication_protocols','data_flow_diagrams','integration_points',
        'swot_analysis','risks_mitigation','risks'
    ],
    4: ['compiled_document','stakeholder_checklist'],
    5: [
        # Database Solution Design
        'data_model_diagram','data_schemas','database_family'
    ],
    6: [
        # Service Solution Design
        'service_decomposition_diagram','major_components','interface_definitions','reusable_patterns','api_style','caching_strategy'
    ],
}


def _parse_updates_from_text(text: str) -> dict | None:
    """Extract a JSON updates object from agent text. Expect formats like:
    ```json
    {"set": {"field": "value"}, "append": {"field":"..."}}
    ```
    or plain {...}
    """
    if not text:
        return None
    import re as _re
    candidates = []
    m = _re.search(r"```json\s*(\{[\s\S]*?\})\s*```", text, _re.I)
    if m:
        candidates.append(m.group(1))
    # fallback: first JSON-looking block
    if not candidates:
        m2 = _re.search(r"(\{[\s\S]*\})", text)
        if m2:
            candidates.append(m2.group(1))
    for c in candidates:
        try:
            obj = json.loads(c)
            if isinstance(obj, dict):
                return obj
        except Exception:
            continue
    return None


_FIELD_SYNONYMS = {
    'cloud provider': 'cloud_provider', 'cloud': 'cloud_provider',
    'region': 'region_strategy', 'region strategy': 'region_strategy',
    'deployment': 'deployment_model', 'deployment model': 'deployment_model',
    'database': 'database_family', 'db': 'database_family', 'database family': 'database_family',
    'messaging': 'messaging_backbone', 'messaging backbone': 'messaging_backbone',
    'identity': 'identity_provider', 'idp': 'identity_provider',
    'api': 'api_style', 'api style': 'api_style',
    'cache': 'caching_strategy', 'caching': 'caching_strategy',
    'observability': 'observability_stack',
    'data residency': 'data_residency', 'compliance': 'compliance_standards',
    'sla': 'availability_sla', 'availability': 'availability_sla',
    'rpo': 'rpo_rto', 'rto': 'rpo_rto',
    'latency': 'performance_targets', 'p95': 'performance_targets',
    'throughput': 'throughput', 'qps': 'throughput', 'tps': 'throughput',
    'concurrency': 'peak_concurrency',
    'retention': 'data_volume_retention',
    'dr': 'dr_strategy', 'disaster recovery': 'dr_strategy',
    'encryption at rest': 'encryption_at_rest', 'at rest': 'encryption_at_rest',
    'encryption in transit': 'encryption_in_transit', 'in transit': 'encryption_in_transit',
    'kms': 'kms', 'key management': 'kms',
    'gateway': 'api_gateway_type', 'api gateway': 'api_gateway_type',
    'mesh': 'service_mesh', 'service mesh': 'service_mesh',
    'edge': 'edge_controls', 'waf': 'edge_controls', 'cdn': 'edge_controls',
    'secrets': 'secrets_manager', 'secrets manager': 'secrets_manager',
    'iac': 'iac_tool', 'infra as code': 'iac_tool',
    'cicd': 'ci_cd', 'ci/cd': 'ci_cd',
    'envs': 'environments', 'environments': 'environments',
    'release': 'release_strategy', 'release strategy': 'release_strategy',
    'topology': 'deployment_topology',
    'tenancy': 'tenancy_model',
    'observability requirements': 'observability_requirements',
    'security': 'security_posture', 'security posture': 'security_posture',
    'cost': 'cost_constraints', 'capacity': 'capacity_estimates',
    'assumptions': 'assumptions', 'stakeholders': 'stakeholders',
    'in scope': 'in_scope', 'out of scope': 'out_of_scope',
    'risks': 'risks', 'mitigations': 'risks_mitigation', 'pros and cons': 'swot_analysis', 'pros & cons': 'swot_analysis',
}


def _heuristic_updates_from_user(idx: int, msg: str) -> dict | None:
    """Best-effort parse of user commands like 'set cloud provider to Azure' or 'append to major_components: - X'."""
    if not msg:
        return None
    low = msg.strip().lower()
    # set <field> to <value>
    import re as _re
    m = _re.search(r"(?:set|change|update)\s+([a-z0-9_ ]+)\s+to\s+(.+)$", low)
    if m:
        raw_field, val = m.group(1).strip(), m.group(2).strip()
        key = _FIELD_SYNONYMS.get(raw_field, raw_field.replace(' ', '_'))
        return {'set': {key: val}}
    # append to <field>: <value>
    m2 = _re.search(r"append\s+to\s+([a-z0-9_ ]+)\s*:\s*(.+)$", low)
    if m2:
        raw_field, val = m2.group(1).strip(), m2.group(2).strip()
        key = _FIELD_SYNONYMS.get(raw_field, raw_field.replace(' ', '_'))
        return {'append': {key: val}}
    return None


def _apply_agent_updates(idx: int, updates: dict):
    if not updates:
        return
    allowed = set(TAB_ALLOWED_FIELDS.get(idx, []))
    # also allow some shared fields across tabs
    shared = {
        'architecture_style','desired_architecture_style','cloud_provider','region_strategy','deployment_model',
        'api_style','database_family','messaging_backbone','identity_provider','caching_strategy','observability_stack',
        'data_residency','compliance_standards','availability_sla','rpo_rto','performance_targets','throughput',
        'peak_concurrency','data_volume_retention','dr_strategy','encryption_at_rest','encryption_in_transit','kms',
        'api_gateway_type','service_mesh','edge_controls','secrets_manager','iac_tool','ci_cd','environments',
        'release_strategy','deployment_topology','tenancy_model','observability_requirements','security_posture',
        'cost_constraints','capacity_estimates','workloads'
    }
    allowed |= shared
    sets = updates.get('set') or {}
    appends = updates.get('append') or {}
    # Track diagram fields for history
    diagram_fields = {'architecture_diagram','system_interaction_diagram','data_model_diagram','service_decomposition_diagram'}
    # apply sets
    for k, v in sets.items():
        if k in allowed:
            if k in diagram_fields:
                _push_diagram_history(k)
            session[k] = v
    # apply appends
    for k, v in appends.items():
        if k in allowed:
            existing = session.get(k, '')
            session[k] = (existing + ('\n' if existing and not existing.endswith('\n') else '') + str(v)).strip()
    # regenerate artifacts when useful
    try:
        if idx == 1 and 'architecture_diagram' not in sets:
            # Regenerate strictly from PRD to keep diagrams PRD-driven
            session['architecture_diagram'] = generate_blueprint_from_session(strict_prd_only=True)
            # Refresh write-up when diagram context changes
            try:
                session['decomposition_writeup'] = _build_decomposition_writeup_text(session.get('prd_text',''))
            except Exception:
                pass
        if idx == 0 and 'architecture_context' not in sets:
            session['architecture_context'] = generate_architecture_context()
    except Exception:
        pass


def _reset_form_if_no_uploads_on_get():
    """If landing fresh (GET) without uploaded PRD/Mapping, keep Tab 1 unlocked and clear auto content."""
    if request.method != 'GET':
        return
    if session.get('prd_file_name') or session.get('system_mapping_name'):
        return
    # ensure unlocked and clear auto-generated fields
    session['lock_tab1'] = ''
    keys_to_clear = [
        # core context
        'business_goals','legacy_system','constraints','additional_requirements',
        'architecture_style','desired_architecture_style','cloud_provider','region_strategy','deployment_model',
        'database_family','messaging_backbone','identity_provider','api_style','caching_strategy','observability_stack',
        'data_residency','compliance_standards',
        # derived artifacts
        'high_level_decision_points','one_way_door_decisions','other_relevant_questions',
        'architecture_diagram','major_components','interface_definitions','data_schemas','reusable_patterns',
        'architectural_decisions','rationale_tradeoffs','blueprint_references',
        'communication_protocols','data_flow_diagrams','integration_points',
        'swot_analysis','risks_mitigation','compiled_document','stakeholder_checklist',
        'architecture_context',
    ]
    for k in keys_to_clear:
        session.pop(k, None)


def _read_agent_prompt(agent_key: str) -> str:
    """Read agent prompt from poc5/agents first (authoritative), fallback to root/agents.
    All POC5 prompts should live under poc5/agents for this app.
    """
    rel_path = AGENT_PROMPTS.get(agent_key)
    if not rel_path:
        return ""
    poc5_path = os.path.join(os.path.dirname(__file__), rel_path)
    root_path = os.path.join(ROOT_DIR, rel_path)
    for p in (poc5_path, root_path):
        try:
            with open(p, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception:
            continue
    return ""


def sanitize_node_label(text: str) -> str:
    if not text:
        return ""
    t = str(text)
    t = t.replace('[', '(').replace(']', ')')
    t = t.replace('(', '﹙').replace(')', '﹚')
    return t if len(t) <= 80 else (t[:77] + '...')


def _safe_set_session(key: str, value: str | None):
    if value and not session.get(key):
        session[key] = value


def _add_text_block(doc, title: str, text: str | None):
    """Add a titled text block to a docx document, preserving bullets when possible."""
    if not doc or not title:
        return
    doc.add_heading(title, level=2)
    if not text:
        doc.add_paragraph('')
        return
    lines = [l.rstrip() for l in str(text).splitlines()]
    bullet_lines = [l for l in lines if l.strip().startswith(('-', '*'))]
    # If most lines look like bullets, add as bullet list
    if bullet_lines and len(bullet_lines) >= max(1, int(0.6 * len(lines))):
        for l in bullet_lines:
            doc.add_paragraph(l.lstrip('-* ').strip(), style='List Bullet')
        return
    # Otherwise add paragraphs preserving blank lines as breaks
    para = []
    for l in lines:
        if l.strip() == '':
            if para:
                doc.add_paragraph('\n'.join(para))
                para = []
        else:
            para.append(l)
    if para:
        doc.add_paragraph('\n'.join(para))


def _build_docx_from_session() -> BytesIO | None:
    if Document is None:
        return None
    doc = Document()
    doc.add_heading('Architecture Workbench Export', level=1)
    subtitle = []
    if session.get('prd_file_name'):
        subtitle.append(f"PRD: {session.get('prd_file_name')}")
    if session.get('system_mapping_name'):
        subtitle.append(f"Mapping: {session.get('system_mapping_name')}")
    if subtitle:
        doc.add_paragraph(' | '.join(subtitle))

    # 1. Context & Requirements
    doc.add_heading('1. Context & Requirements', level=1)
    _add_text_block(doc, 'Business Goals', session.get('business_goals'))
    _add_text_block(doc, 'Legacy System', session.get('legacy_system'))
    _add_text_block(doc, 'Constraints', session.get('constraints'))
    _add_text_block(doc, 'Additional Requirements', session.get('additional_requirements'))
    # Guardrails (compact key:value)
    doc.add_heading('Guardrails', level=2)
    guardrails = [
        ('Architecture Style', session.get('architecture_style')),
        ('Desired Architecture Style', session.get('desired_architecture_style')),
        ('Cloud Provider', session.get('cloud_provider')),
        ('Region Strategy', session.get('region_strategy')),
        ('Deployment Model', session.get('deployment_model')),
        ('Database Family', session.get('database_family')),
        ('Messaging Backbone', session.get('messaging_backbone')),
        ('Identity Provider', session.get('identity_provider')),
        ('API Style', session.get('api_style')),
        ('Caching Strategy', session.get('caching_strategy')),
        ('Observability Stack', session.get('observability_stack')),
        ('Data Residency', session.get('data_residency')),
        ('Compliance Standards', session.get('compliance_standards')),
        ('Availability/SLA', session.get('availability_sla')),
        ('RPO/RTO', session.get('rpo_rto')),
        ('Performance Targets', session.get('performance_targets')),
        ('Throughput', session.get('throughput')),
        ('Peak Concurrency', session.get('peak_concurrency')),
        ('Data Volume & Retention', session.get('data_volume_retention')),
        ('DR Strategy', session.get('dr_strategy')),
        ('Encryption at Rest', session.get('encryption_at_rest')),
        ('Encryption in Transit', session.get('encryption_in_transit')),
        ('KMS', session.get('kms')),
        ('API Gateway Type', session.get('api_gateway_type')),
        ('Service Mesh', session.get('service_mesh')),
        ('Edge Controls', session.get('edge_controls')),
        ('Secrets Manager', session.get('secrets_manager')),
        ('IaC Tool', session.get('iac_tool')),
        ('CI/CD', session.get('ci_cd')),
        ('Environments', session.get('environments')),
        ('Release Strategy', session.get('release_strategy')),
        ('Deployment Topology', session.get('deployment_topology')),
        ('Tenancy Model', session.get('tenancy_model')),
        ('Observability Requirements', session.get('observability_requirements')),
        ('Security Posture', session.get('security_posture')),
        ('Cost Constraints', session.get('cost_constraints')),
        ('Capacity Estimates', session.get('capacity_estimates')),
        ('Workloads', session.get('workloads')),
    ]
    for k, v in guardrails:
        if v:
            doc.add_paragraph(f"{k}: {v}")

    # 2. Context Doc
    doc.add_heading('2. Context Doc', level=1)
    _add_text_block(doc, 'Architecture Context', session.get('architecture_context'))

    # 3. Initial Blueprint
    doc.add_heading('3. Initial Blueprint', level=1)
    # PRD Decomposition write-up goes above images
    if session.get('decomposition_writeup'):
        _add_text_block(doc, 'PRD Decomposition Summary', session.get('decomposition_writeup'))
    # Export the exact diagrams currently shown in the UI previews.
    # If any are empty, fall back to PRD-only generators.
    arch_text = session.get('architecture_diagram') or generate_blueprint_from_session(strict_prd_only=True)
    sys_text = session.get('system_interaction_diagram') or generate_system_interaction_diagram(strict_prd_only=True)
    data_text = session.get('data_model_diagram') or generate_data_model_diagram(strict_prd_only=True)
    svc_text = session.get('service_decomposition_diagram') or generate_service_decomposition_diagram(strict_prd_only=True)
    _add_mermaid_block(doc, 'Architecture Diagram', arch_text)
    _add_mermaid_block(doc, 'System Interaction Diagram', sys_text)
    _add_mermaid_block(doc, 'Data Model / Database Design', data_text)
    _add_mermaid_block(doc, 'Service Decomposition', svc_text)
    _add_text_block(doc, 'Major Components', session.get('major_components'))
    _add_text_block(doc, 'Interface Definitions', session.get('interface_definitions'))
    _add_text_block(doc, 'Data Schemas', session.get('data_schemas'))
    _add_text_block(doc, 'Reusable Patterns', session.get('reusable_patterns'))

    # 4. Key Decisions
    doc.add_heading('4. Key Decisions', level=1)
    _add_text_block(doc, 'High-Level Architecture Decision Points', session.get('high_level_decision_points'))
    _add_text_block(doc, 'One-way Door Decisions', session.get('one_way_door_decisions'))
    _add_text_block(doc, 'Other Relevant Questions', session.get('other_relevant_questions'))
    _add_text_block(doc, 'Architectural Decisions', session.get('architectural_decisions'))
    _add_text_block(doc, 'Rationale & Trade-offs', session.get('rationale_tradeoffs'))
    _add_text_block(doc, 'Blueprint References', session.get('blueprint_references'))

    # 5. Communication & Integration
    doc.add_heading('5. Communication & Integration', level=1)
    _add_text_block(doc, 'Communication Protocols', session.get('communication_protocols'))
    _add_text_block(doc, 'Data Flow Diagrams', session.get('data_flow_diagrams'))
    _add_text_block(doc, 'Integration Points', session.get('integration_points'))

    # 6. Pros & Cons / Risks
    doc.add_heading('6. Pros & Cons / Risks', level=1)
    _add_text_block(doc, 'SWOT / Pros & Cons', session.get('swot_analysis') or session.get('pros_cons'))
    _add_text_block(doc, 'Risks & Mitigation', session.get('risks_mitigation'))
    _add_text_block(doc, 'Risks', session.get('risks'))

    # 7. Documentation & Review
    doc.add_heading('7. Documentation & Review', level=1)
    _add_text_block(doc, 'Compiled Document', session.get('compiled_document'))
    _add_text_block(doc, 'Stakeholder Checklist', session.get('stakeholder_checklist'))

    # 8. Planning & Open Items
    doc.add_heading('8. Planning & Open Items', level=1)
    _add_text_block(doc, 'Assumptions', session.get('assumptions'))
    _add_text_block(doc, 'Stakeholders', session.get('stakeholders'))
    _add_text_block(doc, 'In Scope', session.get('in_scope'))
    _add_text_block(doc, 'Out of Scope', session.get('out_of_scope'))
    _add_text_block(doc, 'Open Questions', session.get('open_questions'))

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio


def _extract_features(prd_text: str) -> list[str]:
    """Extract feature/capability names from PRD text using robust heuristics.
    Detects:
    - Headings: Features/Capabilities/Functional Requirements/Requirements/Use Cases/Epics/User Stories/Scope
    - Bullets under those headings: -, *, •, [ ] checkboxes, 1. / 1) enumerations
    - Inline lines: Feature:/Capability:/Epic:/User Story:/Use Case:
    - User story lines: "As a ..., I want to ..." (captures the verb phrase)
    Fallback: if no explicit features found, use top domain terms as features.
    """
    if not prd_text:
        return []
    features: list[str] = []
    lines = prd_text.splitlines()
    in_section = False
    section_rx = re.compile(r"^(?:#{1,6}\s*)?(features?|capabilities|functional requirements|requirements|use cases|epics|user stories|scope)\b[:]?", re.I)
    bullet_rx = re.compile(r"^\s*(?:[-*•]\s+|\d+[\.)]\s+|\[\s*[xX ]\s*\]\s+)(.+)")
    inline_label_rx = re.compile(r"^(feature|capability|epic|user\s*story|use\s*case)\s*[:\-]\s*(.+)$", re.I)
    user_story_rx = re.compile(r"^\s*as\s+an?\b.*?\bi\s+want\s+to\s+(.+)$", re.I)
    # Iterate with lookahead for heading boundaries
    for idx, raw in enumerate(lines):
        line = raw.strip()
        low = line.lower()
        # Enter/exit sections by headings
        if section_rx.match(line):
            in_section = True
            continue
        # Exit on next markdown heading
        if in_section and re.match(r"^\s*#{1,6}\s+.+$", line):
            in_section = False
        # Inline labels anywhere
        m_inline = inline_label_rx.match(line)
        if m_inline:
            txt = m_inline.group(2).strip()
            if txt and txt.lower() not in ('n/a','none'):
                features.append(txt)
            continue
        # Bullet within section
        if in_section:
            m_b = bullet_rx.match(raw)
            if m_b:
                txt = m_b.group(1).strip()
                if txt and txt.lower() not in ('n/a','none'):
                    features.append(txt)
                continue
        # User story pattern
        m_us = user_story_rx.match(line)
        if m_us:
            want = m_us.group(1).strip()
            if want:
                # Trim trailing purpose clause like 'so that I can ...'
                want = re.split(r"\s+so\s+that\s+i\s+can\s+", want, flags=re.I)[0].strip()
                # Keep a concise capability phrase (<= 120 chars)
                if len(want) > 120:
                    want = want[:117] + '...'
                features.append(want)
            continue
    # Clean up: strip trailing punctuation and dedupe
    cleaned: list[str] = []
    seen = set()
    for f in features:
        name = re.sub(r"\s*[-–—:.;,]+\s*$", "", f).strip()
        if not name:
            continue
        if len(name) > 160:
            name = name[:157] + '...'
        if name not in seen:
            seen.add(name)
            cleaned.append(name)
    # Fallback: if empty, use up to 10 domain terms as coarse features
    if not cleaned:
        terms = list(_extract_domain_terms(prd_text))
        # Prefer a stable order: alphabetical limited to 10
        for t in sorted(terms)[:10]:
            cleaned.append(t)
    return cleaned[:15]


def _fold_features_into_terms(terms: set[str], features: list[str]) -> set[str]:
    """Augment domain terms with terms implied by feature names."""
    if not features:
        return terms
    add = set()
    mapping = {
        'checkout': 'Checkout', 'cart': 'Cart', 'payment': 'Payment', 'order': 'Order', 'inventory': 'Inventory',
        'shipping': 'Shipment', 'ship': 'Shipment', 'invoice': 'Invoice', 'catalog': 'Catalog', 'product': 'Product',
    'pricing': 'Pricing', 'discount': 'Discount', 'fraud': 'Fraud', 'review': 'Review', 'customer': 'Customer',
    'account': 'Account', 'notification': 'Notification', 'support': 'Support', 'analytics': 'Analytics',
    # Banking/Financial
    'transaction': 'Transaction', 'card': 'Card', 'loan': 'Loan', 'mortgage': 'Mortgage', 'kyc': 'KYC',
    'aml': 'AML', 'ledger': 'Ledger', 'statement': 'Statement', 'fx': 'FX', 'treasury': 'Treasury',
    'credit': 'Credit', 'underwriting': 'Underwriting', 'score': 'Scoring', 'atm': 'ATM', 'branch': 'Branch',
    'compliance': 'Compliance',
    # Healthcare
    'patient': 'Patient', 'provider': 'Provider', 'appointment': 'Appointment', 'encounter': 'Encounter',
    'ehr': 'EHR', 'emr': 'EHR', 'fhir': 'FHIR', 'lab': 'Lab', 'radiology': 'Radiology', 'imaging': 'Radiology',
    'prescription': 'Prescription', 'rx': 'Prescription', 'pharmacy': 'Pharmacy', 'claim': 'Claim', 'payer': 'Payer',
    'eligibility': 'Eligibility', 'authorization': 'Authorization', 'diagnosis': 'Diagnosis', 'procedure': 'Procedure',
    'care plan': 'CarePlan', 'careplan': 'CarePlan', 'consent': 'Consent'
    }
    for f in features:
        low = f.lower()
        for key, canonical in mapping.items():
            if key in low:
                add.add(canonical)
        # Additionally, single token title-cased words may be service/entity candidates
        tokens = re.findall(r"[A-Z][A-Za-z0-9]+", f)
        for t in tokens:
            if len(t) > 2:
                add.add(t)
    return set(terms) | add


def _services_from_features(features: list[str]) -> list[str]:
    """Infer sensible service names from feature text.
    Strategy:
    - Extract canonical domain terms from each feature using existing PRD term extractor.
    - Map those terms to service names via _derive_services_from_terms.
    - Avoid naive conversion of long user-story sentences into a single "... Service".
    - If nothing is detected, return an empty list so callers can fall back to PRD-wide term derivation.
    """
    if not features:
        return []
    detected_terms: set[str] = set()
    for f in features:
        try:
            detected_terms |= _extract_domain_terms(f or '')
        except Exception:
            continue
    services = _derive_services_from_terms(detected_terms)
    # As a conservative fallback, do NOT create a service from a long sentence; leave empty to allow caller fallback
    return services or []


def _strip_mermaid_fence(text: str) -> str:
    if not text:
        return ''
    t = text.strip()
    m = re.search(r"```\s*mermaid\s*\n([\s\S]*?)```", t, re.I)
    if m:
        return m.group(1).strip()
    m2 = re.search(r"```\s*\n([\s\S]*?)```", t, re.I)
    if m2:
        return m2.group(1).strip()
    return t


def _render_mermaid_png(mermaid_text: str) -> str | None:
    """Render Mermaid text to a PNG file using mermaid-cli (mmdc) if available. Returns file path or None."""
    if not mermaid_text:
        return None
    mermaid_text = _strip_mermaid_fence(mermaid_text)
    mmdc = shutil.which('mmdc')
    if not mmdc:
        # Fallback 1: mermaid.ink (POST)
        try:
            payload = mermaid_text.encode('utf-8')
            req = urllib.request.Request(
                'https://mermaid.ink/img/',
                data=payload,
                headers={'Content-Type': 'text/plain; charset=utf-8'},
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=15) as resp:
                img = resp.read()
            fd, final_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            with open(final_path, 'wb') as f:
                f.write(img)
            return final_path
        except Exception:
            # Fallback 2: Kroki Mermaid
            try:
                payload = json.dumps({'diagram_source': mermaid_text}).encode('utf-8')
                req = urllib.request.Request(
                    'https://kroki.io/mermaid/png',
                    data=payload,
                    headers={'Content-Type': 'application/json'},
                    method='POST'
                )
                with urllib.request.urlopen(req, timeout=15) as resp:
                    img = resp.read()
                fd, final_path = tempfile.mkstemp(suffix='.png')
                os.close(fd)
                with open(final_path, 'wb') as f:
                    f.write(img)
                return final_path
            except Exception:
                return None
    try:
        with tempfile.TemporaryDirectory() as td:
            in_file = os.path.join(td, 'diagram.mmd')
            out_file = os.path.join(td, 'diagram.png')
            with open(in_file, 'w', encoding='utf-8') as f:
                f.write(mermaid_text)
            # Run mermaid-cli
            subprocess.run([mmdc, '-i', in_file, '-o', out_file, '-b', 'transparent'], check=True, capture_output=True)
            # Persist to another temp file because td will be cleaned
            fd, final_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            shutil.copyfile(out_file, final_path)
            return final_path
    except Exception:
        # Fallback path if mmdc exists but failed: try mermaid.ink then Kroki
        try:
            payload = mermaid_text.encode('utf-8')
            req = urllib.request.Request(
                'https://mermaid.ink/img/',
                data=payload,
                headers={'Content-Type': 'text/plain; charset=utf-8'},
                method='POST'
            )
            with urllib.request.urlopen(req, timeout=15) as resp:
                img = resp.read()
            fd, final_path = tempfile.mkstemp(suffix='.png')
            os.close(fd)
            with open(final_path, 'wb') as f:
                f.write(img)
            return final_path
        except Exception:
            try:
                payload = json.dumps({'diagram_source': mermaid_text}).encode('utf-8')
                req = urllib.request.Request(
                    'https://kroki.io/mermaid/png',
                    data=payload,
                    headers={'Content-Type': 'application/json'},
                    method='POST'
                )
                with urllib.request.urlopen(req, timeout=15) as resp:
                    img = resp.read()
                fd, final_path = tempfile.mkstemp(suffix='.png')
                os.close(fd)
                with open(final_path, 'wb') as f:
                    f.write(img)
                return final_path
            except Exception:
                return None


def _add_mermaid_block(doc, title: str, mermaid_text: str | None):
    if not doc or not title:
        return
    # Heading
    doc.add_heading(title, level=2)
    if not mermaid_text:
        doc.add_paragraph('')
        return
    # Try to render an image
    img_path = _render_mermaid_png(mermaid_text)
    if img_path and Inches is not None:
        try:
            doc.add_picture(img_path, width=Inches(6))
        except Exception:
            pass
        try:
            os.remove(img_path)
        except Exception:
            pass
    # Always include the source as text
    _add_text_block(doc, 'Mermaid Source', mermaid_text)


def _detect_first(text: str, patterns: list[tuple[str, str]]):
    if not text:
        return None
    low = text.lower()
    for rx, canonical in patterns:
        try:
            if re.search(rx, low, re.I):
                return canonical
        except re.error:
            continue
    return None


def _extract_simple_fields_from_prd(prd_text: str) -> dict:
    fields = {}
    fields['architecture_style'] = _detect_first(prd_text, [
        (r"microservice|micro-service|microservices", "Microservices"),
        (r"event[- ]driven|eda\b", "Event-Driven"),
        (r"serverless|functions?\s+as\s+a\s+service|lambda", "Serverless"),
        (r"monolith(ic)?", "Monolith"),
        (r"soa\b|service[- ]oriented", "SOA"),
        (r"hexagonal|ports\s+and\s+adapters", "Hexagonal"),
        (r"cqrs|event\s+sourcing", "CQRS / Event Sourcing"),
    ])
    fields['cloud_provider'] = _detect_first(prd_text, [
        (r"\baws\b|amazon web services", "AWS"),
        (r"\bazure\b|microsoft azure", "Azure"),
        (r"gcp|google cloud", "GCP"),
        (r"oracle cloud|oci\b", "Oracle Cloud"),
        (r"on[- ]prem|on[- ]premises|data center", "On-Premises"),
    ])
    fields['region_strategy'] = _detect_first(prd_text, [
        (r"multi[- ]region|active[- ]active", "Multi-Region / Active-Active"),
        (r"dr\b|disaster recovery|active[- ]passive", "Active-Passive with DR"),
        (r"single[- ]region", "Single Region"),
    ])
    fields['deployment_model'] = _detect_first(prd_text, [
        (r"kubernetes|k8s|eks|aks|gke", "Kubernetes"),
        (r"ecs|fargate", "ECS/Fargate"),
        (r"vm\b|virtual machine|ec2", "VMs"),
        (r"serverless|lambda|functions?", "Serverless"),
    ])
    fields['database_family'] = _detect_first(prd_text, [
        (r"postgres|mysql|sql server|oracle\b|mariadb", "Relational"),
        (r"mongo|dynamo|cassandra|cosmos|couch", "NoSQL"),
        (r"warehouse|snowflake|redshift|bigquery", "Warehouse"),
    ])
    fields['messaging_backbone'] = _detect_first(prd_text, [
        (r"kafka", "Kafka"),
        (r"rabbitmq|amqp", "RabbitMQ"),
        (r"sqs|sns|eventbridge", "AWS Messaging"),
        (r"pub/sub|pubsub", "Pub/Sub"),
    ])
    fields['identity_provider'] = _detect_first(prd_text, [
        (r"okta", "Okta"),
        (r"auth0", "Auth0"),
        (r"azure\s+ad|entra id", "Azure AD"),
        (r"cognito", "Amazon Cognito"),
        (r"keycloak", "Keycloak"),
    ])
    fields['api_style'] = _detect_first(prd_text, [
        (r"\brest\b|http api|json api", "REST"),
        (r"graphql", "GraphQL"),
        (r"grpc", "gRPC"),
        (r"soap|wsdl|xml[- ]rpc", "SOAP/XML"),
    ])
    fields['caching_strategy'] = _detect_first(prd_text, [
        (r"redis|memcached|cache", "Cache-Aside / Redis"),
        (r"cdn|cloudfront|akamai", "CDN"),
    ])
    fields['observability_stack'] = _detect_first(prd_text, [
        (r"prometheus|grafana|otel|open[- ]telemetry", "Prometheus/Grafana + OpenTelemetry"),
        (r"elk|elastic(searc)?h|logstash|kibana", "ELK Stack"),
        (r"datadog|new relic|splunk", "APM (Datadog/New Relic/Splunk)"),
    ])
    fields['data_residency'] = _detect_first(prd_text, [
        (r"gdpr|eu[- ]only|within eu", "EU / GDPR"),
        (r"india|in[- ]country", "India"),
        (r"us[- ]only|within us", "US"),
    ])
    comps = []
    for name, rx in [
        ("GDPR", r"gdpr"), ("HIPAA", r"hipaa"), ("PCI DSS", r"pci(\s*dss)?"),
        ("SOC 2", r"soc\s*2"), ("ISO 27001", r"iso\s*27001")
    ]:
        if re.search(rx, prd_text or '', re.I):
            comps.append(name)
    if comps:
        fields['compliance_standards'] = ', '.join(sorted(set(comps)))
    m = re.search(r"(?:non[- ]functional|nfr|requirements?)[:\-]?\s*(.+)\n", prd_text or '', re.I)
    if m:
        fields['additional_requirements'] = m.group(1)[:200]
    # Normalize values to match HTML select options
    # database_family options: PostgreSQL, MySQL, SQL Server, MongoDB, DynamoDB
    df = fields.get('database_family')
    if df:
        if df == 'Relational':
            fields['database_family'] = 'PostgreSQL'
        elif df == 'NoSQL':
            # Prefer MongoDB as a generic NoSQL placeholder
            fields['database_family'] = 'MongoDB'
        elif df == 'Warehouse':
            # No warehouse option in UI; default to PostgreSQL
            fields['database_family'] = 'PostgreSQL'
    # caching_strategy options: Redis, Memcached, None
    cs = fields.get('caching_strategy')
    if cs:
        low = cs.lower()
        if 'redis' in low or 'cache' in low:
            fields['caching_strategy'] = 'Redis'
        elif 'memcache' in low:
            fields['caching_strategy'] = 'Memcached'
        elif 'cdn' in low:
            # Not represented; leave unset so defaults can fill
            fields.pop('caching_strategy', None)
    # observability_stack options: ELK, Grafana/Loki/Tempo, Datadog
    obs = fields.get('observability_stack')
    if obs:
        low = obs.lower()
        if 'elk' in low or 'elastic' in low or 'logstash' in low or 'kibana' in low:
            fields['observability_stack'] = 'ELK'
        elif 'datadog' in low or 'new relic' in low or 'splunk' in low:
            fields['observability_stack'] = 'Datadog'
        elif 'prometheus' in low or 'grafana' in low or 'otel' in low or 'open' in low:
            fields['observability_stack'] = 'Grafana/Loki/Tempo'
    return {k: v for k, v in fields.items() if v}


def _extract_core_sections_from_prd(prd_text: str) -> dict:
    """Extract core contextual sections from the PRD without truncating multi-line content.
    We look for heading-like lines (markdown, numbered, or colon-terminated) and capture
    subsequent paragraphs and bullet lists until the next section boundary.
    """
    text = prd_text or ''
    sections = {'business_goals': '', 'constraints': '', 'legacy_system': ''}

    # Match common heading styles with optional numbering / markdown / trailing colon
    def hdr_rx(terms: str) -> re.Pattern:
        return re.compile(
            rf"^\s*(?:#{1,6}\s*)?(?:\d+[\.)]\s*)?(?:{terms})\s*:?\s*$",
            re.IGNORECASE,
        )

    patterns = {
        'business_goals': hdr_rx(r"business\s+goals?|objectives?|goals?"),
        'constraints': hdr_rx(r"constraints?|non[- ]?functional\s+requirements|nfrs?"),
        'legacy_system': hdr_rx(r"legacy\s+(?:system|systems|context)|existing\s+system|current\s+state|as[- ]?is"),
    }

    lines = [l.rstrip('\r') for l in text.splitlines()]
    n = len(lines)

    # Generic section heading detector to stop capture at boundaries
    def is_generic_heading(s: str) -> bool:
        s = s or ''
        if s.startswith('- ') or s.startswith('* '):
            return False
        return (
            bool(re.match(r"^\s*#{1,6}\s+\S", s))
            or bool(re.match(r"^\s*(?:\d+|[IVXivx]+)[\.)]\s+\S", s))
            or bool(re.match(r"^\s*[A-Z][-A-Za-z0-9 /&_]{2,60}:?\s*$", s))
        )

    # Find first occurrence index for each section
    idxs: dict[str, int] = {}
    for i, line in enumerate(lines):
        stripped = line.strip()
        for key, rx in patterns.items():
            if key in idxs:
                continue
            if rx.match(stripped):
                idxs[key] = i

    # Helper: capture from start until next boundary
    def grab(start: int) -> str:
        if start is None or start < 0:
            return ''
        buf: list[str] = []
        # Determine the next known target-section heading index after start
        next_known = min((v for v in idxs.values() if v > start), default=n)
        j = start + 1
        while j < min(next_known, n):
            l = lines[j]
            lj = l.strip()
            # Stop at an obvious new section heading (but allow short sub-bullets)
            if is_generic_heading(lj):
                break
            buf.append(l)
            j += 1
        # Trim trailing blank lines
        while buf and buf[-1].strip() == '':
            buf.pop()
        return "\n".join(buf).strip()

    for key, start in idxs.items():
        val = grab(start)
        if val:
            sections[key] = val
    return sections


def _extract_context_fallback(prd_text: str) -> dict:
    """Heuristic fallback to pull Business Goals, Constraints, and Legacy System from raw PRD text.
    Works even when headings are absent or loosely formatted. Captures more lines and stops
    at generic section boundaries to avoid truncation.
    """
    text = (prd_text or '').strip()
    if not text:
        return {}
    lines = [l.rstrip() for l in text.splitlines()]
    n = len(lines)

    def is_generic_heading(s: str) -> bool:
        s = s or ''
        if s.startswith('- ') or s.startswith('* '):
            return False
        return (
            bool(re.match(r"^\s*#{1,6}\s+\S", s))
            or bool(re.match(r"^\s*(?:\d+|[IVXivx]+)[\.)]\s+\S", s))
            or bool(re.match(r"^\s*[A-Z][-A-Za-z0-9 /&_]{2,60}:?\s*$", s))
        )

    def capture_by_keywords(keywords: list[str], max_lines: int = 60) -> str:
        rx = re.compile('|'.join([re.escape(k) for k in keywords]), re.I)
        for i, l in enumerate(lines):
            if rx.search(l):
                # accumulate subsequent lines until a hard break
                buf: list[str] = []
                j = i
                while j < min(i + max_lines, n):
                    lj = lines[j].rstrip()
                    sj = lj.strip()
                    if not sj:
                        # allow single blank lines within the section
                        if buf and (j + 1 < n) and lines[j + 1].strip():
                            buf.append('')
                            j += 1
                            continue
                        else:
                            # stop if we hit multiple blanks
                            if buf:
                                break
                            j += 1
                            continue
                    if is_generic_heading(sj) and j > i:
                        break
                    buf.append(lj)
                    j += 1
                # Trim trailing blanks
                while buf and buf[-1].strip() == '':
                    buf.pop()
                return '\n'.join(buf).strip()
        return ''

    goals = capture_by_keywords(['business goals', 'goals', 'objectives', 'goal:'])
    constraints = capture_by_keywords(['constraints', 'non-functional', 'nfr', 'must', 'should'])
    legacy = capture_by_keywords(['legacy', 'existing system', 'current state', 'as-is'])

    # Final fallback: take a larger excerpt from the start
    def excerpt(prefix: str) -> str:
        head = text[:1200].strip().replace('\r', '')
        return f"{prefix}:\n" + head if head else ''

    return {
        'business_goals': goals or excerpt('Business Goals (excerpt)'),
        'constraints': constraints or excerpt('Constraints (excerpt)'),
        'legacy_system': legacy or excerpt('Legacy System (excerpt)'),
    }


def _top_services_from_decisions(text: str, max_n: int = 5) -> list[str]:
    items = []
    for line in (text or '').splitlines():
        m = re.match(r"^\s*[-*]\s*(.+)$", line)
        if not m:
            continue
        name = re.sub(r"[^A-Za-z0-9 ]+", '', m.group(1)).strip()
        if not name:
            continue
        words = name.split()
        short = ''.join(w.capitalize() for w in words[:3])
        if short and short not in items:
            items.append(short)
            if len(items) >= max_n:
                break
    return items


def _extract_domain_terms(prd_text: str) -> set[str]:
    """Extract canonical product-domain terms from PRD text.
    Returns a set of canonical names: Customer, Account, Product, Catalog, Inventory,
    Order, Cart, Checkout, Payment, Shipment, Invoice, Subscription, Pricing,
    Discount, Return, Review, Notification, Analytics, Fraud, Support, Vendor.
    """
    text = (prd_text or '').lower()
    if not text:
        return set()
    rx = {
        'Customer': r"\bcustomer|end[- ]?user|buyer|member|subscriber\b",
        'Account': r"\baccount(s)?\b",
        'Product': r"\bproduct(s)?|item(s)?|sku(s)?\b",
        'Catalog': r"\bcatalog(ue)?\b",
        'Inventory': r"\binventory|stock\b",
        'Order': r"\border(s)?|purchase(s)?\b",
        'Cart': r"\bcart|basket\b",
        'Checkout': r"\bcheckout\b",
        'Payment': r"\bpayment(s)?|pay\b|credit card|debit card|stripe|paypal|adyen|braintree",
        'Shipment': r"\bshipment|shipping|delivery|fulfillment\b",
        'Invoice': r"\binvoice|billing\b",
        'Subscription': r"\bsubscription(s)?\b",
        'Pricing': r"\bpricing|pricebook|price\b",
        'Discount': r"\bdiscount|promotion|promo|coupon|voucher\b",
        'Return': r"\breturn(s)?|refund(s)?|rma\b",
        'Review': r"\breview(s)?|rating(s)?|feedback\b",
        'Notification': r"\bnotification(s)?|email|sms|push\b",
        'Analytics': r"\banalytics|report(ing)?\b",
        'Fraud': r"\bfraud|chargeback\b",
        'Support': r"\bsupport|helpdesk|ticket(ing)?\b",
        'Vendor': r"\bvendor|supplier\b",
        # Banking/Financial
        'Transaction': r"\btransaction(s)?|txn(s)?\b",
        'Card': r"\bcard(s)?|debit|credit\b",
        'Loan': r"\bloan(s)?\b",
        'Mortgage': r"\bmortgage(s)?\b",
        'KYC': r"\bkyc\b|know your customer",
        'AML': r"\baml\b|anti[- ]money[- ]launder",
        'Ledger': r"\bledger\b|gl\b|general ledger",
        'Statement': r"\bstatement(s)?\b|bank statement",
        'FX': r"\bfx\b|foreign exchange|forex",
        'Treasury': r"\btreasury\b",
        'Credit': r"\bcredit\b|score|scoring",
        'Underwriting': r"\bunderwrit(ing|e)\b",
        'ATM': r"\batm\b",
        'Branch': r"\bbranch\b",
        'Compliance': r"\bcompliance|regulator(y|)|regulation(s)?\b",
        # Healthcare
        'Patient': r"\bpatient(s)?\b",
        'Provider': r"\bprovider(s)?|physician(s)?|doctor(s)?|nurse(s)?\b",
        'Appointment': r"\bappointment(s)?|schedule|scheduling\b",
        'Encounter': r"\bencounter(s)?\b",
        'EHR': r"\behr\b|\bemr\b|electronic health record|electronic medical record",
        'FHIR': r"\bfhir\b|hl7 fhir",
        'Lab': r"\blab(s)?|laborator(y|ies)\b",
        'Radiology': r"\bradiolog(y|ies)|imaging|pac(s)?\b",
        'Prescription': r"\bprescription(s)?|rx\b",
        'Pharmacy': r"\bpharmacy|pharmacies\b",
        'Claim': r"\bclaim(s)?\b",
        'Payer': r"\bpayer(s)?|insurer(s)?|insurance\b",
        'Eligibility': r"\beligibility\b",
        'Authorization': r"\bauthorization(s)?|pre[- ]auth\b",
        'Diagnosis': r"\bdiagnosis|icd[- ]?10|icd[- ]?9\b",
        'Procedure': r"\bprocedure(s)?|cpt\b",
        'CarePlan': r"\bcare\s*plan(s)?\b",
        'Consent': r"\bconsent\b",
    }
    found: set[str] = set()
    for canon, pat in rx.items():
        try:
            if re.search(pat, text, re.I):
                found.add(canon)
        except Exception:
            continue
    # If Products present, imply Catalog
    if 'Product' in found:
        found.add('Catalog')
    return found


def _derive_services_from_terms(terms: set[str]) -> list[str]:
    """Map domain terms to a concise set of service names for diagrams."""
    if not terms:
        return []
    ordering = [
        ('Customer', 'Customer Service'),
        ('Account', 'Account Service'),
        ('Catalog', 'Catalog Service'),
        ('Product', 'Product Service'),
        ('Cart', 'Cart Service'),
        ('Checkout', 'Checkout Service'),
        ('Order', 'Order Service'),
        ('Payment', 'Payment Service'),
        ('Inventory', 'Inventory Service'),
        ('Shipment', 'Shipping Service'),
        ('Invoice', 'Invoice Service'),
        ('Subscription', 'Subscription Service'),
        ('Pricing', 'Pricing Service'),
        ('Discount', 'Discount Service'),
        ('Return', 'Returns Service'),
        ('Review', 'Review Service'),
        ('Notification', 'Notification Service'),
        ('Analytics', 'Analytics Service'),
        ('Fraud', 'Fraud Service'),
        ('Support', 'Support Service'),
        ('Vendor', 'Vendor Service'),
        # Banking/Financial
        ('Transaction', 'Transactions Service'),
        ('Card', 'Card Service'),
        ('Loan', 'Loans Service'),
        ('Mortgage', 'Mortgages Service'),
        ('KYC', 'KYC Service'),
        ('AML', 'AML Service'),
        ('Ledger', 'Ledger Service'),
        ('Statement', 'Statements Service'),
        ('FX', 'FX Service'),
        ('Treasury', 'Treasury Service'),
        ('Credit', 'Credit Service'),
        ('Underwriting', 'Underwriting Service'),
        ('ATM', 'ATM Service'),
        ('Branch', 'Branch Service'),
        ('Compliance', 'Compliance Service'),
        # Healthcare
        ('Patient', 'Patient Service'),
        ('Provider', 'Provider Service'),
        ('Appointment', 'Scheduling Service'),
        ('Encounter', 'Encounter Service'),
        ('EHR', 'EHR Integration Service'),
        ('FHIR', 'FHIR Adapter Service'),
        ('Lab', 'Lab Service'),
        ('Radiology', 'Radiology Service'),
        ('Prescription', 'Prescription Service'),
        ('Pharmacy', 'Pharmacy Service'),
        ('Claim', 'Claims Service'),
        ('Payer', 'Payer Integration Service'),
        ('Eligibility', 'Eligibility Service'),
        ('Authorization', 'Authorization Service'),
        ('Diagnosis', 'Diagnosis Service'),
        ('Procedure', 'Procedure Service'),
        ('CarePlan', 'Care Plan Service'),
        ('Consent', 'Consent Service'),
    ]
    out: list[str] = []
    for key, svc in ordering:
        if key in terms and svc not in out:
            out.append(svc)
    # Ensure we don't overwhelm the diagram
    if len(out) > 12:
        out = out[:12]
    return out


def _compute_prd_decomposition(prd_text: str) -> dict:
    """Compute PRD-only decomposition inputs used for diagram generation."""
    features = _extract_features(prd_text)
    terms = _fold_features_into_terms(_extract_domain_terms(prd_text), features)
    # Prefer feature-derived services; if none, use PRD-wide domain terms
    services = _services_from_features(features) or _derive_services_from_terms(terms)
    simple = _extract_simple_fields_from_prd(prd_text)
    return {
        'features': features,
        'terms': sorted(list(terms)),
        'services': services,
        'simple': simple,
    }


def _build_decomposition_writeup_text(prd_text: str) -> str:
    """Produce a brief, human-readable summary of how diagrams were decomposed from the PRD.
    Avoids 'n/a' by falling back to current selections or sensible defaults when PRD lacks signals.
    """
    prd_text = prd_text or ''
    d = _compute_prd_decomposition(prd_text)
    features = d['features']
    terms = d['terms']
    services = d['services']
    simple = d['simple']

    # Sensible defaults aligned with dropdown defaults
    defaults = {
        'architecture_style': (session.get('architecture_style') or session.get('desired_architecture_style') or 'Microservices'),
        'cloud_provider': session.get('cloud_provider') or 'AWS',
        'region_strategy': session.get('region_strategy') or 'Single Region',
        'deployment_model': session.get('deployment_model') or 'Serverless',
        'database_family': session.get('database_family') or 'PostgreSQL',
        'messaging_backbone': session.get('messaging_backbone') or 'Kafka',
        'identity_provider': session.get('identity_provider') or 'Okta',
        'api_style': (session.get('api_style') or 'REST'),
        'caching_strategy': session.get('caching_strategy') or 'Redis',
        'observability_stack': session.get('observability_stack') or 'Grafana/Loki/Tempo',
    }

    def j(arr):
        return ', '.join(arr) if arr else 'None detected'

    def gv(key: str) -> str:
        # Priority: PRD-derived simple -> current session selection(s) -> defaults
        if key == 'architecture_style':
            return (simple.get('architecture_style')
                    or session.get('architecture_style')
                    or session.get('desired_architecture_style')
                    or defaults['architecture_style'])
        return (simple.get(key) or session.get(key) or defaults.get(key, ''))

    lines = [
        "- Mode: Strict PRD-only decomposition (no mapping enrichment or generic fallbacks)",
        f"- Features detected: {j(features)}",
        f"- Domain terms detected: {j(terms)}",
        f"- Services inferred: {j(services)}",
        "- Guardrails used:",
        f"  - Architecture Style: {gv('architecture_style')}",
        f"  - Cloud Provider: {gv('cloud_provider')}",
        f"  - Region Strategy: {gv('region_strategy')}",
        f"  - Deployment Model: {gv('deployment_model')}",
        f"  - Database Family: {gv('database_family')}",
        f"  - Messaging Backbone: {gv('messaging_backbone')}",
        f"  - Identity Provider: {gv('identity_provider')}",
        f"  - API Style: {gv('api_style')}",
        f"  - Caching Strategy: {gv('caching_strategy')}",
        f"  - Observability Stack: {gv('observability_stack')}",
    ]
    return "\n".join(lines)


def generate_blueprint_from_session(strict_prd_only: bool = False) -> str:
    """Generate the architecture blueprint mermaid.
    When strict_prd_only=True, derive values only from PRD decomposition and
    avoid enrichment from mapping, guesses, or unrelated session fields.
    """
    # In strict mode, re-extract guardrail-like values from PRD only
    prd_text = session.get('prd_text', '')
    if strict_prd_only:
        simple = _extract_simple_fields_from_prd(prd_text)
        style = (simple.get('architecture_style') or '').strip()
        cloud = (simple.get('cloud_provider') or '').strip()
        region = (simple.get('region_strategy') or '').strip()
        deploy = (simple.get('deployment_model') or '').strip()
        dbfam = (simple.get('database_family') or '').strip()
        msg = (simple.get('messaging_backbone') or '').strip()
        idp = (simple.get('identity_provider') or '').strip()
        api_style = (simple.get('api_style') or '').strip()
        caching = (simple.get('caching_strategy') or '').strip()
        observability = (simple.get('observability_stack') or '').strip()
    else:
        style = (session.get('architecture_style') or session.get('desired_architecture_style') or '').strip()
        cloud = (session.get('cloud_provider') or '').strip()
        region = (session.get('region_strategy') or '').strip()
        deploy = (session.get('deployment_model') or '').strip()
        dbfam = (session.get('database_family') or '').strip()
        msg = (session.get('messaging_backbone') or '').strip()
        idp = (session.get('identity_provider') or '').strip()
        api_style = (session.get('api_style') or '').strip()
        caching = (session.get('caching_strategy') or '').strip()
        observability = (session.get('observability_stack') or '').strip()

    # PRD-driven domain → services
    features = _extract_features(prd_text)
    terms = _fold_features_into_terms(_extract_domain_terms(prd_text), features)
    # Prefer services directly implied by features, then domain-derived
    svc_names = _services_from_features(features) or _derive_services_from_terms(terms)
    if not svc_names and not strict_prd_only:
        svc_names = _top_services_from_decisions(session.get('high_level_decision_points', '')) or [
            'CoreServiceA', 'CoreServiceB', 'CoreServiceC'
        ]

    mermaid_lines = []
    mermaid_lines.append("flowchart TB")
    mermaid_lines.append("classDef svc fill:#e3f2fd,stroke:#1976d2,stroke-width:1px,color:#0d47a1")
    mermaid_lines.append("classDef db fill:#fff3cd,stroke:#ff9800,color:#e65100")
    mermaid_lines.append("classDef gw fill:#fce4ec,stroke:#c2185b,color:#880e4f")
    mermaid_lines.append("classDef msg fill:#e8f5e9,stroke:#2e7d32,color:#1b5e20")
    mermaid_lines.append("classDef ext fill:#f3e5f5,stroke:#6a1b9a,color:#4a148c")
    mermaid_lines.append("classDef idp fill:#ede7f6,stroke:#5e35b1,color:#311b92")

    mermaid_lines.append("subgraph Clients")
    mermaid_lines.append("User[User]")
    mermaid_lines.append("end")

    mermaid_lines.append("subgraph Edge")
    apigw_label = "API Gateway"
    if api_style:
        low = api_style.lower()
        if low in ("rest","graphql","grpc"):
            apigw_label = f"{api_style.upper()} Gateway" if low != 'graphql' else 'GraphQL Gateway'
        else:
            apigw_label = sanitize_node_label(api_style)
    mermaid_lines.append(f"APIGW[{sanitize_node_label(apigw_label)}]:::gw")
    if idp:
        mermaid_lines.append("IDP[" + sanitize_node_label(idp) + "]:::idp")
    mermaid_lines.append("end")

    mermaid_lines.append("subgraph Services")
    for i, n in enumerate(svc_names, start=1):
        mermaid_lines.append(f"S{i}[{sanitize_node_label(n)}]:::svc")
    mermaid_lines.append("end")

    mermaid_lines.append("subgraph Data")
    if dbfam:
        db_labels = [dbfam + ' - Primary','Cache','Analytics']
    else:
        db_labels = ['Relational DB','NoSQL DB','Warehouse']
    # Add domain-flavored stores if present
    if 'Payment' in terms:
        db_labels[0] = sanitize_node_label(db_labels[0])
    for i, lbl in enumerate(db_labels, start=1):
        mermaid_lines.append(f"DB{i}(({sanitize_node_label(lbl)})):::db")
    mermaid_lines.append("end")

    if caching and caching.lower() != 'none':
        mermaid_lines.append("subgraph Caching")
        mermaid_lines.append(f"CACHE(({sanitize_node_label(caching)} Cache))")
        mermaid_lines.append("end")
        mermaid_lines.append("classDef cache fill:#e0f7fa,stroke:#00838f,color:#006064")
        mermaid_lines.append("class CACHE cache")

    mermaid_lines.append("subgraph Messaging")
    if msg:
        mnode = 'Kafka' if 'kafka' in msg.lower() else ('RabbitMQ' if 'rabbit' in msg.lower() else msg)
        mermaid_lines.append(f"BUS(({sanitize_node_label(mnode)})):::msg")
    else:
        mermaid_lines.append("BUS((Event Bus)):::msg")
    mermaid_lines.append("end")

    mermaid_lines.append("subgraph External")
    # Always include a stable placeholder so mapping integration can expand reliably
    mermaid_lines.append("EXT_PLACEHOLDER[External Systems]:::ext")
    if 'Payment' in terms:
        mermaid_lines.append("PAYGW[Payment Gateway]:::ext")
    mermaid_lines.append("end")

    # Core flows
    mermaid_lines.append("User-->APIGW")
    if idp:
        mermaid_lines.append("IDP-->|Auth|APIGW")
    for i in range(1, len(svc_names)+1):
        mermaid_lines.append(f"APIGW-->S{i}")
    # Domain-typical data flow hints
    name_to_idx = {n: i for i, n in enumerate(svc_names, start=1)}
    def edge_if(a: str, b: str, label: str | None = None):
        if a in name_to_idx and b in name_to_idx:
            sa = f"S{name_to_idx[a]}"; ta = f"S{name_to_idx[b]}"
            mermaid_lines.append(f"{sa}-->{'|'+sanitize_node_label(label)+'|' if label else ''}{ta}")
    edge_if('Catalog Service','Cart Service','browse/add')
    edge_if('Cart Service','Checkout Service','checkout')
    edge_if('Checkout Service','Payment Service','authorize/capture')
    edge_if('Order Service','Inventory Service','reserve/commit')
    edge_if('Order Service','Shipping Service','fulfillment')
    if 'Payment' in terms and 'Payment Service' in name_to_idx:
        mermaid_lines.append(f"S{name_to_idx['Payment Service']}-->|capture|BUS")
        mermaid_lines.append(f"S{name_to_idx['Payment Service']}-->|settlement|PAYGW")
    for i in range(1, min(len(svc_names), 3)+1):
        mermaid_lines.append(f"S{i}-->DB{i}")
    for i in range(1, len(svc_names)+1):
        mermaid_lines.append(f"S{i}-->|pub/sub|BUS")
    if caching and caching.lower() != 'none':
        for i in range(1, len(svc_names)+1):
            mermaid_lines.append(f"S{i}-. get/set .->CACHE")
    if observability:
        mermaid_lines.append("subgraph Observability")
        mermaid_lines.append(f"OBS[{sanitize_node_label(observability)}]")
        mermaid_lines.append("end")
        mermaid_lines.append("classDef obs fill:#fffde7,stroke:#f9a825,color:#f57f17")
        mermaid_lines.append("class OBS obs")
        for i in range(1, len(svc_names)+1):
            mermaid_lines.append(f"S{i}-. |traces/metrics| .->OBS")

    if cloud or region or deploy:
        mermaid_lines.append(f"%% Cloud: {cloud or 'n/a'} | Region: {region or 'n/a'} | Deploy: {deploy or 'n/a'}")

    mermaid = "\n".join(mermaid_lines)

    # Integrate mapping nodes/edges (always append if available; core remains PRD-driven)
    try:
        mapping = json.loads(session.get('system_mapping', '{}'))
    except Exception:
        mapping = {}
    nodes = mapping.get('nodes') or []
    edges = mapping.get('edges') or []
    if nodes or edges:
        mermaid += "\n%% System Mapping"

    alias_map: dict[str, str] = {}
    shown_nodes_limit = 40
    shown_edges_limit = 40

    def get_alias(name: str) -> str:
        if name in alias_map:
            return alias_map[name]
        alias = f"N{len(alias_map)+1}"
        alias_map[name] = alias
        return alias

    def sanitize_label(text: str) -> str:
        if not text:
            return ""
        t = str(text)
        t = t.replace('|', '/')
        if len(t) > 80:
            t = t[:77] + '...'
        return t

    ext_nodes = nodes[:shown_nodes_limit]
    if ext_nodes:
        extra = []
        for n in ext_nodes:
            alias = get_alias(n)
            extra.append(f"\n{alias}[{sanitize_node_label(n)}]:::ext")
        mermaid = mermaid.replace(
            "EXT_PLACEHOLDER[External Systems]:::ext",
            "EXT_SUMMARY[External Systems]:::ext" + "".join(extra)
        )

    for e in edges[:shown_edges_limit]:
        s = e[0] if len(e) > 0 else None
        t = e[1] if len(e) > 1 else None
        lbl = e[2] if len(e) > 2 else None
        if not s or not t:
            continue
        sa = get_alias(s)
        ta = get_alias(t)
        label = sanitize_label(lbl)
        label_part = f"|{label}|" if label else ""
        mermaid += f"\n{sa}-->{label_part}{ta}"

    if alias_map:
        mermaid += "\nIL[Integration Layer]:::svc"
        for orig, alias in list(alias_map.items())[:min(len(alias_map), 10)]:
            mermaid += f"\n{alias}-->IL"
        for i in range(1, len(svc_names)+1):
            mermaid += f"\nIL-->|adapters|S{i}"

    mermaid += "\n%% Auto-generated initial blueprint\n%% Style: {}".format(style or 'unspecified')
    return mermaid


def generate_system_interaction_diagram(strict_prd_only: bool = False) -> str:
    """Create a Mermaid sequence diagram focused on product-domain interactions.
    When strict_prd_only=True, derive only from PRD decomposition and avoid fallbacks.
    """
    prd_text = session.get('prd_text', '')
    features = _extract_features(prd_text)
    terms = _fold_features_into_terms(_extract_domain_terms(prd_text), features)
    domain_services = _services_from_features(features) or _derive_services_from_terms(terms)
    if strict_prd_only:
        svc_names = domain_services
        simple = _extract_simple_fields_from_prd(prd_text)
        api_style = (simple.get('api_style') or 'REST').upper()
        idp = (simple.get('identity_provider') or 'Identity Provider')
        dbfam = (simple.get('database_family') or 'Database')
        msg = (simple.get('messaging_backbone') or 'Event Bus')
    else:
        fallback_services = _top_services_from_decisions(session.get('high_level_decision_points', '')) or [
            'CoreServiceA', 'CoreServiceB', 'CoreServiceC'
        ]
        svc_names = domain_services or fallback_services
        api_style = (session.get('api_style') or 'REST').upper()
        idp = (session.get('identity_provider') or 'Identity Provider')
        dbfam = (session.get('database_family') or 'Database')
        msg = (session.get('messaging_backbone') or 'Event Bus')

    lines: list[str] = []
    lines.append('sequenceDiagram')
    lines.append('autonumber')
    lines.append('actor User')
    lines.append('participant APIGW as API Gateway')
    for i, n in enumerate(svc_names, 1):
        lines.append(f'participant S{i} as {sanitize_node_label(n)}')
    lines.append(f'participant DB as {sanitize_node_label(dbfam)}')
    lines.append(f'participant BUS as {sanitize_node_label(msg)}')
    lines.append(f'participant IDP as {sanitize_node_label(idp)}')

    lines.append('User->>APIGW: Request (' + api_style + ')')
    lines.append('APIGW->>IDP: Authenticate/Authorize')
    lines.append('IDP-->>APIGW: Token/Claims')

    name_to_idx = {n: i for i, n in enumerate(svc_names, start=1)}
    def route(a: str, b: str, msg: str):
        if a in name_to_idx and b in name_to_idx:
            lines.append(f"APIGW->>S{name_to_idx[a]}: {msg}")
            lines.append(f"S{name_to_idx[a]}->>S{name_to_idx[b]}: {msg}")
            lines.append(f"S{name_to_idx[b]}-->>APIGW: OK")

    # Domain flow if present: browse -> cart -> checkout -> payment -> order -> inventory -> shipping
    if domain_services:
        if 'Catalog Service' in name_to_idx:
            lines.append('APIGW->>S{0}: Browse catalog'.format(name_to_idx['Catalog Service']))
            lines.append('S{0}-->>APIGW: Product list'.format(name_to_idx['Catalog Service']))
        if 'Cart Service' in name_to_idx:
            lines.append('APIGW->>S{0}: Add item to cart'.format(name_to_idx['Cart Service']))
            lines.append('S{0}-->>APIGW: Cart updated'.format(name_to_idx['Cart Service']))
        if 'Checkout Service' in name_to_idx:
            lines.append('APIGW->>S{0}: Start checkout'.format(name_to_idx['Checkout Service']))
            lines.append('S{0}-->>APIGW: Checkout session'.format(name_to_idx['Checkout Service']))
        if 'Payment Service' in name_to_idx:
            lines.append('APIGW->>S{0}: Submit payment details'.format(name_to_idx['Payment Service']))
            lines.append('S{0}->>BUS: PaymentAuthorized event'.format(name_to_idx['Payment Service']))
            lines.append('S{0}-->>APIGW: Payment result'.format(name_to_idx['Payment Service']))
        if 'Order Service' in name_to_idx:
            lines.append('APIGW->>S{0}: Place order'.format(name_to_idx['Order Service']))
            lines.append('S{0}->>BUS: OrderCreated event'.format(name_to_idx['Order Service']))
            lines.append('S{0}-->>APIGW: Order ID'.format(name_to_idx['Order Service']))
        if 'Inventory Service' in name_to_idx:
            lines.append('BUS->>S{0}: Reserve stock'.format(name_to_idx['Inventory Service']))
            lines.append('S{0}->>BUS: StockReserved event'.format(name_to_idx['Inventory Service']))
        if 'Shipping Service' in name_to_idx:
            lines.append('BUS->>S{0}: Create shipment'.format(name_to_idx['Shipping Service']))
            lines.append('S{0}->>BUS: ShipmentCreated event'.format(name_to_idx['Shipping Service']))
    elif not strict_prd_only:
        # generic fallback across services
        for i in range(1, len(svc_names)+1):
            lines.append(f'APIGW->>S{i}: Route request')
            lines.append(f'S{i}-->>APIGW: Response DTO')

    if svc_names:
        for i in range(1, min(3, len(svc_names))+1):
            lines.append(f'S{i}->>DB: CRUD Ops')
            lines.append(f'S{i}->>BUS: Publish events')
    lines.append('APIGW-->>User: Response (200)')
    return "\n".join(lines)


def generate_data_model_diagram(strict_prd_only: bool = False) -> str:
    """Create a Mermaid ER diagram using PRD-derived product-domain entities when available.
    When strict_prd_only=True, avoid decision-based fallbacks; output minimal skeleton if no PRD terms.
    """
    prd_text = session.get('prd_text', '')
    features = _extract_features(prd_text)
    terms = _fold_features_into_terms(_extract_domain_terms(prd_text), features)
    domain = list(terms)

    def ent(name: str, fields: list[str]) -> list[str]:
        out = [f"  {name} {{"]
        for f in fields:
            # crude typing from field name
            t = 'INT' if re.search(r"id|count|qty|quantity|number", f, re.I) else (
                'DECIMAL' if re.search(r"amount|price|total|cost", f, re.I) else (
                'DATETIME' if re.search(r"time|date|at$", f, re.I) else 'STRING'))
            col = re.sub(r"[^A-Za-z0-9_]", '', f) or 'field'
            out.append(f"    {t} {col}")
        out.append("  }")
        return out

    lines: list[str] = ['erDiagram']
    added: set[str] = set()

    # Core user/account
    if 'Customer' in terms:
        lines += ent('Customer', ['customer_id', 'email', 'status', 'created_at'])
        added.add('Customer')
    if 'Account' in terms:
        lines += ent('Account', ['account_id', 'customer_id', 'status', 'created_at'])
        added.add('Account')

    # Catalog/Product/Inventory
    if 'Product' in terms:
        lines += ent('Product', ['product_id', 'sku', 'name', 'price', 'status'])
        added.add('Product')
    if 'Catalog' in terms:
        lines += ent('Catalog', ['catalog_id', 'name'])
        added.add('Catalog')
    if 'Inventory' in terms:
        lines += ent('Inventory', ['inventory_id', 'product_id', 'warehouse_id', 'quantity'])
        added.add('Inventory')

    # Order/Cart/Checkout/Payment/Shipment
    if 'Cart' in terms:
        lines += ent('Cart', ['cart_id', 'customer_id', 'created_at'])
        lines += ent('CartItem', ['cart_item_id', 'cart_id', 'product_id', 'qty', 'price'])
        added.update(['Cart','CartItem'])
    if 'Order' in terms:
        lines += ent('Order', ['order_id', 'customer_id', 'status', 'created_at', 'total_amount'])
        lines += ent('OrderItem', ['order_item_id', 'order_id', 'product_id', 'qty', 'price'])
        added.update(['Order','OrderItem'])
    if 'Payment' in terms:
        lines += ent('Payment', ['payment_id', 'order_id', 'amount', 'method', 'status', 'authorized_at'])
        added.add('Payment')
    if 'Shipment' in terms:
        lines += ent('Shipment', ['shipment_id', 'order_id', 'carrier', 'tracking_number', 'status', 'shipped_at'])
        added.add('Shipment')
    if 'Invoice' in terms:
        lines += ent('Invoice', ['invoice_id', 'order_id', 'total_amount', 'issued_at', 'status'])
        added.add('Invoice')

    # Pricing/Discount/Return/Review
    if 'Pricing' in terms:
        lines += ent('Price', ['price_id', 'product_id', 'amount', 'currency', 'effective_from'])
        added.add('Price')
    if 'Discount' in terms:
        lines += ent('Discount', ['discount_id', 'code', 'amount', 'type', 'expires_at'])
        added.add('Discount')
    if 'Return' in terms:
        lines += ent('Return', ['return_id', 'order_id', 'reason', 'status', 'processed_at'])
        added.add('Return')
    if 'Review' in terms:
        lines += ent('Review', ['review_id', 'product_id', 'customer_id', 'rating', 'comment', 'created_at'])
        added.add('Review')

    # Notification/Fraud/Analytics/Support
    if 'Notification' in terms:
        lines += ent('Notification', ['notification_id', 'customer_id', 'channel', 'template', 'sent_at'])
        added.add('Notification')
    if 'Fraud' in terms:
        lines += ent('FraudCheck', ['fraud_check_id', 'order_id', 'score', 'result', 'checked_at'])
        added.add('FraudCheck')
    if 'Analytics' in terms:
        lines += ent('Event', ['event_id', 'type', 'payload', 'created_at'])
        added.add('Event')
    if 'Support' in terms:
        lines += ent('SupportTicket', ['ticket_id', 'customer_id', 'subject', 'status', 'created_at'])
        added.add('SupportTicket')

    # Relationships (only if both sides added)
    def rel(a: str, b: str, card: str, label: str):
        if a in added and b in added:
            lines.append(f"  {a} {card} {b} : {label}")

    rel('Customer','Account','||--o{','owns')
    rel('Customer','Cart','||--o{','has')
    rel('Cart','CartItem','||--o{','contains')
    rel('Customer','Order','||--o{','places')
    rel('Order','OrderItem','||--o{','contains')
    rel('Order','Payment','||--||','paid by')
    rel('Order','Shipment','||--o{','fulfilled by')
    rel('Order','Invoice','||--||','billed by')
    rel('Product','OrderItem','||--o{','sold as')
    rel('Product','CartItem','||--o{','added as')
    rel('Product','Inventory','||--o{','stocked in')
    rel('Product','Price','||--o{','priced by')
    rel('Discount','Order','o{--||','applies to')
    rel('Return','Order','o{--||','for')
    rel('Review','Product','o{--||','for')
    rel('Review','Customer','o{--||','by')
    # Use a clearer verb to avoid potential parsing confusion with short labels
    rel('Notification','Customer','o{--||','notifies')
    rel('FraudCheck','Order','||--||','checks')
    rel('SupportTicket','Customer','o{--||','opened by')

    # Fallback if nothing added
    if len(lines) == 1:
        if strict_prd_only:
            # Minimal skeleton to keep diagram valid without non-PRD fallbacks
            lines += ent('Entity', ['entity_id', 'name', 'created_at'])
        else:
            svc_names = _top_services_from_decisions(session.get('high_level_decision_points', '')) or ['CoreServiceA']
            base = re.sub(r'[^A-Za-z0-9]', '', (svc_names[0] if svc_names else 'Service')) or 'Service'
            entities = {
                'User': ['user_id', 'email', 'created_at'],
                'Account': ['account_id', 'user_id', 'status', 'created_at'],
                'Event': ['event_id', 'type', 'payload', 'created_at'],
                base: [f'{base.lower()}_id', 'name', 'status']
            }
            for name, cols in entities.items():
                lines += ent(name, cols)
            # Basic relationships using correct cardinalities
            lines.append('  User ||--o{ Account : owns')
            lines.append('  Account ||--o{ Event : emits')
            lines.append(f"  {base} ||--o{{ Event : logs")
    # Ensure trailing newline for Mermaid parser stability
    return "\n".join(lines) + "\n"


def generate_service_decomposition_diagram(strict_prd_only: bool = False) -> str:
    """Flowchart showing services decomposed by domain with key dependencies.
    When strict_prd_only=True, avoid decision-based fallbacks and rely solely on PRD.
    """
    prd_text = session.get('prd_text', '')
    features = _extract_features(prd_text)
    terms = _fold_features_into_terms(_extract_domain_terms(prd_text), features)
    services = _services_from_features(features) or _derive_services_from_terms(terms)
    if not services and not strict_prd_only:
        services = _top_services_from_decisions(session.get('high_level_decision_points', '')) or [
            'CoreServiceA', 'CoreServiceB', 'CoreServiceC'
        ]
    # Buckets by domain keyword
    domains = {
        'Customer': [s for s in services if 'Customer' in s or 'Account' in s or 'Support' in s],
        'Catalog & Product': [s for s in services if 'Catalog' in s or 'Product' in s or 'Pricing' in s or 'Review' in s],
        'Commerce': [s for s in services if 'Cart' in s or 'Checkout' in s or 'Order' in s or 'Discount' in s],
        'Fulfillment': [s for s in services if 'Inventory' in s or 'Shipping' in s or 'Vendor' in s],
        'Finance & Risk': [s for s in services if 'Payment' in s or 'Invoice' in s or 'Fraud' in s],
        'Cross-cutting': [s for s in services if 'Notification' in s or 'Analytics' in s],
    }
    lines: list[str] = [
        'flowchart TD',
        'classDef svc fill:#e3f2fd,stroke:#1976d2,color:#0d47a1',
    ]
    # Subgraphs
    for name, grp in domains.items():
        if not grp:
            continue
        lines.append(f"subgraph {sanitize_node_label(name)}")
        for i, s in enumerate(grp, 1):
            nid = re.sub(r'[^A-Za-z0-9]', '', s)
            lines.append(f"  {nid}[{sanitize_node_label(s)}]:::svc")
        lines.append('end')
    # Dependencies (simple known ones)
    def id_of(svc):
        return re.sub(r'[^A-Za-z0-9]', '', svc)
    sset = set(services)
    def edge(a, b):
        if a in sset and b in sset:
            lines.append(f"{id_of(a)}--> {id_of(b)}")
    edge('Checkout Service','Payment Service')
    edge('Order Service','Inventory Service')
    edge('Order Service','Shipping Service')
    edge('Product Service','Pricing Service')
    edge('Order Service','Notification Service')
    return "\n".join(lines)


def _ensure_default_dropdowns():
    defaults = {
        'cloud_provider': 'AWS',
        'region_strategy': 'Single Region',
    'deployment_model': 'Serverless',
    'database_family': 'PostgreSQL',
        'messaging_backbone': 'Kafka',
        'identity_provider': 'Okta',
        'api_style': 'REST',
    'caching_strategy': 'Redis',
    'observability_stack': 'Grafana/Loki/Tempo',
        'data_residency': 'US',
    }
    for k, v in defaults.items():
        _safe_set_session(k, v)


def _auto_populate_from_sources():
    """Populate as much as possible from PRD and System Mapping if available."""
    prd_text = session.get('prd_text', '')
    mapping = None
    try:
        mapping = json.loads(session.get('system_mapping', '{}')) if session.get('system_mapping') else None
    except Exception:
        mapping = None

    # PRD-driven extraction
    if prd_text:
        hints = {
            'business_goals': session.get('business_goals',''),
            'legacy_system': session.get('legacy_system',''),
            'constraints': session.get('constraints',''),
        }
        extracted = extract_decisions_with_llm(prd_text, hints)
        # Robust parse to avoid leaking raw JSON/irrelevant blocks into UI
        try:
            sections = _parse_decisions_flexible(extracted)
        except Exception:
            sections = {'high_level': [], 'one_way': [], 'style': '', 'other_qs': []}
        # Apply parsed results

        # --- Robust fallback for all three fields ---
        # 1. High-Level Decision Points
        if sections.get('high_level'):
            session['high_level_decision_points'] = '\n'.join(f"- {x}" if not x.strip().startswith('-') else x for x in sections['high_level'])
        else:
            session['high_level_decision_points'] = (
                "- Choice of cloud provider and region for deployment (e.g., AWS in us-east-1)\n"
                "- Adoption of microservices or monolithic architecture pattern\n"
                "- Selection of primary database technology (e.g., PostgreSQL, DynamoDB)\n"
                "- Use of managed messaging backbone (e.g., Kafka, SQS) for service communication\n"
                "- Implementation of CI/CD pipeline with infrastructure-as-code (e.g., Terraform, GitHub Actions)"
            )

        # 2. One-Way Door Decisions
        if sections.get('one_way'):
            session['one_way_door_decisions'] = '\n'.join(f"- {x}" if not x.strip().startswith('-') else x for x in sections['one_way'])
        else:
            session['one_way_door_decisions'] = (
                "- Committing to a specific cloud provider (e.g., AWS, Azure, GCP)\n"
                "- Choosing a database technology that is hard to migrate later\n"
                "- Deciding on a serverless vs. containerized deployment model\n"
                "- Selecting a messaging backbone that impacts integration patterns"
            )

        # 3. Other Relevant Questions
        if sections.get('other_qs'):
            session['other_relevant_questions'] = '\n'.join(f"- {x}" if not x.strip().startswith('-') else x for x in sections['other_qs'])
        else:
            session['other_relevant_questions'] = (
                "- Are there regulatory or compliance requirements impacting architecture?\n"
                "- What are the expected peak workloads and scaling needs?\n"
                "- Are there legacy system integration constraints?\n"
                "- What are the disaster recovery and backup requirements?"
            )

        # 4. Desired Architecture Style (single line) — persist if extracted
        if sections.get('style'):
            style = (sections.get('style') or '').strip()
            if style and style.lower() not in ('tbd','n/a','na'):
                # Prefer setting desired style if explicit
                session['desired_architecture_style'] = style

        for k, v in _extract_simple_fields_from_prd(prd_text).items():
            _safe_set_session(k, v)
        core = _extract_core_sections_from_prd(prd_text)
        for k in ['business_goals','constraints','legacy_system']:
            _safe_set_session(k, core.get(k))
        # Fallback if still empty
        if not session.get('business_goals') or not session.get('constraints') or not session.get('legacy_system'):
            fb = _extract_context_fallback(prd_text)
            for k, v in fb.items():
                _safe_set_session(k, v)
        # Additional requirements: collect strong statements (must/should)
        if not session.get('additional_requirements'):
            req_lines = []
            for l in (prd_text or '').splitlines():
                s = l.strip()
                if re.search(r"\b(must|should|shall|required)\b", s, re.I):
                    req_lines.append(s)
                if len(req_lines) >= 12:
                    break
            if req_lines:
                _safe_set_session('additional_requirements', '\n'.join(f"- {x}" for x in req_lines))

    # Mapping-derived
    if mapping:
        nodes = mapping.get('nodes') or []
        edges = mapping.get('edges') or []
        svc_names = _top_services_from_decisions(session.get('high_level_decision_points','')) or []
        comp_lines = [f"- {n}" for n in svc_names] + [f"- External: {n}" for n in nodes[:20]]
        _safe_set_session('major_components', '\n'.join(comp_lines) or 'Auto from PRD/Mapping')
        iface_lines = []
        for e in edges[:50]:
            s = e[0] if len(e) > 0 else ''
            t = e[1] if len(e) > 1 else ''
            lbl = e[2] if len(e) > 2 else None
            if s and t:
                iface_lines.append(f"- {s} -> {t}" + (f" [{lbl}]" if lbl else ''))
        _safe_set_session('interface_definitions', '\n'.join(iface_lines) or 'Auto from Mapping')
        _safe_set_session('data_schemas', '\n'.join([f"- {n}" for n in nodes[:30]]) or 'Entities from Mapping')
        # Derive workloads string from mapping stats if missing
        if not session.get('workloads'):
            stats = mapping.get('stats') or {}
            node_count = stats.get('node_count') or len(nodes)
            edge_count = stats.get('edge_count') or len(edges)
            _safe_set_session('workloads', f"~{node_count} components, ~{edge_count} integrations; traffic patterns TBD")

    # Patterns and communication
    patt = []
    if (session.get('architecture_style','') or '').lower().startswith('micro'):
        patt += ['- Circuit Breaker', '- Saga / Orchestration', '- API Gateway']
    if session.get('messaging_backbone'):
        patt += ['- Event-Driven (Pub/Sub)', '- Outbox / Idempotency']
    if session.get('caching_strategy'):
        patt += ['- Cache-Aside']
    _safe_set_session('reusable_patterns', '\n'.join(patt) or '- Layered Architecture')

    # Consolidated Architectural Decisions with actual resolved values
    session['architectural_decisions'] = _build_architectural_decisions_markdown()
    style = session.get('architecture_style') or session.get('desired_architecture_style') or 'Unspecified'
    dbfam = session.get('database_family') or 'Relational/NoSQL'
    _safe_set_session('rationale_tradeoffs', f"Chose {style} with {dbfam} storage based on PRD goals and constraints.")
    _safe_set_session('blueprint_references', f"PRD: {session.get('prd_file_name','')} | Mapping: {session.get('system_mapping_name','')}")
    api_style_upper = (session.get('api_style') or '').upper()
    comm = 'HTTP/REST + JSON' if api_style_upper == 'REST' else ('gRPC' if api_style_upper == 'GRPC' else ('GraphQL' if api_style_upper == 'GRAPHQL' else 'HTTP APIs'))
    _safe_set_session('communication_protocols', _humanize_if_json(comm))
    _safe_set_session('data_flow_diagrams', 'See generated Blueprint; flows inferred from Mapping edges.')
    if mapping:
        _safe_set_session('integration_points', _humanize_if_json('\n'.join([f"- {n}" for n in (mapping.get('nodes') or [])[:10]]) or 'Derived from Mapping'))
    # Ensure Initial Blueprint core fields are populated with technical content even without mapping
    svc_names = _top_services_from_decisions(session.get('high_level_decision_points','')) or ['CoreServiceA', 'CoreServiceB', 'CoreServiceC']
    if not session.get('major_components'):
        comps = [f"- {n} (stateless, scalable)" for n in svc_names]
        comps += ['- API Gateway (rate limiting, auth, routing)', '- Identity Provider (OIDC/OAuth2)', '- Message Bus (async pub/sub)', '- Primary Database']
        _safe_set_session('major_components', "\n".join(comps))
    if not session.get('interface_definitions'):
        api_style = (session.get('api_style') or 'REST').upper()
        lines = []
        for n in svc_names:
            base = re.sub(r'[^A-Za-z0-9]', '', n).lower() or 'service'
            if api_style == 'REST':
                lines.append(f"- {n}: GET /{base}s, POST /{base}s, GET /{base}s/{{id}}, PUT /{base}s/{{id}}, DELETE /{base}s/{{id}}")
            elif api_style == 'GRAPHQL':
                lines.append(f"- {n}: Query {base}s, Mutation upsert{base.capitalize()}")
            else:
                lines.append(f"- {n}: rpc List{base.capitalize()}s(), rpc Get{base.capitalize()}(Id), rpc Upsert{base.capitalize()}(Dto)")
        lines += ['- Auth: OIDC token on all requests', '- Versioning: Accept header or URI v1']
        _safe_set_session('interface_definitions', "\n".join(lines))
    if not session.get('data_schemas'):
        dbfam = (session.get('database_family') or '')
        schem_lines = [
            f"- Table: {svc_names[0].lower()} (id PK, name, created_at)",
            "- Table: audit_log (id PK, entity, action, created_at)",
            "- Indexes: ix_audit_log_created_at",
            f"- Storage: {dbfam or 'Relational'} with daily backups"
        ]
        _safe_set_session('data_schemas', "\n".join(schem_lines))
    # Agent 5_3 refinement
    try:
        payload_53 = json.dumps({'prd_excerpt': (prd_text or '')[:1000], 'mapping': mapping or {}}, ensure_ascii=False)
        reply_53 = call_agent('agent_5_3', payload_53)
        refined = _parse_comm_integration(reply_53)
        for k, v in refined.items():
            if k in ('communication_protocols','integration_points','data_flow_diagrams'):
                _safe_set_session(k, _humanize_if_json(v))
            else:
                _safe_set_session(k, v)
    except Exception:
        pass

    # Docs via Agent 5_5
    try:
        payload_55 = json.dumps({'prd_excerpt': (prd_text or '')[:1500], 'mapping_stats': (mapping or {}).get('stats', {})}, ensure_ascii=False)
        reply_55 = call_agent('agent_5_5', payload_55)
        doc_fields = _parse_documentation(reply_55)
        if doc_fields:
            # Keep checklist if provided; compiled doc will be built from session instead of agent text
            if 'stakeholder_checklist' in doc_fields:
                _safe_set_session('stakeholder_checklist', _humanize_if_json(doc_fields['stakeholder_checklist']))
    except Exception:
        pass
    # Always (re)build a clean, sectioned compiled document, never JSON
    session['compiled_document'] = _build_compiled_document_markdown()
    if not session.get('stakeholder_checklist'):
        _safe_set_session('stakeholder_checklist', '\n'.join([
            '- Security sign-off', '- Compliance sign-off', '- SRE runbooks', '- Architecture review'
        ]))

    # Pros & Cons via Agent 5_4
    try:
        payload_54 = json.dumps({'decisions': session.get('architectural_decisions','')}, ensure_ascii=False)
        reply_54 = call_agent('agent_5_4', payload_54)
        pc = _parse_pros_cons(reply_54)
        # Normalize SWOT to ensure all four sections are present; avoid raw JSON in UI
        swot_src = pc.get('swot_analysis') or reply_54 or ''
        if swot_src:
            _safe_set_session('swot_analysis', _build_swot_markdown(swot_src))
        # Humanize risks/mitigations if JSON-like
        if pc.get('risks_mitigation'):
            _safe_set_session('risks_mitigation', _humanize_if_json(pc['risks_mitigation']))
    except Exception:
        pass

    # Artifacts (strict PRD-only for diagrams to ensure fidelity with PRD)
    session['architecture_diagram'] = generate_blueprint_from_session(strict_prd_only=True)
    session['system_interaction_diagram'] = generate_system_interaction_diagram(strict_prd_only=True)
    session['data_model_diagram'] = generate_data_model_diagram(strict_prd_only=True)
    session['service_decomposition_diagram'] = generate_service_decomposition_diagram(strict_prd_only=True)
    # Decomposition write-up for UI and export
    try:
        prd_text = session.get('prd_text', '')
        session['decomposition_writeup'] = _build_decomposition_writeup_text(prd_text) if prd_text else ''
    except Exception:
        session['decomposition_writeup'] = ''
    # Agent 5_1 for Context (fallback to generator)
    try:
        payload_51 = json.dumps({
            'selections': {
                'style': session.get('architecture_style') or session.get('desired_architecture_style'),
                'cloud': session.get('cloud_provider'),
                'region': session.get('region_strategy'),
                'deploy': session.get('deployment_model')
            },
            'decisions': session.get('architectural_decisions',''),
            'mapping_stats': (mapping or {}).get('stats', {}),
        }, ensure_ascii=False)
        reply_51 = call_agent('agent_5_1', payload_51)
        # Always use our generator to avoid agent prompt/JSON leakage in Context Doc
        session['architecture_context'] = generate_architecture_context()
        _normalize_context_doc(force=True)
    except Exception:
        # Fallback if agent call or generation fails
        session['architecture_context'] = generate_architecture_context()
    _ensure_default_dropdowns()


def _fill_educated_guesses():
    """Fill any remaining empty fields with reasonable defaults/guesses."""
    cp = (session.get('cloud_provider') or 'AWS').lower()
    deploy = (session.get('deployment_model') or '').lower()
    # Cloud-specific picks
    kms_guess = 'AWS KMS' if 'aws' in cp else ('Azure Key Vault' if 'azure' in cp else ('GCP KMS' if 'gcp' in cp else 'KMS'))
    secrets_guess = 'AWS Secrets Manager' if 'aws' in cp else ('Azure Key Vault' if 'azure' in cp else ('GCP Secret Manager' if 'gcp' in cp else 'Secrets Manager'))
    gateway_guess = 'Managed API Gateway'
    mesh_guess = 'Istio' if 'kubernetes' in deploy else ''
    edge_guess = 'WAF + CDN'

    guesses = {
        'availability_sla': '99.9%',
        'rpo_rto': 'RPO: 15m, RTO: 1h',
        'performance_targets': 'p95 <= 250ms',
        'throughput': '100 TPS',
        'peak_concurrency': '500',
        'data_volume_retention': '1 TB / 365 days',
        'dr_strategy': 'Warm standby in secondary region',
        'encryption_at_rest': 'Enabled',
        'encryption_in_transit': 'TLS 1.2+',
        'kms': kms_guess,
        'api_gateway_type': gateway_guess,
        'service_mesh': mesh_guess or 'None',
        'edge_controls': edge_guess,
        'secrets_manager': secrets_guess,
        'iac_tool': 'Terraform',
        'ci_cd': 'GitHub Actions',
        'environments': 'Dev, QA, Staging, Prod',
        'release_strategy': 'Blue/Green',
        'deployment_topology': 'VPC with public/private subnets',
        'tenancy_model': 'Single-tenant',
        'observability_requirements': 'Centralized logs, metrics, traces with alerting',
        'security_posture': 'Least privilege IAM, encrypted at rest and in transit, regular security scans',
        'cost_constraints': 'Optimize for cost within budget',
        'capacity_estimates': 'Scale to 2x seasonal peak',
        'migration_strategy': 'Phased strangler migration',
        'risks': '- Compliance gaps\n- Operational complexity\n- Data migration risks',
        'open_questions': '- Finalize data residency\n- Confirm SLOs and capacity',
        'assumptions': '- Access to required cloud accounts\n- Teams available for integration and testing',
        'stakeholders': '- Product Owner\n- Engineering Lead\n- Security\n- SRE',
        'in_scope': '- Core service decomposition\n- API layer\n- Data migration plan',
        'out_of_scope': '- Mobile app redesign\n- Non-core legacy rewrites',
        'pros_cons': 'Pros:\n- Scalability potential\n- Clear service boundaries\n\nCons:\n- Operational complexity\n- Distributed tracing required',
        'risks_mitigation': 'Risks:\n- Compliance and data residency\n- Message ordering/duplication\n\nMitigations:\n- Automate controls & monitoring\n- Idempotency & retries\n- Backup/DR runbooks',
    }
    for k, v in guesses.items():
        _safe_set_session(k, v)
    # Ensure high-level context isn't empty
    prd_text = session.get('prd_text', '')
    if not session.get('business_goals') and prd_text:
        _safe_set_session('business_goals', (prd_text[:400] + '...') if len(prd_text) > 400 else prd_text)
    if not session.get('constraints'):
        _safe_set_session('constraints', 'TBD – capture NFRs, compliance, and technical constraints')
    if not session.get('legacy_system'):
        _safe_set_session('legacy_system', 'TBD – summarize current/legacy system context')


def _diagram_history_key(field: str) -> str:
    return f"history_{field}"


def _push_diagram_history(field: str):
    try:
        cur = session.get(field, '') or ''
        if not cur:
            return
        key = _diagram_history_key(field)
        hist = session.get(key)
        try:
            hist_list = json.loads(hist) if isinstance(hist, str) and hist else (hist or [])
        except Exception:
            hist_list = []
        if not isinstance(hist_list, list):
            hist_list = []
        hist_list.append(cur)
        # cap history
        if len(hist_list) > 10:
            hist_list = hist_list[-10:]
        session[key] = json.dumps(hist_list)
    except Exception:
        pass


def _pop_diagram_history(field: str) -> str | None:
    try:
        key = _diagram_history_key(field)
        hist = session.get(key)
        hist_list = []
        if hist:
            try:
                hist_list = json.loads(hist) if isinstance(hist, str) else (hist or [])
            except Exception:
                hist_list = []
        if hist_list:
            prev = hist_list.pop()
            session[key] = json.dumps(hist_list)
            return prev
        return None
    except Exception:
        return None


def _update_diagram_via_agent(diagram_field: str, mermaid_kind: str, user_prompt: str):
    """Call local Mermaid agent to update the specified diagram field from a natural-language prompt.
    mermaid_kind: 'flowchart'|'sequence'|'er'.
    Updates session[diagram_field] in-place when a result is parsed.
    """
    cur = session.get(diagram_field, '') or ''
    payload = {
        'type': mermaid_kind,
        'current_diagram': cur,
        'prompt': user_prompt or ''
    }
    try:
        reply = call_agent('agent_mermaid', json.dumps(payload, ensure_ascii=False))
    except Exception:
        reply = ''
    # Try JSON updates first
    updates = _parse_updates_from_text(reply)
    if updates and isinstance(updates.get('set'), dict) and diagram_field in updates['set']:
        session[diagram_field] = str(updates['set'][diagram_field])
        return
    # Fallback: parse mermaid fenced block
    try:
        m = re.search(r"```\s*mermaid\s*\n([\s\S]*?)```", reply, re.I)
        if m:
            session[diagram_field] = m.group(1).strip()
            return
    except Exception:
        pass
    # Fallback: if reply looks like mermaid (starts with known tokens)
    first = (reply.strip().splitlines() or [''])[0].strip().lower()
    if first.startswith(('graph','flowchart','sequence','erdiagram','classdiagram','statediagram','gantt','pie','journey','mindmap','timeline','gitgraph')):
        session[diagram_field] = reply.strip()
        return
    # Final fallback: try to interpret user prompt as mermaid changes and apply heuristically
    _apply_prompt_to_mermaid(diagram_field, mermaid_kind, user_prompt)


def _apply_prompt_to_mermaid(diagram_field: str, mermaid_kind: str, prompt: str):
    """Heuristically update the diagram by interpreting the prompt as Mermaid or edge calls.
    Supported minimal patterns:
      - flowchart: lines like A-->B or A --label--> B
      - sequence: lines like A->>B: Message
      - er: lines like A ||--o{ B : relation, and simple entity fields with 'Entity: a,b,c'
    If prompt contains a valid Mermaid block, use it directly.
    """
    text = (prompt or '').strip()
    if not text:
        return
    # If the prompt itself contains mermaid tokens or a fenced block, set directly
    if re.search(r"```\s*mermaid\s*\n([\s\S]*?)```", text, re.I):
        m = re.search(r"```\s*mermaid\s*\n([\s\S]*?)```", text, re.I)
        if m:
            session[diagram_field] = m.group(1).strip()
            return
    tokens = ['flowchart','graph','sequenceDiagram','erDiagram']
    if any(tok.lower() in text.lower() for tok in tokens):
        session[diagram_field] = text
        return
    # Otherwise, append minimal lines to existing diagram
    cur = (session.get(diagram_field, '') or '').strip()
    lines: list[str] = []
    header = ''
    if mermaid_kind == 'flowchart':
        header = 'flowchart TD'
        # Edge: A-->B or A -- label --> B
        for m in re.finditer(r"([A-Za-z0-9_]+)\s*-(?:-+|\.|)?>\s*([A-Za-z0-9_]+)(?::\s*([^\n]+))?", text):
            a, b, lab = m.group(1), m.group(2), m.group(3)
            if lab:
                lines.append(f"{a} -- {lab.strip()} --> {b}")
            else:
                lines.append(f"{a}--> {b}")
        # common intents
        if 'idp' in text.lower():
            lines.append('APIGW-->IDP')
        if 'payments' in text.lower():
            lines.append('APIGW-->PAYMENTS')
    elif mermaid_kind == 'sequence':
        header = 'sequenceDiagram'
        # Calls: A->>B: Message
        calls = re.findall(r"([A-Za-z0-9_]+)\s*-{1,2}>{1,2}\s*([A-Za-z0-9_]+)\s*:\s*([^\n]+)", text)
        for a, b, msg in calls:
            lines.append(f"{a}->>{b}: {msg.strip()}")
        # auth hints
        if 'auth' in text.lower() or 'idp' in text.lower():
            lines += ['APIGW->>IDP: Authenticate', 'IDP-->>APIGW: Token']
    elif mermaid_kind == 'er':
        header = 'erDiagram'
        # Relationships: A ||--o{ B : rel
        rels = re.findall(r"([A-Za-z0-9_]+)\s*\|\|--o\{\s*([A-Za-z0-9_]+)\s*:\s*([^\n]+)", text)
        for a, b, rel in rels:
            lines.append(f"{a} ||--o{{ {b} : {rel.strip()}")
        # Entities: Name: a,b,c
        ents = re.findall(r"([A-Za-z][A-Za-z0-9_]*)\s*:\s*([A-Za-z0-9_, ]+)", text)
        for name, fields in ents:
            cols = [c.strip() for c in fields.split(',') if c.strip()]
            if cols:
                lines.append(f"{name} {{")
                for c in cols:
                    t = 'INT' if re.search(r"id|count|num", c, re.I) else 'STRING'
                    safe = re.sub(r"[^A-Za-z0-9_]", '', c) or 'field'
                    lines.append(f"  {t} {safe}")
                lines.append("}")
    if not lines and not header:
        return
    # Build combined
    if not cur:
        new = header + "\n" + "\n".join(lines)
        session[diagram_field] = new.strip()
        return
    # Ensure header exists
    cur_lines = cur.splitlines()
    if header and not any(cur_lines and cur_lines[0].strip().lower().startswith(header.split()[0].lower()) for _ in [0]):
        cur = header + "\n" + cur
        cur_lines = cur.splitlines()
    # Add participants for sequence if missing
    if mermaid_kind == 'sequence':
        existing_parts = set()
        for l in cur_lines:
            m = re.match(r"\s*participant\s+([A-Za-z0-9_]+)", l)
            if m:
                existing_parts.add(m.group(1))
        for l in lines:
            m = re.match(r"\s*([A-Za-z0-9_]+)-+>{1,2}\s*([A-Za-z0-9_]+)", l)
            if m:
                a, b = m.group(1), m.group(2)
                if a not in existing_parts:
                    cur_lines.insert(1, f"participant {a}")
                    existing_parts.add(a)
                if b not in existing_parts:
                    cur_lines.insert(1, f"participant {b}")
                    existing_parts.add(b)
        cur = "\n".join(cur_lines)
    # Append new lines, avoiding duplicates
    cur_set = set([l.strip() for l in cur.splitlines() if l.strip()])
    appended = False
    for l in lines:
        if l.strip() and l.strip() not in cur_set:
            cur += "\n" + l
            cur_set.add(l.strip())
            appended = True
    if appended:
        session[diagram_field] = cur.strip()


def call_agent(agent_key, context):
    prompt = _read_agent_prompt(agent_key)
    return f"{prompt}\n\n{context}"


def read_text_safely(file_path: str) -> str:
    _, ext = os.path.splitext(file_path.lower())
    try:
        if ext in ('.txt', '.md', '.json'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        if ext == '.docx':
            try:
                import docx as docx_mod
                d = docx_mod.Document(file_path)
                return '\n'.join(p.text for p in d.paragraphs)
            except Exception:
                pass
        if ext == '.pdf':
            try:
                import importlib
                pdf_high = importlib.import_module('pdfminer.high_level')
                return getattr(pdf_high, 'extract_text')(file_path) or ''
            except Exception:
                pass
    except Exception:
        pass
    try:
        with open(file_path, 'rb') as f:
            data = f.read()
        return data.decode('utf-8', errors='ignore')
    except Exception:
        return ''


def parse_system_mapping(file_path: str):
    nodes = set()
    edges = []
    _, ext = os.path.splitext(file_path.lower())
    try:
        if ext == '.csv':
            with open(file_path, newline='', encoding='utf-8', errors='ignore') as f:
                reader = csv.DictReader(f)
                if reader.fieldnames is None:
                    f.seek(0)
                    reader = csv.reader(f)
                    for row in reader:
                        if len(row) >= 2:
                            s, t = row[0].strip(), row[1].strip()
                            if s and t:
                                nodes.update([s, t])
                                edges.append((s, t, None))
                else:
                    for row in reader:
                        s = (row.get('source') or row.get('from') or '').strip()
                        t = (row.get('target') or row.get('to') or '').strip()
                        if not s and len(row) >= 1:
                            vals = list(row.values())
                            s = (vals[0] or '').strip()
                            t = (vals[1] or '').strip() if len(vals) > 1 else ''
                        if s and t:
                            label = (row.get('label') or '').strip() or None
                            nodes.update([s, t])
                            edges.append((s, t, label))
        elif ext == '.json':
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
            for n in data.get('nodes', []):
                nid = (n.get('id') if isinstance(n, dict) else str(n)).strip()
                if nid:
                    nodes.add(nid)
            for e in data.get('edges', []):
                s = (e.get('source') or e.get('from') or '').strip()
                t = (e.get('target') or e.get('to') or '').strip()
                if s and t:
                    edges.append((s, t, (e.get('label') or None)))
                    nodes.update([s, t])
        else:
            text = read_text_safely(file_path)
            for line in text.splitlines():
                parts = [p.strip() for p in line.replace(';', ',').split(',') if p.strip()]
                for token in parts:
                    if '->' in token:
                        s, t = [x.strip() for x in token.split('->', 1)]
                        if s and t:
                            nodes.update([s, t])
                            edges.append((s, t, None))
    except Exception:
        pass
    truncated = False
    MAX_EDGES = 300
    if len(edges) > MAX_EDGES:
        edges = edges[:MAX_EDGES]
        truncated = True
    return {
        'nodes': sorted(nodes),
        'edges': edges,
        'stats': {
            'node_count': len(nodes),
            'edge_count': len(edges),
            'truncated': truncated,
        }
    }


def extract_decisions_with_llm(prd_text: str, context_hints: dict | None = None) -> str:
    """Use local Agents (agent_5_2) to extract decisions; avoid external API calls."""
    system_prompt = (
        "You are an Architecture Decision Extractor. Read the PRD text and extract:\n"
        "- High-Level Architecture Decision Points (bulleted)\n"
        "- One-way door decisions (irreversible or costly-to-reverse) (bulleted)\n"
        "- Desired Architecture Style (single line)\n"
        "- Other Relevant Questions (bulleted)\n"
        "Return sections with clear, markdown headings: '# High-Level Architecture Decision Points', '# One-way Door Decisions', '# Desired Architecture Style', '# Other Relevant Questions'."
    )
    user_payload = prd_text or ''
    if context_hints:
        try:
            user_payload = json.dumps({'prd': prd_text, 'hints': context_hints}, ensure_ascii=False)
        except Exception:
            pass
    try:
        return call_agent('agent_5_2', system_prompt + "\n\n" + user_payload)
    except Exception:
        head = (prd_text or '')[:2000]
        return (
            "# High-Level Architecture Decision Points\n"
            "- Review data domains\n- Identify SLAs and latency requirements\n\n"
            "# One-way Door Decisions\n- Choose database family\n- Pick messaging backbone\n\n"
            "# Desired Architecture Style\n- TBD\n\n"
            "# Other Relevant Questions\n- Compliance constraints?\n- Multi-region needs?\n\n---\nExcerpt:\n" + head
        )


def parse_extracted_decisions(md: str) -> dict:
    sections = {'high_level': [], 'one_way': [], 'style': '', 'other_qs': []}
    if not md:
        return sections
    cur = None
    for line in md.splitlines():
        l = line.strip()
        if re.match(r"^#\s*High-Level Architecture Decision Points", l, re.I):
            cur = 'high_level'; continue
        if re.match(r"^#\s*One-way Door Decisions", l, re.I):
            cur = 'one_way'; continue
        if re.match(r"^#\s*Desired Architecture Style", l, re.I):
            cur = 'style'; continue
        if re.match(r"^#\s*Other Relevant Questions", l, re.I):
            cur = 'other_qs'; continue
        if cur in ('high_level','one_way','other_qs'):
            m = re.match(r"^[-*]\s*(.+)$", l)
            if m:
                sections[cur].append(m.group(1).strip())
        elif cur == 'style' and l:
            if not sections['style']:
                sections['style'] = l
    return sections

def _try_parse_json_decisions(text: str) -> dict | None:
    """Attempt to parse decisions from a JSON object embedded in text.
    Accepts variations in key names and value shapes (array or string).
    Returns sections dict like parse_extracted_decisions.
    """
    if not text:
        return None
    # Extract a JSON object from fenced block or the first curly braces block
    m = re.search(r"```json\s*(\{[\s\S]*?\})\s*```", text, re.I)
    blob = None
    if m:
        blob = m.group(1)
    if not blob:
        m2 = re.search(r"(\{[\s\S]*\})", text)
        if m2:
            blob = m2.group(1)
    if not blob:
        return None
    try:
        data = json.loads(blob)
    except Exception:
        return None
    if not isinstance(data, dict):
        return None
    # Normalize keys
    def get_any(d: dict, keys: list[str]):
        for k in keys:
            if k in d:
                return d[k]
            # support case-insensitive and camelCase lookup
            for dk in d.keys():
                if dk.lower() == k.lower():
                    return d[dk]
        return None

    out = {'high_level': [], 'one_way': [], 'style': '', 'other_qs': []}
    high = get_any(data, ['high_level_decision_points','highLevelDecisionPoints','high_level','decisions','decision_points'])
    one = get_any(data, ['one_way_door_decisions','oneWayDoorDecisions','one_way','oneWay'])
    style = get_any(data, ['desired_architecture_style','architecture_style','style'])
    qs = get_any(data, ['other_relevant_questions','otherQuestions','questions'])

    def to_list(v) -> list[str]:
        if v is None:
            return []
        if isinstance(v, list):
            return [str(x).strip() for x in v if str(x).strip()]
        if isinstance(v, str):
            lines = []
            for line in v.splitlines():
                m = re.match(r"^[-*]\s*(.+)$", line.strip())
                if m:
                    lines.append(m.group(1).strip())
            if not lines and v.strip():
                # split on semicolons or periods as a weak fallback
                parts = re.split(r"[;•\u2022]|\.(?:\s|$)", v)
                lines = [p.strip() for p in parts if p and p.strip()]
            return lines
        return []

    out['high_level'] = to_list(high)
    out['one_way'] = to_list(one)
    out['style'] = (style or '').strip() if isinstance(style, str) else (str(style).strip() if style is not None else '')
    out['other_qs'] = to_list(qs)
    # If at least one field is populated, return
    if out['high_level'] or out['one_way'] or out['style'] or out['other_qs']:
        return out
    return None


def _parse_decisions_flexible(text: str) -> dict:
    """Parse decisions robustly from agent output text.
    Order: JSON block → Markdown sections → heuristic bullets.
    """
    # 1) JSON
    js = _try_parse_json_decisions(text)
    if js:
        return js
    # 2) Markdown
    md = parse_extracted_decisions(text)
    if md and (md.get('high_level') or md.get('one_way') or md.get('style') or md.get('other_qs')):
        return md
    # 3) Heuristic fallback: pull any bulleted lines
    hl: list[str] = []
    for line in (text or '').splitlines():
        m = re.match(r"^[-*]\s*(.+)$", line.strip())
        if m:
            item = m.group(1).strip()
            if item and item not in hl:
                hl.append(item)
        if len(hl) >= 12:
            break
    return {'high_level': hl, 'one_way': [], 'style': '', 'other_qs': []}


def _parse_comm_integration(md: str) -> dict:
    """Parse Communication & Integration content from agent text into fields."""
    if not md:
        return {}
    out: dict = {}
    comm = None
    integrations: list[str] = []
    flows: list[str] = []
    proto_rx = re.compile(r"\b(grpc|rest|http|kafka|amqp|mq|pub/sub|eventbridge|sns|sqs)\b", re.I)
    in_integration = False
    for line in md.splitlines():
        l = line.strip()
        if not l:
            continue
        if comm is None and proto_rx.search(l):
            comm = l
        if re.search(r"integration", l, re.I):
            in_integration = True
        m = re.match(r"^[-*]\s*(.+)$", l)
        if m and in_integration:
            integrations.append(m.group(1).strip())
        else:
            flows.append(l)
    if comm:
        out['communication_protocols'] = comm
    if integrations:
        out['integration_points'] = "\n".join(f"- {x}" for x in integrations)
    if flows:
        out['data_flow_diagrams'] = "\n".join(flows[:50])
    return out


def _parse_pros_cons(md: str) -> dict:
    """Parse SWOT and risks/mitigations from agent text."""
    if not md:
        return {}
    out: dict = {}
    swot_lines: list[str] = []
    risk_lines: list[str] = []
    cur = None
    for line in md.splitlines():
        l = line.strip()
        if re.search(r"swot|pros|cons", l, re.I):
            cur = 'swot'; continue
        if re.search(r"risk|mitigation", l, re.I):
            cur = 'risk'; continue
        if not l:
            continue
        if cur == 'swot':
            swot_lines.append(l)
        elif cur == 'risk':
            risk_lines.append(l)
    if swot_lines:
        out['swot_analysis'] = "\n".join(swot_lines)
    else:
        out['swot_analysis'] = md
    if risk_lines:
        out['risks_mitigation'] = "\n".join(risk_lines)
    return out


def _humanize_if_json(text: str) -> str:
    """If text looks like JSON, convert to human-readable bullets; else return as-is."""
    if not text:
        return ''
    t = str(text).strip()
    if not (t.startswith('{') or t.startswith('[')):
        return t
    try:
        data = json.loads(t)
    except Exception:
        return t
    lines: list[str] = []
    def emit(prefix: str, val):
        if isinstance(val, dict):
            lines.append(f"- {prefix}:")
            for k, v in val.items():
                emit(f"  {k}", v)
        elif isinstance(val, list):
            lines.append(f"- {prefix}:")
            for i, item in enumerate(val, 1):
                emit(f"  {i}", item)
        else:
            sval = str(val)
            lines.append(f"- {prefix}: {sval}")
    if isinstance(data, dict):
        for k, v in data.items():
            emit(str(k), v)
    elif isinstance(data, list):
        for i, item in enumerate(data, 1):
            emit(str(i), item)
    return "\n".join(lines)


def _build_swot_markdown(raw: str | None) -> str:
    """Normalize any given text (or empty) into a clear SWOT markdown with all four sections."""
    s_lines: list[str] = []
    w_lines: list[str] = []
    o_lines: list[str] = []
    t_lines: list[str] = []
    text = (raw or '').strip()
    if text:
        # Try to parse simple headings or prefixes
        cur = None
        for line in text.splitlines():
            l = line.strip()
            if not l:
                continue
            if re.search(r"strengths", l, re.I):
                cur = 'S'; continue
            if re.search(r"weaknesses", l, re.I):
                cur = 'W'; continue
            if re.search(r"opportunities", l, re.I):
                cur = 'O'; continue
            if re.search(r"threats", l, re.I):
                cur = 'T'; continue
            m = re.match(r"^[-*]\s*(.+)$", l)
            if m:
                item = m.group(1).strip()
                if cur == 'S': s_lines.append(item)
                elif cur == 'W': w_lines.append(item)
                elif cur == 'O': o_lines.append(item)
                elif cur == 'T': t_lines.append(item)
                else:
                    # Distribute uncategorized bullets heuristically
                    if len(s_lines) <= len(w_lines) and len(s_lines) <= len(o_lines):
                        s_lines.append(item)
                    elif len(w_lines) <= len(o_lines):
                        w_lines.append(item)
                    else:
                        o_lines.append(item)
    # Provide sensible defaults if any section is empty
    if not s_lines:
        s_lines = ['Clear service boundaries', 'Horizontal scalability', 'Tech stack flexibility']
    if not w_lines:
        w_lines = ['Operational complexity', 'Distributed tracing and debugging', 'Increased infra footprint']
    if not o_lines:
        o_lines = ['Faster feature delivery via independent deployments', 'Improved resilience via isolation', 'Observability-driven optimization']
    if not t_lines:
        t_lines = ['Compliance/regulatory risk', 'Vendor lock-in', 'Integration fragility with legacy systems']
    parts = [
        '# Strengths',
        *[f"- {x}" for x in s_lines],
        '',
        '# Weaknesses',
        *[f"- {x}" for x in w_lines],
        '',
        '# Opportunities',
        *[f"- {x}" for x in o_lines],
        '',
        '# Threats',
        *[f"- {x}" for x in t_lines],
    ]
    return "\n".join(parts)


def _build_compiled_document_markdown() -> str:
    """Compose a detailed, sectioned Architecture Document in markdown from current session.
    Always returns human-readable markdown (no JSON).
    """
    def val(key: str, default: str = 'TBD') -> str:
        v = (session.get(key) or '').strip()
        return v if v else default

    def bullets(text: str) -> list[str]:
        out = []
        for line in (text or '').splitlines():
            l = line.strip()
            if not l:
                continue
            l = re.sub(r"^[-*]\s*", '', l)
            out.append(l)
        return out

    # Guardrails list
    guardrails = [
        ("Architecture Style", val('architecture_style', val('desired_architecture_style','TBD'))),
        ("Cloud Provider", val('cloud_provider','TBD')),
        ("Region Strategy", val('region_strategy','TBD')),
        ("Deployment Model", val('deployment_model','TBD')),
        ("Database Family", val('database_family','TBD')),
        ("Messaging Backbone", val('messaging_backbone','TBD')),
        ("Identity Provider", val('identity_provider','TBD')),
        ("API Style", val('api_style','TBD')),
        ("Caching Strategy", val('caching_strategy','TBD')),
        ("Observability Stack", val('observability_stack','TBD')),
        ("Compliance Standards", val('compliance_standards','TBD')),
        ("Data Residency", val('data_residency','TBD')),
    ]

    # Decisions content
    decisions_md = session.get('architectural_decisions','')
    one_way_md = session.get('one_way_door_decisions','')
    other_q_md = session.get('other_relevant_questions','')
    # Communication & Integration (humanize JSON if needed)
    comm = _humanize_if_json(session.get('communication_protocols',''))
    integ = _humanize_if_json(session.get('integration_points',''))
    flows = session.get('data_flow_diagrams','')

    parts: list[str] = []
    parts += [
        '# Compiled Architecture Document',
        '',
        '## 1. Overview',
        f"Desired Style: {val('architecture_style', val('desired_architecture_style','TBD'))}",
        f"Cloud / Region / Deploy: {val('cloud_provider','TBD')} / {val('region_strategy','TBD')} / {val('deployment_model','TBD')}",
    ]
    parts += [
        '',
        '## 2. Business Context',
        '### 2.1 Goals',
        val('business_goals','TBD'),
        '',
        '### 2.2 Legacy System Context',
        val('legacy_system','TBD'),
        '',
        '### 2.3 Constraints & Guardrails',
        *[f"- {k}: {v}" for k, v in guardrails],
    ]
    parts += [
        '',
        '## 3. Architecture Overview',
        'See Initial Blueprint, System Interaction, Data Model, and Service Decomposition diagrams in their respective tabs. Key selections are listed above.',
    ]
    parts += [
        '',
        '## 4. Key Decisions',
        '### 4.1 High-Level Architecture Decision Points',
        *(f"- {x}" for x in bullets(decisions_md) or ['TBD']),
        '',
        '### 4.2 One-way Door Decisions',
        *(f"- {x}" for x in bullets(one_way_md) or ['TBD']),
        '',
        '### 4.3 Desired Architecture Style',
        val('architecture_style', val('desired_architecture_style','TBD')),
        '',
        '### 4.4 Other Relevant Questions',
        *(f"- {x}" for x in bullets(other_q_md) or ['TBD']),
    ]
    parts += [
        '',
        '## 5. Communication & Integration',
        '### 5.1 Communication Protocols',
        comm or 'TBD',
        '',
        '### 5.2 Integration Points',
        integ or 'TBD',
        '',
        '### 5.3 Data Flows',
        flows or 'TBD',
    ]
    parts += [
        '',
        '## 6. Non-Functional Requirements',
        f"- Availability/SLA: {val('availability_sla','TBD')}",
        f"- RPO/RTO: {val('rpo_rto','TBD')}",
        f"- Performance Targets: {val('performance_targets','TBD')}",
        f"- Throughput: {val('throughput','TBD')}",
        f"- Peak Concurrency: {val('peak_concurrency','TBD')}",
        f"- Data Volume & Retention: {val('data_volume_retention','TBD')}",
        f"- Security Posture: {val('security_posture','TBD')}",
    ]
    parts += [
        '',
        '## 7. Risks & Mitigations',
        _humanize_if_json(session.get('risks','')) or '- TBD',
        '',
        '### 7.1 Mitigation Strategies',
        _humanize_if_json(session.get('risks_mitigation','')) or '- TBD',
    ]
    parts += [
        '',
        '## 8. Environments & Release',
        f"- Environments: {val('environments','TBD')}",
        f"- Release Strategy: {val('release_strategy','TBD')}",
        '',
        '## 9. Observability',
        val('observability_requirements','- TBD'),
        '',
        '## 10. Migration Strategy',
        val('migration_strategy','TBD'),
        '',
        '## 11. Open Questions',
        val('open_questions','- TBD'),
        '',
        '## 12. References',
        f"PRD: {val('prd_file_name','n/a')} | Mapping: {val('system_mapping_name','n/a')}",
    ]
    return "\n".join(str(p) for p in parts)


def _normalize_context_doc(force: bool = False):
    """Ensure Context Doc is human-readable markdown (not JSON or agent prompt). Rebuild when needed."""
    try:
        cur = (session.get('architecture_context') or '').strip()
        bad = False
        if force:
            bad = True
        if not bad and cur:
            # JSON-like block at start
            if cur.startswith('{') or cur.startswith('['):
                bad = True
            # Agent prompt content or embedded JSON payload
            if not bad and re.search(r"\bYou are an?\b|Return sections|^\{\s*\"selections\"\s*:|\"decisions\"\s*:\s*\[", cur, re.I | re.M):
                bad = True
        if bad:
            session['architecture_context'] = generate_architecture_context()
    except Exception:
        try:
            session['architecture_context'] = generate_architecture_context()
        except Exception:
            pass


def _parse_documentation(md: str) -> dict:
    """Split compiled document vs stakeholder checklist if headings are present."""
    if not md:
        return {}
    parts = re.split(r"(?im)^#\s*Stakeholder\s*Checklist\s*$", md, maxsplit=1)
    out = {'compiled_document': md.strip()}
    if len(parts) >= 1 and parts[0].strip():
        out['compiled_document'] = parts[0].strip()
    if len(parts) == 2:
        out['stakeholder_checklist'] = parts[1].strip()
    return out


def _append_chat(idx: int, role: str, msg: str):
    key = f"chat_tab{idx}"
    existing = session.get(key, '')
    prefix = "User" if role == 'user' else "Agent"
    new = (existing + f"\n\n{prefix}: " + (msg or '').strip()).strip()
    session[key] = new
    _trim_chat(idx)


def _trim_chat(idx: int, max_chars: int = 4000):
    """Trim stored chat transcript for tab idx to avoid unbounded growth."""
    key = f"chat_tab{idx}"
    txt = session.get(key, '')
    if not txt:
        return
    if len(txt) > max_chars:
        # Keep the tail within limit, but try to cut at a paragraph boundary
        tail = txt[-max_chars:]
        cut = tail.find("\n\n")
        session[key] = tail[cut+2:] if cut != -1 else tail


def _build_architectural_decisions_markdown() -> str:
    """Compose a single Architectural Decisions markdown with actual resolved values.
    Sections:
    - High-Level Architecture Decision Points (bulleted) including resolved guardrail selections
    - One-way Door Decisions (bulleted)
    - Desired Architecture Style (single line)
    - Other Relevant Questions (bulleted)
    Never returns empty placeholders; uses existing session defaults/guesses where needed.
    """
    def _val(key: str, default: str = 'TBD') -> str:
        v = (session.get(key) or '').strip()
        return v if v else default

    def _bullets(text: str) -> list[str]:
        out: list[str] = []
        for line in (text or '').splitlines():
            l = line.strip()
            if not l:
                continue
            out.append(re.sub(r"^[-*]\s*", '', l))
        return out

    # Core resolved selections we want to surface as actual decisions
    resolved_pairs = [
        ('Architecture Style', _val('architecture_style', _val('desired_architecture_style','TBD'))),
        ('Cloud Provider', _val('cloud_provider','TBD')),
        ('Region Strategy', _val('region_strategy','TBD')),
        ('Deployment Model', _val('deployment_model','TBD')),
        ('Database Family', _val('database_family','TBD')),
        ('Messaging Backbone', _val('messaging_backbone','TBD')),
        ('Identity Provider', _val('identity_provider','TBD')),
        ('API Style', _val('api_style','TBD')),
        ('Caching Strategy', _val('caching_strategy','TBD')),
        ('Observability Stack', _val('observability_stack','TBD')),
        ('Compliance Standards', _val('compliance_standards','TBD')),
        ('Data Residency', _val('data_residency','TBD')),
    ]

    high_level_src = session.get('high_level_decision_points','')
    one_way_src = session.get('one_way_door_decisions','')
    other_q_src = session.get('other_relevant_questions','')
    style_val = _val('architecture_style', _val('desired_architecture_style','TBD'))

    high_level_bullets = [f"{k}: {v}" for k, v in resolved_pairs]
    # Add explicit pattern/trade-off bullets derived from selections
    style = style_val
    deploy = _val('deployment_model','TBD')
    dbfam = _val('database_family','TBD')
    msg = _val('messaging_backbone','TBD')
    if style != 'TBD':
        high_level_bullets.append(f"Pattern: {style} — trade-off of agility vs. operational complexity")
    if deploy != 'TBD':
        high_level_bullets.append(f"Deployment: {deploy} — trade-off of cost vs. control")
    if dbfam != 'TBD':
        high_level_bullets.append(f"Data: {dbfam} — trade-off of consistency, scalability, and operational overhead")
    if msg != 'TBD':
        high_level_bullets.append(f"Integration: {msg} — trade-off of latency vs. decoupling and reliability")
    # Append any user/agent extracted bullets (dedup)
    seen = set(b.lower() for b in high_level_bullets)
    for b in _bullets(high_level_src):
        if b and b.lower() not in seen:
            high_level_bullets.append(b)
            seen.add(b.lower())

    one_way_bullets = _bullets(one_way_src) or [
        'Commit to cloud provider and region (data residency, DR topology)',
        'Select primary database family (migration is costly)',
        'Pick messaging backbone (ecosystem and semantics impact)',
        'Choose deployment model (Serverless vs. K8s/VMs)'
    ]
    other_q_bullets = _bullets(other_q_src) or [
        'Clarify compliance scope (GDPR/HIPAA/PCI) and data boundaries',
        'Confirm SLOs (p95 latency, throughput) and capacity planning',
        'Define DR targets and multi-region requirements',
        'List legacy integration constraints and migration cutover plan'
    ]

    parts: list[str] = []
    parts.append('# High-Level Architecture Decision Points')
    parts += [f"- {x}" for x in high_level_bullets]
    parts.append('')
    parts.append('# One-way Door Decisions')
    parts += [f"- {x}" for x in one_way_bullets]
    parts.append('')
    parts.append('# Desired Architecture Style')
    parts.append(style_val)
    parts.append('')
    parts.append('# Other Relevant Questions')
    parts += [f"- {x}" for x in other_q_bullets]
    return "\n".join(parts).strip()


def _get_tab_context_payload(idx: int) -> dict:
    # Build a lightweight context from current session relevant to each tab
    s = session
    if idx == 0:
        return {
            'business_goals': s.get('business_goals',''),
            'legacy_system': s.get('legacy_system',''),
            'constraints': s.get('constraints',''),
            'guardrails': {
                'cloud_provider': s.get('cloud_provider',''),
                'region_strategy': s.get('region_strategy',''),
                'deployment_model': s.get('deployment_model',''),
                'database_family': s.get('database_family',''),
                'messaging_backbone': s.get('messaging_backbone',''),
                'identity_provider': s.get('identity_provider',''),
                'api_style': s.get('api_style',''),
                'caching_strategy': s.get('caching_strategy',''),
                'observability_stack': s.get('observability_stack',''),
            },
            'additional_requirements': s.get('additional_requirements',''),
        }
    if idx == 6:
        return {
            'architecture_context': s.get('architecture_context','')
        }
    if idx == 1:
        return {
            'architecture_diagram': s.get('architecture_diagram',''),
            'system_interaction_diagram': s.get('system_interaction_diagram',''),
            'data_model_diagram': s.get('data_model_diagram',''),
            'major_components': s.get('major_components',''),
            'interface_definitions': s.get('interface_definitions',''),
            'data_schemas': s.get('data_schemas',''),
            'reusable_patterns': s.get('reusable_patterns',''),
        }
    if idx == 2:
        return {
            'architectural_decisions': s.get('architectural_decisions',''),
            'rationale_tradeoffs': s.get('rationale_tradeoffs',''),
            'blueprint_references': s.get('blueprint_references',''),
        }
    if idx == 3:
        return {
            'communication_protocols': s.get('communication_protocols',''),
            'data_flow_diagrams': s.get('data_flow_diagrams',''),
            'integration_points': s.get('integration_points',''),
        }
    if idx == 4:
        return {
            'swot_analysis': s.get('swot_analysis',''),
            'risks_mitigation': s.get('risks_mitigation',''),
        }
    if idx == 5:
        return {
            'compiled_document': s.get('compiled_document',''),
            'stakeholder_checklist': s.get('stakeholder_checklist',''),
        }
    return {}


def generate_architecture_context() -> str:
    """Build a comprehensive Architecture Context markdown using PRD + session + mapping.
    Includes product overview, features, decisions summary, constraints/guardrails, exhaustive assumptions,
    NFRs, integration/data landscape, environments, risks/mitigations, open questions, and references.
    """
    prd_text = session.get('prd_text', '')
    # Mapping snapshot
    try:
        mapping = json.loads(session.get('system_mapping', '{}'))
    except Exception:
        mapping = {}
    nodes = mapping.get('nodes') or []
    edges = mapping.get('edges') or []
    node_count = len(nodes)
    edge_count = len(edges)
    sample_nodes = nodes[:12]
    sample_edges = [f"{e[0]} -> {e[1]}" + (f" [{e[2]}]" if len(e) > 2 and e[2] else '') for e in edges[:12] if len(e) >= 2]

    def val(key: str, default: str = 'TBD') -> str:
        return (session.get(key) or '').strip() or default

    # PRD-derived features and domain terms
    features = _extract_features(prd_text)
    terms = sorted(list(_fold_features_into_terms(_extract_domain_terms(prd_text), features)))

    # Product overview from Business Goals or PRD intro
    product_overview = (session.get('business_goals') or '').strip()
    if not product_overview and prd_text:
        # Use the first 2-3 sentences as a coarse overview
        snippet = re.split(r"(?<=[.!?])\s+", prd_text.strip())[:3]
        product_overview = ' '.join(snippet).strip()[:600]

    # Decisions summary bullets
    high_level = (session.get('high_level_decision_points') or '').strip()
    one_way = (session.get('one_way_door_decisions') or '').strip()
    decision_bullets: list[str] = []
    def split_bullets(text: str) -> list[str]:
        out = []
        for line in (text or '').splitlines():
            line = line.strip()
            if not line:
                continue
            line = re.sub(r"^[-*]\s*", '', line)
            out.append(line)
        return out
    decision_bullets += split_bullets(high_level)
    # Mark one-way doors clearly
    decision_bullets += [f"[One-way] {x}" for x in split_bullets(one_way)]
    # Add key guardrails as explicit decisions
    guardrail_pairs = [
        ('Architecture Style', val('architecture_style', val('desired_architecture_style','TBD'))),
        ('Cloud Provider', val('cloud_provider','TBD')),
        ('Region Strategy', val('region_strategy','TBD')),
        ('Deployment Model', val('deployment_model','TBD')),
        ('Database Family', val('database_family','TBD')),
        ('Messaging Backbone', val('messaging_backbone','TBD')),
        ('Identity Provider', val('identity_provider','TBD')),
        ('API Style', val('api_style','TBD')),
        ('Caching Strategy', val('caching_strategy','TBD')),
        ('Observability Stack', val('observability_stack','TBD')),
    ]
    decision_bullets += [f"{k}: {v}" for k, v in guardrail_pairs]

    # Build exhaustive assumptions list
    base_assumptions = split_bullets(val('assumptions',''))
    extra_assumptions = [
        'Cloud accounts and required subscriptions are provisioned with baseline guardrails',
        'CI/CD is available (e.g., GitHub Actions) with infrastructure-as-code (Terraform) pipelines',
        'Security controls: least-privilege IAM, TLS 1.2+, encryption at rest; secrets managed in a KMS-backed store',
        'Environments exist (dev, test, stage, prod) with promotion and change control',
        'SLA/SLOs are monitored with alerting; on-call is staffed',
        'Data classification and residency requirements are identified and enforceable',
        'Disaster Recovery targets (RPO/RTO) are feasible given chosen region strategy',
        'Access to legacy systems and external integrations is available for migration and testing',
    ]
    # Include specific selections as assumptions if not already captured
    for k, label in [
        ('cloud_provider','Cloud'),('region_strategy','Region Strategy'),('deployment_model','Deployment'),
        ('database_family','Database'),('messaging_backbone','Messaging'),('identity_provider','Identity'),
        ('api_style','API'),('caching_strategy','Caching'),('observability_stack','Observability')
    ]:
        v = val(k, '')
        if v and all(v not in a for a in base_assumptions):
            extra_assumptions.append(f"Selection confirmed: {label} = {v}")
    assumptions_md = '\n'.join([f"- {a}" for a in (base_assumptions + extra_assumptions)]) or '- TBD'

    # Read Agent 5_1 guidance (optional)
    guidance = _read_agent_prompt('agent_5_1')

    lines: list[str] = []
    lines.append('# Architecture Context')
    if guidance:
        lines.append(f"> Guidance: {guidance.strip()}")
    # Overview
    lines += [
        '',
        '## Overview',
        f"Desired Style: {val('architecture_style', val('desired_architecture_style', 'TBD'))}",
        f"Cloud/Region/Deploy: {val('cloud_provider','TBD')} / {val('region_strategy','TBD')} / {val('deployment_model','TBD')}",
    ]
    # Product overview
    lines += [
        '',
        '## Product Overview',
        product_overview or 'TBD',
    ]
    # Features
    lines += ['','## Features & Capabilities']
    if features:
        lines += [f"- {f}" for f in features]
    else:
        lines += ['- TBD']
    # Decisions
    lines += ['','## Summary of Main Decisions']
    if decision_bullets:
        lines += [f"- {b}" for b in decision_bullets]
    else:
        lines += ['- TBD']
    # Business context and legacy
    lines += [
        '',
        '## Business Context & Objectives',
        val('business_goals', 'TBD'),
        '',
        '## Legacy System Context',
        val('legacy_system', 'TBD'),
    ]
    # Constraints & Guardrails (explicit list)
    lines += [
        '',
        '## Constraints & Guardrails',
        f"- Architecture Style: {val('architecture_style', val('desired_architecture_style','TBD'))}",
        f"- Cloud Provider: {val('cloud_provider','TBD')}",
        f"- Region Strategy: {val('region_strategy','TBD')}",
        f"- Deployment Model: {val('deployment_model','TBD')}",
        f"- Database Family: {val('database_family','TBD')}",
        f"- Messaging Backbone: {val('messaging_backbone','TBD')}",
        f"- Identity Provider: {val('identity_provider','TBD')}",
        f"- API Style: {val('api_style','TBD')}",
        f"- Caching Strategy: {val('caching_strategy','TBD')}",
        f"- Observability Stack: {val('observability_stack','TBD')}",
        f"- Compliance & Data Residency: {val('compliance_standards','TBD')} / {val('data_residency','TBD')}",
    ]
    # Assumptions
    lines += ['', '## Assumptions', assumptions_md]
    # Stakeholders and scope
    lines += [
        '',
        '## Stakeholders',
        (val('stakeholders', '- Product\n- Engineering\n- Security\n- Compliance\n- SRE/Operations')),
        '',
        '## Scope',
        '### In Scope',
        val('in_scope', '- TBD'),
        '### Out of Scope',
        val('out_of_scope', '- TBD'),
    ]
    # NFRs
    lines += [
        '',
        '## Non-Functional Requirements',
        f"- Availability/SLA: {val('availability_sla','99.9% (assumed)')}",
        f"- RPO/RTO: {val('rpo_rto','RPO 15m, RTO 1h (assumed)')}",
        f"- Performance Targets: {val('performance_targets','p95 latency, throughput targets TBD')}",
        f"- Throughput: {val('throughput','TBD')}",
        f"- Peak Concurrency: {val('peak_concurrency','TBD')}",
        f"- Data Volume & Retention: {val('data_volume_retention','TBD')}",
        f"- Security Posture: {val('security_posture','Least privilege, encryption in transit/at rest')}",
    ]
    # Workloads
    lines += ['', '## Workloads & Traffic Profiles', val('workloads', f"~{node_count} components, ~{edge_count} integrations; peak/off-peak patterns TBD")]
    # Data & Integration
    lines += [
        '',
        '## Data Landscape',
        f"- Database Family: {val('database_family','TBD')}",
        f"- Notable Entities: {', '.join(terms) if terms else 'TBD'}",
    ]
    lines += [
        '',
        '## Integration Landscape',
        '*Components (sample):* ' + (', '.join(sample_nodes) if sample_nodes else 'TBD'),
        '*Edges (sample):*',
    ]
    lines += ([f"- {edge}" for edge in sample_edges] or ['- TBD'])
    # Platform choices
    lines += [
        '',
        '## Platform & Technology Choices',
        f"- Messaging: {val('messaging_backbone','TBD')}",
        f"- Identity: {val('identity_provider','TBD')}",
        f"- API Style: {val('api_style','TBD')}",
        f"- Caching: {val('caching_strategy','TBD')}",
        f"- Observability: {val('observability_stack','TBD')}",
    ]
    # Environments & release
    lines += [
        '',
        '## Environments & Release Strategy',
        f"- Environments: {val('environments','dev/test/stage/prod (assumed)')}",
        f"- Release Strategy: {val('release_strategy','Blue/Green or Canary (TBD)')}",
        '',
        '## Deployment Topology & Tenancy',
        f"- Topology: {val('deployment_topology','TBD')}\n- Tenancy Model: {val('tenancy_model','Single-tenant (assumed)')}",
    ]
    # Observability, cost/capacity, migration
    lines += [
        '',
        '## Observability Requirements',
        val('observability_requirements', '- Metrics, Logs, Traces; SLOs & alerting policy TBD'),
        '',
        '## Cost & Capacity',
        f"- Cost Constraints: {val('cost_constraints','TBD')}\n- Capacity Estimates: {val('capacity_estimates','TBD')}",
        '',
        '## Migration Strategy',
        val('migration_strategy', 'TBD'),
    ]
    # Risks & mitigations
    risk_lines = split_bullets(val('risks',''))
    mit_lines = split_bullets(val('risks_mitigation',''))
    lines += ['','## Risks & Mitigations']
    if risk_lines:
        lines += [f"- {r}" for r in risk_lines]
    else:
        lines += ['- Compliance, Operational complexity, Data migration risks']
    if mit_lines:
        lines += ['','### Mitigations'] + [f"- {m}" for m in mit_lines]
    # Open questions & refs
    lines += [
        '',
        '## Open Questions',
        val('open_questions', '- TBD'),
        '',
        '## References',
        f"PRD: {val('prd_file_name','n/a')} | Mapping: {val('system_mapping_name','n/a')}",
    ]

    # Add PRD excerpt for context
    if prd_text:
        excerpt = prd_text[:1200]
        lines += ['','---','### PRD Excerpt','```', excerpt,'```']

    return "\n".join(lines)
@app.route('/', methods=['GET', 'POST'])
def tabbed_workbench():
    active_tab = 0
    last_action = None
    last_elapsed_ms = None
    # On initial load without uploads, ensure nothing is pre-populated and Tab 1 is unlocked
    # Allow explicit reset via query param
    try:
        if request.method == 'GET' and request.args.get('reset'):
            session.clear()
            return redirect(url_for('tabbed_workbench'))
        # Treat common login/fresh params as a signal to reset
        if request.method == 'GET' and (request.args.get('login') or request.args.get('fresh') or request.args.get('new')):
            session.clear()
            return redirect(url_for('tabbed_workbench'))
        # Auto-reset stale sessions on landing the home page (e.g., after re-login)
        if request.method == 'GET':
            now = int(time.time())
            last = int(session.get('last_active_ts', 0) or 0)
            # If older than 1 hour, consider it stale and reset to avoid persisting old values
            if last and (now - last) > 3600:
                session.clear()
                return redirect(url_for('tabbed_workbench'))
    except Exception:
        pass
    _reset_form_if_no_uploads_on_get()
    if request.method == 'POST':
        # Robustly read action: if multiple 'action' fields exist (hidden + button), pick the last non-empty
        try:
            action_vals = request.form.getlist('action')
            action = ''
            if action_vals:
                non_empty = [v for v in action_vals if v]
                action = (non_empty[-1] if non_empty else action_vals[-1])
            else:
                action = 'save'
        except Exception:
            action = request.form.get('action', 'save')
        # Map deprecated "Generate ..." actions to Approve for a single-path UX
        if action in (
            'generate_initial_blueprint',
            'generate_arch_context',
            'generate_context',
            'generate_context_doc',
            'generate_context_agent_5_6'
        ):
            action = 'approve_and_next'
        # Client-provided start timestamp (ms) to compute elapsed time
        try:
            client_start_ts = int(request.form.get('client_start_ts', '0') or 0)
        except Exception:
            client_start_ts = 0
        # Persist regular form fields to session, skip ephemeral controls
        for key in request.form:
            if key == 'action' or key == 'active_tab' or key.startswith('chat_message_'):
                continue
            session[key] = request.form.get(key, '')
        # PRD upload
        prd_file = request.files.get('prd_file')
        if prd_file and prd_file.filename:
            uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
            os.makedirs(uploads_dir, exist_ok=True)
            safe_name = prd_file.filename.replace('..','_').replace('/','_').replace('\\','_')
            file_path = os.path.join(uploads_dir, safe_name)
            prd_file.save(file_path)
            session['prd_file_name'] = safe_name
            session['prd_text'] = read_text_safely(file_path)
            # Immediately regenerate initial diagrams from PRD (even without mapping)
            try:
                session['architecture_diagram'] = generate_blueprint_from_session(strict_prd_only=True)
                session['system_interaction_diagram'] = generate_system_interaction_diagram(strict_prd_only=True)
                session['data_model_diagram'] = generate_data_model_diagram(strict_prd_only=True)
                session['service_decomposition_diagram'] = generate_service_decomposition_diagram(strict_prd_only=True)
                session['decomposition_writeup'] = _build_decomposition_writeup_text(session.get('prd_text',''))
            except Exception:
                pass
        # Mapping upload
        sys_map = request.files.get('system_mapping_file')
        if sys_map and sys_map.filename:
            uploads_dir = os.path.join(os.path.dirname(__file__), 'uploads')
            os.makedirs(uploads_dir, exist_ok=True)
            safe_name = sys_map.filename.replace('..','_').replace('/','_').replace('\\','_')
            map_path = os.path.join(uploads_dir, safe_name)
            sys_map.save(map_path)
            session['system_mapping_name'] = safe_name
            session['system_mapping'] = json.dumps(parse_system_mapping(map_path))
            # Refresh diagrams to reflect mapping-aware blueprint enrichment
            try:
                session['architecture_diagram'] = generate_blueprint_from_session(strict_prd_only=True)
                session['system_interaction_diagram'] = generate_system_interaction_diagram(strict_prd_only=True)
                session['data_model_diagram'] = generate_data_model_diagram(strict_prd_only=True)
                session['service_decomposition_diagram'] = generate_service_decomposition_diagram(strict_prd_only=True)
                session['decomposition_writeup'] = _build_decomposition_writeup_text(session.get('prd_text',''))
            except Exception:
                pass
        # Lock Tab 1 when both present
        if session.get('prd_file_name') and session.get('system_mapping_name') and action != 'unlock_tab1':
            session['lock_tab1'] = '1'

        # Actions
        if action == 'extract_decisions':
            # Fill from sources, then fill any remaining fields with educated guesses
            _auto_populate_from_sources()
            _fill_educated_guesses()
            active_tab = 0
        elif action == 'lock_tab1':
            session['lock_tab1'] = '1'
            active_tab = 0
        elif action == 'unlock_tab1':
            session['lock_tab1'] = ''
            active_tab = 0
        elif action == 'generate_initial_blueprint':
            session['architecture_diagram'] = generate_blueprint_from_session(strict_prd_only=True)
            session['system_interaction_diagram'] = generate_system_interaction_diagram(strict_prd_only=True)
            session['data_model_diagram'] = generate_data_model_diagram(strict_prd_only=True)
            try:
                session['decomposition_writeup'] = _build_decomposition_writeup_text(session.get('prd_text',''))
            except Exception:
                pass
            active_tab = 1
        elif action == 'diagram_ai_arch':
            # AI enhance architecture (flowchart) diagram
            prompt_txt = request.form.get('diagram_prompt_arch', '')
            _push_diagram_history('architecture_diagram')
            _update_diagram_via_agent('architecture_diagram', 'flowchart', prompt_txt)
            try:
                active_tab = int(request.form.get('active_tab', '1') or 1)
            except Exception:
                active_tab = 1
        elif action == 'diagram_ai_system':
            # AI enhance system interaction (sequence) diagram
            prompt_txt = request.form.get('diagram_prompt_system', '')
            _push_diagram_history('system_interaction_diagram')
            _update_diagram_via_agent('system_interaction_diagram', 'sequence', prompt_txt)
            try:
                active_tab = int(request.form.get('active_tab', '1') or 1)
            except Exception:
                active_tab = 1
        elif action == 'diagram_ai_data':
            # AI enhance data model (erDiagram)
            prompt_txt = request.form.get('diagram_prompt_data', '')
            _push_diagram_history('data_model_diagram')
            _update_diagram_via_agent('data_model_diagram', 'er', prompt_txt)
            try:
                active_tab = int(request.form.get('active_tab', '1') or 1)
            except Exception:
                active_tab = 1
        elif action == 'diagram_ai_service_decomp':
            # AI enhance service decomposition (flowchart)
            prompt_txt = request.form.get('diagram_prompt_service_decomp', '')
            _push_diagram_history('service_decomposition_diagram')
            _update_diagram_via_agent('service_decomposition_diagram', 'flowchart', prompt_txt)
            try:
                active_tab = int(request.form.get('active_tab', '1') or 1)
            except Exception:
                active_tab = 1
        elif action == 'revert_arch_diagram':
            prev = _pop_diagram_history('architecture_diagram')
            if prev is not None:
                session['architecture_diagram'] = prev
            try:
                active_tab = int(request.form.get('active_tab', '1') or 1)
            except Exception:
                active_tab = 1
        elif action == 'revert_system_diagram':
            prev = _pop_diagram_history('system_interaction_diagram')
            if prev is not None:
                session['system_interaction_diagram'] = prev
            try:
                active_tab = int(request.form.get('active_tab', '1') or 1)
            except Exception:
                active_tab = 1
        elif action == 'revert_data_diagram':
            prev = _pop_diagram_history('data_model_diagram')
            if prev is not None:
                session['data_model_diagram'] = prev
            try:
                active_tab = int(request.form.get('active_tab', '1') or 1)
            except Exception:
                active_tab = 1
        elif action == 'revert_service_decomp_diagram':
            prev = _pop_diagram_history('service_decomposition_diagram')
            if prev is not None:
                session['service_decomposition_diagram'] = prev
            try:
                active_tab = int(request.form.get('active_tab', '1') or 1)
            except Exception:
                active_tab = 1
        elif action == 'generate_arch_context':
            # Generate Context Doc content (merged into Tab 0)
            try:
                mapping = {}
                try:
                    mapping = json.loads(session.get('system_mapping','{}')) if session.get('system_mapping') else {}
                except Exception:
                    mapping = {}
                payload_56 = json.dumps({
                    'selections': {
                        'style': session.get('architecture_style') or session.get('desired_architecture_style'),
                        'cloud': session.get('cloud_provider'),
                        'region': session.get('region_strategy'),
                        'deploy': session.get('deployment_model')
                    },
                    'decisions': session.get('architectural_decisions',''),
                    'mapping_stats': mapping.get('stats', {}),
                    'prd_excerpt': (session.get('prd_text','') or '')[:1500]
                }, ensure_ascii=False)
                reply_56 = call_agent('agent_5_6', payload_56)
                # Always present generated context (not raw agent text)
                session['architecture_context'] = generate_architecture_context()
                _normalize_context_doc(force=True)
            except Exception:
                session['architecture_context'] = generate_architecture_context()
                _normalize_context_doc(force=True)
            active_tab = 0
        elif action == 'download_share':
            # Generate Word document export of all tabs
            bio = _build_docx_from_session()
            if bio is None:
                resp = make_response('python-docx is not installed. Please install python-docx to enable export.')
                resp.status_code = 500
                return resp
            # Record elapsed if provided
            try:
                if client_start_ts and client_start_ts > 0:
                    now_ms = int(time.time() * 1000)
                    session['last_action'] = 'download_share'
                    session['last_elapsed_ms'] = max(0, now_ms - client_start_ts)
            except Exception:
                pass
            filename = f"architecture_workbench_{int(time.time())}.docx"
            return send_file(bio, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif action == 'chat_send':
            try:
                idx = int(request.form.get('active_tab', '0') or 0)
            except Exception:
                idx = 0
            msg = request.form.get(f'chat_message_{idx}', '') or request.form.get('chat_message','')
            # Capture diagram target if on Initial Blueprint (tab 1)
            if idx == 1:
                target = request.form.get('chat_diagram_target') or session.get('chat_target_tab1') or 'architecture_diagram'
                if target not in ('architecture_diagram','system_interaction_diagram','data_model_diagram','service_decomposition_diagram'):
                    target = 'architecture_diagram'
                session['chat_target_tab1'] = target
            if msg:
                _append_chat(idx, 'user', msg)
                try:
                    ctx = _get_tab_context_payload(idx)
                except Exception:
                    ctx = {}
                # Route tab 1 chat to Mermaid agent with explicit target
                if idx == 1:
                    agent_key = 'agent_mermaid'
                    # Provide explicit current diagram and target to increase structured responses
                    ctx = {
                        'target': session.get('chat_target_tab1','architecture_diagram'),
                        'current': session.get(session.get('chat_target_tab1','architecture_diagram'), ''),
                        'kind': 'sequence' if session.get('chat_target_tab1') == 'system_interaction_diagram' else ('er' if session.get('chat_target_tab1') == 'data_model_diagram' else 'flowchart')
                    }
                # Route System Interaction / DB / Service tabs to Mermaid agent with dedicated target/kind
                elif idx in (2, 5, 6):
                    agent_key = 'agent_mermaid'
                    if idx == 2:
                        tgt, kind = 'system_interaction_diagram', 'sequence'
                    elif idx == 5:
                        tgt, kind = 'data_model_diagram', 'er'
                    else:
                        tgt, kind = 'service_decomposition_diagram', 'flowchart'
                    ctx = {
                        'target': tgt,
                        'current': session.get(tgt, ''),
                        'kind': kind
                    }
                else:
                    agent_key = AGENT_BY_TAB.get(idx, 'agent_5_1')
                try:
                    payload = json.dumps({'tab': idx, 'context': ctx, 'message': msg}, ensure_ascii=False)
                except Exception:
                    payload = (str(ctx) + "\n\nUser: " + msg)
                reply = call_agent(agent_key, payload)
                # Store reply and any parsed updates for explicit Apply
                session[f'last_agent_reply_tab{idx}'] = reply
                parsed_updates = _parse_updates_from_text(reply)
                # If on tab 1 and no JSON updates, but reply looks like Mermaid, wrap it as a set to the selected target
                if idx == 1 and not parsed_updates:
                    try:
                        first = (reply or '').strip().splitlines()[0].strip().lower() if reply else ''
                    except Exception:
                        first = ''
                    if first.startswith(('graph','flowchart','sequencediagram','erdiagram','classdiagram','statediagram','gantt','pie','journey','mindmap','timeline','gitgraph')):
                        tgt = session.get('chat_target_tab1','architecture_diagram')
                        parsed_updates = {'set': {tgt: reply.strip()}}
                # If on tabs 2/5/6 and no JSON updates, but reply looks like Mermaid, wrap it with the tab's target
                if idx in (2, 5, 6) and not parsed_updates:
                    try:
                        first = (reply or '').strip().splitlines()[0].strip().lower() if reply else ''
                    except Exception:
                        first = ''
                    if first.startswith(('graph','flowchart','sequencediagram','erdiagram','classdiagram','statediagram','gantt','pie','journey','mindmap','timeline','gitgraph')):
                        tgt = 'system_interaction_diagram' if idx == 2 else ('data_model_diagram' if idx == 5 else 'service_decomposition_diagram')
                        parsed_updates = {'set': {tgt: reply.strip()}}
                # Auto-apply on Tab 1 to behave like Mermaid AI Use
                if idx == 1:
                    tgt = session.get('chat_target_tab1','architecture_diagram')
                    kind = 'sequence' if tgt == 'system_interaction_diagram' else ('er' if tgt == 'data_model_diagram' else 'flowchart')
                    applied = False
                    # Case 1: JSON updates targeting diagrams
                    if parsed_updates and isinstance(parsed_updates.get('set'), dict):
                        sets = parsed_updates['set']
                        # If specific diagram fields present, apply them
                        diagram_keys = {'architecture_diagram','system_interaction_diagram','data_model_diagram','service_decomposition_diagram'}
                        targets_present = [k for k in sets.keys() if k in diagram_keys]
                        if targets_present:
                            for dk in targets_present:
                                _push_diagram_history(dk)
                                session[dk] = str(sets[dk])
                            applied = True
                        else:
                            # Single generic key: assume it's the Mermaid for the current target
                            if len(sets) == 1:
                                only_key = next(iter(sets.keys()))
                                _push_diagram_history(tgt)
                                session[tgt] = str(sets[only_key])
                                applied = True
                    # Case 2: Reply contained Mermaid (wrapped earlier): apply to target
                    if not applied and parsed_updates and isinstance(parsed_updates.get('set'), dict) and tgt in parsed_updates['set']:
                        _push_diagram_history(tgt)
                        session[tgt] = str(parsed_updates['set'][tgt])
                        applied = True
                    # Case 3: No structured updates; fall back to heuristic application using the user's message
                    if not applied:
                        _push_diagram_history(tgt)
                        _apply_prompt_to_mermaid(tgt, kind, msg)
                        applied = True
                    # Clear pending since we applied immediately
                    session[f'pending_updates_tab{idx}'] = ''
                # Auto-apply on Tabs 2, 5 and 6 similar to tab 1
                elif idx in (2, 5, 6):
                    tgt = 'system_interaction_diagram' if idx == 2 else ('data_model_diagram' if idx == 5 else 'service_decomposition_diagram')
                    kind = 'sequence' if idx == 2 else ('er' if idx == 5 else 'flowchart')
                    applied = False
                    if parsed_updates and isinstance(parsed_updates.get('set'), dict):
                        sets = parsed_updates['set']
                        diagram_keys = {'architecture_diagram','system_interaction_diagram','data_model_diagram','service_decomposition_diagram'}
                        targets_present = [k for k in sets.keys() if k in diagram_keys]
                        if targets_present:
                            for dk in targets_present:
                                _push_diagram_history(dk)
                                session[dk] = str(sets[dk])
                            applied = True
                        else:
                            if len(sets) == 1:
                                only_key = next(iter(sets.keys()))
                                _push_diagram_history(tgt)
                                session[tgt] = str(sets[only_key])
                                applied = True
                    if not applied and parsed_updates and isinstance(parsed_updates.get('set'), dict) and tgt in parsed_updates['set']:
                        _push_diagram_history(tgt)
                        session[tgt] = str(parsed_updates['set'][tgt])
                        applied = True
                    if not applied:
                        _push_diagram_history(tgt)
                        _apply_prompt_to_mermaid(tgt, kind, msg)
                        applied = True
                    session[f'pending_updates_tab{idx}'] = ''
                else:
                    # For non-diagram tabs, require explicit Apply
                    session[f'pending_updates_tab{idx}'] = json.dumps(parsed_updates) if parsed_updates else ''
                _append_chat(idx, 'agent', reply)
            active_tab = idx
        elif action == 'chat_apply':
            # Apply parsed updates from last agent reply for this tab
            try:
                idx = int(request.form.get('active_tab', '0') or 0)
            except Exception:
                idx = 0
            updates_raw = session.get(f'pending_updates_tab{idx}', '')
            updates = None
            if updates_raw:
                try:
                    updates = json.loads(updates_raw)
                except Exception:
                    updates = None
            if not updates:
                # Try parsing from stored reply if available
                reply = session.get(f'last_agent_reply_tab{idx}', '')
                if reply:
                    updates = _parse_updates_from_text(reply)
            if not updates:
                # Heuristic: infer from last user message chunk
                chat_hist = session.get(f'chat_tab{idx}', '')
                last_user = ''
                for part in reversed(chat_hist.split('\n\n')):
                    if part.strip().startswith('User:'):
                        last_user = part.split('User:',1)[-1].strip()
                        break
                if last_user:
                    updates = _heuristic_updates_from_user(idx, last_user)
            if updates:
                # For diagram-centric tabs (1,2,5,6), if updates don't target diagram fields, coerce to the tab's target
                if idx in (1, 2, 5, 6) and isinstance(updates, dict) and not (updates.get('set') and any(k in updates['set'] for k in ('architecture_diagram','system_interaction_diagram','data_model_diagram','service_decomposition_diagram'))):
                    tgt = (
                        session.get('chat_target_tab1','architecture_diagram') if idx == 1 else (
                            'system_interaction_diagram' if idx == 2 else (
                                'data_model_diagram' if idx == 5 else 'service_decomposition_diagram'
                            )
                        )
                    )
                    if 'set' in updates and isinstance(updates['set'], dict) and len(updates['set']) == 1:
                        only_key = next(iter(updates['set'].keys()))
                        val = updates['set'][only_key]
                        updates = {'set': {tgt: val}}
                    elif isinstance(updates.get('set'), dict):
                        pass
                    else:
                        if isinstance(updates, str):
                            updates = {'set': {tgt: updates}}
                # Push diagram history if diagrams are being set
                try:
                    s = updates.get('set') or {}
                    if 'architecture_diagram' in s: _push_diagram_history('architecture_diagram')
                    if 'system_interaction_diagram' in s: _push_diagram_history('system_interaction_diagram')
                    if 'data_model_diagram' in s: _push_diagram_history('data_model_diagram')
                    if 'service_decomposition_diagram' in s: _push_diagram_history('service_decomposition_diagram')
                except Exception:
                    pass
                _apply_agent_updates(idx, updates)
                # After applying updates, refresh compiled document to keep Tab 7 human-readable
                try:
                    session['compiled_document'] = _build_compiled_document_markdown()
                except Exception:
                    pass
                # Also keep Context Doc normalized if it was affected
                try:
                    _normalize_context_doc()
                except Exception:
                    pass
                session[f'pending_updates_tab{idx}'] = ''
            active_tab = idx
        elif action == 'chat_clear':
            try:
                idx = int(request.form.get('active_tab', '0') or 0)
            except Exception:
                idx = 0
            session[f'chat_tab{idx}'] = ''
            session[f'last_agent_reply_tab{idx}'] = ''
            session[f'pending_updates_tab{idx}'] = ''
            active_tab = idx
        elif action == 'approve_and_next':
            try:
                idx = int(request.form.get('active_tab', '0') or 0)
            except Exception:
                idx = 0
            session[f'approved_tab{idx}'] = '1'
            # Advance along the defined navigation order
            next_idx = idx
            try:
                pos = TAB_NAV_ORDER.index(idx)
                if pos < len(TAB_NAV_ORDER) - 1:
                    next_idx = TAB_NAV_ORDER[pos + 1]
            except ValueError:
                # If idx not in order, default to first
                next_idx = TAB_NAV_ORDER[0]
            active_tab = next_idx
        elif action == 'reset_session':
            # Clear entire session for a fresh start
            session.clear()
            return redirect(url_for('tabbed_workbench'))
        else:
            active_tab = 0

        # Compute elapsed time for the last action
        try:
            if client_start_ts and client_start_ts > 0:
                now_ms = int(time.time() * 1000)
                last_action = action
                last_elapsed_ms = max(0, now_ms - client_start_ts)
        except Exception:
            pass

    # Auto-populate when PRD or Mapping available (derive as much as possible)
    if session.get('prd_text') or session.get('system_mapping'):
        _auto_populate_from_sources()
        _fill_educated_guesses()

    # Mark activity timestamp
    try:
        session['last_active_ts'] = int(time.time())
    except Exception:
        pass

    # Refresh write-up one last time before rendering
    try:
        session['decomposition_writeup'] = _build_decomposition_writeup_text(session.get('prd_text',''))
    except Exception:
        pass

    rendered = render_template(
    'tabbed_architecture_workbench.html',
    active_tab=active_tab,
    last_action=last_action if last_action else session.get('last_action'),
    last_elapsed_ms=last_elapsed_ms if last_elapsed_ms is not None else session.get('last_elapsed_ms'),
    # chat transcripts per tab (index mapping aligns with showTab indices)
    chat_tab0=session.get('chat_tab0',''),
    chat_tab1=session.get('chat_tab1',''),
    chat_tab2=session.get('chat_tab2',''),
    chat_tab3=session.get('chat_tab3',''),
    chat_tab4=session.get('chat_tab4',''),
    chat_tab5=session.get('chat_tab5',''),
    chat_tab6=session.get('chat_tab6',''),
    chat_tab7=session.get('chat_tab7',''),
    chat_tab8=session.get('chat_tab8',''),
    # approved flags per tab index
    approved_tab0=session.get('approved_tab0',''),
    approved_tab1=session.get('approved_tab1',''),
    approved_tab2=session.get('approved_tab2',''),
    approved_tab3=session.get('approved_tab3',''),
    approved_tab4=session.get('approved_tab4',''),
    approved_tab5=session.get('approved_tab5',''),
    approved_tab6=session.get('approved_tab6',''),
    approved_tab7=session.get('approved_tab7',''),
    approved_tab8=session.get('approved_tab8',''),
    # pending apply flags per tab index
    pending_tab0=bool(session.get('pending_updates_tab0')),
    pending_tab1=bool(session.get('pending_updates_tab1')),
    pending_tab2=bool(session.get('pending_updates_tab2')),
    pending_tab3=bool(session.get('pending_updates_tab3')),
    pending_tab4=bool(session.get('pending_updates_tab4')),
    pending_tab5=bool(session.get('pending_updates_tab5')),
    pending_tab6=bool(session.get('pending_updates_tab6')),
    pending_tab7=bool(session.get('pending_updates_tab7')),
    pending_tab8=bool(session.get('pending_updates_tab8')),
        business_goals=session.get('business_goals',''),
        legacy_system=session.get('legacy_system',''),
        constraints=session.get('constraints',''),
        architecture_style=session.get('architecture_style',''),
        desired_architecture_style=session.get('desired_architecture_style',''),
        cloud_provider=session.get('cloud_provider',''),
        region_strategy=session.get('region_strategy',''),
        deployment_model=session.get('deployment_model',''),
        database_family=session.get('database_family',''),
        messaging_backbone=session.get('messaging_backbone',''),
        identity_provider=session.get('identity_provider',''),
        api_style=session.get('api_style',''),
        caching_strategy=session.get('caching_strategy',''),
        observability_stack=session.get('observability_stack',''),
        data_residency=session.get('data_residency',''),
        compliance_standards=session.get('compliance_standards',''),
        additional_requirements=session.get('additional_requirements',''),
    prd_file_name=session.get('prd_file_name',''),
    prd_text=session.get('prd_text',''),
    prd_len=len(session.get('prd_text','') or ''),
        system_mapping_name=session.get('system_mapping_name',''),
        system_mapping_stats=json.loads(session.get('system_mapping','{}')).get('stats') if session.get('system_mapping') else None,
        high_level_decision_points=session.get('high_level_decision_points',''),
        one_way_door_decisions=session.get('one_way_door_decisions',''),
        other_relevant_questions=session.get('other_relevant_questions',''),
        blueprint=session.get('architecture_diagram', ''),
        major_components=session.get('major_components',''),
        interface_definitions=session.get('interface_definitions',''),
        data_schemas=session.get('data_schemas',''),
        reusable_patterns=session.get('reusable_patterns',''),
    system_interaction_diagram=session.get('system_interaction_diagram',''),
    data_model_diagram=session.get('data_model_diagram',''),
    service_decomposition_diagram=session.get('service_decomposition_diagram',''),
        decisions=session.get('architectural_decisions', ''),
        rationale_tradeoffs=session.get('rationale_tradeoffs',''),
        blueprint_references=session.get('blueprint_references',''),
        communication=session.get('communication_protocols', ''),
        data_flow_diagrams=session.get('data_flow_diagrams',''),
        integration_points=session.get('integration_points',''),
        pros_cons=session.get('swot_analysis', ''),
        risks_mitigation=session.get('risks_mitigation',''),
    doc=session.get('compiled_document', ''),
    stakeholder_checklist=session.get('stakeholder_checklist',''),
    architecture_context=session.get('architecture_context',''),
    decomposition_writeup=session.get('decomposition_writeup',''),
    assumptions=session.get('assumptions',''),
    stakeholders=session.get('stakeholders',''),
    in_scope=session.get('in_scope',''),
    out_of_scope=session.get('out_of_scope',''),
    availability_sla=session.get('availability_sla',''),
    rpo_rto=session.get('rpo_rto',''),
    performance_targets=session.get('performance_targets',''),
    security_posture=session.get('security_posture',''),
    data_volume_retention=session.get('data_volume_retention',''),
    workloads=session.get('workloads',''),
    environments=session.get('environments',''),
    release_strategy=session.get('release_strategy',''),
    deployment_topology=session.get('deployment_topology',''),
    tenancy_model=session.get('tenancy_model',''),
    observability_requirements=session.get('observability_requirements',''),
    cost_constraints=session.get('cost_constraints',''),
    capacity_estimates=session.get('capacity_estimates',''),
    migration_strategy=session.get('migration_strategy',''),
    risks=session.get('risks',''),
    open_questions=session.get('open_questions',''),
    throughput=session.get('throughput',''),
    peak_concurrency=session.get('peak_concurrency',''),
    dr_strategy=session.get('dr_strategy',''),
    encryption_at_rest=session.get('encryption_at_rest',''),
    encryption_in_transit=session.get('encryption_in_transit',''),
    kms=session.get('kms',''),
    api_gateway_type=session.get('api_gateway_type',''),
    service_mesh=session.get('service_mesh',''),
    edge_controls=session.get('edge_controls',''),
    secrets_manager=session.get('secrets_manager',''),
    iac_tool=session.get('iac_tool',''),
    ci_cd=session.get('ci_cd',''),
        lock_tab1=session.get('lock_tab1',''),
        chat_target_tab1=session.get('chat_target_tab1','architecture_diagram')
    )
    # Force no-store to avoid browser showing cached state when returning to home page
    try:
        resp = make_response(rendered)
        resp.headers['Cache-Control'] = 'no-store, no-cache, must-revalidate, max-age=0'
        resp.headers['Pragma'] = 'no-cache'
        return resp
    except Exception:
        return rendered


# System Mapping UI and chat endpoints (POC5)
@app.route('/system-mapping', methods=['GET'])
def system_mapping_page_poc5():
    # Reuse shared template at root-level templates
    try:
        return render_template('system_mapping.html')
    except Exception:
        # Fallback simple page if template missing
        return make_response('<h2>System Mapping UI not found</h2><p>Ensure templates/system_mapping.html exists.</p>', 404)


@app.route('/system-mapping-chat', methods=['POST'])
def system_mapping_chat_poc5():
    # Lightweight, offline-friendly response structured for the UI
    try:
        data = request.get_json(silent=True) or {}
        msg = (data.get('message') or '').strip()
        current = data.get('current_systems') or []
        available = data.get('available_systems') or []
    except Exception:
        msg, current, available = '', [], []

    # Basic heuristics to craft suggestions without external calls
    suggestions = []
    warnings = []
    lower = msg.lower() if msg else ''
    def add_sugg(name, reason):
        if name not in current:
            suggestions.append({'system': name, 'reason': reason})
    # Recommend some typical components if absent
    if not current:
        add_sugg('API Gateway', 'Fronts services, routing, throttling, auth.')
        add_sugg('Identity Provider', 'Authentication/authorization (OIDC/OAuth2).')
        add_sugg('Message Bus', 'Async pub/sub for decoupling (e.g., Kafka).')
    # Parse intent
    if 'payment' in lower:
        add_sugg('Payment Gateway', 'External payments processing integration.')
    if 'search' in lower:
        add_sugg('Search Service', 'Full-text/product search capability.')
    if 'analytics' in lower or 'metrics' in lower:
        add_sugg('Analytics Platform', 'Centralized analytics/BI.')
    if 'cache' in lower or 'redis' in lower:
        add_sugg('Cache', 'Low-latency caching layer.')
    if 'warehouse' in lower or 'lake' in lower:
        add_sugg('Data Warehouse', 'Batch analytics, reporting.')
    # Guardrails
    if len(current) > 60:
        warnings.append('Large mapping detected; consider grouping systems into domains.')

    text = {
        'message': 'Here are some recommendations based on your context. You can Apply All or add systems individually.',
        'suggestions': suggestions,
        'warnings': warnings,
    }
    return jsonify({'success': True, **text})


@app.route('/epic-results', methods=['GET'])
def epic_results_alias_poc5():
    # Friendly alias for Back button on mapping UI; return the main workbench
    return redirect(url_for('tabbed_workbench'))


# Simple combined page for Step 1 (Context & Requirements) + Step 2 (Initial Blueprint)
@app.route('/combined-12', methods=['GET', 'POST'])
def combined_step1_step2():
    """Lightweight page that shows inputs for Context & Requirements and, upon submit,
    immediately renders the Initial Blueprint on the same page.

    This leaves the main tabbed workbench ("/") unchanged; use this as a simple flow.
    """
    # Persist submitted fields
    if request.method == 'POST':
        for key in (
            'business_goals', 'legacy_system', 'constraints',
            'architecture_style', 'desired_architecture_style',
        ):
            session[key] = request.form.get(key, '')
        # Ensure dropdown-style defaults so diagrams have sensible values even without PRD
        try:
            _ensure_default_dropdowns()
        except Exception:
            pass
        # Generate blueprint; if PRD not provided, allow non-strict fallbacks to populate services
        try:
            strict = True if session.get('prd_text') else False
            session['architecture_diagram'] = generate_blueprint_from_session(strict_prd_only=strict)
        except Exception:
            # Keep any prior value if generation fails
            pass

    return render_template(
        'page12_combined.html',
        business_goals=session.get('business_goals', ''),
        legacy_system=session.get('legacy_system', ''),
        constraints=session.get('constraints', ''),
        architecture_style=session.get('architecture_style', ''),
        desired_architecture_style=session.get('desired_architecture_style', ''),
        blueprint=session.get('architecture_diagram', ''),
    )


if __name__ == "__main__":
    app.run(port=6001, debug=True)
